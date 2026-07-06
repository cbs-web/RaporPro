import datetime
import os
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


APP_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_TEMPLATE_PATH = os.path.join(APP_DIR, "Tutanak Örnek.docx")


def _clean(value, fallback=""):
    text = "" if value is None else str(value).strip()
    return text if text else fallback


def _safe_name(value, fallback="Proje"):
    text = _clean(value, fallback)
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return "_".join(text.split()) or fallback


def _today():
    return datetime.datetime.now().strftime("%d.%m.%Y")


def _template_path(veri=None):
    ayarlar = (veri or {}).get("ayarlar", {})
    custom = _clean(ayarlar.get("tutanak_sablon_path"))
    if custom and os.path.exists(custom):
        return custom
    if os.path.exists(DEFAULT_TEMPLATE_PATH):
        return DEFAULT_TEMPLATE_PATH
    candidates = list(Path(APP_DIR).glob("Tutanak*.docx"))
    if candidates:
        return str(candidates[0])
    return DEFAULT_TEMPLATE_PATH


def _project_name(veri):
    kunye = (veri or {}).get("kunye", {})
    mah = _clean(kunye.get("mah"))
    ada = _clean(kunye.get("ada"))
    parsel = _clean(kunye.get("par"))
    if mah and ada and parsel:
        return f"{mah} {ada} Ada {parsel} Parsel"
    if ada and parsel:
        return f"{ada} Ada {parsel} Parsel"
    return _clean(kunye.get("sahibi"), "Proje")


def _fmt_coord(value):
    text = _clean(value, "-")
    if text == "-" or text.endswith("°"):
        return text
    return f"{text}°"


def _fmt_date(value):
    text = _clean(value)
    return text or _today()


def _set_cell(cell, value):
    if len(cell.paragraphs) > 1:
        for paragraph in list(cell.paragraphs)[1:]:
            paragraph._element.getparent().remove(paragraph._element)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(_clean(value, "-"))


def _set_row_value(table, row_idx, value, col_start=2):
    row = table.rows[row_idx]
    for col in range(col_start, len(row.cells)):
        _set_cell(row.cells[col], value)


def _clear_body(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _append_table_clone(doc, source_table):
    tbl = deepcopy(source_table._tbl)
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(tbl)
    else:
        body.append(tbl)
    return Table(tbl, doc)


def _add_page_break(doc):
    paragraph = doc.add_paragraph()
    _format_tutanak_paragraph(paragraph)
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _format_tutanak_paragraph(paragraph, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1
    fmt.keep_together = True
    fmt.keep_with_next = bool(keep_next)


def _format_tutanak_table(table, keep_with_next=False):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _format_tutanak_paragraph(paragraph, keep_next=keep_with_next)


def _compact_output_page_layout(doc):
    for section in doc.sections:
        if section.top_margin is None or section.top_margin.cm > 0.8:
            section.top_margin = Cm(0.75)
        if section.bottom_margin is None or section.bottom_margin.cm > 0.8:
            section.bottom_margin = Cm(0.75)
        if section.left_margin is None or section.left_margin.cm > 1.8:
            section.left_margin = Cm(1.6)
        if section.right_margin is None or section.right_margin.cm > 1.8:
            section.right_margin = Cm(1.6)


def _count_nonempty_rows(rows):
    count = 0
    for row in rows or []:
        values = row.values() if isinstance(row, dict) else row
        if any(_clean(value) for value in values):
            count += 1
    return count


def _safe_float(value):
    try:
        return float(str(value).replace(",", ".").strip())
    except Exception:
        return 0.0


def _orselenmis_numune_sayisi(sondaj):
    derinlik = _safe_float((sondaj or {}).get("der"))
    if derinlik <= 0:
        return 0
    return int((derinlik + 1e-6) // 1.5)


def _ud_numune_sayisi(sondaj):
    count = 0
    for row in (sondaj or {}).get("numuneler", []) or []:
        if isinstance(row, dict):
            values = row.values()
        else:
            values = row
        text = " ".join(_clean(value) for value in values).upper()
        tokens = text.replace("-", " ").replace("_", " ").replace("/", " ").split()
        if any(token == "UD" or (token.startswith("UD") and token[2:].isdigit()) for token in tokens):
            count += 1
    return count


def _sondaj_turu(veri, sondaj):
    ayarlar = (veri or {}).get("ayarlar", {})
    text = _clean(ayarlar.get("sondaj_turu") or (sondaj or {}).get("sondaj_turu")).lower()
    if text in ("kaya", "rock"):
        return "Kaya"
    if text in ("zemin", "soil"):
        return "Zemin"
    return "Kaya" if (sondaj or {}).get("kaya") else "Zemin"


def _delgi_capi(veri, sondaj):
    ayarlar = (veri or {}).get("ayarlar", {})
    text = _clean(ayarlar.get("delgi_capi") or (sondaj or {}).get("delgi_capi"), "76mm").replace(" ", "")
    if text.lower() in ("76", "76mm"):
        return "76mm"
    if text.lower() in ("89", "89mm"):
        return "89mm"
    return "76mm"


def _yass_text(sondaj):
    for key in ("yass_d2", "yass_d1"):
        value = _clean(sondaj.get(key))
        if value:
            return value
    return "-"


def _statement_date(sondaj):
    return _fmt_date(sondaj.get("bit_tar") or sondaj.get("bas_tar"))


def _sondaj_statement(veri, sondaj):
    ayarlar = (veri or {}).get("ayarlar", {})
    firma = _clean(ayarlar.get("tutanak_sondaj_firma"), "Kale Detay Sondaj")
    tarih = _statement_date(sondaj)
    return (
        f"Yukarıda belirtilen sondaj kuyusu {firma} tarafından {tarih} tarihinde açılarak gerekli "
        "tespit ve deneyler yapılmış olup, iş bu tutanak 1 nüsha olarak tanzim ve imza edilmiştir."
    )


def _fill_sondaj_table(table, veri, sondaj):
    ayarlar = (veri or {}).get("ayarlar", {})
    _set_row_value(table, 1, _project_name(veri), col_start=1)
    _set_row_value(table, 2, _clean(sondaj.get("no"), "SK-"), col_start=2)
    _set_row_value(table, 3, _clean(sondaj.get("k"), "-"), col_start=2)
    _set_row_value(table, 4, _sondaj_turu(veri, sondaj), col_start=2)
    _set_row_value(table, 5, _clean(ayarlar.get("tutanak_uygulama_sekli"), "Burgusuz/Sulu"), col_start=2)
    _set_row_value(table, 6, _clean(ayarlar.get("tutanak_sondaj_makinesi"), "SMK-500"), col_start=2)
    _set_row_value(table, 7, _clean(sondaj.get("bas_tar"), "-"), col_start=2)
    _set_row_value(table, 8, _clean(sondaj.get("bit_tar"), "-"), col_start=2)
    _set_row_value(table, 9, _clean(sondaj.get("der"), "-"), col_start=2)
    _set_cell(table.rows[10].cells[2], _fmt_coord(sondaj.get("y")))
    _set_cell(table.rows[10].cells[3], _fmt_coord(sondaj.get("x")))
    _set_row_value(table, 11, _delgi_capi(veri, sondaj), col_start=2)
    _set_row_value(table, 12, str(_orselenmis_numune_sayisi(sondaj)), col_start=2)
    _set_row_value(table, 13, str(_ud_numune_sayisi(sondaj)), col_start=2)
    _set_row_value(table, 14, str(_count_nonempty_rows(sondaj.get("spt", []))), col_start=2)
    _set_row_value(table, 15, str(_count_nonempty_rows(sondaj.get("pmt", []))), col_start=2)
    _set_row_value(table, 16, "-", col_start=2)
    _set_row_value(table, 17, "-", col_start=2)
    _set_row_value(table, 18, "-", col_start=2)
    _set_row_value(table, 19, _yass_text(sondaj), col_start=2)


def _image_target_cell(table):
    for row in table.rows:
        for cell in row.cells:
            if list(cell._tc.iter(qn("w:drawing"))) or list(cell._tc.iter(qn("w:pict"))):
                return cell
    if len(table.rows) >= 1 and len(table.rows[0].cells) >= 6:
        return table.rows[0].cells[5]
    return table.rows[0].cells[0]


def _clear_cell_content(cell):
    try:
        cell._tc.clear_content()
        return
    except Exception:
        pass
    tc_pr = cell._tc.tcPr
    for child in list(cell._tc):
        if child is not tc_pr:
            cell._tc.remove(child)


def _replace_location_image(table, image_path):
    if not image_path or not os.path.exists(image_path):
        return
    cell = _image_target_cell(table)
    _clear_cell_content(cell)
    paragraph = cell.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(11.5))


def _fill_jeofizik_table(table, veri, ss, idx):
    ayarlar = (veri or {}).get("ayarlar", {})
    jeofizik = (veri or {}).get("jeofizik", {})
    firma = _clean(ayarlar.get("firma_adi"), "UB ZEMİN MÜHENDİSLİK")
    tarih = _fmt_date(jeofizik.get("tarih"))
    title = f"JF – SİSMİK ÇALIŞMALAR KABUL TUTANAĞI (Serim {idx})"
    _set_row_value(table, 0, title, col_start=0)
    _set_row_value(table, 1, firma, col_start=1)
    _set_row_value(table, 2, _project_name(veri), col_start=1)
    _set_row_value(table, 3, tarih, col_start=1)
    coords = list(ss.get("coords", []) or [])
    while len(coords) < 6:
        coords.append("")
    _set_cell(table.rows[5].cells[2], _fmt_coord(coords[0]))
    _set_cell(table.rows[5].cells[3], _fmt_coord(coords[0]))
    _set_cell(table.rows[5].cells[4], _fmt_coord(coords[4] or coords[2]))
    _set_cell(table.rows[5].cells[5], _fmt_coord(coords[4] or coords[2]))
    _set_cell(table.rows[6].cells[2], _fmt_coord(coords[1]))
    _set_cell(table.rows[6].cells[3], _fmt_coord(coords[1]))
    _set_cell(table.rows[6].cells[4], _fmt_coord(coords[5] or coords[3]))
    _set_cell(table.rows[6].cells[5], _fmt_coord(coords[5] or coords[3]))
    defaults = {
        8: _clean(ayarlar.get("tutanak_jeofizik_cihaz"), "GEODE"),
        9: _clean(ayarlar.get("tutanak_jeofon"), "3,0m - 4,5 Hz"),
        10: _clean(ayarlar.get("tutanak_offset"), "3,0m"),
        11: _clean(ayarlar.get("tutanak_kanal_sayisi"), "12"),
        14: _clean(ayarlar.get("tutanak_kaynak"), "Balyoz"),
    }
    for row_idx, value in defaults.items():
        _set_row_value(table, row_idx, value, col_start=1)
    statement = (
        f"Yukarıda belirtilen jeofizik ölçüm {firma} tarafından {tarih} tarihinde yapılmış olup, "
        f"iş bu tutanak 1 nüsha olarak tanzim ve imza edilmiştir. {tarih}"
    )
    _set_row_value(table, 15, statement, col_start=0)


def _export_docx_to_pdf(docx_path, pdf_path):
    try:
        import pythoncom
        import win32com.client
    except Exception as exc:
        raise RuntimeError(f"Tutanak PDF aktarımı için pywin32 bulunamadı: {exc}") from exc
    app = None
    doc = None
    pythoncom.CoInitialize()
    try:
        app = win32com.client.DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        doc = app.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    finally:
        if doc is not None:
            doc.Close(False)
        if app is not None:
            app.Quit()
        pythoncom.CoUninitialize()


def tutanak_dosya_adi(veri, ext=".docx"):
    return f"{_safe_name(_project_name(veri), 'Proje')}_Tutanaklar{ext}"


def tutanaklari_olustur(veri, output_path, lokasyon_haritasi=None):
    template_path = _template_path(veri)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Tutanak şablonu bulunamadı: {template_path}")
    template = Document(template_path)
    if len(template.tables) < 4:
        raise ValueError("Tutanak şablonunda beklenen tablolar bulunamadı.")

    ext = os.path.splitext(output_path)[1].lower()
    if ext not in (".docx", ".pdf"):
        output_path = f"{output_path}.docx"
        ext = ".docx"

    out_doc = Document(template_path)
    _clear_body(out_doc)
    _compact_output_page_layout(out_doc)
    sondajlar = list((veri or {}).get("sondaj", []) or [])
    ss_list = list(((veri or {}).get("jeofizik", {}) or {}).get("ss_list", []) or [])

    first_block = True
    for sondaj in sondajlar:
        if not first_block:
            _add_page_break(out_doc)
        first_block = False
        sondaj_table = _append_table_clone(out_doc, template.tables[0])
        _fill_sondaj_table(sondaj_table, veri, sondaj)
        _format_tutanak_table(sondaj_table, keep_with_next=True)
        statement = out_doc.add_paragraph(_sondaj_statement(veri, sondaj))
        _format_tutanak_paragraph(statement, keep_next=True)
        location_table = _append_table_clone(out_doc, template.tables[1])
        _replace_location_image(location_table, lokasyon_haritasi)
        _format_tutanak_table(location_table, keep_with_next=True)
        signature_table = _append_table_clone(out_doc, template.tables[2])
        _format_tutanak_table(signature_table, keep_with_next=False)

    for idx, ss in enumerate(ss_list, start=1):
        if not first_block:
            _add_page_break(out_doc)
        first_block = False
        jeofizik_table = _append_table_clone(out_doc, template.tables[3])
        _fill_jeofizik_table(jeofizik_table, veri, ss, idx)
        _format_tutanak_table(jeofizik_table, keep_with_next=False)

    if first_block:
        out_doc.add_paragraph("Tutanak oluşturulacak sondaj veya jeofizik verisi bulunamadı.")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    if ext == ".docx":
        out_doc.save(output_path)
        return {"path": output_path, "sondaj_count": len(sondajlar), "jeofizik_count": len(ss_list)}

    with tempfile.TemporaryDirectory(prefix="raporpro_tutanak_") as tmp_dir:
        docx_path = os.path.join(tmp_dir, tutanak_dosya_adi(veri, ".docx"))
        out_doc.save(docx_path)
        _export_docx_to_pdf(docx_path, output_path)
    return {"path": output_path, "sondaj_count": len(sondajlar), "jeofizik_count": len(ss_list)}

import os
import re
import shutil
import tempfile
from copy import deepcopy
from io import BytesIO

import fitz
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


APP_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_EK_NORMAL = os.path.join(APP_DIR, "sablonlar", "ekler", "EK-Yeni-Sondaj-Tutanakli.docx")
DEFAULT_EK_ARAZI_DENEYLI = os.path.join(APP_DIR, "sablonlar", "ekler", "EK-Yeni-Sondaj-Tutanakli-Arazi-Deneyli.docx")

EK_SET_NORMAL = "normal"
EK_SET_ARAZI_DENEYLI = "arazi_deneyli"
EK_SET_LABELS = {
    EK_SET_NORMAL: "Normal Ekler",
    EK_SET_ARAZI_DENEYLI: "Arazi Deneyli Ekler",
}

EKLER_TAGS = ("[EKLER]", "[EKLER_AUTO]", "[EKLER_BOLUMU]")
EK_HEADING_RE = re.compile(r"^EK\s*[-:]?\s*(\d+)\s*$", re.IGNORECASE)
A4_WIDTH = 595.0
A4_HEIGHT = 842.0


def _com_guvenli_temizle(pythoncom, com_initialized=False, belge=None, uygulama=None):
    """COM temizligini, bir adimdaki hata digerlerini engellemeden tamamla."""
    if belge is not None:
        try:
            belge.Close(False)
        except Exception:
            pass
    if uygulama is not None:
        try:
            uygulama.Quit()
        except Exception:
            pass
    if com_initialized:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _com_ozelligini_ayarla(nesne, ad, deger):
    try:
        setattr(nesne, ad, deger)
    except Exception:
        pass


def _a4_size_for(width, height):
    if float(width or 0) > float(height or 0):
        return A4_HEIGHT, A4_WIDTH
    return A4_WIDTH, A4_HEIGHT


def _fit_rect(src_width, src_height, page_width, page_height, margin=28):
    box = fitz.Rect(margin, margin, page_width - margin, page_height - margin)
    scale = min(box.width / max(src_width, 1), box.height / max(src_height, 1))
    target_width = src_width * scale
    target_height = src_height * scale
    x0 = box.x0 + (box.width - target_width) / 2
    y0 = box.y0 + (box.height - target_height) / 2
    return fitz.Rect(x0, y0, x0 + target_width, y0 + target_height)


def _clean(value):
    return "" if value is None else str(value).strip()


def _safe_name(value, fallback="Proje"):
    text = _clean(value) or fallback
    text = re.sub(r"[^\w\-\.]+", "_", text, flags=re.UNICODE).strip("._")
    return text or fallback


def _row_has_value(row):
    if isinstance(row, dict):
        values = row.values()
    elif isinstance(row, (list, tuple)):
        values = row
    else:
        values = [row]
    for value in values:
        text = _clean(value)
        if text and text not in ("-", "0.0-0.0"):
            return True
    return False


def proje_presiyometre_var_mi(veri):
    for sondaj in (veri or {}).get("sondaj", []) or []:
        for row in sondaj.get("pmt", []) or []:
            if _row_has_value(row):
                return True
    return False


def uygun_ek_seti(veri):
    return EK_SET_ARAZI_DENEYLI if proje_presiyometre_var_mi(veri) else EK_SET_NORMAL


def ek_sablon_yollari(veri=None):
    ayarlar = (veri or {}).get("ayarlar", {})
    normal = _clean(ayarlar.get("ek_tutanak_path")) or DEFAULT_EK_NORMAL
    arazi_deneyli = _clean(ayarlar.get("ek_arazi_deneyli_path")) or DEFAULT_EK_ARAZI_DENEYLI
    return {EK_SET_NORMAL: normal, EK_SET_ARAZI_DENEYLI: arazi_deneyli}


def ek_set_sablonu(veri, set_key=None):
    set_key = set_key if set_key in EK_SET_LABELS else uygun_ek_seti(veri)
    paths = ek_sablon_yollari(veri)
    if set_key == EK_SET_ARAZI_DENEYLI:
        return set_key, "Arazi Deneyli", paths[EK_SET_ARAZI_DENEYLI]
    return EK_SET_NORMAL, "Tutanaklı", paths[EK_SET_NORMAL]


def uygun_ek_sablonu(veri):
    _, label, source = ek_set_sablonu(veri)
    return label, source


def ek_dosya_adi(veri, source_path=None, label=None):
    kunye = (veri or {}).get("kunye", {})
    owner = _safe_name(kunye.get("sahibi"), "Proje")
    suffix = "Arazi_Deneyli" if (label or "").casefold().startswith("arazi") else "Tutanakli"
    ext = os.path.splitext(source_path or "")[1] or ".docx"
    return f"{owner}_EK_Sondaj_{suffix}{ext}"


def ek_pdf_dosya_adi(veri, set_key=None):
    kunye = (veri or {}).get("kunye", {})
    owner = _safe_name(kunye.get("sahibi"), "Proje")
    set_key = set_key if set_key in EK_SET_LABELS else uygun_ek_seti(veri)
    suffix = "Arazi_Deneyli" if set_key == EK_SET_ARAZI_DENEYLI else "Normal"
    return f"{owner}_Ekler_{suffix}.pdf"


def _docx_source_path(path):
    if not path:
        return ""
    if path.lower().endswith(".docx") and os.path.exists(path):
        return path
    root, _ = os.path.splitext(path)
    converted = f"{root}.docx"
    if os.path.exists(converted):
        return converted
    return ""


def _body_elements(doc):
    for child in doc._body._element:
        if child.tag != qn("w:sectPr"):
            yield child


def _element_text(element):
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def ek_bloklari_oku(path):
    source_path = _docx_source_path(path)
    if not source_path:
        raise FileNotFoundError(f"Ek dosyası okunabilmek için .docx olmalı: {path}")
    source_doc = Document(source_path)
    blocks = []
    current_no = None
    current_elements = []
    for element in _body_elements(source_doc):
        match = EK_HEADING_RE.match(_element_text(element))
        if match:
            if current_no and current_elements:
                blocks.append({"no": current_no, "elements": current_elements})
            current_no = match.group(1)
            current_elements = []
        if current_no:
            current_elements.append(deepcopy(element))
    if current_no and current_elements:
        blocks.append({"no": current_no, "elements": current_elements})
    if not blocks:
        blocks = [{"no": "", "elements": [deepcopy(element) for element in _body_elements(source_doc)]}]
    return blocks


def _block_titles(block):
    titles = []
    for element in block.get("elements", []):
        text = _element_text(element)
        if not text or EK_HEADING_RE.match(text):
            continue
        titles.append(text)
    return titles


def ek_basliklari(veri, set_key=None):
    set_key, label, source = ek_set_sablonu(veri, set_key)
    blocks = ek_bloklari_oku(source)
    items = []
    for block in blocks:
        titles = _block_titles(block)
        items.append(
            {
                "no": block["no"],
                "titles": titles,
                "title": " / ".join(titles),
                "set_key": set_key,
                "set_label": label,
            }
        )
    return items


def ek_icerik_haritasi(veri, set_key=None):
    set_key = set_key if set_key in EK_SET_LABELS else uygun_ek_seti(veri)
    return (veri or {}).setdefault("ek_icerikleri", {}).setdefault(set_key, {})


def _page_break_element():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def _elements_for_blocks(blocks):
    elements = []
    for block in blocks:
        elements.append(_page_break_element())
        elements.extend(deepcopy(element) for element in block["elements"])
    return elements


def _iter_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _remove_tags_from_paragraph(paragraph, tags):
    for run in paragraph.runs:
        for tag in tags:
            if tag in run.text:
                run.text = run.text.replace(tag, "")
    if any(tag in paragraph.text for tag in tags):
        paragraph.text = ""


def _find_tag_paragraph(doc, tags):
    for paragraph in _iter_paragraphs(doc):
        if any(tag in paragraph.text for tag in tags):
            return paragraph
    return None


def _insert_elements_after(paragraph, elements):
    anchor = paragraph._p
    for element in reversed(elements):
        anchor.addnext(deepcopy(element))


def _append_elements(doc, elements):
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))
    for element in elements:
        if sect_pr is not None:
            sect_pr.addprevious(deepcopy(element))
        else:
            body.append(deepcopy(element))


def _individual_tags(no):
    return (f"[EK-{no}]", f"[EK{no}]", f"[EK_{no}]", f"[EKLER_{no}]")


def ekleri_rapora_ekle(doc, veri):
    label, source = uygun_ek_sablonu(veri)
    blocks = ek_bloklari_oku(source)
    inserted = set()

    for idx, block in enumerate(blocks):
        if not block["no"]:
            continue
        tags = _individual_tags(block["no"])
        paragraph = _find_tag_paragraph(doc, tags)
        if paragraph:
            _remove_tags_from_paragraph(paragraph, tags)
            _insert_elements_after(paragraph, _elements_for_blocks([block]))
            inserted.add(idx)

    remaining = [block for idx, block in enumerate(blocks) if idx not in inserted]
    if remaining:
        paragraph = _find_tag_paragraph(doc, EKLER_TAGS)
        if paragraph:
            _remove_tags_from_paragraph(paragraph, EKLER_TAGS)
            _insert_elements_after(paragraph, _elements_for_blocks(remaining))
        else:
            _append_elements(doc, _elements_for_blocks(remaining))

    return {"label": label, "source": source, "count": len(blocks)}


def _font_path():
    for path in (
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
    ):
        if os.path.exists(path):
            return path
    return None


def _pil_font(size):
    font_path = _font_path()
    if font_path:
        try:
            return ImageFont.truetype(font_path, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def _draw_centered_lines(draw, lines, font, y, image_width, fill=(20, 20, 20), line_gap=18):
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        height = bbox[3] - bbox[1]
        draw.text(((image_width - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap
    return y


def _append_cover_page(pdf_doc, no, titles):
    image_width, image_height = 1240, 1754
    image = Image.new("RGB", (image_width, image_height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _pil_font(92)
    body_font = _pil_font(56)
    y = 570
    y = _draw_centered_lines(draw, [f"EK-{no}"], title_font, y, image_width, fill=(0, 0, 0), line_gap=44)
    title_lines = [line for line in titles if line]
    if title_lines:
        _draw_centered_lines(draw, title_lines, body_font, y + 40, image_width, fill=(35, 35, 35), line_gap=26)
    stream = BytesIO()
    image.save(stream, format="PNG")
    page = pdf_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
    page.insert_image(page.rect, stream=stream.getvalue())


def _append_message_page(pdf_doc, title, lines):
    image_width, image_height = 1240, 1754
    image = Image.new("RGB", (image_width, image_height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _pil_font(58)
    body_font = _pil_font(34)
    y = 260
    y = _draw_centered_lines(draw, [title], title_font, y, image_width, fill=(0, 0, 0), line_gap=34)
    wrapped = []
    for line in lines:
        text = str(line)
        while len(text) > 74:
            wrapped.append(text[:74])
            text = text[74:]
        wrapped.append(text)
    _draw_centered_lines(draw, wrapped, body_font, y + 30, image_width, fill=(55, 55, 55), line_gap=18)
    stream = BytesIO()
    image.save(stream, format="PNG")
    page = pdf_doc.new_page(width=A4_WIDTH, height=A4_HEIGHT)
    page.insert_image(page.rect, stream=stream.getvalue())


def _append_pdf(pdf_doc, path):
    with fitz.open(path) as source:
        for page_index in range(source.page_count):
            src_page = source[page_index]
            src_rect = src_page.rect
            page_width, page_height = _a4_size_for(src_rect.width, src_rect.height)
            page = pdf_doc.new_page(width=page_width, height=page_height)
            target = _fit_rect(src_rect.width, src_rect.height, page_width, page_height, margin=0)
            page.show_pdf_page(target, source, page_index)


def _append_image(pdf_doc, path):
    with Image.open(path) as image:
        width, height = image.size
    page_width, page_height = _a4_size_for(width, height)
    page = pdf_doc.new_page(width=page_width, height=page_height)
    page.insert_image(_fit_rect(width, height, page_width, page_height), filename=path)


def _office_to_pdf(path, tmp_dir):
    ext = os.path.splitext(path)[1].lower()
    output_path = os.path.join(tmp_dir, f"{_safe_name(os.path.basename(path))}.pdf")
    if ext in (".doc", ".docx"):
        import pythoncom
        import win32com.client

        app = None
        doc = None
        com_initialized = False
        try:
            pythoncom.CoInitialize()
            com_initialized = True
            app = win32com.client.DispatchEx("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
            _com_ozelligini_ayarla(app, "AutomationSecurity", 3)
            try:
                _com_ozelligini_ayarla(app.Options, "UpdateLinksAtOpen", False)
            except Exception:
                pass
            doc = app.Documents.Open(
                os.path.abspath(path),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
                OpenAndRepair=False,
                NoEncodingDialog=True,
            )
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)
        finally:
            _com_guvenli_temizle(
                pythoncom,
                com_initialized=com_initialized,
                belge=doc,
                uygulama=app,
            )
        return output_path
    if ext in (".xls", ".xlsx"):
        import pythoncom
        import win32com.client

        app = None
        workbook = None
        com_initialized = False
        try:
            pythoncom.CoInitialize()
            com_initialized = True
            app = win32com.client.DispatchEx("Excel.Application")
            app.Visible = False
            app.DisplayAlerts = False
            _com_ozelligini_ayarla(app, "AutomationSecurity", 3)
            _com_ozelligini_ayarla(app, "AskToUpdateLinks", False)
            _com_ozelligini_ayarla(app, "EnableEvents", False)
            workbook = app.Workbooks.Open(
                os.path.abspath(path),
                UpdateLinks=0,
                ReadOnly=True,
                IgnoreReadOnlyRecommended=True,
                AddToMru=False,
                Notify=False,
            )
            workbook.ExportAsFixedFormat(0, os.path.abspath(output_path))
        finally:
            _com_guvenli_temizle(
                pythoncom,
                com_initialized=com_initialized,
                belge=workbook,
                uygulama=app,
            )
        return output_path
    raise ValueError(f"Bu dosya türü PDF'e çevrilemiyor: {ext}")


def _append_attachment(pdf_doc, path, tmp_dir):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        _append_pdf(pdf_doc, path)
        return
    if ext in (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"):
        _append_image(pdf_doc, path)
        return
    if ext in (".doc", ".docx", ".xls", ".xlsx"):
        converted = _office_to_pdf(path, tmp_dir)
        _append_pdf(pdf_doc, converted)
        return
    raise ValueError(f"Desteklenmeyen ek dosyası türü: {ext}")


def _pdf_ciktisini_dogrula(path):
    if not os.path.isfile(path) or os.path.getsize(path) <= 0:
        raise RuntimeError("Ekler PDF çıktısı boş veya oluşturulamadı.")
    kontrol = None
    try:
        kontrol = fitz.open(path)
        if kontrol.page_count <= 0:
            raise RuntimeError("Ekler PDF çıktısında sayfa bulunamadı.")
    finally:
        if kontrol is not None:
            try:
                kontrol.close()
            except Exception:
                pass


def ekler_pdf_olustur(veri, output_path, set_key=None):
    set_key, label, _ = ek_set_sablonu(veri, set_key)
    basliklar = ek_basliklari(veri, set_key)
    icerikler = ek_icerik_haritasi(veri, set_key)
    warnings = []
    attached_count = 0
    output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    os.makedirs(output_dir, exist_ok=True)
    temp_fd, temp_output_path = tempfile.mkstemp(
        prefix=f".{os.path.basename(output_path)}.",
        suffix=".tmp.pdf",
        dir=output_dir,
    )
    os.close(temp_fd)
    try:
        with tempfile.TemporaryDirectory(prefix="raporpro_ekler_") as tmp_dir:
            pdf_doc = None
            try:
                pdf_doc = fitz.open()
                for item in basliklar:
                    no = str(item["no"])
                    _append_cover_page(pdf_doc, no, item.get("titles", []))
                    for path in list(icerikler.get(no, []) or []):
                        if not path or not os.path.exists(path):
                            warnings.append(f"EK-{no}: dosya bulunamadı: {path}")
                            continue
                        try:
                            _append_attachment(pdf_doc, path, tmp_dir)
                            attached_count += 1
                        except Exception as exc:
                            warnings.append(f"EK-{no}: {os.path.basename(path)} eklenemedi: {exc}")
                            _append_message_page(
                                pdf_doc,
                                f"EK-{no} Dosya Eklenemedi",
                                [os.path.basename(path), str(exc)],
                            )
                if pdf_doc.page_count == 0:
                    _append_message_page(pdf_doc, "Ekler", ["Ek kapak sayfası bulunamadı."])
                pdf_doc.save(temp_output_path, garbage=4, deflate=True)
            finally:
                if pdf_doc is not None:
                    try:
                        pdf_doc.close()
                    except Exception:
                        pass
        _pdf_ciktisini_dogrula(temp_output_path)
        os.replace(temp_output_path, output_path)
    finally:
        if os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except OSError:
                pass
    return {
        "path": output_path,
        "set_key": set_key,
        "label": label,
        "cover_count": len(basliklar),
        "attached_count": attached_count,
        "warnings": warnings,
    }


def ek_olustur(veri, folder):
    label, source = uygun_ek_sablonu(veri)
    if not source or not os.path.exists(source):
        raise FileNotFoundError(f"{label} ek şablonu bulunamadı: {source}")
    os.makedirs(folder, exist_ok=True)
    target = os.path.join(folder, ek_dosya_adi(veri, source, label))
    shutil.copy2(source, target)
    return target

# Dosya: RaporPro/raporlama_parsel.py
"""Parsel bazli Word raporundaki sabit tablo ve kosullu bolumleri yonetir."""

from __future__ import annotations

import math
import re
import unicodedata

from docx.text.paragraph import Paragraph

from rapor_parsel_bilgileri import (
    aktif_fay_satirlari,
    rapor_bilgilerini_normalize_et,
)
from raporlama_tablo import apply_report_table_style, set_cell_text_clean


def _normalized(value):
    text = unicodedata.normalize("NFKD", str(value or "").casefold())
    return "".join(char for char in text if not unicodedata.combining(char))


def _table_header_text(table):
    if not table.rows:
        return ""
    return " | ".join(cell.text for cell in table.rows[0].cells)


def _find_table(doc, header_fragment):
    needle = _normalized(header_fragment)
    for table in doc.tables:
        if needle in _normalized(_table_header_text(table)):
            return table
    return None


def _remove_table(table):
    element = getattr(table, "_element", None)
    parent = element.getparent() if element is not None else None
    if parent is not None:
        parent.remove(element)
        return True
    return False


def _remove_paragraph(paragraph):
    element = getattr(paragraph, "_element", None)
    parent = element.getparent() if element is not None else None
    if parent is not None:
        parent.remove(element)
        return True
    return False


def _remove_paragraphs_with_prefix(doc, prefixes):
    normalized_prefixes = tuple(_normalized(prefix) for prefix in prefixes)
    removed = 0
    for paragraph in list(doc.paragraphs):
        text = _normalized(paragraph.text).strip()
        if any(text.startswith(prefix) for prefix in normalized_prefixes):
            removed += int(_remove_paragraph(paragraph))
    return removed


def _replace_table_rows(table, rows, widths_cm=None):
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if index < len(cells):
                set_cell_text_clean(cells[index], value, font_size=10)
    apply_report_table_style(
        table,
        header_rows=1,
        text_cols={0},
        widths_cm=widths_cm,
    )
    return len(rows)


def _number(value):
    text = str(value or "").strip().replace(",", ".")
    if not text:
        return None
    try:
        number = float(text)
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def _distance_m(coords):
    values = [_number(value) for value in (coords or [])]
    if len(values) < 4:
        return None
    pairs = [
        (values[index], values[index + 1])
        for index in range(0, len(values) - 1, 2)
        if values[index] is not None and values[index + 1] is not None
    ]
    if len(pairs) < 2:
        return None
    y1, x1 = pairs[0]
    y2, x2 = pairs[-1]
    if all((-90 <= y <= 90 and -180 <= x <= 180) for y, x in ((y1, x1), (y2, x2))):
        radius = 6_371_008.8
        lat1 = math.radians(y1)
        lat2 = math.radians(y2)
        dlat = lat2 - lat1
        dlon = math.radians(x2 - x1)
        a = (
            math.sin(dlat / 2) ** 2
            + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2) ** 2
        )
        return radius * 2 * math.atan2(math.sqrt(a), math.sqrt(max(0.0, 1 - a)))
    return math.hypot(y2 - y1, x2 - x1)


def _format_distance(coords):
    distance = _distance_m(coords)
    if distance is None:
        return "-"
    if abs(distance - round(distance)) < 0.05:
        return str(int(round(distance)))
    return f"{distance:.1f}".replace(".", ",")


def _with_unit(value, unit):
    text = str(value or "").strip()
    if not text:
        return "-"
    if re.search(r"[A-Za-zçğıöşüÇĞİÖŞÜ]", text):
        return text
    return f"{text} {unit}"


def _serim_rows(ss_list, data, *, masw=False):
    rows = []
    shot_key = "masw_vurus_sayisi" if masw else "sismik_vurus_sayisi"
    record_key = "masw_kayit_uzunlugu" if masw else "sismik_kayit_uzunlugu"
    for index, serim in enumerate(ss_list or [], start=1):
        if not isinstance(serim, dict):
            continue
        name = str(serim.get("ad") or f"Serim {index}").strip()
        rows.append(
            [
                name,
                str(data.get(shot_key) or "-").strip(),
                _with_unit(data.get(record_key), "sn"),
                _format_distance(serim.get("coords")),
            ]
        )
    return rows


def rapor_sabit_tablolarini_uygula(doc, veri):
    """Sablondaki sabit iklim, fay ve yontem tablolarini degistirmeden birak."""
    return {
        "iklim_tablolari": 0,
        "aktif_fay_satiri": 0,
        "sismik_satiri": 0,
        "masw_satiri": 0,
    }


def _heading_number(paragraph):
    if not str(getattr(paragraph.style, "name", "")).startswith("Heading"):
        return None
    match = re.match(r"\s*(\d+(?:\.\d+)*)", paragraph.text or "")
    if not match:
        return None
    return tuple(int(part) for part in match.group(1).split("."))


def numarali_bolumu_kaldir(doc, heading_number):
    """Numarali basligi ve bir sonraki esit/ust basliga kadarki icerigi kaldir."""
    target = tuple(int(part) for part in str(heading_number).strip(".").split("."))
    body = doc.element.body
    children = list(body.iterchildren())
    start_index = None
    for index, child in enumerate(children):
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, doc._body)
            if _heading_number(paragraph) == target:
                start_index = index
                break
    if start_index is None:
        return 0

    end_index = len(children)
    for index in range(start_index + 1, len(children)):
        child = children[index]
        if not child.tag.endswith("}p"):
            continue
        number = _heading_number(Paragraph(child, doc._body))
        if number is not None and len(number) <= len(target):
            end_index = index
            break
    for child in children[start_index:end_index]:
        body.remove(child)
    return end_index - start_index


def _heading_text_update(doc, number, text):
    target = tuple(int(part) for part in str(number).strip(".").split("."))
    for paragraph in doc.paragraphs:
        if _heading_number(paragraph) == target:
            paragraph.text = text
            return True
    return False


def rapor_kosullu_bolumlerini_uygula(doc, veri):
    """Veri bulunmayan arazi ve jeofizik alt bolumlerini rapordan kaldir."""
    veri = veri if isinstance(veri, dict) else {}
    jeofizik = veri.get("jeofizik", {})
    jeofizik = jeofizik if isinstance(jeofizik, dict) else {}
    sondajlar = veri.get("sondaj", [])
    sondajlar = sondajlar if isinstance(sondajlar, list) else []
    has_ss = bool(jeofizik.get("ss_list"))
    has_mt = bool(jeofizik.get("mt_list"))
    has_spt = any(
        isinstance(sondaj, dict) and bool(sondaj.get("spt"))
        for sondaj in sondajlar
    )
    has_pmt = any(
        isinstance(sondaj, dict) and bool(sondaj.get("pmt"))
        for sondaj in sondajlar
    )
    has_rock = any(
        isinstance(sondaj, dict) and bool(sondaj.get("kaya"))
        for sondaj in sondajlar
    )
    removed = []

    if not has_ss and not has_mt:
        if numarali_bolumu_kaldir(doc, "3.1"):
            removed.append("3.1")
    else:
        if not has_ss:
            for number in ("3.1.1", "3.1.2", "3.1.4"):
                if numarali_bolumu_kaldir(doc, number):
                    removed.append(number)
        if not has_mt and numarali_bolumu_kaldir(doc, "3.1.3"):
            removed.append("3.1.3")

    if not sondajlar:
        for number in ("3.3", "3.4"):
            if numarali_bolumu_kaldir(doc, number):
                removed.append(number)
    elif not (has_spt or has_pmt or has_rock):
        if numarali_bolumu_kaldir(doc, "3.4"):
            removed.append("3.4")
    elif not has_spt and not has_rock:
        if numarali_bolumu_kaldir(doc, "3.4.1"):
            removed.append("3.4.1")
    elif has_rock and not has_spt:
        _heading_text_update(doc, "3.4.1", "3.4.1. Karot Değerlendirmeleri")

    return removed


__all__ = [
    "numarali_bolumu_kaldir",
    "rapor_kosullu_bolumlerini_uygula",
    "rapor_sabit_tablolarini_uygula",
]

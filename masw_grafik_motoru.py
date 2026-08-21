# Dosya: RaporPro/masw_grafik_motoru.py
"""Jeofizik degerlendirme Word'lerinden MASW hiz grafiklerini rapora aktarir."""

from __future__ import annotations

import os
import re
import unicodedata
from copy import deepcopy
from dataclasses import dataclass
from functools import lru_cache

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.shared import Cm


MASW_GRAFIK_ETIKETI = "[RESIM_MASW]"
MASW_GRAFIK_GENISLIK_CM = 13.2
MASW_GRAFIK_AZAMI_YUKSEKLIK_CM = 9.0
MASW_SAYFA_BASINA_GRAFIK = 2


@dataclass(frozen=True)
class MaswGrafikKaydi:
    """Kaynak Word'den alinmis tek bir MASW hiz grafigi."""

    kaynak_yolu: str
    paragraf_no: int
    drawing: object
    iliskili_resimler: tuple[tuple[str, bytes, str, str], ...]
    genislik_emu: int
    yukseklik_emu: int


@dataclass(frozen=True)
class MaswGrafikEklemeSonucu:
    """Rapor ekleme isleminin ozetini tasir."""

    eklenen: int
    kaldirilan_sabit_gorsel: int
    hatalar: tuple[str, ...] = ()


def _metin_normalize(value):
    text = str(value or "").replace("ı", "i").replace("İ", "I")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"\s+", " ", text.casefold()).strip()


def _dogal_siralama_anahtari(path):
    name = _metin_normalize(os.path.basename(str(path or "")))
    return tuple(int(part) if part.isdigit() else part for part in re.split(r"(\d+)", name))


def masw_word_yollari_normalize(paths):
    """Yollari temizler, tekrarlarini kaldirir ve dosya adina gore siralar."""

    unique = []
    seen = set()
    for value in paths or []:
        path = str(value or "").strip().strip('"')
        if not path:
            continue
        path = os.path.normpath(os.path.expandvars(os.path.expanduser(path)))
        key = os.path.normcase(os.path.abspath(path))
        if key in seen:
            continue
        seen.add(key)
        unique.append(path)
    return sorted(unique, key=_dogal_siralama_anahtari)


def _dispersiyon_basligi_mi(text):
    normalized = _metin_normalize(text)
    return "dispersiyon" in normalized and "egrisi" in normalized


def _masw_grafik_basligi_mi(text):
    normalized = _metin_normalize(text).replace(" ", "")
    return "masw" in normalized and "olcum" in normalized and "grafik" in normalized


def _masw_dispersiyon_sonrasi_cizimleri(paragraphs, marker_index):
    """Bir dispersiyon basligindan sonraki tek grafik blogunu belge sirasiyla toplar."""

    candidates = []
    for index in range(marker_index + 1, len(paragraphs)):
        paragraph = paragraphs[index]
        drawings = paragraph._p.xpath(".//w:drawing")
        if drawings:
            candidates.extend((index, drawing) for drawing in drawings)
            continue

        if not paragraph.text.strip():
            continue

        break

    return candidates


def _drawing_resimlerini_oku(document, drawing):
    images = []
    seen = set()
    for blip in drawing.xpath(".//a:blip"):
        source_rid = blip.get(qn("r:embed"))
        if not source_rid or source_rid in seen:
            continue
        source_part = document.part.related_parts.get(source_rid)
        if source_part is None or not getattr(source_part, "blob", None):
            continue
        extension = os.path.splitext(str(source_part.partname))[1].lower() or ".bin"
        images.append(
            (
                source_rid,
                bytes(source_part.blob),
                str(source_part.content_type),
                extension,
            )
        )
        seen.add(source_rid)
    return tuple(images)


def _drawing_boyutu(drawing):
    extents = drawing.xpath(".//wp:extent")
    if not extents:
        return 1, 1
    try:
        width = max(1, int(extents[0].get("cx") or 1))
        height = max(1, int(extents[0].get("cy") or 1))
    except (TypeError, ValueError):
        return 1, 1
    return width, height


def masw_grafik_kaydi_oku(path):
    """Kaynak Word'deki dispersiyon egrisini izleyen S-hizi grafigini okur."""

    source_path = os.path.abspath(os.fspath(path))
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Dosya bulunamadi: {source_path}")
    if os.path.splitext(source_path)[1].lower() != ".docx":
        raise ValueError("MASW grafik kaynagi DOCX biciminde olmalidir.")
    stat = os.stat(source_path)
    return _masw_grafik_kaydi_oku_cached(
        source_path,
        int(stat.st_mtime_ns),
        int(stat.st_size),
    )


@lru_cache(maxsize=32)
def _masw_grafik_kaydi_oku_cached(source_path, _mtime_ns, _size):
    """Degismeyen kaynak Word'u tekrar tekrar acmadan grafik kaydini dondurur."""

    document = Document(source_path)
    paragraphs = list(document.paragraphs)
    marker_indices = [
        index
        for index, paragraph in enumerate(paragraphs)
        if _dispersiyon_basligi_mi(paragraph.text)
    ]

    selected_index = None
    selected_drawing = None
    for marker_index in reversed(marker_indices):
        candidates = _masw_dispersiyon_sonrasi_cizimleri(paragraphs, marker_index)
        if candidates:
            selected_index, selected_drawing = candidates[-1]
            break

    if selected_drawing is None:
        for index in range(len(paragraphs) - 1, -1, -1):
            drawings = paragraphs[index]._p.xpath(".//w:drawing")
            if drawings:
                selected_index = index
                selected_drawing = drawings[-1]
                break

    if selected_drawing is None:
        raise ValueError("Word dosyasinda aktarilabilecek MASW grafigi bulunamadi.")

    images = _drawing_resimlerini_oku(document, selected_drawing)
    if not images:
        raise ValueError("MASW grafiginin gomulu resim verisi okunamadi.")
    width, height = _drawing_boyutu(selected_drawing)
    return MaswGrafikKaydi(
        kaynak_yolu=source_path,
        paragraf_no=int(selected_index),
        drawing=deepcopy(selected_drawing),
        iliskili_resimler=images,
        genislik_emu=width,
        yukseklik_emu=height,
    )


def masw_word_kaynaklarini_dogrula(paths):
    """Gecerli grafik kayitlarini ve kullaniciya gosterilecek hatalari dondurur."""

    records = []
    errors = []
    for path in masw_word_yollari_normalize(paths):
        try:
            records.append(masw_grafik_kaydi_oku(path))
        except Exception as exc:
            errors.append(f"{os.path.basename(path) or path}: {exc}")
    return records, errors


def _benzersiz_part_adi(package, index, extension):
    used = {str(part.partname) for part in package.iter_parts()}
    candidate_index = max(1, int(index))
    while True:
        candidate = f"/word/media/masw_graph_{candidate_index}{extension}"
        if candidate not in used:
            return PackURI(candidate)
        candidate_index += 1


def _sonraki_docpr_id(document):
    values = []
    for element in document.element.xpath(".//wp:docPr"):
        try:
            values.append(int(element.get("id") or 0))
        except (TypeError, ValueError):
            continue
    return max(values, default=0) + 1


def _drawing_hedefe_kopyala(document, record, image_index, docpr_id):
    drawing = deepcopy(record.drawing)
    rid_map = {}
    for offset, (source_rid, blob, content_type, extension) in enumerate(
        record.iliskili_resimler
    ):
        part = Part(
            _benzersiz_part_adi(
                document.part.package,
                image_index + offset,
                extension,
            ),
            content_type,
            blob,
            document.part.package,
        )
        rid_map[source_rid] = document.part.relate_to(part, RT.IMAGE)

    for blip in drawing.xpath(".//a:blip"):
        source_rid = blip.get(qn("r:embed"))
        if source_rid in rid_map:
            blip.set(qn("r:embed"), rid_map[source_rid])

    width_limit = int(Cm(MASW_GRAFIK_GENISLIK_CM))
    height_limit = int(Cm(MASW_GRAFIK_AZAMI_YUKSEKLIK_CM))
    scale = min(
        width_limit / max(1, record.genislik_emu),
        height_limit / max(1, record.yukseklik_emu),
    )
    width = max(1, int(record.genislik_emu * scale))
    height = max(1, int(record.yukseklik_emu * scale))
    for extent in drawing.xpath(".//wp:extent") + drawing.xpath(".//a:xfrm/a:ext"):
        extent.set("cx", str(width))
        extent.set("cy", str(height))
    for properties in drawing.xpath(".//wp:docPr"):
        properties.set("id", str(docpr_id))
        properties.set("name", f"MASW Hiz Grafigi {image_index}")
        properties.set("descr", os.path.basename(record.kaynak_yolu))
    return drawing


def _paragrafi_sil(paragraph):
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _masw_hedefini_bul(document):
    paragraphs = list(document.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if MASW_GRAFIK_ETIKETI in paragraph.text:
            caption = None
            for previous in reversed(paragraphs[:index]):
                if not previous.text.strip():
                    continue
                if _masw_grafik_basligi_mi(previous.text):
                    caption = previous
                break
            return paragraph, caption, [paragraph]

    caption = next(
        (paragraph for paragraph in paragraphs if _masw_grafik_basligi_mi(paragraph.text)),
        None,
    )
    if caption is None:
        return None, None, []

    caption_index = paragraphs.index(caption)
    image_block = []
    for paragraph in paragraphs[caption_index + 1 :]:
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if has_drawing or not paragraph.text.strip():
            image_block.append(paragraph)
            continue
        break
    anchor = image_block[0] if image_block else None
    return anchor, caption, image_block


def _paragrafi_sonrasina_ekle(document, anchor_element):
    paragraph = document.add_paragraph()
    anchor_element.addnext(paragraph._p)
    return paragraph


def masw_grafiklerini_rapora_ekle(document, paths):
    """Sabit ornekleri temizleyip secilen Word'lerin hiz grafiklerini yerlestirir."""

    records, errors = masw_word_kaynaklarini_dogrula(paths)
    anchor, caption, old_block = _masw_hedefini_bul(document)
    removed_count = sum(
        len(paragraph._p.xpath(".//w:drawing")) for paragraph in old_block
    )

    if anchor is None and caption is not None:
        anchor = _paragrafi_sonrasina_ekle(document, caption._p)
        old_block = [anchor]

    if anchor is None:
        if records:
            errors.append(
                f"Rapor şablonunda {MASW_GRAFIK_ETIKETI} etiketi veya MASW ölçüm grafikleri başlığı bulunamadı."
            )
        return MaswGrafikEklemeSonucu(0, removed_count, tuple(errors))

    anchor.clear()
    for paragraph in old_block:
        if paragraph is not anchor:
            _paragrafi_sil(paragraph)

    if not records:
        _paragrafi_sil(anchor)
        if caption is not None:
            _paragrafi_sil(caption)
        return MaswGrafikEklemeSonucu(0, removed_count, tuple(errors))

    if caption is not None:
        caption.paragraph_format.keep_with_next = True

    current = anchor
    next_docpr_id = _sonraki_docpr_id(document)
    for index, record in enumerate(records, start=1):
        if index > 1:
            current = _paragrafi_sonrasina_ekle(document, current._p)
        current.alignment = WD_ALIGN_PARAGRAPH.CENTER
        current.paragraph_format.space_after = Cm(0)
        current.paragraph_format.keep_with_next = index % MASW_SAYFA_BASINA_GRAFIK == 1 and index < len(records)
        if index > MASW_SAYFA_BASINA_GRAFIK and (index - 1) % MASW_SAYFA_BASINA_GRAFIK == 0:
            current.paragraph_format.page_break_before = True
        drawing = _drawing_hedefe_kopyala(
            document,
            record,
            image_index=index,
            docpr_id=next_docpr_id,
        )
        next_docpr_id += 1
        current.add_run()._r.append(drawing)

    return MaswGrafikEklemeSonucu(len(records), removed_count, tuple(errors))


__all__ = [
    "MASW_GRAFIK_ETIKETI",
    "MaswGrafikEklemeSonucu",
    "MaswGrafikKaydi",
    "masw_grafik_kaydi_oku",
    "masw_grafiklerini_rapora_ekle",
    "masw_word_kaynaklarini_dogrula",
    "masw_word_yollari_normalize",
]

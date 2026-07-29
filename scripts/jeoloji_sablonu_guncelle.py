# Dosya: RaporPro/scripts/jeoloji_sablonu_guncelle.py
"""Dahili Word şablonundaki sabit jeoloji metinlerini etiketlere dönüştürür."""

from __future__ import annotations

import copy
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "sablonlar" / "rapor" / "varsayilan_rapor_sablonu.docx"


def _find_paragraph(doc, predicate):
    return next((paragraph for paragraph in doc.paragraphs if predicate(paragraph.text.strip())), None)


def _replace_text_preserving_format(paragraph, text):
    run_properties = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = copy.deepcopy(paragraph.runs[0]._r.rPr)
    paragraph.clear()
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)
    return paragraph


def _insert_after(paragraph, text):
    new_element = OxmlElement("w:p")
    if paragraph._p.pPr is not None:
        new_element.append(copy.deepcopy(paragraph._p.pPr))
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    _replace_text_preserving_format(new_paragraph, text)
    return new_paragraph


def _remove_until(doc, start_paragraph, end_predicate):
    current = start_paragraph._p.getnext()
    while current is not None:
        next_element = current.getnext()
        paragraph = Paragraph(current, start_paragraph._parent)
        if end_predicate(paragraph.text.strip()):
            break
        current.getparent().remove(current)
        current = next_element


def _paragraph_after_heading(doc, heading_prefix, predicate):
    paragraphs = doc.paragraphs
    start = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip().startswith(heading_prefix)
        ),
        None,
    )
    if start is None:
        return None
    return next(
        (
            paragraph
            for paragraph in paragraphs[start + 1 :]
            if predicate(paragraph.text.strip())
        ),
        None,
    )


def update_template(template_path=DEFAULT_TEMPLATE):
    path = Path(template_path)
    doc = Document(path)
    changed = False

    regional_intro = _find_paragraph(
        doc,
        lambda text: text.startswith(
            "Çalışma alanı ve yakın çevresinde literatürde Çanakkale Formasyonuna"
        ),
    )
    if regional_intro is not None:
        _replace_text_preserving_format(regional_intro, "[BOLGESEL_JEOLOJI]")
        changed = True

    regional_units = _find_paragraph(
        doc,
        lambda text: text == "Çanakkale Formasyonu (Tmçk)",
    )
    if regional_units is not None:
        _replace_text_preserving_format(
            regional_units,
            "[BOLGESEL_JEOLOJI_BIRIMLERI]",
        )
        _remove_until(
            doc,
            regional_units,
            lambda text: text.startswith("Bölgenin stratigrafik kesiti"),
        )
        changed = True

    mt_text = _find_paragraph(
        doc,
        lambda text: text.startswith(
            "İnceleme alanında yapılan mikrotremör ölçümlerinde Çamrakdere Üyesi"
        ),
    )
    if mt_text is not None:
        _replace_text_preserving_format(mt_text, "[MT_BIRIM_METNI]")
        changed = True

    engineering = _paragraph_after_heading(
        doc,
        "6. İNCELEME ALANI MÜHENDİSLİK JEOLOJİSİ",
        lambda text: "[ZEMIN_OZET]" in text,
    )
    if engineering is not None:
        _replace_text_preserving_format(
            engineering,
            "[MUHENDISLIK_JEOLOJISI]",
        )
        map_text = _insert_after(
            engineering,
            "Çalışma alanına ait mühendislik jeolojisi haritası Şekil 10'da verilmiştir.",
        )
        _insert_after(engineering, "[ZEMIN_OZET]")
        # Değişken yalnız ekleme sırasını açık eder; paragraf sırası:
        # dinamik jeoloji, zemin özeti, harita cümlesi.
        _ = map_text
        changed = True

    section_text = _paragraph_after_heading(
        doc,
        "7. JEOLOJİK KESİT",
        lambda text: text.startswith(
            "Çalışma alanı ve yakın çevresinde Alçıtepe Üyesine ait"
        ),
    )
    if section_text is not None:
        _replace_text_preserving_format(
            section_text,
            "[JEOLOJIK_KESIT_ACIKLAMA]",
        )
        _remove_until(
            doc,
            section_text,
            lambda text: text.startswith("Çalışma alanında;"),
        )
        changed = True

    conclusion = _paragraph_after_heading(
        doc,
        "8. SONUÇ VE ÖNERİLER",
        lambda text: text.startswith(
            "İnceleme alanı literatürde Üst Miyosen yaşlı"
        ),
    )
    if conclusion is not None:
        _replace_text_preserving_format(conclusion, "[JEOLOJI_SONUC]")
        changed = True

    required_tags = {
        "[BOLGESEL_JEOLOJI]",
        "[BOLGESEL_JEOLOJI_BIRIMLERI]",
        "[MUHENDISLIK_JEOLOJISI]",
        "[JEOLOJIK_KESIT_ACIKLAMA]",
        "[JEOLOJI_SONUC]",
        "[MT_BIRIM_METNI]",
    }
    actual_tags = {
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("[")
    }
    missing = required_tags - actual_tags
    if missing:
        raise RuntimeError(
            "Şablonda oluşturulamayan jeoloji etiketleri: "
            + ", ".join(sorted(missing))
        )

    if changed:
        with tempfile.NamedTemporaryFile(
            suffix=".docx",
            delete=False,
            dir=path.parent,
        ) as handle:
            temp_path = Path(handle.name)
        try:
            doc.save(temp_path)
            os.replace(temp_path, path)
        finally:
            if temp_path.exists():
                temp_path.unlink()
    return changed


if __name__ == "__main__":
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_TEMPLATE
    print("güncellendi" if update_template(target) else "zaten güncel")

# Dosya: RaporPro/scripts/rapor_parsel_sablonu_guncelle.py
"""Dahili rapor sablonundaki yalniz proje-ozel metinleri etiketlere donusturur."""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

from docx import Document


DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parents[1]
    / "sablonlar"
    / "rapor"
    / "varsayilan_rapor_sablonu.docx"
)

PARAGRAPH_REPLACEMENTS = (
    ("Çanakkale Belediyesi", "[ILGILI_IDARE]"),
    ("Aralık 2025", "[RAPOR_AY_YIL]"),
    ("Bu çalışma; [IL] ili", "[ETUT_AMAC_KAPSAM]"),
    ("Hazırlanan bu rapor kapsamında", "[RAPOR_KAPSAM]"),
    ("İnceleme alanı; [IL] ili", "[PARSEL_TANITIM]"),
    ("İnceleme alanı, ‘’15.10.2017", "[IMAR_PLANI_ACIKLAMA]"),
    ("Çalışma alanı ‘‘[IMAR_ALANI]’’", "[IMAR_ADASI_ACIKLAMA]"),
    ("Çalışma alanında [SONDAJ_BILGISI] Sondajlar", "[SONDAJ_ARAZI_GIRIS]"),
    ("Ayrıca jeofizik çalışmalar kapsamında", "[JEOFIZIK_ARAZI_GIRIS]"),
    ("P atışlarından elde edilen", "[VP_ACIKLAMA]"),
    ("Arazide ölçülmüş Vs hızları", "[MASW_SONUC_ACIKLAMA]"),
    ("İnceleme alanında yapılması planlanan yapının", "[MT_REZONANS_ACIKLAMA]"),
    ("Çalışma alanında [SONDAJ_BILGISI] Sondaj derinlik", "[SONDAJ_BOLUM_GIRIS]"),
    ("İnceleme alanında açılmış olan temel sondaj kuyularındaki numuneler, Atterberg", "[LAB_FIZIK_GIRIS]"),
    ("İnceleme alanında açılmış olan temel sondaj kuyularındaki numunelerin belirli", "[LAB_MEKANIK_GIRIS]"),
    ("Çalışma alanında sondaj noktalarından geçen kesit", "[KESIT_GIRIS]"),
    ("[IL] ili, [ILCE] ilçesi", "[SONUC_GIRIS]"),
    ("İnceleme alanı; Çanakkale ili", "[SONUC_KONUM]"),
    ("İnceleme alanı, “05.10.2017", "[SONUC_IMAR]"),
    ("Çalışma alanında Türkiye Heyelan Envanter", "[SONUC_AFET]"),
)


def _paragraph_text_set(paragraph, text):
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    target = next((run for run in runs if run.text.strip()), runs[0])
    target.text = text
    for run in runs:
        if run is not target:
            paragraph._p.remove(run._r)


def _header_tables_dynamic(doc):
    header_text = (
        "Proje Adı: [S3_PROJE_ADI]\n"
        "İmar Bilgileri: [S3_IL] İli, [S3_ILCE] İlçesi, [S3_MAHALLE] Mahallesi, "
        "[S3_PAFTA] Pafta, [S3_ADA] Ada, [S3_PARSEL] Parsel"
    )
    changed = 0
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for table in header.tables:
                if table.rows and len(table.rows[0].cells) > 1:
                    cell = table.rows[0].cells[1]
                    if "Proje Adı:" in cell.text and cell.text != header_text:
                        cell.text = header_text
                        changed += 1
    return changed


def update_template(path):
    path = Path(path).resolve()
    doc = Document(path)
    changed = 0
    matched = set()

    for paragraph in doc.paragraphs:
        current = " ".join(paragraph.text.split())
        for prefix, replacement in PARAGRAPH_REPLACEMENTS:
            if current.startswith(prefix):
                _paragraph_text_set(paragraph, replacement)
                matched.add(prefix)
                changed += 1
                break

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == ".":
            paragraph._element.getparent().remove(paragraph._element)
            changed += 1

    for paragraph in doc.paragraphs:
        current = " ".join(paragraph.text.split())
        if current == "Bölgenin stratigrafik kesiti Şekil 5’ te verişmiştir.":
            _paragraph_text_set(
                paragraph,
                "Bölgenin stratigrafik kesiti Şekil 5'te verilmiştir.",
            )
            changed += 1
        elif current == "www.afat.go.tr, MTA 2010":
            _paragraph_text_set(
                paragraph,
                "https://www.afad.gov.tr/; https://yerbilimleri.mta.gov.tr/",
            )
            changed += 1
        elif current.startswith("Çalışma alanında yapılan sondajlarda karotyüzdeleri"):
            _paragraph_text_set(
                paragraph,
                "Çalışma alanındaki karot yüzdeleri proje verilerine göre raporda sunulacaktır.",
            )
            changed += 1
        elif current.startswith("Çalışma alanında killi birimler bulunduğundan"):
            _paragraph_text_set(
                paragraph,
                "Presiyometre deney sonuçları proje verilerine göre raporda sunulacaktır.",
            )
            changed += 1

    changed += _header_tables_dynamic(doc)
    missing = [prefix for prefix, _replacement in PARAGRAPH_REPLACEMENTS if prefix not in matched]
    if missing:
        raise RuntimeError(
            "Sablonda eslesmeyen sabit metinler var: " + ", ".join(missing)
        )

    with tempfile.NamedTemporaryFile(
        prefix=path.stem + "_",
        suffix=".docx",
        dir=path.parent,
        delete=False,
    ) as handle:
        temp_path = Path(handle.name)
    try:
        doc.save(temp_path)
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()
    return changed


def main(argv=None):
    argv = list(sys.argv[1:] if argv is None else argv)
    target = Path(argv[0]) if argv else DEFAULT_TEMPLATE
    changed = update_template(target)
    print(f"Sablon guncellendi: {target} ({changed} degisiklik)")


if __name__ == "__main__":
    main()

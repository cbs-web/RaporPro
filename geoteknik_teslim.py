# Dosya: RaporPro/geoteknik_teslim.py
"""Geoteknik muhendisine verilecek duzenlenebilir veri paketlerini olusturur."""

from __future__ import annotations

import math
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.pagebreak import Break

from jeofizik_sheet_motoru import jeofizik_sheet_rows_to_ss_list


GIRILMEDI = "Girilmedi"
MAVI = "20364B"
ACIK_MAVI = "D9E5F2"
ACIK_GRI = "F3F5F7"
KIRMIZI = "C00000"

_INCE_KENAR = Side(style="thin", color="9AA7B2")
_KALIN_KENAR = Side(style="medium", color=MAVI)
_SAYI_RE = re.compile(r"^[+-]?\d+(?:[.,]\d+)?$")


def _metin(value, bos=GIRILMEDI):
    if value is None:
        return bos
    text = str(value).strip()
    return text if text else bos


def _excel_degeri(value, bos=GIRILMEDI):
    """Sayilari sayi tutar; formulu andiran metinleri guvenli metne cevirir."""
    if value is None or str(value).strip() == "":
        return bos
    if isinstance(value, bool):
        return "Evet" if value else "Hayir"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        if isinstance(value, float) and not math.isfinite(value):
            return bos
        return value
    text = str(value).strip()
    if _SAYI_RE.fullmatch(text):
        number = float(text.replace(",", "."))
        return int(number) if number.is_integer() else number
    if text == "-":
        return text
    if text.startswith(("=", "+", "-", "@")):
        return "'" + text
    return text


def _sayi(value):
    try:
        number = float(str(value).strip().replace(",", "."))
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def _refu_mu(value):
    text = str(value or "").upper().replace(" ", "")
    return text in {"R", "REF", "REFU", "REFÜ"} or "50/" in text


def _n30_hesapla(v30, v45, mevcut=""):
    if str(mevcut or "").strip():
        return "R" if _refu_mu(mevcut) else str(mevcut).strip()
    if _refu_mu(v30) or _refu_mu(v45):
        return "R"
    n30, n45 = _sayi(v30), _sayi(v45)
    if n30 is None or n45 is None:
        return GIRILMEDI
    total = n30 + n45
    return str(int(total)) if total.is_integer() else str(total)


def _dict_degeri(row, *keys):
    if not isinstance(row, dict):
        return ""
    for key in keys:
        if key in row and str(row.get(key) or "").strip():
            return row.get(key)
    return ""


def _liste_degeri(row, index):
    if isinstance(row, (list, tuple)) and index < len(row):
        return row[index]
    return ""


def _sondajlar(veri):
    return [item for item in (veri or {}).get("sondaj", []) or [] if isinstance(item, dict)]


def _satirlar(sondaj, key):
    return [row for row in sondaj.get(key, []) or [] if isinstance(row, (dict, list, tuple))]


def _spt_satiri(row):
    if isinstance(row, dict):
        der = _dict_degeri(row, "der", "derinlik", "depth")
        v15 = _dict_degeri(row, "v15", "15", "n15")
        v30 = _dict_degeri(row, "v30", "30")
        v45 = _dict_degeri(row, "v45", "45")
        n30 = _dict_degeri(row, "n30", "N30", "n30_toplam")
    else:
        der, v15, v30, v45, n30 = (_liste_degeri(row, idx) for idx in range(5))
    n30 = _n30_hesapla(v30, v45, n30)
    durum = "Refü" if any(_refu_mu(value) for value in (v15, v30, v45, n30)) else "Geçerli"
    return der, v15, v30, v45, n30, durum


def _derinlik_anahtari(value):
    number = _sayi(value)
    return f"{number:.2f}" if number is not None else str(value or "").strip().casefold()


def _spt_kaynak_haritasi(sondaj):
    result = {}
    for item in sondaj.get("spt_kaynaklari", []) or []:
        if not isinstance(item, dict):
            continue
        der = _derinlik_anahtari(item.get("derinlik"))
        source = item.get("kaynak") or item.get("kaynak_yolu")
        if der and source:
            result[der] = os.path.basename(str(source))
    return result


def etkin_jeofizik_serimleri(veri):
    """Raporla ayni oncelikle Sheet veya manuel sismik serimleri dondurur."""
    veri = veri if isinstance(veri, dict) else {}
    rows = (veri.get("jeofizik_sheet") or {}).get("rows", [])
    if rows:
        try:
            parsed = jeofizik_sheet_rows_to_ss_list(rows)
        except Exception:
            parsed = []
        parsed = [item for item in parsed if isinstance(item, dict) and item.get("layers")]
        if parsed:
            return parsed
    return [
        item
        for item in (veri.get("jeofizik") or {}).get("ss_list", []) or []
        if isinstance(item, dict) and item.get("layers")
    ]


def sondaj_teslim_ozeti(veri):
    sondajlar = _sondajlar(veri)
    return {
        "sondaj": len(sondajlar),
        "litoloji": sum(len(_satirlar(item, "litoloji")) for item in sondajlar),
        "spt": sum(len(_satirlar(item, "spt")) for item in sondajlar),
        "pmt": sum(len(_satirlar(item, "pmt")) for item in sondajlar),
        "kaya": sum(len(_satirlar(item, "kaya")) for item in sondajlar),
        "numune": sum(len(_satirlar(item, "numuneler")) for item in sondajlar),
    }


def _workbook_sayfasi(workbook, title, headers, widths, tab_color=MAVI):
    sheet = workbook.create_sheet(title)
    sheet.sheet_properties.tabColor = tab_color
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A2"
    sheet.append(headers)
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor=MAVI)
        cell.font = Font(name="Aptos", size=10, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=_KALIN_KENAR, bottom=_KALIN_KENAR, left=_INCE_KENAR, right=_INCE_KENAR)
    sheet.row_dimensions[1].height = 30
    for idx, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + idx)].width = width
    sheet.auto_filter.ref = f"A1:{chr(64 + len(headers))}1"
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.print_title_rows = "1:1"
    sheet.oddFooter.center.text = "RaporPro - Geoteknik Teslim Paketi"
    return sheet


def _workbook_satiri_ekle(sheet, values):
    sheet.append([_excel_degeri(value) for value in values])
    row_idx = sheet.max_row
    fill = PatternFill("solid", fgColor=ACIK_GRI if row_idx % 2 == 0 else "FFFFFF")
    for cell in sheet[row_idx]:
        cell.fill = fill
        cell.font = Font(name="Aptos", size=10, color="1F2933")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(top=_INCE_KENAR, bottom=_INCE_KENAR, left=_INCE_KENAR, right=_INCE_KENAR)


def _sondaj_excel_olustur(veri, path):
    sondajlar = _sondajlar(veri)
    ayarlar = (veri or {}).get("ayarlar", {}) or {}
    workbook = Workbook()
    workbook.remove(workbook.active)
    workbook.properties.title = "RaporPro Sondaj Verileri"
    workbook.properties.subject = "Geoteknik muhendisine teslim icin duzenlenebilir sondaj verileri"

    summary = _workbook_sayfasi(
        workbook,
        "Sondaj Özeti",
        [
            "Sondaj No", "Derinlik (m)", "Kot (m)", "Enlem", "Boylam",
            "Başlangıç Tarihi", "Bitiş Tarihi", "YASS İlk (m)", "YASS İlk Tarih",
            "YASS Son (m)", "YASS Son Tarih", "Sondaj Türü", "Delgi Çapı",
            "Litoloji Satırı", "SPT", "PMT", "Kaya/Karot", "Numune",
        ],
        [14, 12, 11, 15, 15, 15, 15, 12, 15, 12, 15, 13, 12, 13, 9, 9, 12, 10],
    )
    lith = _workbook_sayfasi(
        workbook, "Litoloji", ["Sondaj No", "Başlangıç (m)", "Bitiş (m)", "Nihai Litoloji Tanımı"],
        [15, 15, 15, 48], "6B8E23",
    )
    spt = _workbook_sayfasi(
        workbook, "SPT", ["Sondaj No", "Derinlik (m)", "0-15", "15-30", "30-45", "N30", "Durum", "Kaynak Fotoğraf"],
        [15, 14, 10, 10, 10, 10, 12, 28], "C55A11",
    )
    pmt = _workbook_sayfasi(
        workbook, "Presiyometre", ["Sondaj No", "Derinlik (m)", "Em (kg/cm²)", "Pl (kg/cm²)", "Em/Pl"],
        [15, 16, 18, 18, 14], "8064A2",
    )
    rock = _workbook_sayfasi(
        workbook, "Kaya ve Karot", ["Sondaj No", "Derinlik Aralığı (m)", "TCR (%)", "SCR (%)", "RQD (%)"],
        [15, 22, 15, 15, 15], "7F6000",
    )
    samples = _workbook_sayfasi(
        workbook, "Numuneler", ["Sondaj No", "Derinlik/Aralık (m)", "Numune Türü/No"],
        [15, 24, 32], "008C95",
    )

    for index, sondaj in enumerate(sondajlar, start=1):
        sondaj_no = sondaj.get("no") or f"SK-{index}"
        _workbook_satiri_ekle(summary, [
            sondaj_no, sondaj.get("der"), sondaj.get("k"), sondaj.get("y"), sondaj.get("x"),
            sondaj.get("bas_tar"), sondaj.get("bit_tar"), sondaj.get("yass_d1"), sondaj.get("yass_t1"),
            sondaj.get("yass_d2"), sondaj.get("yass_t2"),
            ayarlar.get("sondaj_turu"), ayarlar.get("delgi_capi"),
            len(_satirlar(sondaj, "litoloji")), len(_satirlar(sondaj, "spt")),
            len(_satirlar(sondaj, "pmt")), len(_satirlar(sondaj, "kaya")),
            len(_satirlar(sondaj, "numuneler")),
        ])

        for row in _satirlar(sondaj, "litoloji"):
            if isinstance(row, dict):
                values = (
                    _dict_degeri(row, "bas", "başlangıç", "baslangic", "start"),
                    _dict_degeri(row, "bit", "bitiş", "bitis", "end"),
                    _dict_degeri(row, "tanim", "tanım", "litoloji", "description"),
                )
            else:
                values = tuple(_liste_degeri(row, idx) for idx in range(3))
            _workbook_satiri_ekle(lith, [sondaj_no, *values])

        source_map = _spt_kaynak_haritasi(sondaj)
        for row in _satirlar(sondaj, "spt"):
            der, v15, v30, v45, n30, durum = _spt_satiri(row)
            _workbook_satiri_ekle(
                spt,
                [sondaj_no, der, v15, v30, v45, n30, durum, source_map.get(_derinlik_anahtari(der), GIRILMEDI)],
            )

        for row in _satirlar(sondaj, "pmt"):
            if isinstance(row, dict):
                der = _dict_degeri(row, "der", "derinlik")
                em = _dict_degeri(row, "em", "Em", "EM")
                pl = _dict_degeri(row, "pl", "Pl", "PL")
            else:
                der, em, pl = (_liste_degeri(row, idx) for idx in range(3))
            em_num, pl_num = _sayi(em), _sayi(pl)
            ratio = GIRILMEDI if em_num is None or pl_num in (None, 0) else round(em_num / pl_num, 2)
            _workbook_satiri_ekle(pmt, [sondaj_no, der, em, pl, ratio])

        for row in _satirlar(sondaj, "kaya"):
            if isinstance(row, dict):
                values = (
                    _dict_degeri(row, "der", "derinlik", "aralik"),
                    _dict_degeri(row, "tcr", "TCR"),
                    _dict_degeri(row, "scr", "SCR"),
                    _dict_degeri(row, "rqd", "RQD"),
                )
            else:
                values = tuple(_liste_degeri(row, idx) for idx in range(4))
            _workbook_satiri_ekle(rock, [sondaj_no, *values])

        for row in _satirlar(sondaj, "numuneler"):
            if isinstance(row, dict):
                values = (
                    _dict_degeri(row, "der", "derinlik", "aralik"),
                    _dict_degeri(row, "tur", "tür", "no", "numune"),
                )
            else:
                values = tuple(_liste_degeri(row, idx) for idx in range(2))
            _workbook_satiri_ekle(samples, [sondaj_no, *values])

    workbook.save(path)


def _docx_hucre_dolgu(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _docx_hucre_genislik(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))


def _docx_satir_bolunmesin(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _docx_baslik_tekrarla(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _docx_tablo(document, headers, rows, widths, font_size=8.5):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    _docx_baslik_tekrarla(header)
    _docx_satir_bolunmesin(header)
    for idx, (cell, text) in enumerate(zip(header.cells, headers)):
        _docx_hucre_dolgu(cell, MAVI)
        _docx_hucre_genislik(cell, widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(text))
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, values in enumerate(rows, start=1):
        row = table.add_row()
        _docx_satir_bolunmesin(row)
        for idx, (cell, value) in enumerate(zip(row.cells, values)):
            _docx_hucre_genislik(cell, widths[idx])
            if row_index % 2 == 0:
                _docx_hucre_dolgu(cell, ACIK_GRI)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(_metin(value))
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
    return table


def _docx_belge_hazirla(title, veri, landscape=True):
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(5)
    run = title_p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(MAVI)
    kunye = (veri or {}).get("kunye", {}) or {}
    project = _metin(kunye.get("sahibi"), "Adsız proje")
    location = " / ".join(
        item for item in (_metin(kunye.get("mah"), ""), _metin(kunye.get("ilce"), ""), _metin(kunye.get("il"), "")) if item
    )
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run(f"{project}" + (f"  |  {location}" if location else ""))
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(88, 99, 110)
    return document


def _docx_bolum_basligi(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(MAVI)


def _sondaj_docx_olustur(veri, path):
    sondajlar = _sondajlar(veri)
    ayarlar = (veri or {}).get("ayarlar", {}) or {}
    document = _docx_belge_hazirla("Sondaj Özet Tabloları", veri, landscape=True)

    general_rows = []
    count_rows = []
    water_rows = []
    for index, sondaj in enumerate(sondajlar, start=1):
        no = sondaj.get("no") or f"SK-{index}"
        general_rows.append([
            no, sondaj.get("der"), sondaj.get("k"), sondaj.get("y"), sondaj.get("x"),
            sondaj.get("bas_tar"), sondaj.get("bit_tar"), ayarlar.get("sondaj_turu"), ayarlar.get("delgi_capi"),
        ])
        count_rows.append([
            no, len(_satirlar(sondaj, "litoloji")), len(_satirlar(sondaj, "spt")),
            len(_satirlar(sondaj, "pmt")), len(_satirlar(sondaj, "kaya")), len(_satirlar(sondaj, "numuneler")),
        ])
        water_rows.append([
            no, sondaj.get("yass_d1"), sondaj.get("yass_t1"), sondaj.get("yass_d2"), sondaj.get("yass_t2"),
        ])

    _docx_bolum_basligi(document, "1. Genel Sondaj Bilgileri")
    _docx_tablo(
        document,
        ["Sondaj", "Derinlik (m)", "Kot (m)", "Enlem", "Boylam", "Başlangıç", "Bitiş", "Tür", "Delgi Çapı"],
        general_rows,
        [2.0, 2.2, 1.8, 3.2, 3.2, 2.7, 2.7, 2.4, 2.2],
    )
    _docx_bolum_basligi(document, "2. Veri ve Deney Sayıları")
    _docx_tablo(
        document,
        ["Sondaj", "Litoloji Satırı", "SPT", "Presiyometre", "Kaya/Karot", "Numune"],
        count_rows,
        [3.0, 4.0, 3.0, 4.0, 4.0, 3.5],
        font_size=9,
    )
    _docx_bolum_basligi(document, "3. Yeraltı Suyu Gözlemleri")
    _docx_tablo(
        document,
        ["Sondaj", "İlk Ölçüm (m)", "İlk Ölçüm Tarihi", "Son Ölçüm (m)", "Son Ölçüm Tarihi"],
        water_rows,
        [3.0, 4.0, 5.0, 4.0, 5.0],
        font_size=9,
    )
    document.save(path)


def sondaj_veri_paketi_olustur(veri, hedef_klasor):
    if not _sondajlar(veri):
        raise ValueError("Sondaj veri paketi için sondaj kaydı bulunamadı.")
    os.makedirs(hedef_klasor, exist_ok=True)
    xlsx_path = os.path.join(hedef_klasor, "Sondaj_Verileri.xlsx")
    docx_path = os.path.join(hedef_klasor, "Sondaj_Ozet_Tablolari.docx")
    _sondaj_excel_olustur(veri, xlsx_path)
    _sondaj_docx_olustur(veri, docx_path)
    return [xlsx_path, docx_path]


_JEO_PARAMETRELER = (
    ("Kalınlık (m)", "h"),
    ("Vp (m/s)", "vp"),
    ("Vs (m/s)", "vs"),
    ("Yoğunluk (g/cm³)", "rho"),
    ("Poisson Oranı", "nu"),
    ("Elastisite Mod. (kg/cm²)", "E"),
    ("Kayma Mod. (kg/cm²)", "G"),
    ("Bulk Mod. (kg/cm²)", "K"),
)


def _jeo_degeri(key, value, son_tabaka=False):
    if key == "h" and son_tabaka:
        return "-"
    if value is None or str(value).strip() == "":
        return "-"
    if key in {"E", "G", "K"}:
        number = _sayi(value)
        return int(round(number)) if number is not None else str(value).strip()
    number = _sayi(value)
    if number is None:
        return str(value).strip()
    return int(number) if number.is_integer() else round(number, 3)


def _jeo_gruplari(veri):
    serimler = etkin_jeofizik_serimleri(veri)
    return [serimler[start:start + 3] for start in range(0, len(serimler), 3)]


def _jeo_excel_olustur(veri, path):
    groups = _jeo_gruplari(veri)
    if not groups:
        raise ValueError("Jeofizik parametre tablosu için sismik serim verisi bulunamadı.")
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Jeofizik Parametreleri"
    sheet.sheet_properties.tabColor = "8064A2"
    sheet.sheet_view.showGridLines = False
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.paperSize = sheet.PAPERSIZE_A4
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.column_dimensions["A"].width = 31
    current_row = 1
    max_end_column = 1

    for group_index, group in enumerate(groups):
        if group_index:
            current_row += 2
            sheet.row_breaks.append(Break(id=current_row - 1))
        start_row = current_row
        sheet.cell(current_row, 1, "Parametre")
        sheet.merge_cells(start_row=current_row, start_column=1, end_row=current_row + 1, end_column=1)
        column = 2
        flat_layers = []
        for serim_index, serim in enumerate(group, start=1):
            layers = [layer for layer in serim.get("layers", []) or [] if isinstance(layer, dict)]
            if not layers:
                continue
            start_column = column
            name = serim.get("ad") or f"Serim {group_index * 3 + serim_index}"
            for layer_index, layer in enumerate(layers, start=1):
                flat_layers.append((layer, layer_index == len(layers)))
                sheet.cell(current_row + 1, column, f"Tab. {layer_index}")
                sheet.column_dimensions[get_column_letter(column)].width = 13
                column += 1
            sheet.cell(current_row, start_column, str(name))
            if column - start_column > 1:
                sheet.merge_cells(start_row=current_row, start_column=start_column, end_row=current_row, end_column=column - 1)

        if not flat_layers:
            continue
        end_column = 1 + len(flat_layers)
        max_end_column = max(max_end_column, end_column)
        for row in range(current_row, current_row + 2):
            for col in range(1, end_column + 1):
                cell = sheet.cell(row, col)
                cell.fill = PatternFill("solid", fgColor=MAVI)
                cell.font = Font(name="Aptos", size=10, bold=True, color="FFFFFF")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = Border(top=_KALIN_KENAR, bottom=_KALIN_KENAR, left=_INCE_KENAR, right=_INCE_KENAR)
        current_row += 2
        for parameter_index, (label, key) in enumerate(_JEO_PARAMETRELER):
            sheet.cell(current_row, 1, label)
            for layer_index, (layer, is_last) in enumerate(flat_layers, start=2):
                sheet.cell(current_row, layer_index, _excel_degeri(_jeo_degeri(key, layer.get(key), is_last), bos="-"))
            for col in range(1, end_column + 1):
                cell = sheet.cell(current_row, col)
                cell.fill = PatternFill("solid", fgColor=ACIK_GRI if parameter_index % 2 else "FFFFFF")
                cell.font = Font(
                    name="Aptos", size=10, bold=col == 1,
                    color=KIRMIZI if key in {"vp", "vs"} else "1F2933",
                )
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = Border(top=_INCE_KENAR, bottom=_INCE_KENAR, left=_INCE_KENAR, right=_INCE_KENAR)
            current_row += 1
        sheet.row_dimensions[start_row].height = 24
        sheet.row_dimensions[start_row + 1].height = 24
    sheet.print_area = f"A1:{sheet.cell(current_row - 1, max_end_column).coordinate}"
    sheet.oddFooter.center.text = "RaporPro - Jeofizik Parametreleri"
    workbook.save(path)


def _jeo_docx_hucre_yaz(cell, text, bold=False, color="1F2933", font_size=8.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(color)


def _jeo_docx_tablosu(document, group, group_offset):
    flat_layers = []
    serim_groups = []
    for index, serim in enumerate(group, start=1):
        layers = [layer for layer in serim.get("layers", []) or [] if isinstance(layer, dict)]
        if not layers:
            continue
        name = serim.get("ad") or f"Serim {group_offset + index}"
        serim_groups.append((name, len(layers)))
        flat_layers.extend((layer, layer_index == len(layers) - 1) for layer_index, layer in enumerate(layers))
    if not flat_layers:
        return None

    table = document.add_table(rows=2, cols=1 + len(flat_layers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].merge(table.rows[1].cells[0])
    _docx_hucre_dolgu(table.rows[0].cells[0], MAVI)
    _jeo_docx_hucre_yaz(table.rows[0].cells[0], "Parametre", bold=True, color="FFFFFF")
    _docx_hucre_genislik(table.rows[0].cells[0], 4.3)
    layer_width = max(1.35, min(2.5, 21.5 / len(flat_layers)))
    column = 1
    for name, count in serim_groups:
        start = table.rows[0].cells[column]
        if count > 1:
            start = start.merge(table.rows[0].cells[column + count - 1])
        _docx_hucre_dolgu(start, MAVI)
        _docx_hucre_genislik(start, layer_width * count)
        _jeo_docx_hucre_yaz(start, name, bold=True, color="FFFFFF")
        for layer_no in range(1, count + 1):
            cell = table.rows[1].cells[column]
            _docx_hucre_dolgu(cell, MAVI)
            _docx_hucre_genislik(cell, layer_width)
            _jeo_docx_hucre_yaz(cell, f"Tab. {layer_no}", bold=True, color="FFFFFF")
            column += 1
    for header in table.rows[:2]:
        _docx_baslik_tekrarla(header)
        _docx_satir_bolunmesin(header)

    for parameter_index, (label, key) in enumerate(_JEO_PARAMETRELER):
        row = table.add_row()
        _docx_satir_bolunmesin(row)
        if parameter_index % 2:
            for cell in row.cells:
                _docx_hucre_dolgu(cell, ACIK_GRI)
        _docx_hucre_genislik(row.cells[0], 4.3)
        _jeo_docx_hucre_yaz(
            row.cells[0], label, bold=True,
            color=KIRMIZI if key in {"vp", "vs"} else "1F2933",
        )
        for index, (layer, is_last) in enumerate(flat_layers, start=1):
            _docx_hucre_genislik(row.cells[index], layer_width)
            _jeo_docx_hucre_yaz(
                row.cells[index], _jeo_degeri(key, layer.get(key), is_last),
                color=KIRMIZI if key in {"vp", "vs"} else "1F2933",
            )
    return table


def _jeo_docx_olustur(veri, path):
    groups = _jeo_gruplari(veri)
    if not groups:
        raise ValueError("Jeofizik parametre tablosu için sismik serim verisi bulunamadı.")
    document = _docx_belge_hazirla("Jeofizik Parametreleri", veri, landscape=True)
    for group_index, group in enumerate(groups):
        if group_index:
            document.add_page_break()
        names = [str(item.get("ad") or f"Serim {group_index * 3 + idx + 1}") for idx, item in enumerate(group)]
        _docx_bolum_basligi(document, " - ".join(names))
        _jeo_docx_tablosu(document, group, group_index * 3)
    document.save(path)


def jeofizik_parametre_paketi_olustur(veri, hedef_klasor):
    if not etkin_jeofizik_serimleri(veri):
        raise ValueError("Jeofizik parametre paketi için sismik serim verisi bulunamadı.")
    os.makedirs(hedef_klasor, exist_ok=True)
    xlsx_path = os.path.join(hedef_klasor, "Jeofizik_Parametreleri.xlsx")
    docx_path = os.path.join(hedef_klasor, "Jeofizik_Parametreleri.docx")
    _jeo_excel_olustur(veri, xlsx_path)
    _jeo_docx_olustur(veri, docx_path)
    return [xlsx_path, docx_path]


def ham_veri_kaynaklari(veri, proje_dosyasi=""):
    """Proje klasorleri, bagli Excel'ler ve SPT kaynak fotograflarini listeler."""
    veri = veri if isinstance(veri, dict) else {}
    sources = []
    project_root = Path(proje_dosyasi).expanduser().resolve().parent if proje_dosyasi else None
    folder_candidates = (
        ("Evraklar", "EVRAKLAR"),
        ("Laboratuvar", "LAB"),
        ("Jeofizik", "JEOFİZİK"),
        ("Jeofizik", "JEOFIZIK"),
        ("Presiyometre", "Presiyometre"),
    )
    if project_root and project_root.is_dir():
        for category, folder_name in folder_candidates:
            candidate = project_root / folder_name
            if candidate.is_dir():
                sources.append((category, str(candidate)))

    def mevcut_dosya(path_value):
        if not path_value:
            return None
        candidate = Path(str(path_value)).expanduser()
        if not candidate.is_absolute() and project_root:
            candidate = project_root / candidate
        return candidate.resolve() if candidate.is_file() else None

    files = veri.get("dosyalar", {}) or {}
    for category, key in (("Laboratuvar", "lab_excel_path"), ("Jeofizik", "jeo_excel_path")):
        path = mevcut_dosya(files.get(key))
        if path:
            sources.append((category, str(path)))

    for sondaj in _sondajlar(veri):
        for item in sondaj.get("spt_kaynaklari", []) or []:
            if not isinstance(item, dict):
                continue
            path = mevcut_dosya(item.get("kaynak_yolu"))
            if path:
                sources.append(("SPT_Fotograflari", str(path)))

    unique = []
    seen = set()
    for category, source in sources:
        key = os.path.normcase(os.path.abspath(source))
        if key in seen:
            continue
        seen.add(key)
        unique.append((category, source))
    return unique


def _benzersiz_hedef(path):
    if not path.exists():
        return path
    counter = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{counter}{path.suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def ham_verileri_kopyala(kaynaklar, hedef_klasor):
    """Ham kaynaklari kategori klasorlerine kopyalar; dosya bazinda sonuc dondurur."""
    target_root = Path(hedef_klasor).resolve()
    target_root.mkdir(parents=True, exist_ok=True)
    copied = []
    warnings = []
    copied_sources = set()

    for category, source_text in kaynaklar or []:
        source = Path(source_text)
        if not source.exists():
            warnings.append(f"Kaynak bulunamadı: {source}")
            continue
        candidates = [source] if source.is_file() else [item for item in source.rglob("*") if item.is_file()]
        for candidate in candidates:
            try:
                candidate_resolved = candidate.resolve()
                if target_root == candidate_resolved or target_root in candidate_resolved.parents:
                    continue
                source_key = os.path.normcase(str(candidate_resolved))
                if source_key in copied_sources:
                    continue
                copied_sources.add(source_key)
                category_root = target_root / str(category)
                if source.is_dir():
                    relative = candidate.relative_to(source)
                    target = category_root / source.name / relative
                else:
                    target = category_root / candidate.name
                target.parent.mkdir(parents=True, exist_ok=True)
                target = _benzersiz_hedef(target)
                shutil.copy2(candidate, target)
                copied.append(str(target))
            except Exception as exc:
                warnings.append(f"{candidate}: {exc}")
    return copied, warnings


__all__ = [
    "GIRILMEDI",
    "etkin_jeofizik_serimleri",
    "ham_veri_kaynaklari",
    "ham_verileri_kopyala",
    "jeofizik_parametre_paketi_olustur",
    "sondaj_teslim_ozeti",
    "sondaj_veri_paketi_olustur",
]

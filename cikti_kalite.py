# Dosya: RaporPro/cikti_kalite.py
"""Uretilmis RaporPro dosyalari icin bagimsiz cikti kalite denetimi."""

from __future__ import annotations

import datetime
import hashlib
import json
import os
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import fitz
from docx import Document
from openpyxl import load_workbook
from PIL import Image

from yardimcilar import atomic_json_dump


CIKTI_KALITE_SURUMU = 1
_ETIKET_RE = re.compile(
    r"\[(?=[^\]\r\n]{0,80}(?:[A-ZÇĞİÖŞÜ_]))[^\]\r\n]{2,80}\]"
    r"|RESIM:[A-Za-z0-9_ÇĞİÖŞÜçğıöşü:-]+"
)
_VOLATILE_SETTING_PREFIXES = (
    "varsayilan_cikti_",
    "cikti_merkezi_",
    "log_export_",
)
_VOLATILE_SETTING_KEYS = {
    "yedek_sayisi",
    "surum_gecmisi_sayisi",
    "spt_guven_esigi",
    "spt_auto_pro",
}


def _bulgu(report, level, label, detail, path=""):
    report.setdefault("findings", []).append({
        "id": f"cikti.{len(report.get('findings', [])) + 1}",
        "level": level,
        "category": "Çıktı kalitesi",
        "label": label,
        "detail": detail,
        "path": str(path or ""),
        "target": "cikti",
        "suggestion": "Çıktıyı güncel proje verisiyle yeniden oluşturun." if level != "info" else "",
    })


def _raporu_tamamla(report):
    findings = report.get("findings", [])
    errors = [item for item in findings if item.get("level") == "error"]
    warnings = [item for item in findings if item.get("level") == "warning"]
    report["errors"] = errors
    report["warnings"] = warnings
    report["info"] = [item for item in findings if item.get("level") == "info"]
    report["counts"] = {
        "error": len(errors),
        "warning": len(warnings),
        "info": len(report["info"]),
    }
    report["state"] = "HATA" if errors else ("UYARI" if warnings else "TEMİZ")
    return report


def _docx_paragraflari(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from container.paragraphs
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def _docx_denetle(path, report):
    try:
        with zipfile.ZipFile(path) as archive:
            bad_member = archive.testzip()
            if bad_member:
                _bulgu(report, "error", "Bozuk Word paketi", f"DOCX içindeki {bad_member} okunamıyor.", path)
                return
            media_count = sum(1 for name in archive.namelist() if name.startswith("word/media/"))
    except Exception as exc:
        _bulgu(report, "error", "Word dosyası", f"DOCX paketi açılamadı: {exc}", path)
        return

    try:
        document = Document(path)
        text = "\n".join(paragraph.text or "" for paragraph in _docx_paragraflari(document))
        tags = sorted(set(_ETIKET_RE.findall(text)))
        if tags:
            preview = ", ".join(tags[:8])
            suffix = " ..." if len(tags) > 8 else ""
            _bulgu(
                report,
                "error",
                "Tamamlanmamış Word etiketi",
                f"Çıktıda {len(tags)} işlenmemiş etiket kaldı: {preview}{suffix}",
                path,
            )
        meaningful_text = re.sub(r"\s+", "", text)
        if not meaningful_text and not document.tables and media_count == 0:
            _bulgu(report, "error", "Boş Word çıktısı", "Word belgesinde metin, tablo veya görsel bulunamadı.", path)
        if not document.sections:
            _bulgu(report, "error", "Word sayfa yapısı", "Word belgesinde sayfa bölümü bulunamadı.", path)
        for idx, section in enumerate(document.sections, start=1):
            if int(section.page_width or 0) <= 0 or int(section.page_height or 0) <= 0:
                _bulgu(report, "error", "Word sayfa boyutu", f"{idx}. bölümün sayfa boyutu geçersiz.", path)
    except Exception as exc:
        _bulgu(report, "error", "Word içeriği", f"Word içeriği okunamadı: {exc}", path)


def _pdf_denetle(path, report):
    document = None
    try:
        document = fitz.open(path)
        if document.page_count <= 0:
            _bulgu(report, "error", "Boş PDF", "PDF dosyasında sayfa bulunamadı.", path)
            return
        blank_pages = []
        for page_no, page in enumerate(document, start=1):
            text = (page.get_text("text") or "").strip()
            has_images = bool(page.get_images(full=True))
            has_drawings = bool(page.get_drawings())
            if not text and not has_images and not has_drawings:
                blank_pages.append(page_no)
            if page.rect.width <= 0 or page.rect.height <= 0:
                _bulgu(report, "error", "PDF sayfa boyutu", f"{page_no}. sayfanın boyutu geçersiz.", path)
        if blank_pages:
            pages = ", ".join(str(item) for item in blank_pages[:12])
            _bulgu(report, "warning", "Boş PDF sayfası", f"İçeriksiz görünen sayfalar: {pages}", path)
    except Exception as exc:
        _bulgu(report, "error", "PDF dosyası", f"PDF açılamadı: {exc}", path)
    finally:
        if document is not None:
            document.close()


def _gorsel_denetle(path, report):
    try:
        with Image.open(path) as image:
            width, height = image.size
            image.load()
            if width < 600 or height < 400:
                _bulgu(
                    report,
                    "warning",
                    "Düşük görsel çözünürlüğü",
                    f"Görsel {width}x{height} px; rapor çıktısı için küçük olabilir.",
                    path,
                )
            gray = image.convert("L")
            gray.thumbnail((96, 96))
            low, high = gray.getextrema()
            if high - low <= 2:
                _bulgu(report, "warning", "Boş görsel", "Görsel neredeyse tamamen tek renk görünüyor.", path)
    except Exception as exc:
        _bulgu(report, "error", "Görsel dosyası", f"Görsel açılamadı: {exc}", path)


def _xlsx_denetle(path, report):
    workbook = None
    try:
        workbook = load_workbook(path, read_only=True, data_only=False)
        if not workbook.sheetnames:
            _bulgu(report, "error", "Boş Excel çıktısı", "Excel dosyasında çalışma sayfası bulunamadı.", path)
            return
        if not any(
            sheet.max_row > 1 or sheet.max_column > 1 or sheet.cell(1, 1).value not in (None, "")
            for sheet in workbook.worksheets
        ):
            _bulgu(report, "warning", "Boş Excel çıktısı", "Excel sayfalarında veri bulunamadı.", path)
    except Exception as exc:
        _bulgu(report, "error", "Excel dosyası", f"Excel açılamadı: {exc}", path)
    finally:
        if workbook is not None:
            workbook.close()


def _svg_denetle(path, report):
    try:
        root = ElementTree.parse(path).getroot()
        if not list(root) and not (root.text or "").strip():
            _bulgu(report, "warning", "Boş SVG", "SVG dosyasında çizim öğesi bulunamadı.", path)
    except Exception as exc:
        _bulgu(report, "error", "SVG dosyası", f"SVG açılamadı: {exc}", path)


def dosya_kalite_raporu(path):
    report = {"path": str(path or ""), "findings": []}
    if not path or not os.path.isfile(path):
        _bulgu(report, "error", "Eksik çıktı", "Çıktı dosyası bulunamadı.", path)
        return _raporu_tamamla(report)
    if os.path.getsize(path) <= 0:
        _bulgu(report, "error", "Boş çıktı dosyası", "Dosya boyutu sıfır.", path)
        return _raporu_tamamla(report)

    ext = Path(path).suffix.casefold()
    if ext == ".docx":
        _docx_denetle(path, report)
    elif ext == ".pdf":
        _pdf_denetle(path, report)
    elif ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        _gorsel_denetle(path, report)
    elif ext in {".xlsx", ".xlsm"}:
        _xlsx_denetle(path, report)
    elif ext == ".svg":
        _svg_denetle(path, report)
    elif os.path.getsize(path) < 128:
        _bulgu(report, "warning", "Çok küçük çıktı", "Dosya beklenenden çok küçük görünüyor.", path)
    return _raporu_tamamla(report)


def _parmak_izi_verisi(veri):
    payload = dict(veri or {})
    payload.pop("proje_durumu", None)
    settings = dict(payload.get("ayarlar") or {})
    for key in list(settings):
        if key in _VOLATILE_SETTING_KEYS or key.startswith(_VOLATILE_SETTING_PREFIXES):
            settings.pop(key, None)
    payload["ayarlar"] = settings
    return payload


def proje_parmak_izi(veri):
    text = json.dumps(
        _parmak_izi_verisi(veri),
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    )
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()


def dosya_parmak_izi(path, chunk_size=1024 * 1024):
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def cikti_dosyalari_denetle(paths, veri=None):
    report = {
        "quality_schema": CIKTI_KALITE_SURUMU,
        "created_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "project_fingerprint": proje_parmak_izi(veri) if veri is not None else "",
        "files": [],
        "findings": [],
    }
    seen = set()
    for path in paths or []:
        key = os.path.normcase(os.path.abspath(str(path))) if path else ""
        if not key or key in seen:
            continue
        seen.add(key)
        file_report = dosya_kalite_raporu(path)
        report["files"].append({
            "path": str(path),
            "state": file_report["state"],
            "size": os.path.getsize(path) if os.path.isfile(path) else 0,
            "sha256": dosya_parmak_izi(path) if os.path.isfile(path) else "",
        })
        report["findings"].extend(file_report["findings"])
    if not report["files"]:
        _bulgu(report, "warning", "Çıktı listesi", "Denetlenecek çıktı dosyası bulunamadı.")
    return _raporu_tamamla(report)


def kalite_manifestosu_yaz(path, report, veri=None):
    payload = dict(report or {})
    payload["quality_schema"] = CIKTI_KALITE_SURUMU
    payload["project_fingerprint"] = proje_parmak_izi(veri) if veri is not None else payload.get("project_fingerprint", "")
    atomic_json_dump(payload, path, indent=2, ensure_ascii=False)
    return path


def kalite_manifestosu_dogrula(path, veri=None):
    report = {"path": str(path or ""), "findings": []}
    try:
        with open(path, "r", encoding="utf-8") as handle:
            manifest = json.load(handle)
    except Exception as exc:
        _bulgu(report, "error", "Kalite manifestosu", f"Manifesto okunamadı: {exc}", path)
        return _raporu_tamamla(report)

    if veri is not None and manifest.get("project_fingerprint") != proje_parmak_izi(veri):
        _bulgu(
            report,
            "warning",
            "Güncelliğini yitirmiş çıktı",
            "Proje verisi, bu çıktılar denetlendikten sonra değişmiş.",
            path,
        )
    for item in manifest.get("files", []):
        file_path = item.get("path")
        if not file_path or not os.path.isfile(file_path):
            _bulgu(report, "error", "Silinmiş çıktı", f"Manifestodaki dosya bulunamadı: {file_path or '-'}", file_path)
        elif item.get("sha256") and dosya_parmak_izi(file_path) != item.get("sha256"):
            _bulgu(report, "warning", "Değiştirilmiş çıktı", f"Dosya denetimden sonra değiştirilmiş: {file_path}", file_path)
    return _raporu_tamamla(report)


__all__ = [
    "CIKTI_KALITE_SURUMU",
    "cikti_dosyalari_denetle",
    "dosya_kalite_raporu",
    "dosya_parmak_izi",
    "kalite_manifestosu_dogrula",
    "kalite_manifestosu_yaz",
    "proje_parmak_izi",
]

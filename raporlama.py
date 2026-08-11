# Dosya: RaporPro/raporlama.py
import os
import datetime
import re
import time
import tempfile
import unicodedata
import copy
from types import SimpleNamespace
from tkinter import filedialog
import pandas as pd
import traceback
import numpy as np
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from yardimcilar import (
    atomic_docx_save,
    docx_metadata_nortrle,
    temizle_baslik,
    zemin_sinifi_cevir,
    safe_float,
)
from motor import GeoEngine
from hidrojeoloji_raporu import (
    hidrojeoloji_durum_metni,
    hidrojeoloji_word_paragrafini_uygula,
)
from jeoloji_raporu import (
    jeoloji_kisa_formasyon_metni,
    jeoloji_rapor_bloklari,
)
from jeofizik_sheet_motoru import jeofizik_sheet_rows_to_ss_list, jeofizik_ss_koordinatlarini_koru
from performans import log_exception, perf_log, perf_timer
from rapor_etiketleri import DUZELTME_ETIKET_ADLARI, DUZELTME_ETIKET_GRUPLARI
from rapor_sablonu import proje_rapor_sablon_profili, rapor_sablonu_durumu
from rapor_revizyon import revizyon_isaretleri_ekle
from raporlama_deger import clean_val, fmt_jeo, jeofizik_vp_layers_sadelestir, read_table_file
from raporlama_arazi import arazi_deney_rapor_verileri, arazi_deney_word_bolumlerini_uygula
from rapor_parsel_bilgileri import (
    rapor_metin_degerleri,
    rapor_proje_adi,
)
from raporlama_litoloji import (
    INCE_DANELILER,
    IRI_DANELILER,
    LITOLOJI_DAGILIM_BIRIMLERI,
    litoloji_dagilim_birimi,
    litoloji_dagilim_paragraflari,
)
from raporlama_tablo import (
    TABLE_ALT_FILL,
    TABLE_BORDER_COLOR,
    TABLE_HEADER_FILL,
    TABLE_LABEL_FILL,
    TABLE_TEXT_COLOR,
    apply_report_table_style,
    create_word_table,
    keep_table_together,
    repeat_table_header,
    set_cell_border,
    set_cell_margins,
    set_cell_shading,
    set_cell_text_clean,
    set_cell_width,
    set_table_fit_to_window,
    set_vertical_cell_alignment,
    style_cell_text,
    style_report_table_row,
)
from raporlama_parsel import (
    rapor_kosullu_bolumlerini_uygula,
    rapor_sabit_tablolarini_uygula,
)


def rapor_baglami_olustur(kaynak, *, word_path=None, veri=None, durum_bildir=False):
    """Rapor motoruna yalniz ihtiyac duydugu alanlarin anlik goruntusunu ver."""

    kaynak_veri = getattr(kaynak, "veri", {}) if veri is None else veri
    status_callback = getattr(kaynak, "set_status", None) if durum_bildir else None
    if not callable(status_callback):
        status_callback = lambda *_args, **_kwargs: None
    return SimpleNamespace(
        word_path=word_path if word_path is not None else getattr(kaynak, "word_path", None),
        veri=copy.deepcopy(kaynak_veri),
        jeo_excel_path=getattr(kaynak, "jeo_excel_path", None),
        lab_excel_path=getattr(kaynak, "lab_excel_path", None),
        img_yer=getattr(kaynak, "img_yer", None),
        img_tkgm=getattr(kaynak, "img_tkgm", None),
        img_pga=getattr(kaynak, "img_pga", None),
        img_mjh=getattr(kaynak, "img_mjh", None),
        word_img_jeofizik=getattr(kaynak, "word_img_jeofizik", None),
        word_img_sondaj=getattr(kaynak, "word_img_sondaj", None),
        set_status=status_callback,
    )


def _log_silent(name, exc):
    log_exception(f"raporlama.{name}", exc_value=exc)

def lab_sheet_satirlari(app_instance):
    try:
        rows = app_instance.veri.get("lab_sheet", {}).get("rows", [])
    except Exception:
        return []
    clean_rows = []
    for row in rows or []:
        cells = ["" if cell is None else str(cell) for cell in (row or [])]
        while cells and not str(cells[-1]).strip():
            cells.pop()
        clean_rows.append(cells)
    while clean_rows and not any(str(cell).strip() for cell in clean_rows[-1]):
        clean_rows.pop()
    return clean_rows

def lab_sheet_verisi_var_mi(app_instance):
    return any(any(str(cell).strip() for cell in row) for row in lab_sheet_satirlari(app_instance))

def duzeltme_etiketleri_temizle(tags):
    selected = []
    seen = set()
    for tag in tags or []:
        clean = str(tag or "").strip()
        if not clean or clean in seen:
            continue
        selected.append(clean)
        seen.add(clean)
    return selected

def duzeltme_etiket_sablonu_olustur(tags, template_path):
    selected = duzeltme_etiketleri_temizle(tags)
    if not selected:
        raise ValueError("En az bir etiket seçilmelidir.")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    title = doc.add_heading("RAPOR DÜZELTME ETİKET ÇIKTISI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph(
        "Bu dosya yalnızca seçilen etiketlerin güncel proje verileriyle yeniden üretilmiş halidir."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for tag in selected:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        heading.paragraph_format.keep_with_next = True
        heading_run = heading.add_run(DUZELTME_ETIKET_ADLARI.get(tag, tag))
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        placeholder = doc.add_paragraph(tag)
        placeholder.paragraph_format.keep_with_next = False
        if "RESIM" in tag:
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(template_path)
    return selected

def duzeltme_etiket_ciktisi_olustur(app_instance, tags, final_path):
    if not final_path:
        return False, "Kaydedilecek dosya seçilmedi."
    selected = duzeltme_etiketleri_temizle(tags)
    if not selected:
        return False, "En az bir etiket seçilmelidir."

    with tempfile.TemporaryDirectory(prefix="raporpro_duzeltme_") as tmp:
        tmp_template = os.path.join(tmp, "duzeltme_etiket_sablonu.docx")
        duzeltme_etiket_sablonu_olustur(selected, tmp_template)
        context = rapor_baglami_olustur(app_instance, word_path=tmp_template)
        success, msg = raporla(context, final_path=final_path, autosave=False)
        if success:
            return True, f"Düzeltme etiket çıktısı oluşturuldu: {len(selected)} etiket."
        return False, msg

def clean_word_tags(doc):
    for p in iter_all_paragraphs(doc):
        if "AUTO" in p.text:
            if "AUTO :" in p.text: 
                for run in p.runs:
                    if "AUTO :" in run.text:
                        run.text = run.text.replace("AUTO :", "AUTO:")

def iter_all_paragraphs(doc):
    for p in doc.paragraphs: yield p
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs: yield p
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer]:
            if hf:
                for p in hf.paragraphs: yield p
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs: yield p

def word_bosluk_paragrafi(space_before_pt=6, space_after_pt=6):
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(space_before_pt * 20)))
    spacing.set(qn("w:after"), str(int(space_after_pt * 20)))
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    paragraph.append(p_pr)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = " "
    run.append(text)
    paragraph.append(run)
    return paragraph

def word_sayfa_sonu_paragrafi():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph

def _plain_heading_text(text):
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    text = text.strip(" .:-–—")
    return text

def _normalize_style_name(name):
    text = unicodedata.normalize("NFKD", str(name or "").lower())
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return text.replace("ı", "i")

def _numbered_heading_level(text):
    clean = _plain_heading_text(text)
    match = re.match(r"^(\d+(?:\.\d+)*)\.?\s+.+", clean)
    if not match:
        return 0
    return len(match.group(1).split("."))

def _heading_style_level(paragraph):
    style_name = _normalize_style_name(getattr(getattr(paragraph, "style", None), "name", ""))
    match = re.match(r"^(?:heading|baslik|başlık)\s*(\d+)", style_name)
    if not match:
        return 0
    return int(match.group(1))

def _major_heading_mi(paragraph):
    text = _plain_heading_text(paragraph.text)
    if not text or "[" in text or "]" in text:
        return False
    if len(text) > 120:
        return False

    style_level = _heading_style_level(paragraph)
    if style_level == 1:
        return True
    if style_level > 1:
        return False

    if _numbered_heading_level(text) != 1:
        return False
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(1 for ch in letters if ch.upper() == ch) / max(1, len(letters))
    word_count = len(text.split())
    return upper_ratio >= 0.72 and word_count <= 12

def buyuk_basliklari_yeni_sayfaya_al(doc):
    count = 0
    heading_indexes = []
    for idx, paragraph in enumerate(doc.paragraphs):
        numbered_level = _numbered_heading_level(paragraph.text)
        style_level = _heading_style_level(paragraph)
        if numbered_level > 1 or style_level > 1:
            paragraph.paragraph_format.page_break_before = False
            paragraph.paragraph_format.keep_with_next = True
            heading_indexes.append(idx)
            continue
        if idx == 0 or not _major_heading_mi(paragraph):
            continue
        paragraph.paragraph_format.page_break_before = True
        paragraph.paragraph_format.keep_with_next = True
        heading_indexes.append(idx)
        count += 1

    paragraphs = doc.paragraphs
    for heading_index in heading_indexes:
        next_index = heading_index + 1
        while next_index < len(paragraphs) and not paragraphs[next_index].text.strip():
            blank = paragraphs[next_index]
            blank.paragraph_format.page_break_before = False
            blank.paragraph_format.keep_with_next = True
            next_index += 1
    return count

def jeo_parametre_degeri_formatla(anahtar, deger, son_tabaka=False):
    if anahtar == "h" and son_tabaka:
        return "-"
    if anahtar in ("E", "G", "K"):
        text = clean_val(deger)
        if text == "-":
            return "-"
        try:
            return str(int(round(float(text.replace(",", ".")))))
        except Exception as exc:
            _log_silent("jeo_parametre_degeri_formatla", exc)
            return fmt_jeo(deger)
    return fmt_jeo(deger)

BINA_FIELDS_MAP = [
    ("Bina Kullanım Amacı", "kul"),
    ("Bina Kullanım Sınıfı", "sinif"),
    ("Bina Önem Katsayısı", "onem"),
    ("Yapı Malzemesi", "malz"),
    ("Bodrum Kat Adedi / Toplam Kat Adedi", ("bod", "kat")),
    ("Plan Boyutları", "plan"),
    ("Yapı Yüksekliği (Hn)", "yukseklik"),
    ("Bina Yükseklik Sınıfı", "yukseklik_sinif"),
    ("Temel Alanı / Toplam İnşaat Alanı", ("temel_alan", "ins")),
    ("Olası Kazı Derinliği", "der"),
]


def bina_bilgisi_tablo_degeri(blok, key):
    """Birleşik bina satırlarını soldaki başlık sırasıyla biçimlendir."""
    if not isinstance(key, tuple):
        return clean_val(blok.get(key, ""))

    values = [clean_val(blok.get(item, "")) for item in key]
    if key == ("temel_alan", "ins"):
        values = [
            value if value == "-" or "m²" in value or re.search(r"\bm\s*2\b", value, re.I)
            else f"{value} m²"
            for value in values
        ]
        values = [re.sub(r"\bm\s*2\b", "m²", value, flags=re.I) for value in values]
    return " / ".join(values)


def bina_bloklari_rapor(bina):
    if not isinstance(bina, dict) or not bina.get("coklu_blok"):
        return []
    bloklar = []
    for idx, blok in enumerate(bina.get("bloklar", []) or []):
        if not isinstance(blok, dict):
            continue
        row = {key: clean_val(value) for key, value in blok.items()}
        if not any(value != "-" for key, value in row.items() if key != "blok_adi"):
            continue
        if row.get("blok_adi", "-") == "-":
            row["blok_adi"] = f"Blok {idx + 1}"
        bloklar.append(row)
    return bloklar

def bina_bilgileri_tablo_kayitlari(bina):
    bloklar = bina_bloklari_rapor(bina)
    if bloklar:
        return [(blok.get("blok_adi", f"Blok {idx + 1}"), blok) for idx, blok in enumerate(bloklar)]
    return [(None, bina)]

def bina_bilgileri_dikey_tablo_olustur(doc, bina):
    kayitlar = bina_bilgileri_tablo_kayitlari(bina)
    coklu_blok = len(kayitlar) > 1
    total_cols = 1 + (len(kayitlar) * 3)
    table = doc.add_table(rows=0, cols=total_cols)
    table.style = 'Table Grid'

    def merge_block_cells(row, block_idx):
        start = 1 + block_idx * 3
        return row.cells[start].merge(row.cells[start + 2])

    if coklu_blok:
        row = table.add_row()
        set_cell_text_clean(row.cells[0], "Bina Bilgileri", bold=True)
        for block_idx, (blok_adi, _) in enumerate(kayitlar):
            cell = merge_block_cells(row, block_idx)
            set_cell_text_clean(cell, clean_val(blok_adi), bold=True)
            set_vertical_cell_alignment(cell, "center")

    for label, key in BINA_FIELDS_MAP:
        row = table.add_row()
        set_cell_text_clean(row.cells[0], label, bold=True)
        for block_idx, (_, blok) in enumerate(kayitlar):
            cell = merge_block_cells(row, block_idx)
            set_cell_text_clean(cell, bina_bilgisi_tablo_degeri(blok, key), bold=False)
            set_vertical_cell_alignment(cell, "center")

    header_row = table.add_row()
    set_cell_text_clean(header_row.cells[0], "Binadan Temel Zeminine Aktarılan En Yükler (t/m2)", bold=True)
    if coklu_blok:
        for block_idx, (blok_adi, _) in enumerate(kayitlar):
            cell = merge_block_cells(header_row, block_idx)
            set_cell_text_clean(cell, clean_val(blok_adi), bold=True)
            set_vertical_cell_alignment(cell, "center")
    else:
        for offset, h in enumerate(("Min", "Ortalama", "Maks"), start=1):
            set_cell_text_clean(header_row.cells[offset], h, bold=True)
            set_vertical_cell_alignment(header_row.cells[offset], "center")

    sub_header_row = None
    if coklu_blok:
        sub_header_row = table.add_row()
        set_cell_text_clean(sub_header_row.cells[0], "", bold=True)
        for block_idx, _ in enumerate(kayitlar):
            start = 1 + block_idx * 3
            for offset, h in enumerate(("Min", "Ortalama", "Maks")):
                set_cell_text_clean(sub_header_row.cells[start + offset], h, bold=True)
                set_vertical_cell_alignment(sub_header_row.cells[start + offset], "center")

    row_gqe = table.add_row()
    set_cell_text_clean(row_gqe.cells[0], "(G+Q+E)", bold=True)
    for block_idx, (_, blok) in enumerate(kayitlar):
        start = 1 + block_idx * 3
        set_cell_text_clean(row_gqe.cells[start], clean_val(blok.get("gqe_min", "")))
        set_cell_text_clean(row_gqe.cells[start + 1], clean_val(blok.get("gqe_ort", "")))
        set_cell_text_clean(row_gqe.cells[start + 2], clean_val(blok.get("gqe_max", "")))

    row_comb = table.add_row()
    set_cell_text_clean(row_comb.cells[0], "1.4G+1.6Q", bold=True)
    for block_idx, (_, blok) in enumerate(kayitlar):
        start = 1 + block_idx * 3
        set_cell_text_clean(row_comb.cells[start], clean_val(blok.get("comb_min", "")))
        set_cell_text_clean(row_comb.cells[start + 1], clean_val(blok.get("comb_ort", "")))
        set_cell_text_clean(row_comb.cells[start + 2], clean_val(blok.get("comb_max", "")))

    for r in [row_gqe, row_comb]:
        for i in range(1, total_cols):
            set_vertical_cell_alignment(r.cells[i], "center")
    widths_cm = [5.2] + [1.55] * (total_cols - 1)
    apply_report_table_style(table, header_rows=0, label_cols={0}, widths_cm=widths_cm)
    if coklu_blok:
        style_report_table_row(table.rows[0])
    style_report_table_row(header_row)
    if sub_header_row is not None:
        style_report_table_row(sub_header_row)
    return table

def bina_bilgileri_tablolari_olustur(doc, bina):
    table = bina_bilgileri_dikey_tablo_olustur(doc, bina)
    return [table]

def _paragraph_source(doc, paragraphs=None):
    return paragraphs if paragraphs is not None else iter_all_paragraphs(doc)

def replace_text(doc, tag, value, paragraphs=None, paragraph_index=None):
    val_str = str(value)
    candidates = [paragraph_index.get(tag)] if paragraph_index is not None else _paragraph_source(doc, paragraphs)
    for p in candidates:
        if p is None:
            continue
        if tag in p.text:
            replaced = False
            for run in p.runs:
                if tag in run.text:
                    run.text = run.text.replace(tag, val_str)
                    replaced = True
            if not replaced:
                full_text = "".join(r.text for r in p.runs)
                if tag in full_text and p.runs:
                    p.runs[0].text = full_text.replace(tag, val_str)
                    for r in p.runs[1:]:
                        r.text = ""

def replace_many_text(doc, replacements, paragraphs=None):
    items = [(str(tag), str(value)) for tag, value in (replacements or {}).items() if tag]
    if not items:
        return 0
    changed = 0
    for p in _paragraph_source(doc, paragraphs):
        paragraph_text = p.text
        if not paragraph_text or not any(tag in paragraph_text for tag, _ in items):
            continue
        for run in p.runs:
            if not run.text:
                continue
            original = run.text
            updated = original
            for tag, value in items:
                if tag in updated:
                    changed += updated.count(tag)
                    updated = updated.replace(tag, value)
            if updated != original:
                run.text = updated

        full_text = "".join(r.text for r in p.runs)
        if any(tag in full_text for tag, _ in items):
            updated = full_text
            for tag, value in items:
                if tag in updated:
                    changed += updated.count(tag)
                    updated = updated.replace(tag, value)
            if p.runs:
                p.runs[0].text = updated
                for run in p.runs[1:]:
                    run.text = ""
            else:
                p.text = updated
    return changed

def find_paragraph_with_tag(doc, tag, paragraphs=None):
    for p in _paragraph_source(doc, paragraphs):
        if tag in p.text:
            return p
    return None

def build_paragraph_tag_index(paragraphs, tags):
    index = {}
    remaining = set(tags or [])
    for p in paragraphs or []:
        if not remaining:
            break
        text = p.text or ""
        matched = [tag for tag in remaining if tag in text]
        for tag in matched:
            index[tag] = p
            remaining.remove(tag)
    return index

def islem_tablo_yerlestir(doc, tag, headers, data_list, paragraphs=None, paragraph_index=None):
    candidates = [paragraph_index.get(tag)] if paragraph_index is not None else _paragraph_source(doc, paragraphs)
    for p in candidates:
        if p is None:
            continue
        if tag in p.text: 
            p.text = p.text.replace(tag, "")
            if data_list:
                table = create_word_table(doc, headers, data_list)
                p._p.addnext(table._tbl)
            return

def doc_replace_text_everywhere(doc, old_text, new_text, paragraphs=None, paragraph_index=None):
    replace_text(doc, old_text, new_text, paragraphs=paragraphs, paragraph_index=paragraph_index)

def replace_tag_with_paragraphs(doc, tag, text_list, paragraphs=None, paragraph_index=None):
    candidates = [paragraph_index.get(tag)] if paragraph_index is not None else _paragraph_source(doc, paragraphs)
    for p in candidates:
        if p is None:
            continue
        if tag in p.text:
            p.text = p.text.replace(tag, "")
            for text in reversed(text_list):
                if not text or not text.strip(): continue
                new_p = OxmlElement("w:p"); new_r = OxmlElement("w:r"); new_t = OxmlElement("w:t"); new_t.text = text; new_r.append(new_t); new_p.append(new_r); p._p.addnext(new_p)
            return


def replace_tag_with_report_blocks(
    doc,
    tag,
    blocks,
    paragraphs=None,
    paragraph_index=None,
):
    """Etiketi, yer tutucunun biçimini koruyan rapor paragraflarıyla değiştir."""
    candidates = (
        [paragraph_index.get(tag)]
        if paragraph_index is not None
        else _paragraph_source(doc, paragraphs)
    )
    for paragraph in candidates:
        if paragraph is None or tag not in paragraph.text:
            continue

        paragraph_properties = (
            copy.deepcopy(paragraph._p.pPr)
            if paragraph._p.pPr is not None
            else None
        )
        run_properties = None
        if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
            run_properties = copy.deepcopy(paragraph.runs[0]._r.rPr)
        paragraph.text = paragraph.text.replace(tag, "")
        placeholder_empty = not paragraph.text.strip()
        inserted = False

        for block in reversed(blocks or []):
            text = str(block.get("metin", "") if isinstance(block, dict) else block)
            if not text.strip():
                continue
            block_type = (
                str(block.get("tur", "metin"))
                if isinstance(block, dict)
                else "metin"
            )
            new_paragraph = OxmlElement("w:p")
            if paragraph_properties is not None:
                new_paragraph.append(copy.deepcopy(paragraph_properties))
            p_pr = new_paragraph.get_or_add_pPr()
            if p_pr.find(qn("w:keepLines")) is None:
                p_pr.append(OxmlElement("w:keepLines"))
            if block_type == "birim_basligi":
                if p_pr.find(qn("w:keepNext")) is None:
                    p_pr.append(OxmlElement("w:keepNext"))

            new_run = OxmlElement("w:r")
            if run_properties is not None:
                new_run.append(copy.deepcopy(run_properties))
            r_pr = new_run.get_or_add_rPr()
            for bold_tag in ("w:b", "w:bCs"):
                bold_node = r_pr.find(qn(bold_tag))
                if bold_node is None:
                    bold_node = OxmlElement(bold_tag)
                    r_pr.append(bold_node)
                bold_node.set(
                    qn("w:val"),
                    "1" if block_type == "birim_basligi" else "0",
                )

            new_text = OxmlElement("w:t")
            new_text.text = text
            new_run.append(new_text)
            new_paragraph.append(new_run)
            paragraph._p.addnext(new_paragraph)
            inserted = True
        if placeholder_empty and inserted:
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)
        return True
    return False


def doc_replace_img(doc, keyword, img_path, paragraphs=None, paragraph_index=None):
    label = _safe_perf_label(keyword)
    if not img_path or not os.path.exists(img_path):
        perf_log(f"report.image.{label}.skip", detail=f"{keyword}|missing")
        return
    replaced = 0
    with perf_timer(f"report.image.{label}", _file_perf_detail(img_path)):
        candidates = [paragraph_index.get(keyword)] if paragraph_index is not None else _paragraph_source(doc, paragraphs)
        for p in candidates:
            if p is None:
                continue
            if keyword in p.text:
                p.text = ""
                run = p.add_run()
                run.add_picture(img_path, width=Cm(16))
                replaced += 1
    perf_log(f"report.image.{label}.matches", detail=_report_detail(count=replaced, keyword=keyword))

def first_existing_path(*paths):
    for path in paths:
        if path and os.path.exists(path):
            return path
    return None

def mjh_resim_yolu(app_instance):
    """Yalnız mühendislik jeolojisi haritasını döndür.

    Farklı harita türlerini yedek olarak kullanmak, raporda doğru etikete yanlış
    paftanın yerleşmesine neden olur.
    """
    return first_existing_path(getattr(app_instance, 'img_mjh', None))

def _report_detail(**items):
    parts = []
    for key, value in items.items():
        try:
            parts.append(f"{key}={value}")
        except Exception:
            parts.append(f"{key}=?")
    return " ".join(parts)

def _safe_perf_label(text):
    label = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(text or "").strip("[]"))
    return (label or "tag")[:80]

def _file_perf_detail(path):
    if not path:
        return "path=missing"
    try:
        size_kb = os.path.getsize(path) / 1024
        return _report_detail(path=os.path.basename(path), size_kb=f"{size_kb:.1f}")
    except Exception:
        return _report_detail(path=os.path.basename(str(path)))

def _doc_perf_stats(doc):
    try:
        table_rows = sum(len(table.rows) for table in doc.tables)
        return _report_detail(
            paragraphs=sum(1 for _ in iter_all_paragraphs(doc)),
            tables=len(doc.tables),
            table_rows=table_rows,
        )
    except Exception as exc:
        _log_silent("doc_perf_stats", exc)
        return ""

def raporla(app_instance, final_path=None, autosave=True):
    template_info = rapor_sablonu_durumu(
        getattr(app_instance, "word_path", None),
        proje_rapor_sablon_profili(getattr(app_instance, "veri", {})),
    )
    template_path = template_info.get("path", "")
    if not template_path:
        return False, "Dahili rapor şablonu bulunamadı. Geçerli bir özel Word şablonu seçin."
    app_instance.set_status("Rapor oluşturuluyor...", level="warning")
    if final_path is None and hasattr(app_instance, "root"):
        app_instance.root.update()
    if autosave and hasattr(app_instance, "veri_kaydet"):
        app_instance.veri_kaydet()
    step_time = [time.perf_counter()]

    def report_step(name, detail=""):
        now = time.perf_counter()
        perf_log(f"report.step.{name}", now - step_time[0], detail)
        step_time[0] = now
    
    try:
        with perf_timer("report.open_template", template_path):
            doc = Document(template_path)
        with perf_timer("report.clean_tags"):
            clean_word_tags(doc)
        report_paragraphs = list(iter_all_paragraphs(doc))
        structural_tags = [
            "[BOLGESEL_JEOLOJI]", "[BOLGESEL_JEOLOJI_BIRIMLERI]",
            "[MUHENDISLIK_JEOLOJISI]",
            "[JEOLOJIK_KESIT_ACIKLAMA]", "[JEOLOJI_SONUC]", "[MT_BIRIM_METNI]",
            "[BINA_BILGILERI]", "[Sondaj]", "[YASS_TABLO]", "[LAB_FIZIK]", "[LAB_MEKANIK]",
            "[ZEMIN_OZET]", "[LITOLOJI_DAGILIM]", "[SPT]", "[PMT]", "[KAYA_TABLO]",
            "[JEO_PARAMETRE]", "[MASW]", "[VP]", "[JEO_KOOR]", "[MT_TABLO]",
            "[JEO_SONUC]", "[YASS_ONERI]", "[HIDROJEOLOJI_DURUM]",
            "RESIM:Yerbuldurur", "[RESIM_YERBULDURUR]", "RESIM:TKGM", "RESIM:PGA",
            "[RESIM_JEOFIZIK]", "RESIM:MJH", "[RESIM_MJH]", "[RESIM:MJH]",
            "[RESIM_SONDAJ]", "[RESIM:SONDAJ]",
        ]
        report_tag_index = build_paragraph_tag_index(report_paragraphs, structural_tags)
        with perf_timer("report.revision_markers"):
            marker_count = revizyon_isaretleri_ekle(doc, report_tag_index, structural_tags)
        report_step("revision_markers", _report_detail(count=marker_count))
        report_step("template_ready", _doc_perf_stats(doc))
        kunye = app_instance.veri["kunye"]
        jeofizik = app_instance.veri["jeofizik"]
        arazi = app_instance.veri["arazi"]
        sondajlar = app_instance.veri["sondaj"]
        
        ss_list = jeofizik.get("ss_list", [])
        mt_list = jeofizik.get("mt_list", [])
        
        jeo_excel_path = getattr(app_instance, 'jeo_excel_path', None)
        param_ss_list = []
        jeo_sheet_rows = app_instance.veri.get("jeofizik_sheet", {}).get("rows", []) if isinstance(getattr(app_instance, "veri", None), dict) else []
        jeo_sheet_ready = any(any(str(cell).strip() for cell in row) for row in jeo_sheet_rows or [])
        
        if jeo_sheet_ready:
            try:
                with perf_timer("report.read_jeofizik_sheet", "internal"):
                    param_ss_list = jeofizik_sheet_rows_to_ss_list(jeo_sheet_rows)
                    jeofizik_ss_koordinatlarini_koru(param_ss_list, ss_list)
            except Exception as e:
                _log_silent("jeofizik_sheet_parse", e)
                traceback.print_exc()
                param_ss_list = []

        if not param_ss_list and jeo_excel_path:
            try:
                with perf_timer("report.read_jeofizik_excel", jeo_excel_path):
                    df_jeo = read_table_file(jeo_excel_path, header=None)
                current_serim = None
                for idx, row in df_jeo.iterrows():
                    row_str = [str(x).strip() for x in row if pd.notna(x)]
                    if not row_str: continue
                    if "Sismik Ölçü ve Hesaplarının Sahibi" in str(row.iloc[0]):
                        s_name = "SS"
                        for cell in row_str:
                            if "Serim" in str(cell) or "SS" in str(cell): s_name = str(cell).strip(); break
                        current_serim = {"ad": s_name, "layers": []}
                        param_ss_list.append(current_serim)
                    elif current_serim is not None:
                        r0 = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ""
                        if "VP =" in r0 or "Boyuna Dalga" in r0: current_serim["raw_vp"] = row.tolist()
                        elif "VS =" in r0 or "Enine Dalga" in r0: current_serim["raw_vs"] = row.tolist()
                        elif "Tabaka Kalınlığı" in r0: current_serim["raw_h"] = row.tolist()
                        elif "Tabaka Yoğunluğu" in r0: current_serim["raw_rho"] = row.tolist()
                        elif "Poisson Oranı" in r0: current_serim["raw_nu"] = row.tolist()
                        elif "Elastisite" in r0 or "Young" in r0: current_serim["raw_E"] = row.tolist()
                        elif "Kayma Modülü" in r0 or "Gmax" in r0: current_serim["raw_G"] = row.tolist()
                        elif "Bulk" in r0 or "Sıkışmazlık" in r0: current_serim["raw_K"] = row.tolist() 
                        elif "Vs30 =" in r0 or "Vs30" in r0: current_serim["raw_vs30"] = row.tolist()
                
                for s in param_ss_list:
                    if "raw_vp" not in s: continue
                    vs30_val = "-"
                    if "raw_vs30" in s:
                        for val in s["raw_vs30"][2:]: 
                            if pd.notna(val) and str(val).strip() != "": vs30_val = val; break
                    for col_idx in range(2, len(s["raw_vp"])):
                        vp_val = s["raw_vp"][col_idx]
                        if pd.isna(vp_val) or str(vp_val).strip() == "": continue
                        layer = {}
                        layer["vp"] = vp_val
                        layer["vs"] = s.get("raw_vs", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_vs", [])) else "-"
                        layer["h"] = s.get("raw_h", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_h", [])) else "-"
                        layer["rho"] = s.get("raw_rho", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_rho", [])) else "-"
                        layer["nu"] = s.get("raw_nu", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_nu", [])) else "-"
                        layer["E"] = s.get("raw_E", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_E", [])) else "-"
                        layer["G"] = s.get("raw_G", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_G", [])) else "-"
                        layer["K"] = s.get("raw_K", [np.nan]*(col_idx+1))[col_idx] if col_idx < len(s.get("raw_K", [])) else "-" 
                        layer["vs30"] = vs30_val if col_idx == 2 else "-"
                        s["layers"].append(layer)
            except Exception as e:
                _log_silent("jeofizik_excel_parse", e)
                traceback.print_exc()
                param_ss_list = ss_list 
        elif not param_ss_list:
            param_ss_list = ss_list
        param_layer_count = sum(len(ss.get("layers", []) or []) for ss in param_ss_list)
        report_step(
            "jeofizik_data_ready",
            _report_detail(ss=len(ss_list), mt=len(mt_list), param_ss=len(param_ss_list), layers=param_layer_count),
        )

        tum_vs30 = []
        for ss in param_ss_list:
            layers = ss.get("layers", [])
            if not layers: continue
            vs30_raw = layers[0].get("vs30", 0)
            try: v = float(str(vs30_raw).replace(",",".")); 
            except Exception as exc:
                _log_silent("vs30_parse", exc)
                v=0
            if v > 0: tum_vs30.append(v)
            
        # MT VERİLERİNDEN T0 (BASKIN PERİYOT) ÇEKİMİ
        tum_t0 = []
        for mt in mt_list:
            to_raw = mt.get("to", "")
            if str(to_raw).strip() == "":
                continue
            try:
                v_to = float(str(to_raw).replace(",", "."))
                if v_to > 0: tum_t0.append(v_to)
            except Exception as exc:
                _log_silent("mt_t0_parse", exc)
            
        yass_seviyeleri = []
        for s in sondajlar:
            try:
                d1 = safe_float(s.get("yass_d1")); d2 = safe_float(s.get("yass_d2"))
                if d1 > 0: yass_seviyeleri.append(d1)
                if d2 > 0: yass_seviyeleri.append(d2)
            except Exception as exc:
                _log_silent("yass_parse", exc)

        basic_tag_start = time.perf_counter()
        prefixes = ["", "S1_", "S2_", "S3_", "S4_", "S5_"]
        kunye_map = [("sahibi", "PROJE_ADI"), ("il", "IL"), ("ilce", "ILCE"), ("mah", "MAHALLE"), ("mev", "MEVKI"), ("paf", "PAFTA"), ("ada", "ADA"), ("par", "PARSEL")]
        basic_replacements = {}
        for key, tag_base in kunye_map:
            val = (
                rapor_proje_adi(app_instance.veri)
                if tag_base == "PROJE_ADI"
                else kunye.get(key, "")
            )
            for pre in prefixes: basic_replacements[f"[{pre}{tag_base}]"] = val
        
        basic_replacements["[KATEGORI]"] = arazi.get("kategori", "-")
        basic_replacements["[KATEGORI_ZEMIN]"] = arazi.get("zemin", "-")
        basic_replacements["[PGA]"] = arazi.get("pga", "-")
        basic_replacements["[JEO_TARIH]"] = jeofizik.get("tarih", "-")
        basic_replacements["[SAYI_SS]"] = str(len(ss_list))
        basic_replacements["[SAYI_MT]"] = str(len(mt_list))
        basic_replacements["[YEREL_ZEMIN]"] = app_instance.veri["bina"].get("ysinif", "-")
        basic_replacements["[KOT_ORT]"] = arazi.get("ort", "-")
        basic_replacements["[KOT_MAX]"] = arazi.get("max", "-")
        basic_replacements["[KOT_MIN]"] = arazi.get("min", "-")
        basic_replacements["[EGIM_YUZDE]"] = arazi.get("egim", "-")
        basic_replacements["[EGIM_YONU]"] = arazi.get("yon", "-")
        
        imar_alani_ham = arazi.get("imar_alani", "").strip(); imar_alani_final = f"({imar_alani_ham})" if imar_alani_ham else "-"
        basic_replacements["[IMAR_ALANI]"] = imar_alani_final
        basic_replacements["[IMAR_DURUMU]"] = arazi.get("imar_durumu", "-")

        ay = clean_val(app_instance.veri["arazi"].get("alan_y", "-")); ax = clean_val(app_instance.veri["arazi"].get("alan_x", "-"))
        basic_replacements["[ALAN_ENLEM]"] = ay
        basic_replacements["[ALAN_BOYLAM]"] = ax
        
        if sondajlar:
            ozet_parca = ", ".join([f"{s['no']}: {s['der']}m" for s in sondajlar])
            sondaj_metni = f"Sahada toplam {len(sondajlar)} adet sondaj kuyusu ({ozet_parca}) açılmıştır."
        else: sondaj_metni = "Sahada sondaj çalışması yapılmamıştır."
        basic_replacements["[SONDAJ_BILGISI]"] = sondaj_metni
        replaced_basic_tags = replace_many_text(doc, basic_replacements, paragraphs=report_paragraphs)
        parcel_text_replacements = rapor_metin_degerleri(app_instance.veri)
        replaced_parcel_tags = replace_many_text(
            doc,
            parcel_text_replacements,
            paragraphs=report_paragraphs,
        )
        fixed_table_result = rapor_sabit_tablolarini_uygula(
            doc,
            app_instance.veri,
        )
        basic_tag_detail = _report_detail(tags=len(basic_replacements), replaced=replaced_basic_tags, sondaj=len(sondajlar))
        basic_tag_detail += " " + _report_detail(
            parcel_tags=replaced_parcel_tags,
            fault_rows=fixed_table_result.get("aktif_fay_satiri", 0),
            seismic_rows=fixed_table_result.get("sismik_satiri", 0),
        )
        perf_log("report.basic_tags.replace", time.perf_counter() - basic_tag_start, basic_tag_detail)
        report_step("basic_tags", basic_tag_detail)

        jeoloji_blocks = jeoloji_rapor_bloklari(app_instance.veri)
        jeoloji_tag_map = {
            "[BOLGESEL_JEOLOJI]": "bolgesel_giris",
            "[BOLGESEL_JEOLOJI_BIRIMLERI]": "bolgesel_birimler",
            "[MUHENDISLIK_JEOLOJISI]": "muhendislik",
            "[JEOLOJIK_KESIT_ACIKLAMA]": "kesit",
            "[JEOLOJI_SONUC]": "sonuc",
            "[MT_BIRIM_METNI]": "mt",
        }
        jeoloji_replaced = 0
        for tag, block_key in jeoloji_tag_map.items():
            if replace_tag_with_report_blocks(
                doc,
                tag,
                jeoloji_blocks[block_key],
                paragraph_index=report_tag_index,
            ):
                jeoloji_replaced += 1
        report_step(
            "geology_texts",
            _report_detail(
                tags=jeoloji_replaced,
                units=len(app_instance.veri.get("jeoloji", {}).get("birimler", [])),
            ),
        )

        bina = app_instance.veri["bina"]
        bina_table_start = time.perf_counter()
        bina_detail = _report_detail(blocks=len(bina_bilgileri_tablo_kayitlari(bina)), coklu=bool(bina.get("coklu_blok")))
        p = report_tag_index.get("[BINA_BILGILERI]")
        if p is not None and "[BINA_BILGILERI]" in p.text:
            p.text = p.text.replace("[BINA_BILGILERI]", "")
            tables = bina_bilgileri_tablolari_olustur(doc, bina)
            anchor = p._p
            for idx, table in enumerate(tables):
                anchor.addnext(table._tbl)
                anchor = table._tbl
                if idx < len(tables) - 1:
                    spacer = OxmlElement("w:p")
                    anchor.addnext(spacer)
                    anchor = spacer
        perf_log("report.table.bina", time.perf_counter() - bina_table_start, bina_detail)
        report_step("bina_table", bina_detail)
        
        sondaj_table_start = time.perf_counter()
        sondaj_lit_count = sum(len(s.get("litoloji", []) or []) for s in app_instance.veri["sondaj"])
        sondaj_table_detail = _report_detail(sondaj=len(app_instance.veri["sondaj"]), litoloji=sondaj_lit_count)
        p = report_tag_index.get("[Sondaj]")
        if p is not None and "[Sondaj]" in p.text:
                p.text = p.text.replace("[Sondaj]", ""); headers = ["Kuyu No", "Başlangıç", "Bitiş", "Enlem", "Boylam", "Kot", "Derinlik", "Litoloji"]; table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
                for i, h in enumerate(headers): set_cell_text_clean(table.rows[0].cells[i], h, bold=True); set_vertical_cell_alignment(table.rows[0].cells[i], "center")
                for s in app_instance.veri["sondaj"]:
                    lits = s.get("litoloji", []); kuyu_no = clean_val(s["no"]); bas_tar = clean_val(s["bas_tar"]); bit_tar = clean_val(s["bit_tar"]); enlem = clean_val(s["y"]); boylam = clean_val(s["x"]); kot = clean_val(s["k"])
                    if not lits: 
                        row = table.add_row(); set_cell_text_clean(row.cells[0], kuyu_no)
                        for i, v in enumerate([bas_tar, bit_tar, enlem, boylam, kot, "-", "-"]): set_cell_text_clean(row.cells[i+1], v)
                        continue
                    start_idx = len(table.rows)
                    for lit in lits:
                        row = table.add_row(); lit_depth = f"{clean_val(lit[0])}-{clean_val(lit[1])}"; set_cell_text_clean(row.cells[6], lit_depth); set_cell_text_clean(row.cells[7], clean_val(lit[2]), alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    end_idx = len(table.rows) - 1; vals = [kuyu_no, bas_tar, bit_tar, enlem, boylam, kot]
                    if end_idx > start_idx:
                        for col_idx, val in enumerate(vals): first = table.rows[start_idx].cells[col_idx]; last = table.rows[end_idx].cells[col_idx]; first.merge(last); set_cell_text_clean(first, val); set_vertical_cell_alignment(first, "center")
                    else:
                        for col_idx, val in enumerate(vals): set_cell_text_clean(table.rows[start_idx].cells[col_idx], val); set_vertical_cell_alignment(table.rows[start_idx].cells[col_idx], "center")
                apply_report_table_style(table, header_rows=1, text_cols={7}, widths_cm=[1.5, 1.8, 1.8, 2.2, 2.2, 1.4, 1.7, 5.4])
                p._p.addnext(table._tbl)
        perf_log("report.table.sondaj", time.perf_counter() - sondaj_table_start, sondaj_table_detail)
        report_step("sondaj_table", sondaj_table_detail)
        
        yass_data = []
        for s in app_instance.veri["sondaj"]: v1 = f"{clean_val(s.get('yass_d1'))} ({clean_val(s.get('yass_t1'))})" if s.get('yass_d1') else "-"; v2 = f"{clean_val(s.get('yass_d2'))} ({clean_val(s.get('yass_t2'))})" if s.get('yass_d2') else "-"; yass_data.append([s["no"], v1, v2])
        yass_table_start = time.perf_counter()
        islem_tablo_yerlestir(doc, "[YASS_TABLO]", ["Kuyu No", "1. Ölçüm (Delgi Sonu)", "2. Ölçüm (Statik)"], yass_data, paragraph_index=report_tag_index)
        yass_detail = _report_detail(rows=len(yass_data))
        perf_log("report.table.yass", time.perf_counter() - yass_table_start, yass_detail)
        report_step("yass_table", yass_detail)

        lab_fizik_headers = ["Birim", "Değer", "Çakıl (%)", "Kum (%)", "Silt+Kil (%)", "Kil (%)", "LL (%)", "PL (%)", "PI (%)", "Wn (%)", "γn (g/cm³)", "γk (g/cm³)"]
        lab_mekanik_headers = ["Birim", "Değer", "İçsel Sürtünme (ϕ)", "Kohezyon (c)"]
        
        zemin_ozet_yazildi = False
        lab_birim_isimleri = []
        lito_groups = {}
        unit_spt_values = {}
        lab_section_start = time.perf_counter()

        lab_rows = lab_sheet_satirlari(app_instance)
        lab_sheet_ready = any(any(str(cell).strip() for cell in row) for row in lab_rows)
        lab_excel_path = getattr(app_instance, "lab_excel_path", None)

        if lab_sheet_ready or lab_excel_path:
            try:
                if lab_sheet_ready:
                    with perf_timer("report.read_lab_sheet", "internal"):
                        df_lab = pd.DataFrame(lab_rows).replace(r'^\s*$', np.nan, regex=True)
                else:
                    with perf_timer("report.read_lab_excel", lab_excel_path):
                        df_lab = pd.read_excel(lab_excel_path, header=None)
                h_idx = 0
                for r, row in df_lab.head(30).iterrows():
                    if any("Sondaj No" in str(x) for x in row): h_idx = r; break
                hb = df_lab.iloc[h_idx:h_idx+5].copy().replace(r'^\s*$', np.nan, regex=True); hb.iloc[0] = hb.iloc[0].ffill()
                col_sigs = [" ".join([str(x).strip().upper() for x in hb.iloc[:, c] if pd.notna(x)]) for c in range(hb.shape[1])]
                cols = {}
                candidates_c = []
                for i, s in enumerate(col_sigs):
                    s_clean = temizle_baslik(s)
                    if "siniflama" in s_clean or "uscs" in s_clean or "sinif" in s_clean: cols['sinif']=i
                    if "wn" in s_clean or "su muhtevasi" in s_clean or "nem" in s_clean: cols['Wn']=i
                    if "ll" in s_clean or "likit" in s_clean: cols['LL']=i
                    if "pl" in s_clean or "plastik" in s_clean: cols['PL']=i
                    if "pi" in s_clean or "plastisite" in s_clean: cols['PI']=i
                    if "cakil" in s_clean or "gravel" in s_clean: cols['cakil']=i
                    if "kum" in s_clean or "sand" in s_clean: cols['kum']=i
                    if "kil" in s_clean or "clay" in s_clean: cols['kil']=i
                    if "silt" in s_clean: cols['silt']=i
                    if "dogal" in s_clean or "gn" in s_clean: cols['gn']=i
                    if "kuru" in s_clean or "gk" in s_clean: cols['gk']=i
                    if "direkt" in s_clean and ("c" in s_clean or "kohezyon" in s_clean): candidates_c.append(i)
                if candidates_c:
                    best_c = -1; max_count = -1; data_start = h_idx + 5
                    for cand_idx in candidates_c:
                        count = df_lab.iloc[data_start:, cand_idx].count()
                        if count > max_count: max_count = count; best_c = cand_idx
                    if best_c != -1: cols['c'] = best_c; potential_phi = best_c + 1; 
                    if potential_phi < df_lab.shape[1]: cols['phi'] = potential_phi

                if 'sinif' in cols:
                    data_start = h_idx + 5; df_c = df_lab.iloc[data_start:].copy(); df_c.columns = range(df_c.shape[1])
                    df_c.rename(columns={v:k for k,v in cols.items()}, inplace=True)
                    df_c['GRUP'] = df_c['sinif'].apply(zemin_sinifi_cevir); df_gruplanmis = df_c.groupby('GRUP')

                    for n, _ in df_gruplanmis:
                        if str(n).lower() != "tanımsız":
                            lab_birim_isimleri.append(str(n))
                    lab_birim_isimleri.sort(key=len, reverse=True)

                    lito_groups = {name: {} for name in lab_birim_isimleri}
                    unit_spt_values = {name: [] for name in lab_birim_isimleri}

                    for s in app_instance.veri["sondaj"]:
                        kuyu_no = s["no"]
                        for lit in s.get("litoloji", []):
                            top_val = safe_float(lit[0])
                            bot_val = safe_float(lit[1])
                            desc = str(lit[2]).strip().lower()
                            
                            matched_unit = None
                            for lab_name in lab_birim_isimleri:
                                if lab_name.lower() in desc:
                                    matched_unit = lab_name
                                    break 
                            
                            if matched_unit:
                                if kuyu_no not in lito_groups[matched_unit]: 
                                    lito_groups[matched_unit][kuyu_no] = []
                                lito_groups[matched_unit][kuyu_no].append((top_val, bot_val))

                    for s in app_instance.veri["sondaj"]:
                        kuyu_no = s["no"]
                        for spt in s.get("spt", []):
                            try:
                                spt_depth = safe_float(spt[0])
                                val_str = str(spt[4]).strip().upper()
                                if val_str == "R" or "REF" in val_str: n30 = 51.0
                                else: n30 = float(spt[4])

                                for unit_name, kuyu_dict in lito_groups.items():
                                    if kuyu_no in kuyu_dict:
                                        for (top, bot) in kuyu_dict[kuyu_no]:
                                            if top <= spt_depth < bot:
                                                unit_spt_values[unit_name].append(n30)
                                                break 
                            except Exception as exc:
                                _log_silent("spt_unit_summary", exc)

                    ozet_metinler = []
                    
                    for grup_adi, grup_df in df_gruplanmis:
                        if str(grup_adi).lower() == "tanımsız": continue

                        val_cakil = pd.to_numeric(grup_df.get('cakil', pd.Series()), errors='coerce')
                        ort_cakil = val_cakil.mean() if not val_cakil.empty else 0.0
                        
                        val_kum = pd.to_numeric(grup_df.get('kum', pd.Series()), errors='coerce')
                        ort_kum = val_kum.mean() if not val_kum.empty else 0.0
                        
                        silt_val = pd.to_numeric(grup_df.get('silt', pd.Series()), errors='coerce').fillna(0)
                        kil_val = pd.to_numeric(grup_df.get('kil', pd.Series()), errors='coerce').fillna(0)
                        val_ince = silt_val + kil_val
                        ort_ince = val_ince.mean() if not val_ince.empty else 0.0
                        
                        g_clean = str(grup_adi).strip()
                        is_cohesive = False
                        
                        if g_clean in INCE_DANELILER:
                            is_cohesive = True
                        elif g_clean in IRI_DANELILER:
                            is_cohesive = False
                        else:
                            g_low = g_clean.lower()
                            is_cohesive = True if "kil" in g_low or "silt" in g_low else False
                        
                        found_labels = set()
                        spt_source = unit_spt_values.get(g_clean, []) 

                        if spt_source:
                            for val in spt_source:
                                if is_cohesive:
                                    if val < 2: found_labels.add("çok yumuşak")
                                    elif val < 5: found_labels.add("yumuşak")
                                    elif val < 9: found_labels.add("orta katı")
                                    elif val < 16: found_labels.add("katı")
                                    elif val < 31: found_labels.add("çok katı")
                                    else: found_labels.add("sert")
                                else:
                                    if val < 5: found_labels.add("çok gevşek")
                                    elif val < 11: found_labels.add("gevşek")
                                    elif val < 31: found_labels.add("orta sıkı")
                                    elif val < 51: found_labels.add("sıkı")
                                    else: found_labels.add("çok sıkı")
                            
                            order_map = {"çok yumuşak": 1, "yumuşak": 2, "orta katı": 3, "katı": 4, "çok katı": 5, "sert": 6, "çok gevşek": 1, "gevşek": 2, "orta sıkı": 3, "sıkı": 4, "çok sıkı": 5}
                            sorted_labels = sorted(list(found_labels), key=lambda x: order_map.get(x, 99))
                            durum_text = " - ".join(sorted_labels)
                        else:
                            durum_text = ""

                        if durum_text:
                            if is_cohesive: giris_cumlesi = f"{grup_adi} birimleri {durum_text} olup"
                            else: giris_cumlesi = f"{grup_adi} birimleri sıkılığı {durum_text} olup"
                        else: giris_cumlesi = f"{grup_adi} birimleri"

                        lab_kisim = f"laboratuvar sonuçlarına göre içeriğinde ortalama olarak %{ort_cakil:.2f} Çakıl, %{ort_kum:.2f} Kum ve %{ort_ince:.2f} Silt-Kil barındırmaktadır"
                        ozet_metinler.append(f"{giris_cumlesi}, {lab_kisim}.")
                    
                    if ozet_metinler:
                        doc_replace_text_everywhere(doc, "[ZEMIN_OZET]", " ".join(ozet_metinler), paragraph_index=report_tag_index)
                        zemin_ozet_yazildi = True

                    p = report_tag_index.get("[LAB_FIZIK]")
                    if p is not None and "[LAB_FIZIK]" in p.text:
                            p.text = p.text.replace("[LAB_FIZIK]", ""); t = create_word_table(doc, lab_fizik_headers, [])
                            table_has_data = False
                            for n, g in df_gruplanmis:
                                if n == "Tanımsız": continue
                                
                                s_exists = 'silt' in g
                                k_exists = 'kil' in g
                                row_min = t.add_row(); row_max = t.add_row(); row_avg = t.add_row()
                                row_min.cells[0].text = n; row_min.cells[1].text = "Min"; row_max.cells[1].text = "Max"; row_avg.cells[1].text = "Ort"
                                row_min.cells[0].merge(row_avg.cells[0]); set_cell_text_clean(row_min.cells[0], n, bold=False); set_vertical_cell_alignment(row_min.cells[0], "center")
                                set_cell_text_clean(row_min.cells[1], "Min"); set_cell_text_clean(row_max.cells[1], "Max"); set_cell_text_clean(row_avg.cells[1], "Ort")
                                
                                if s_exists or k_exists:
                                    s_val = pd.to_numeric(g['silt'] if s_exists else pd.Series(0, index=g.index), errors='coerce').fillna(0)
                                    k_val = pd.to_numeric(g['kil'] if k_exists else pd.Series(0, index=g.index), errors='coerce').fillna(0)
                                    fines = s_val + k_val
                                    set_cell_text_clean(row_min.cells[4], f"{fines.min():.2f}")
                                    set_cell_text_clean(row_max.cells[4], f"{fines.max():.2f}")
                                    set_cell_text_clean(row_avg.cells[4], f"{fines.mean():.2f}")

                                pmap = [('cakil', 2), ('kum', 3), ('kil', 5), ('LL', 6), ('PL', 7), ('PI', 8), ('Wn', 9), ('gn', 10), ('gk', 11)]
                                for key, ci in pmap:
                                    if key in g:
                                        safe_series = g[key].astype(str).str.replace(',', '.', regex=False)
                                        v = pd.to_numeric(safe_series, errors='coerce').dropna()
                                        v = v[v >= 0]
                                        if not v.empty: 
                                            set_cell_text_clean(row_min.cells[ci], f"{v.min():.2f}"); set_cell_text_clean(row_max.cells[ci], f"{v.max():.2f}"); set_cell_text_clean(row_avg.cells[ci], f"{v.mean():.2f}"); table_has_data = True
                                for r_obj in [row_min, row_max, row_avg]: 
                                    for c_obj in r_obj.cells: set_vertical_cell_alignment(c_obj, "center")
                            if table_has_data:
                                apply_report_table_style(t, header_rows=1, text_cols={0}, widths_cm=[2.1, 1.2, 1.35, 1.35, 1.45, 1.35, 1.2, 1.2, 1.2, 1.25, 1.45, 1.45])
                                p._p.addnext(t._tbl)

                    
                    p = report_tag_index.get("[LAB_MEKANIK]")
                    if p is not None and "[LAB_MEKANIK]" in p.text:
                            p.text = p.text.replace("[LAB_MEKANIK]", ""); t = create_word_table(doc, lab_mekanik_headers, [])
                            table_has_data = False
                            for n, g in df_gruplanmis:
                                if n == "Tanımsız": continue
                                vals_c = pd.to_numeric(g.get('c', pd.Series()), errors='coerce').dropna()
                                vals_phi = pd.to_numeric(g.get('phi', pd.Series()), errors='coerce').dropna()
                                if vals_c.empty and vals_phi.empty: continue
                                row_min = t.add_row(); row_max = t.add_row(); row_avg = t.add_row()
                                row_min.cells[0].merge(row_avg.cells[0]); set_cell_text_clean(row_min.cells[0], n, bold=False); set_vertical_cell_alignment(row_min.cells[0], "center")
                                set_cell_text_clean(row_min.cells[1], "Min"); set_cell_text_clean(row_max.cells[1], "Max"); set_cell_text_clean(row_avg.cells[1], "Ort")
                                pmap = [('phi', 2), ('c', 3)]
                                for key, ci in pmap:
                                    if key in g:
                                        safe_series = g[key].astype(str).str.replace(',', '.', regex=False)
                                        v = pd.to_numeric(safe_series, errors='coerce').dropna()
                                        v = v[v >= 0]
                                        if not v.empty: 
                                            set_cell_text_clean(row_min.cells[ci], f"{v.min():.2f}"); set_cell_text_clean(row_max.cells[ci], f"{v.max():.2f}"); set_cell_text_clean(row_avg.cells[ci], f"{v.mean():.2f}"); table_has_data = True
                                for r_obj in [row_min, row_max, row_avg]: 
                                    for c_obj in r_obj.cells: set_vertical_cell_alignment(c_obj, "center")
                            if table_has_data:
                                apply_report_table_style(t, header_rows=1, text_cols={0}, widths_cm=[3.0, 1.4, 3.0, 3.0])
                                p._p.addnext(t._tbl)

            except Exception as exc:
                _log_silent("zemin_ozet_table", exc)

        if not zemin_ozet_yazildi: doc_replace_text_everywhere(doc, "[ZEMIN_OZET]", "Laboratuvar verisi girilmediği için zemin özeti oluşturulamadı.", paragraph_index=report_tag_index)
        islem_tablo_yerlestir(doc, "[LAB_FIZIK]", lab_fizik_headers, [], paragraph_index=report_tag_index)
        islem_tablo_yerlestir(doc, "[LAB_MEKANIK]", lab_mekanik_headers, [], paragraph_index=report_tag_index)

        lito_paragraphs = litoloji_dagilim_paragraflari(app_instance.veri.get("sondaj", []))
        replace_tag_with_paragraphs(doc, "[LITOLOJI_DAGILIM]", lito_paragraphs, paragraph_index=report_tag_index)
        lab_detail = _report_detail(
            lab_sheet=lab_sheet_ready,
            lab_excel=bool(lab_excel_path),
            lab_units=len(lab_birim_isimleri),
            lito_paragraphs=len(lito_paragraphs),
        )
        perf_log("report.lab_and_lithology.detail", time.perf_counter() - lab_section_start, lab_detail)
        report_step("lab_and_lithology", lab_detail)

        field_report_data = arazi_deney_rapor_verileri(app_instance.veri["sondaj"])
        spt_data = field_report_data["spt_data"]
        pmt_data = field_report_data["pmt_data"]
        kaya_data = field_report_data["kaya_data"]
        arazi_deney_word_bolumlerini_uygula(doc, report_tag_index, field_report_data)
        field_tables_start = time.perf_counter()
        p = report_tag_index.get("[SPT]")
        if p is not None and "[SPT]" in p.text:
                p.text = p.text.replace("[SPT]", ""); headers = ["Kuyu No", "Derinlik", "0-15", "15-30", "30-45", "N30"]; table = create_word_table(doc, headers, [])
                for s in app_instance.veri["sondaj"]:
                    spt_rows = s.get("spt", []); valid_spt = [row for row in spt_rows if len(row) >= 5]
                    if not valid_spt: continue
                    start_idx = len(table.rows)
                    for idx, row in enumerate(valid_spt):
                        r = table.add_row(); kuyu_text = clean_val(s["no"]) if idx == 0 else ""
                        set_cell_text_clean(r.cells[0], kuyu_text); set_cell_text_clean(r.cells[1], clean_val(row[0])); set_cell_text_clean(r.cells[2], clean_val(row[1])); set_cell_text_clean(r.cells[3], clean_val(row[2])); set_cell_text_clean(r.cells[4], clean_val(row[3])); set_cell_text_clean(r.cells[5], clean_val(row[4]))
                        for c in r.cells: set_vertical_cell_alignment(c, "center")
                    end_idx = len(table.rows) - 1
                    if end_idx > start_idx: first = table.rows[start_idx].cells[0]; first.merge(table.rows[end_idx].cells[0]); set_cell_text_clean(first, clean_val(s["no"])); set_vertical_cell_alignment(first, "center")
                apply_report_table_style(table, header_rows=1, widths_cm=[2.0, 2.0, 1.5, 1.5, 1.5, 1.5])
                p._p.addnext(table._tbl)

        if pmt_data:
            islem_tablo_yerlestir(doc, "[PMT]", ["Kuyu No", "Derinlik", "Em (kg/cm2)", "Pl (kg/cm2)"], pmt_data, paragraph_index=report_tag_index)
        if kaya_data:
            islem_tablo_yerlestir(doc, "[KAYA_TABLO]", ["Kuyu No", "Derinlik", "TCR (%)", "SCR (%)", "RQD (%)"], kaya_data, paragraph_index=report_tag_index)
        field_detail = _report_detail(spt=len(spt_data), pmt=len(pmt_data), kaya=len(kaya_data))
        perf_log("report.table.field_tests", time.perf_counter() - field_tables_start, field_detail)
        report_step("field_test_tables", field_detail)
        jeo_parametre_start = time.perf_counter()
        valid_serimler = []
        jeo_parametre_table_count = 0
        
        # --- JEO_PARAMETRE ÇİFT BAŞLIKLI DİKEY TABLO (TRANSPOZE) ---
        p = report_tag_index.get("[JEO_PARAMETRE]")
        if p is not None and "[JEO_PARAMETRE]" in p.text:
                p.text = p.text.replace("[JEO_PARAMETRE]", "")

                param_etiketleri = [
                    ("Kalınlık (m)", "h"), ("Vp (m/s)", "vp"), ("Vs (m/s)", "vs"), 
                    ("Yoğunluk (g/cm³)", "rho"), ("Poisson Oranı", "nu"), 
                    ("Elastisite Mod. (kg/cm²)", "E"), ("Kayma Mod. (kg/cm²)", "G"), 
                    ("Bulk Mod. (kg/cm²)", "K")
                ]

                def jeo_parametre_tablosu_olustur(serimler):
                    all_layers_flat = []
                    serim_gruplari = []
                    for ss in serimler:
                        layers = ss.get("layers", [])
                        if not layers:
                            continue
                        serim_ad = clean_val(ss.get("ad", "-"))
                        all_layers_flat.extend([(layer, idx == len(layers) - 1) for idx, layer in enumerate(layers)])
                        serim_gruplari.append((serim_ad, len(layers)))

                    total_cols = 1 + len(all_layers_flat)
                    if total_cols <= 1:
                        return None

                    table = doc.add_table(rows=2, cols=total_cols)
                    table.style = 'Table Grid'
                    set_cell_text_clean(table.rows[0].cells[0], "Parametre", bold=True)

                    current_col = 1
                    for serim_ad, layer_count in serim_gruplari:
                        start_cell = table.rows[0].cells[current_col]
                        end_cell = table.rows[0].cells[current_col + layer_count - 1]
                        if layer_count > 1:
                            start_cell.merge(end_cell)
                        set_cell_text_clean(start_cell, serim_ad, bold=True)
                        current_col += layer_count

                    table.rows[0].cells[0].merge(table.rows[1].cells[0])

                    current_col = 1
                    for _, layer_count in serim_gruplari:
                        for layer_idx in range(1, layer_count + 1):
                            set_cell_text_clean(table.rows[1].cells[current_col], f"Tab. {layer_idx}", bold=True)
                            current_col += 1

                    kirmizi_parametre_satirlari = []
                    for baslik, anahtar in param_etiketleri:
                        row_cells = table.add_row().cells
                        set_cell_text_clean(row_cells[0], baslik, bold=True)
                        for col_idx, (layer, is_last) in enumerate(all_layers_flat):
                            val = jeo_parametre_degeri_formatla(anahtar, layer.get(anahtar, "-"), is_last)
                            set_cell_text_clean(row_cells[col_idx + 1], val)
                        if anahtar in ("vp", "vs"):
                            kirmizi_parametre_satirlari.append(row_cells)

                    for row in table.rows:
                        for cell in row.cells:
                            set_vertical_cell_alignment(cell, "center")
                    apply_report_table_style(table, header_rows=2, label_cols={0}, repeat_headers=False)
                    for row_cells in kirmizi_parametre_satirlari:
                        for cell in row_cells:
                            style_cell_text(cell, font_color="C00000")
                    keep_table_together(table)
                    return table

                valid_serimler = [ss for ss in param_ss_list if ss.get("layers", [])]
                tables = []
                for start in range(0, len(valid_serimler), 3):
                    table = jeo_parametre_tablosu_olustur(valid_serimler[start:start + 3])
                    if table is not None:
                        tables.append(table)
                if tables:
                    jeo_parametre_table_count = len(tables)
                    anchor = p._p
                    for idx, table in enumerate(tables):
                        anchor.addnext(table._tbl)
                        anchor = table._tbl
                        if idx < len(tables) - 1:
                            page_break = word_sayfa_sonu_paragrafi()
                            anchor.addnext(page_break)
                            anchor = page_break
        jeo_parametre_detail = _report_detail(serim=len(valid_serimler), tables=jeo_parametre_table_count)
        perf_log("report.table.jeo_parametre", time.perf_counter() - jeo_parametre_start, jeo_parametre_detail)
        report_step("jeo_parametre_table", jeo_parametre_detail)

        p = report_tag_index.get("[MASW]")
        if p is not None and "[MASW]" in p.text:
                p.text = p.text.replace("[MASW]", "")
                masw_headers = ["Serim No", "Ortam No", "Vs(m/sn)", "Kalınlık h (m)", "Vs30(m/sn)"]
                table = create_word_table(doc, masw_headers, [])
                for ss in param_ss_list:
                    layers = ss.get("layers", [])
                    if not layers: continue
                    start_idx = len(table.rows)
                    vs30_val = fmt_jeo(layers[0].get("vs30", "-"))
                    serim_ad = clean_val(ss.get("ad", "-"))
                    
                    for idx, layer in enumerate(layers):
                        r = table.add_row()
                        set_cell_text_clean(r.cells[0], serim_ad if idx == 0 else "")
                        set_cell_text_clean(r.cells[1], str(idx + 1))
                        set_cell_text_clean(r.cells[2], fmt_jeo(layer.get("vs", "-")))
                        h_val = "-" if idx == len(layers) - 1 else fmt_jeo(layer.get("h", "-"))
                        set_cell_text_clean(r.cells[3], h_val)
                        set_cell_text_clean(r.cells[4], vs30_val if idx == 0 else "")
                        for c in r.cells: set_vertical_cell_alignment(c, "center")
                        
                    end_idx = len(table.rows) - 1
                    if end_idx > start_idx:
                        first_c0 = table.rows[start_idx].cells[0]
                        first_c0.merge(table.rows[end_idx].cells[0])
                        set_cell_text_clean(first_c0, serim_ad)
                        set_vertical_cell_alignment(first_c0, "center")
                        
                        first_c4 = table.rows[start_idx].cells[4]
                        first_c4.merge(table.rows[end_idx].cells[4])
                        set_cell_text_clean(first_c4, vs30_val)
                        set_vertical_cell_alignment(first_c4, "center")
                apply_report_table_style(table, header_rows=1, widths_cm=[2.2, 2.0, 2.3, 2.3, 2.3])
                p._p.addnext(table._tbl)
        masw_detail = _report_detail(serim=sum(1 for ss in param_ss_list if ss.get("layers", [])))
        report_step("masw_table", masw_detail)

        p = report_tag_index.get("[VP]")
        if p is not None and "[VP]" in p.text:
                p.text = p.text.replace("[VP]", "")
                vp_headers = ["Ölçü No", "Ortam No", "Vp (m/sn)"]
                table = create_word_table(doc, vp_headers, [])
                for ss in param_ss_list:
                    layers = jeofizik_vp_layers_sadelestir(ss.get("layers", []))
                    if not layers: continue
                    start_idx = len(table.rows)
                    serim_ad = clean_val(ss.get("ad", "-"))
                    
                    for idx, layer in enumerate(layers):
                        r = table.add_row()
                        set_cell_text_clean(r.cells[0], serim_ad if idx == 0 else "")
                        set_cell_text_clean(r.cells[1], str(idx + 1))
                        set_cell_text_clean(r.cells[2], fmt_jeo(layer.get("vp", "-")))
                        for c in r.cells: set_vertical_cell_alignment(c, "center")
                        
                    end_idx = len(table.rows) - 1
                    if end_idx > start_idx:
                        first_c0 = table.rows[start_idx].cells[0]
                        first_c0.merge(table.rows[end_idx].cells[0])
                        set_cell_text_clean(first_c0, serim_ad)
                        set_vertical_cell_alignment(first_c0, "center")
                apply_report_table_style(table, header_rows=1, widths_cm=[2.8, 2.0, 2.4])
                p._p.addnext(table._tbl)
        vp_detail = _report_detail(serim=sum(1 for ss in param_ss_list if ss.get("layers", [])))
        report_step("vp_table", vp_detail)

        p = report_tag_index.get("[JEO_KOOR]")
        if p is not None and "[JEO_KOOR]" in p.text:
                p.text = p.text.replace("[JEO_KOOR]", "")
                table = doc.add_table(rows=0, cols=7)
                table.style = 'Table Grid'
                
                def f_coord(v):
                    c = clean_val(v)
                    return f"{c}°" if c != "-" and not c.endswith("°") else c

                r0 = table.add_row()
                r0.cells[1].merge(r0.cells[6])
                set_cell_text_clean(r0.cells[1], "Koordinatlar(WGS84)", bold=True)
                
                r1 = table.add_row()
                r1.cells[1].merge(r1.cells[2]); set_cell_text_clean(r1.cells[1], "Düz Atış", bold=True)
                r1.cells[3].merge(r1.cells[4]); set_cell_text_clean(r1.cells[3], "Orta Atış", bold=True)
                r1.cells[5].merge(r1.cells[6]); set_cell_text_clean(r1.cells[5], "Ters Atış", bold=True)
                
                r2 = table.add_row()
                for idx, txt in enumerate(["", "Enlem", "Boylam", "Enlem", "Boylam", "Enlem", "Boylam"]):
                    if txt: set_cell_text_clean(r2.cells[idx], txt, bold=True)
                
                r0.cells[0].merge(r2.cells[0])
                set_cell_text_clean(r0.cells[0], "Çalışma No", bold=True)
                
                for ss in ss_list: 
                    r = table.add_row()
                    set_cell_text_clean(r.cells[0], clean_val(ss.get("ad", "-")))
                    c_list = ss.get("coords", ["-", "-", "-", "-", "-", "-"])
                    while len(c_list) < 6: c_list.append("-")
                    for i in range(6): set_cell_text_clean(r.cells[i+1], f_coord(c_list[i]))
                
                if mt_list:
                    rmt0 = table.add_row()
                    rmt0.cells[1].merge(rmt0.cells[6])
                    set_cell_text_clean(rmt0.cells[1], "Koordinatlar (WGS84)", bold=True)
                    
                    rmt1 = table.add_row()
                    rmt1.cells[1].merge(rmt1.cells[3]); set_cell_text_clean(rmt1.cells[1], "Enlem", bold=True)
                    rmt1.cells[4].merge(rmt1.cells[6]); set_cell_text_clean(rmt1.cells[4], "Boylam", bold=True)
                    
                    rmt0.cells[0].merge(rmt1.cells[0])
                    set_cell_text_clean(rmt0.cells[0], "Çalışma No", bold=True)
                    
                    for mt in mt_list:
                        r = table.add_row()
                        set_cell_text_clean(r.cells[0], clean_val(mt.get("no", "-")))
                        r.cells[1].merge(r.cells[3])
                        set_cell_text_clean(r.cells[1], f_coord(mt.get("y", "-")))
                        r.cells[4].merge(r.cells[6])
                        set_cell_text_clean(r.cells[4], f_coord(mt.get("x", "-")))

                for row in table.rows:
                    for cell in row.cells: set_vertical_cell_alignment(cell, "center")
                apply_report_table_style(table, header_rows=3, widths_cm=[2.2, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0])
                p._p.addnext(table._tbl)

        mt_table_data = []; mt_headers = ["Ölçü No", "Baskın Frekans (Hz)", "Baskın Periyot (To) (sn)", "Ta (sn)", "Tb (sn)", "H/V Oranı", "Kayıt Süresi (dk)", "Formasyon"]
        mt_formasyon = jeoloji_kisa_formasyon_metni(app_instance.veri) or "-"
        for mt in mt_list:
            row = [mt.get("no", "-"), clean_val(mt.get("freq", "-")), clean_val(mt.get("to", "-")), clean_val(mt.get("ta", "-")), clean_val(mt.get("tb", "-")), clean_val(mt.get("hv", "-")), clean_val(mt.get("sure", "-")), mt_formasyon]
            mt_table_data.append(row)
        islem_tablo_yerlestir(doc, "[MT_TABLO]", mt_headers, mt_table_data, paragraph_index=report_tag_index)
        report_step("jeofizik_coord_mt_tables", _report_detail(ss=len(ss_list), mt=len(mt_list)))

        # DİNAMİK JEOFİZİK SONUÇ CÜMLESİ OLUŞTURMA
        sonuc_parcalari = []
        if tum_vs30:
            vs30_min = int(min(tum_vs30)); vs30_max = int(max(tum_vs30))
            vs30_range_str = f"{vs30_min}-{vs30_max}" if vs30_min != vs30_max else f"{vs30_min}"
            sonuc_parcalari.append(f"Vs30={vs30_range_str} m/sn")
            
        if tum_t0:
            t0_min = min(tum_t0); t0_max = max(tum_t0)
            if t0_min == t0_max:
                t0_str = "{:.2f}".format(t0_min).replace('.', ',')
            else:
                t0_str = "{:.2f}-{:.2f}".format(t0_min, t0_max).replace('.', ',')
            sonuc_parcalari.append(f"zemin hakim titreşim periyodu {t0_str}sn")

        if sonuc_parcalari:
            birlestirilmis = " olarak, ".join(sonuc_parcalari)
            jeo_sonuc_cumlesi = f"Çalışma alanında yapılan jeofizik çalışmalar sonucunda {birlestirilmis} olarak bulunmuştur."
        else: 
            jeo_sonuc_cumlesi = ""
        
        if not yass_seviyeleri:
            yass_oneri = "Yapılan sondaj çalışmaları sonucunda çalışma alanında yeraltı suyuna rastlanmamıştır. Ancak, olası yüzey ve atık sularının yapı temeline ve temelin oturacağı zemine sızarak meydana getirebileceği olumsuz etkiler göz önüne alınarak; su geçirgenliğini önlemek amacıyla standartlara uygun bir yalıtım projelendirilmeli ve suları temelden uzak tutacak etkin bir drenaj sistemi oluşturulmalıdır."
        else:
            min_y = min(yass_seviyeleri)
            max_y = max(yass_seviyeleri)
            if min_y == max_y:
                r_str = f"-{min_y}m derinlikte"
            else:
                r_str = f"-{min_y}m ila -{max_y}m derinlikleri arasında"
            
            yass_oneri = f"Yapılan sondaj çalışmaları sonucunda çalışma alanında {r_str} yeraltı suyuna rastlanmıştır. Yeraltı, yüzey ve atık sularının yapı temeline ve temelin oturacağı zemine sızarak meydana getirebileceği olumsuz etkiler göz önüne alınarak; su geçirgenliğini önlemek amacıyla standartlara uygun bir yalıtım projelendirilmeli ve suları temelden uzak tutacak etkin bir drenaj sistemi oluşturulmalıdır."
            
        replace_text(doc, "[JEO_SONUC]", jeo_sonuc_cumlesi, paragraph_index=report_tag_index)
        replace_text(doc, "[YASS_ONERI]", yass_oneri, paragraph_index=report_tag_index)
        hidrojeoloji_metni = hidrojeoloji_durum_metni(arazi, sondajlar)
        hidrojeoloji_word_paragrafini_uygula(
            doc,
            report_tag_index,
            hidrojeoloji_metni,
        )
        report_step(
            "result_texts",
            _report_detail(
                vs30=len(tum_vs30),
                t0=len(tum_t0),
                yass=len(yass_seviyeleri),
                hidrojeoloji=bool(hidrojeoloji_metni),
            ),
        )

        mjh_path = mjh_resim_yolu(app_instance)
        image_paths = [
            app_instance.img_yer,
            app_instance.img_tkgm,
            app_instance.img_pga,
            getattr(app_instance, 'word_img_jeofizik', None),
            mjh_path,
            getattr(app_instance, 'word_img_sondaj', None),
        ]
        image_detail = _report_detail(paths=sum(1 for path in image_paths if path and os.path.exists(path)), tags=10)
        with perf_timer("report.replace_images", image_detail):
            doc_replace_img(doc, "RESIM:Yerbuldurur", app_instance.img_yer, paragraphs=report_paragraphs)
            doc_replace_img(doc, "[RESIM_YERBULDURUR]", app_instance.img_yer, paragraphs=report_paragraphs)

            doc_replace_img(doc, "RESIM:TKGM", app_instance.img_tkgm, paragraphs=report_paragraphs)
            doc_replace_img(doc, "RESIM:PGA", app_instance.img_pga, paragraphs=report_paragraphs)

            doc_replace_img(doc, "[RESIM_JEOFIZIK]", getattr(app_instance, 'word_img_jeofizik', None), paragraphs=report_paragraphs)
            doc_replace_img(doc, "RESIM:MJH", mjh_path, paragraphs=report_paragraphs)
            doc_replace_img(doc, "[RESIM_MJH]", mjh_path, paragraphs=report_paragraphs)
            doc_replace_img(doc, "[RESIM:MJH]", mjh_path, paragraphs=report_paragraphs)
            doc_replace_img(doc, "[RESIM_SONDAJ]", getattr(app_instance, 'word_img_sondaj', None), paragraphs=report_paragraphs)
            doc_replace_img(doc, "[RESIM:SONDAJ]", getattr(app_instance, 'word_img_sondaj', None), paragraphs=report_paragraphs)
        report_step("images", image_detail)

        removed_sections = rapor_kosullu_bolumlerini_uygula(
            doc,
            app_instance.veri,
        )
        report_step(
            "conditional_sections",
            _report_detail(removed=",".join(removed_sections) or "none"),
        )
        
        final = final_path
        if final is None:
            cikti_klasor = app_instance.veri.get("ayarlar", {}).get("varsayilan_cikti_klasor", "")
            save_opts = {"defaultextension": ".docx", "filetypes": [("Word Dosyası", "*.docx")]}
            if cikti_klasor and os.path.isdir(cikti_klasor):
                save_opts["initialdir"] = cikti_klasor
            final = filedialog.asksaveasfilename(**save_opts)
        if final:
            heading_page_break_enabled = str(app_instance.veri.get("ayarlar", {}).get("rapor_buyuk_baslik_yeni_sayfa", "1")).lower() not in ("0", "false", "no", "off", "hayir", "hayır")
            if heading_page_break_enabled:
                with perf_timer("report.major_headings_page_break"):
                    heading_count = buyuk_basliklari_yeni_sayfaya_al(doc)
                report_step("major_headings_page_break", _report_detail(headings=heading_count))
            docx_metadata_nortrle(doc)
            with perf_timer("report.save_docx", final):
                atomic_docx_save(doc, final)
            report_step("save_docx", _file_perf_detail(final))
            return True, "Rapor oluşturuldu!"
        return False, "İptal edildi."

    except Exception as e:
        traceback.print_exc(); return False, f"Hata: {str(e)}"

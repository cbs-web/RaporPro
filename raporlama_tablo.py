# Dosya: RaporPro/raporlama_tablo.py
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from performans import log_exception
from raporlama_deger import clean_val


def _log_silent(name, exc):
    log_exception(f"raporlama_tablo.{name}", exc_value=exc)

TABLE_HEADER_FILL = "D9E2F3"
TABLE_LABEL_FILL = "F2F5F9"
TABLE_ALT_FILL = "FAFBFC"
TABLE_BORDER_COLOR = "B7C1CC"
TABLE_TEXT_COLOR = "1F2937"

def set_vertical_cell_alignment(cell, align="center"):
    try: tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcValign = OxmlElement('w:vAlign'); tcValign.set(qn('w:val'), align); tcPr.append(tcValign)
    except Exception as exc: _log_silent("set_vertical_cell_alignment", exc)

def _xml_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child

def _rgb_from_hex(value):
    value = str(value or "").strip().lstrip("#")
    if len(value) != 6:
        return None
    try:
        return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))
    except ValueError:
        return None

def set_cell_shading(cell, fill):
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = _xml_child(tc_pr, "w:shd")
        shd.set(qn("w:fill"), fill)
    except Exception as exc:
        _log_silent("set_cell_shading", exc)

def set_cell_margins(cell, top=70, left=80, bottom=70, right=80):
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = _xml_child(tc_pr, "w:tcMar")
        for key, value in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
            node = _xml_child(tc_mar, f"w:{key}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
    except Exception as exc:
        _log_silent("set_cell_margins", exc)

def set_cell_border(cell, color=TABLE_BORDER_COLOR, size="6"):
    try:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = _xml_child(tc_pr, "w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            node = _xml_child(borders, f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)
    except Exception as exc:
        _log_silent("set_cell_border", exc)

def set_cell_width(cell, width_cm):
    try:
        if not width_cm:
            return
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = _xml_child(tc_pr, "w:tcW")
        tc_w.set(qn("w:w"), str(int(float(width_cm) * 567)))
        tc_w.set(qn("w:type"), "dxa")
    except Exception as exc:
        _log_silent("set_cell_width", exc)

def repeat_table_header(row):
    try:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))
    except Exception as exc:
        _log_silent("repeat_table_header", exc)

def set_table_fit_to_window(table):
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception as exc:
        _log_silent("set_table_fit_to_window.alignment", exc)
    try:
        table.autofit = True
    except Exception as exc:
        _log_silent("set_table_fit_to_window.autofit", exc)
    try:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        tbl_w = _xml_child(tbl_pr, "w:tblW")
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")
        layout = _xml_child(tbl_pr, "w:tblLayout")
        layout.set(qn("w:type"), "autofit")
    except Exception as exc:
        _log_silent("set_table_fit_to_window.tblpr", exc)

def style_cell_text(cell, bold=None, font_size=None, font_color=None, alignment=None):
    rgb = _rgb_from_hex(font_color)
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            if bold is not None:
                run.font.bold = bold
            if font_size is not None:
                run.font.size = Pt(font_size)
            if rgb is not None:
                run.font.color.rgb = rgb

def style_report_table_row(row, fill=TABLE_HEADER_FILL, bold=True, font_size=10):
    for cell in row.cells:
        set_cell_shading(cell, fill)
        set_vertical_cell_alignment(cell, "center")
        set_cell_margins(cell)
        set_cell_border(cell)
        style_cell_text(cell, bold=bold, font_size=font_size, font_color=TABLE_TEXT_COLOR, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def apply_report_table_style(table, header_rows=1, label_cols=None, text_cols=None, widths_cm=None):
    label_cols = set(label_cols or [])
    text_cols = set(text_cols or [])
    set_table_fit_to_window(table)
    for row_idx, row in enumerate(table.rows):
        is_header = row_idx < header_rows
        if is_header:
            repeat_table_header(row)
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell)
            set_cell_border(cell)
            set_vertical_cell_alignment(cell, "center")
            if widths_cm and col_idx < len(widths_cm):
                set_cell_width(cell, widths_cm[col_idx])
            alignment = WD_ALIGN_PARAGRAPH.LEFT if (not is_header and col_idx in text_cols) else WD_ALIGN_PARAGRAPH.CENTER
            style_cell_text(cell, bold=None, font_size=10, font_color=TABLE_TEXT_COLOR, alignment=alignment)
            if is_header:
                set_cell_shading(cell, TABLE_HEADER_FILL)
                style_cell_text(cell, bold=True, font_size=10, font_color=TABLE_TEXT_COLOR, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            elif col_idx in label_cols:
                set_cell_shading(cell, TABLE_LABEL_FILL)
                style_cell_text(cell, bold=True, font_size=10, font_color=TABLE_TEXT_COLOR, alignment=alignment)
            elif row_idx % 2 == 0:
                set_cell_shading(cell, TABLE_ALT_FILL)

def set_cell_text_clean(cell, text, font_name="Times New Roman", font_size=10, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    if len(cell.paragraphs) > 1:
        for i in range(len(cell.paragraphs) - 1, 0, -1): p_element = cell.paragraphs[i]._element; p_element.getparent().remove(p_element)
    p = cell.paragraphs[0]; p.clear(); p.alignment = alignment
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1
    set_cell_margins(cell); set_cell_border(cell)
    run = p.add_run(clean_val(text)); run.font.name = font_name; run.font.size = Pt(font_size); run.font.bold = bold; run.font.italic = italic
    rgb = _rgb_from_hex(TABLE_TEXT_COLOR)
    if rgb is not None:
        run.font.color.rgb = rgb
    return run

def create_word_table(doc, headers, data, text_cols=None, widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers): set_cell_text_clean(hdr_cells[i], h, font_size=11, bold=True); set_vertical_cell_alignment(hdr_cells[i], "center")
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data): set_cell_text_clean(row_cells[i], val, font_size=10, bold=False); set_vertical_cell_alignment(row_cells[i], "center")
    apply_report_table_style(table, header_rows=1, text_cols=text_cols, widths_cm=widths_cm)
    return table

# Dosya: RaporPro/yonetmelik_motoru.py
import datetime
from html.parser import HTMLParser
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import tempfile
import unicodedata
from urllib.parse import unquote, urljoin, urlparse
from urllib.request import Request, urlopen

from yardimcilar import atomic_json_dump, temizle_baslik


RAPORPRO_CONFIG_DIR = Path(os.environ.get("APPDATA") or (Path.home() / "AppData" / "Roaming")) / "RaporPro"
YONETMELIK_DIR = RAPORPRO_CONFIG_DIR / "yonetmelikler"
YONETMELIK_INDEX_PATH = YONETMELIK_DIR / "index.json"
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".html", ".htm")
OFFICIAL_USER_AGENT = "RaporPro/1.0 (+https://yapiisleri.csb.gov.tr)"
RESMI_YONETMELIK_KAYNAKLARI = {
    "zemin_temel_etudu_2019": {
        "id": "zemin_temel_etudu_2019",
        "title": "Zemin ve Temel Etüdü Uygulama Esasları ve Rapor Formatı",
        "page_url": "https://yapiisleri.csb.gov.tr/zemin-ve-temel-etudu-uygulama-esaslari-ve-rapor-formati-haber-238674",
        "expected_extensions": (".docx",),
        "link_text_hints": ("eki indir", "eki indirmek", "tıklayınız", "tiklayiniz"),
        "description": "Çevre, Şehircilik ve İklim Değişikliği Bakanlığı resmi DOCX eki",
    },
}
RESMI_YONETMELIK_YERLESIK_METINLER = {
    "zemin_temel_etudu_2019": """
ZEMİN VE TEMEL ETÜDÜ UYGULAMA ESASLARI VE RAPOR FORMATI
Yerleşik RaporPro referans metni

Kaynak kimliği:
Çevre, Şehircilik ve İklim Değişikliği Bakanlığı tarafından duyurulan
Zemin ve Temel Etüdü Uygulama Esasları ve Rapor Formatı, 09.03.2019 tarihli
ve 30709 sayılı Resmî Gazete yayımı ile yürürlüğe girmiştir. Bu yerleşik
metin, program internet bağlantısı olmadan çalışırken düzeltme notlarını
ilgili zemin etüdü başlıklarıyla eşleştirmek için kullanılan sabit başvuru
özetidir. Tam resmi DOCX/PDF dosyası elde edildiğinde Yönetmelik Merkezi'ne
ayrıca eklenebilir.

Kapsam:
Parsel bazında hazırlanacak zemin ve temel etüt raporlarında arazi
araştırmaları, laboratuvar çalışmaları, jeofizik çalışmalar, yeraltı suyu
verileri, yerel deprem etkileri, zemin veya kaya birimlerinin mühendislik
özellikleri ve rapor formatı birlikte değerlendirilir.

Arazi araştırmaları ve sondajlar:
Sondajlar, yapı özellikleri, temel sistemi, zemin koşulları ve araştırma
amacına uygun olacak şekilde planlanır. Sondaj derinliği ve sayısı sahadaki
zemin veya kaya koşullarını temsil edecek yeterlilikte olmalıdır. SPT,
presiyometre, yerinde deneyler, karot, örselenmiş numune ve örselenmemiş
numune bilgileri raporda tutarlı şekilde gösterilir. Sondaj loglarında
litoloji sürekliliği, deney derinlikleri, numune derinlikleri ve yeraltı suyu
ölçümleri kontrol edilir.

Laboratuvar deneyleri:
Laboratuvar deneyleri zemin veya kaya birimlerini temsil edecek seviyelerde
yapılır. Eksik deney aralıkları, numune alınmayan seviyeler, sınıflama
deneyleri, doğal birim hacim ağırlık, su muhtevası, Atterberg limitleri,
granülometri, kesme kutusu, üç eksenli basınç, konsolidasyon, tek eksenli
basınç, nokta yükleme, TCR, SCR ve RQD gibi deneylerin rapor tablolarıyla
uyumu denetlenir. Belediye veya idare ek laboratuvar deneyi istediğinde rapor
tabloları ve zemin parametre özetleri yeni verilere göre yenilenmelidir.

Jeofizik ve deprem verileri:
MASW, sismik kırılma, mikrotremor ve benzeri jeofizik çalışmalar varsa
serim, tabaka, Vs, Vp, zemin hakim periyodu ve benzeri değerler rapor
tablolarıyla uyumlu verilmelidir. Türkiye Bina Deprem Yönetmeliği ile ilişkili
yerel zemin sınıfı, tasarım spektrumu ve deprem parametreleri raporda tutarlı
olmalıdır.

Harita, koordinat ve görseller:
Araştırma noktaları vaziyet planı, sondaj lokasyon haritası, jeofizik
lokasyon haritası, mühendislik jeolojisi haritası ve koordinat tabloları
birbiriyle uyumlu olmalıdır. Sondaj, sismik serim ve mikrotremor noktalarının
etiketleri ve koordinatları rapordaki tablo ve haritalarla çelişmemelidir.

Rapor formatı ve düzeltme kontrolü:
Rapor metni, çizelgeler, loglar, kesitler, haritalar ve ekler aynı proje
verisine dayanmalıdır. İdare düzeltme talebinde eksik sondaj, eksik deney,
yanlış derinlik, yanlış koordinat, güncel olmayan harita, eksik jeofizik veri,
eksik bina bilgisi veya tutarsız zemin parametresi belirtilirse ilgili veri
girişi tamamlanmalı ve sadece etkilenen rapor bölümleri yenilenmelidir.
""".strip(),
}


class _LinkParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.links = []
        self._active = None

    def handle_starttag(self, tag, attrs):
        if tag.lower() != "a":
            return
        attrs = dict(attrs)
        href = attrs.get("href")
        if href:
            self._active = {"href": href, "text": ""}

    def handle_data(self, data):
        if self._active is not None:
            self._active["text"] += data or ""

    def handle_endtag(self, tag):
        if tag.lower() == "a" and self._active is not None:
            self.links.append(self._active)
            self._active = None


class _TextHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag.lower() in {"script", "style", "noscript"}:
            self._skip_depth += 1

    def handle_endtag(self, tag):
        if tag.lower() in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if self._skip_depth:
            return
        text = re.sub(r"\s+", " ", str(data or "")).strip()
        if text:
            self.parts.append(text)

    def text(self):
        return "\n".join(self.parts)


def _now_iso():
    return datetime.datetime.now().replace(microsecond=0).isoformat()


def _normalize(text):
    text = str(text or "").lower()
    text = text.replace("ı", "i").replace("İ", "i")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^a-z0-9çğıöşü\s\-.]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _tokens(text):
    stop = {
        "ve", "veya", "ile", "icin", "için", "olan", "olarak", "gore", "göre",
        "bir", "bu", "su", "şu", "da", "de", "ile", "ise", "var", "yok",
        "rapor", "talep", "duzeltme", "düzeltme", "belediye", "kontrolor", "kontrolör",
    }
    result = []
    for token in re.findall(r"[a-z0-9çğıöşü\-.]{3,}", _normalize(text)):
        if token not in stop:
            result.append(token)
    return result


def _safe_stem(path, title=""):
    base = title or Path(path).stem
    cleaned = temizle_baslik(base).strip("._- ") or "yonetmelik"
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", cleaned)[:80] or "yonetmelik"


def _read_text_file(path):
    for encoding in ("utf-8", "cp1254", "latin-1"):
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def _read_html(path):
    html = _read_text_file(path)
    parser = _TextHTMLParser()
    parser.feed(html)
    return parser.text()


def _read_docx(path):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"python-docx yüklenemedi: {exc}") from exc
    doc = Document(path)
    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(path):
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError(f"PDF metni için PyMuPDF/fitz yüklenemedi: {exc}") from exc
    text_parts = []
    doc = fitz.open(path)
    try:
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                text_parts.append(text.strip())
    finally:
        doc.close()
    return "\n\n".join(text_parts)


def yonetmelik_metni_oku(path):
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Desteklenmeyen yönetmelik dosyası: {ext}")
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext in (".html", ".htm"):
        return _read_html(path)
    return _read_text_file(path)


def _baslik_mi(line):
    stripped = str(line or "").strip()
    if not stripped:
        return False
    if len(stripped) > 140:
        return False
    patterns = [
        r"^(madde|MADDE)\s+\d+",
        r"^\d+(\.\d+)*[\)\.-]\s+",
        r"^(bölüm|BÖLÜM|kısım|KISIM)\s+",
        r"^[A-ZÇĞİÖŞÜ0-9\s\-/]{8,}$",
    ]
    return any(re.search(pattern, stripped) for pattern in patterns)


def _chunk_paragraphs(text, max_chars=2200):
    lines = [re.sub(r"\s+", " ", line).strip() for line in str(text or "").splitlines()]
    chunks = []
    title = "Genel"
    current = []

    def flush():
        nonlocal current
        body = "\n".join(item for item in current if item).strip()
        if body:
            chunks.append({"title": title, "text": body})
        current = []

    for line in lines:
        if not line:
            continue
        if _baslik_mi(line) and current:
            flush()
            title = line[:160]
            continue
        if _baslik_mi(line) and not current:
            title = line[:160]
            continue
        current.append(line)
        if sum(len(item) + 1 for item in current) >= max_chars:
            flush()
    flush()

    if not chunks and str(text or "").strip():
        raw = re.sub(r"\s+", " ", str(text)).strip()
        chunks = [{"title": "Genel", "text": raw[i:i + max_chars]} for i in range(0, len(raw), max_chars)]

    cleaned = []
    for idx, chunk in enumerate(chunks, start=1):
        body = chunk["text"].strip()
        if len(body) < 30:
            continue
        cleaned.append({
            "chunk_id": f"c{idx}",
            "title": chunk.get("title") or f"Bölüm {idx}",
            "text": body,
            "norm": _normalize(f"{chunk.get('title', '')} {body}"),
        })
    return cleaned


def _load_index(base_dir=None):
    index_path = Path(base_dir) / "index.json" if base_dir else YONETMELIK_INDEX_PATH
    if not index_path.exists():
        return {"version": 1, "documents": []}
    try:
        data = json.loads(index_path.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            data.setdefault("version", 1)
            data.setdefault("documents", [])
            return data
    except Exception:
        pass
    return {"version": 1, "documents": []}


def _save_index(index, base_dir=None):
    base = Path(base_dir) if base_dir else YONETMELIK_DIR
    base.mkdir(parents=True, exist_ok=True)
    atomic_json_dump(index, base / "index.json", ensure_ascii=False, indent=2)


def _doc_chunks_path(doc_id, base_dir=None):
    base = Path(base_dir) if base_dir else YONETMELIK_DIR
    return base / f"{doc_id}.chunks.json"


def resmi_yonetmelik_kaynaklari():
    return [dict(item) for item in RESMI_YONETMELIK_KAYNAKLARI.values()]


def resmi_yonetmelik_baglanti_bul(html, page_url, expected_extensions=None, link_text_hints=None):
    expected_extensions = tuple(ext.lower() for ext in (expected_extensions or SUPPORTED_EXTENSIONS))
    link_text_hints = tuple(_normalize(item) for item in (link_text_hints or ()))
    parser = _LinkParser()
    parser.feed(str(html or ""))
    candidates = []
    for link in parser.links:
        href = str(link.get("href") or "").strip()
        if not href or href.startswith("#") or href.lower().startswith("javascript:"):
            continue
        absolute_url = urljoin(page_url, href)
        url_lower = absolute_url.lower()
        text_norm = _normalize(link.get("text", ""))
        score = 0
        if any(url_lower.split("?", 1)[0].endswith(ext) for ext in expected_extensions):
            score += 60
        if "webdosya" in url_lower or "dosya" in url_lower:
            score += 15
        if link_text_hints and any(hint and hint in text_norm for hint in link_text_hints):
            score += 30
        if score > 0:
            candidates.append((score, absolute_url))
    if not candidates:
        return ""
    candidates.sort(key=lambda item: (-item[0], item[1]))
    return candidates[0][1]


def _filename_from_url(url, fallback="yonetmelik.docx"):
    path = unquote(urlparse(url).path or "")
    name = Path(path).name or fallback
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" ._") or fallback
    if Path(name).suffix.lower() not in SUPPORTED_EXTENSIONS:
        suffix = Path(fallback).suffix or ".docx"
        name = f"{Path(name).stem or 'yonetmelik'}{suffix}"
    return name


def _download_bytes(url, timeout=45):
    request = Request(url, headers={"User-Agent": OFFICIAL_USER_AGENT})
    with urlopen(request, timeout=timeout) as response:
        return response.read()


def _update_doc_metadata(doc_id, metadata, base_dir=None):
    if not metadata:
        return
    index = _load_index(base_dir)
    changed = False
    for doc in index.get("documents", []):
        if doc.get("id") == doc_id:
            doc.update(metadata)
            changed = True
            break
    if changed:
        _save_index(index, base_dir)


def yonetmelikleri_listele(base_dir=None):
    index = _load_index(base_dir)
    return list(index.get("documents") or [])


def yonetmelik_ekle(path, title=None, base_dir=None, metadata=None):
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError("Yönetmelik dosyası bulunamadı.")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Desteklenmeyen dosya türü: {ext}")

    base = Path(base_dir) if base_dir else YONETMELIK_DIR
    base.mkdir(parents=True, exist_ok=True)
    text = yonetmelik_metni_oku(path)
    if len(str(text or "").strip()) < 50:
        raise RuntimeError("Yönetmelik metni okunamadı veya çok kısa görünüyor.")

    digest = hashlib.sha1((str(path.resolve()) + str(path.stat().st_mtime_ns) + text[:5000]).encode("utf-8", errors="ignore")).hexdigest()[:12]
    doc_id = f"ym_{digest}"
    doc_title = str(title or path.stem).strip() or path.stem
    stored_name = f"{doc_id}_{_safe_stem(path, doc_title)}{ext}"
    stored_path = base / stored_name
    if path.resolve() != stored_path.resolve():
        shutil.copy2(path, stored_path)

    chunks = _chunk_paragraphs(text)
    chunks_path = _doc_chunks_path(doc_id, base)
    atomic_json_dump({"doc_id": doc_id, "chunks": chunks}, chunks_path, ensure_ascii=False, indent=2)

    index = _load_index(base)
    docs = [doc for doc in index.get("documents", []) if doc.get("id") != doc_id]
    record = {
        "id": doc_id,
        "title": doc_title,
        "source_path": str(path),
        "stored_path": str(stored_path),
        "added_at": _now_iso(),
        "ext": ext,
        "char_count": len(text),
        "chunk_count": len(chunks),
        "active": True,
    }
    if metadata:
        record.update(metadata)
    docs.append(record)
    docs.sort(key=lambda item: str(item.get("title", "")).lower())
    index["documents"] = docs
    _save_index(index, base)
    return record


def yerlesik_yonetmelik_ekle(source_id, base_dir=None, refresh=False, download_error=None):
    source = RESMI_YONETMELIK_KAYNAKLARI.get(source_id)
    text = RESMI_YONETMELIK_YERLESIK_METINLER.get(source_id, "")
    if not source or not text:
        raise KeyError(f"Yerleşik yönetmelik kaynağı bulunamadı: {source_id}")
    if not refresh:
        for doc in yonetmelikleri_listele(base_dir):
            if doc.get("official_id") == source_id:
                return {
                    "record": doc,
                    "already_exists": True,
                    "embedded_fallback": bool(doc.get("embedded_fallback")),
                    "download_url": doc.get("official_download_url", ""),
                }

    with tempfile.TemporaryDirectory() as tmp:
        temp_path = Path(tmp) / f"{source_id}_yerlesik.txt"
        temp_path.write_text(text, encoding="utf-8")
        record = yonetmelik_ekle(
            temp_path,
            title=f"{source.get('title')} (Yerleşik)",
            base_dir=base_dir,
            metadata={
                "official_id": source_id,
                "official_page_url": source.get("page_url", ""),
                "official_download_url": "",
                "description": source.get("description", ""),
                "embedded_fallback": True,
                "download_error": str(download_error or ""),
            },
        )
    _update_doc_metadata(record["id"], {"source_path": "RaporPro yerleşik kaynak"}, base_dir=base_dir)
    record["source_path"] = "RaporPro yerleşik kaynak"
    return {"record": record, "already_exists": False, "embedded_fallback": True, "download_url": ""}


def varsayilan_yonetmelikleri_hazirla(base_dir=None):
    results = []
    for source_id in RESMI_YONETMELIK_KAYNAKLARI:
        results.append(yerlesik_yonetmelik_ekle(source_id, base_dir=base_dir))
    return results


def resmi_yonetmelik_indir_ve_ekle(source_id, base_dir=None, timeout=45, refresh=False, allow_embedded_fallback=True):
    source = RESMI_YONETMELIK_KAYNAKLARI.get(source_id)
    if not source:
        raise KeyError(f"Resmi yönetmelik kaynağı bulunamadı: {source_id}")
    if not refresh:
        for doc in yonetmelikleri_listele(base_dir):
            if doc.get("official_id") == source_id:
                return {"record": doc, "already_exists": True, "download_url": doc.get("official_download_url", "")}

    page_url = source.get("page_url", "")
    download_url = source.get("download_url", "")
    try:
        if not download_url:
            html = _download_bytes(page_url, timeout=timeout).decode("utf-8", errors="ignore")
            download_url = resmi_yonetmelik_baglanti_bul(
                html,
                page_url,
                expected_extensions=source.get("expected_extensions"),
                link_text_hints=source.get("link_text_hints"),
            )
        if not download_url:
            raise RuntimeError("Resmi sayfada indirilecek yönetmelik eki bulunamadı.")
        data = _download_bytes(download_url, timeout=timeout)
    except Exception as exc:
        if allow_embedded_fallback:
            return yerlesik_yonetmelik_ekle(source_id, base_dir=base_dir, refresh=refresh, download_error=exc)
        raise

    filename = _filename_from_url(download_url, fallback=f"{source_id}.docx")
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        suffix = ".docx"

    with tempfile.TemporaryDirectory() as tmp:
        temp_path = Path(tmp) / filename
        if temp_path.suffix.lower() != suffix:
            temp_path = temp_path.with_suffix(suffix)
        temp_path.write_bytes(data)
        record = yonetmelik_ekle(
            temp_path,
            title=source.get("title"),
            base_dir=base_dir,
            metadata={
                "official_id": source_id,
                "official_page_url": page_url,
                "official_download_url": download_url,
                "description": source.get("description", ""),
            },
        )
    _update_doc_metadata(record["id"], {"source_path": download_url}, base_dir=base_dir)
    record["source_path"] = download_url
    return {"record": record, "already_exists": False, "download_url": download_url}


def yonetmelik_sil(doc_id, base_dir=None):
    base = Path(base_dir) if base_dir else YONETMELIK_DIR
    index = _load_index(base)
    docs = []
    removed = None
    for doc in index.get("documents", []):
        if doc.get("id") == doc_id:
            removed = doc
            continue
        docs.append(doc)
    if not removed:
        return False
    index["documents"] = docs
    _save_index(index, base)
    for path in [removed.get("stored_path"), str(_doc_chunks_path(doc_id, base))]:
        try:
            if path and Path(path).exists():
                Path(path).unlink()
        except Exception:
            pass
    return True


def _load_chunks(doc_id, base_dir=None):
    path = _doc_chunks_path(doc_id, base_dir)
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return list(data.get("chunks") or [])
    except Exception:
        return []


def _score_chunk(query_tokens, query_norm, chunk_norm):
    if not query_tokens or not chunk_norm:
        return 0
    score = 0
    token_set = set(query_tokens)
    for token in token_set:
        if token in chunk_norm:
            score += 3 + min(4, chunk_norm.count(token))
    phrases = [
        "laboratuvar deneyi", "sondaj derinliği", "spt", "presiyometre", "jeofizik",
        "yeraltı suyu", "zemin sınıfı", "rapor formatı", "numune", "karot",
    ]
    for phrase in phrases:
        norm_phrase = _normalize(phrase)
        if norm_phrase in query_norm and norm_phrase in chunk_norm:
            score += 8
    return score


def _excerpt(text, query_tokens, limit=420):
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(text) <= limit:
        return text
    lower = _normalize(text)
    positions = [lower.find(_normalize(token)) for token in query_tokens if lower.find(_normalize(token)) >= 0]
    pos = min(positions) if positions else 0
    start = max(0, pos - limit // 3)
    end = min(len(text), start + limit)
    snippet = text[start:end].strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(text):
        snippet += "..."
    return snippet


def yonetmelik_ara(query, limit=8, base_dir=None):
    query = str(query or "").strip()
    if not query:
        return []
    query_norm = _normalize(query)
    query_tokens = _tokens(query)
    results = []
    for doc in yonetmelikleri_listele(base_dir):
        if not doc.get("active", True):
            continue
        for chunk in _load_chunks(doc.get("id"), base_dir):
            score = _score_chunk(query_tokens, query_norm, chunk.get("norm") or _normalize(chunk.get("text", "")))
            if score <= 0:
                continue
            results.append({
                "score": score,
                "doc_id": doc.get("id"),
                "doc_title": doc.get("title", ""),
                "chunk_id": chunk.get("chunk_id", ""),
                "chunk_title": chunk.get("title", ""),
                "excerpt": _excerpt(chunk.get("text", ""), query_tokens),
                "stored_path": doc.get("stored_path", ""),
            })
    results.sort(key=lambda item: (-item.get("score", 0), item.get("doc_title", ""), item.get("chunk_title", "")))
    return results[:limit]


def duzeltme_yonetmelik_dayanaklari(text, limit=6, base_dir=None):
    text = str(text or "").strip()
    if not text:
        return {"items": [], "documents": yonetmelikleri_listele(base_dir), "warnings": ["Düzeltme metni boş."]}
    parts = []
    current = []
    for line in text.splitlines():
        if re.match(r"^\s*\d+\s*[\-\)\.]", line) and current:
            parts.append(" ".join(current).strip())
            current = [line.strip()]
        elif line.strip():
            current.append(line.strip())
    if current:
        parts.append(" ".join(current).strip())
    if not parts:
        parts = [text]

    items = []
    seen = set()
    per_part_limit = max(2, min(4, limit))
    for part in parts:
        for hit in yonetmelik_ara(part, limit=per_part_limit, base_dir=base_dir):
            key = (hit.get("doc_id"), hit.get("chunk_id"))
            if key in seen:
                continue
            seen.add(key)
            item = dict(hit)
            item["source_text"] = part
            items.append(item)
            if len(items) >= limit:
                break
        if len(items) >= limit:
            break

    warnings = []
    docs = yonetmelikleri_listele(base_dir)
    if not docs:
        warnings.append("Yönetmelik Merkezi'ne henüz dosya eklenmemiş.")
    elif not items:
        warnings.append("Düzeltme metniyle eşleşen yönetmelik dayanağı bulunamadı.")
    return {"items": items, "documents": docs, "warnings": warnings}

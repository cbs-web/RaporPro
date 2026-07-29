# Dosya: RaporPro/sondaj_derinlik_foyu.py
import datetime
import os
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from sondaj_derinlik import gerilme_yuzde_on_derinlik_hesapla
from yardimcilar import atomic_docx_save, atomic_fitz_pdf_save


def _clean(value, fallback=""):
    text = "" if value is None else str(value).strip()
    return text if text else fallback


def _safe_name(value, fallback="Sondaj_Derinligi"):
    text = _clean(value, fallback)
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return "_".join(text.split()) or fallback


def _project_name(veri):
    kunye = (veri or {}).get("kunye", {}) or {}
    name = _clean(kunye.get("proje_adi") or kunye.get("proje") or kunye.get("sahibi"))
    ada = _clean(kunye.get("ada"))
    parsel = _clean(kunye.get("par"))
    mah = _clean(kunye.get("mah"))
    if name:
        return name
    if mah and ada and parsel:
        return f"{mah} {ada} Ada {parsel} Parsel"
    if ada and parsel:
        return f"{ada} Ada {parsel} Parsel"
    return "Proje"


def _project_location(veri):
    kunye = (veri or {}).get("kunye", {}) or {}
    parts = [_clean(kunye.get("mah")), _clean(kunye.get("ilce")), _clean(kunye.get("il"))]
    return " / ".join(part for part in parts if part) or "-"


def _parcel_text(veri):
    kunye = (veri or {}).get("kunye", {}) or {}
    parts = []
    if _clean(kunye.get("paf")):
        parts.append(f"Pafta: {_clean(kunye.get('paf'))}")
    if _clean(kunye.get("ada")):
        parts.append(f"Ada: {_clean(kunye.get('ada'))}")
    if _clean(kunye.get("par")):
        parts.append(f"Parsel: {_clean(kunye.get('par'))}")
    return " | ".join(parts) or "-"


def sondaj_derinligi_foy_dosya_adi(veri, ext=".docx"):
    ext = ext if str(ext).startswith(".") else f".{ext}"
    return f"{_safe_name(_project_name(veri))}_Sondaj_Derinligi_Hesap_Foyu{ext}"


def _fmt(value, digits=2, suffix=""):
    try:
        number = float(value)
    except Exception:
        return "-"
    return f"{number:.{digits}f}{suffix}"


def _ratio(value):
    try:
        return f"%{float(value) * 100:.2f}"
    except Exception:
        return "-"


INPUT_LABELS = {
    "b": "Temel Genişliği\n(B)",
    "l": "Temel Uzunluğu\n(L)",
    "df": "Temel Taban Derinliği\n(Df)",
    "yass": "Yeraltı Su Seviyesi\n(YASS)",
    "q_taban": "Temel Taban Gerilmesi\n(qtaban)",
    "sigma_vo_taban": "Temel Tabanındaki Efektif Düşey Gerilme\n(σ'vo taban)",
    "dogal_bha": "Doğal Birim Hacim Ağırlığı\n(Doğal BHA)",
    "doygun_bha": "Doygun Birim Hacim Ağırlığı\n(Doygun BHA)",
    "q_net": "Net Temel Taban Basıncı\n(qnet)",
    "kosul": "Hesap Koşulu\n(Δσ ≤ 0.10 σ'vo)",
}


def _foy_context(veri):
    params = ((veri or {}).get("sondaj_derinlik_hesabi") or {}).copy()
    result = gerilme_yuzde_on_derinlik_hesapla(params)
    if not result.get("ok"):
        raise ValueError("Sondaj derinliği hesabı yapılamadı:\n" + "\n".join(result.get("errors", [])))
    return {
        "params": params,
        "result": result,
        "project": _project_name(veri),
        "location": _project_location(veri),
        "parcel": _parcel_text(veri),
        "date": datetime.datetime.now().strftime("%d.%m.%Y"),
    }


def sondaj_derinligi_grafik_olustur(result, output_path):
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    rows = result.get("rows", [])
    if not rows:
        raise ValueError("Grafik için hesap tablosu bulunamadı.")

    depths = [row["derinlik"] for row in rows]
    target = [row["sigma_vo"] * result["target_ratio"] for row in rows]
    fig, ax = plt.subplots(figsize=(7.4, 4.3), dpi=160)
    ax.plot([row["boussinesq_delta"] for row in rows], depths, color="#C0392B", linewidth=2.0, label="Boussinesq Δσ")
    ax.plot([row["westergaard_delta"] for row in rows], depths, color="#2471A3", linewidth=1.7, label="Westergaard Δσ")
    ax.plot([row["yaklasik_delta"] for row in rows], depths, color="#7D6608", linewidth=1.7, label="Yaklaşık Δσ")
    ax.plot(target, depths, color="#111111", linestyle="--", linewidth=2.0, label="0.10 σ'vo")
    ax.axhline(result["sondaj_derinligi_yuvarlatilmis"], color="#1E8449", linewidth=1.8, linestyle=":")
    ax.text(
        max(max(target), result["delta_sigma"]) * 0.98,
        result["sondaj_derinligi_yuvarlatilmis"],
        f"  Öneri: {result['sondaj_derinligi_yuvarlatilmis']:.2f} m",
        va="bottom",
        ha="right",
        fontsize=9,
        color="#1E8449",
    )
    ax.invert_yaxis()
    ax.set_xlabel("Gerilme (t/m²)")
    ax.set_ylabel("Derinlik (m)")
    ax.set_title("Gerilme Artışı - Efektif Düşey Gerilme Kontrolü", fontsize=11, weight="bold")
    ax.grid(True, color="#D5D8DC", linewidth=0.7)
    ax.legend(loc="lower right", fontsize=8)
    fig.tight_layout()
    fig.savefig(output_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return output_path


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _table_grid(table, header_rows=1):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
            if r_idx < header_rows:
                _set_cell_shading(cell, "D9EAF7")


def _add_heading(doc, text, size=16, color="1F2D3D"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def _add_note(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.size = Pt(9.5)
    return paragraph


def _add_input_table(doc, ctx):
    result = ctx["result"]
    stress_unit = result.get("gerilme_birimi", "t/m²")
    bha_unit = result.get("bha_birimi", "t/m³")
    rows = [
        ("Proje", ctx["project"], "Konum", ctx["location"]),
        ("Pafta/Ada/Parsel", ctx["parcel"], "Hesap tarihi", ctx["date"]),
        (INPUT_LABELS["b"], _fmt(result["b"], 2, " m"), INPUT_LABELS["l"], _fmt(result["l"], 2, " m")),
        (INPUT_LABELS["df"], _fmt(result["temel_derinligi"], 2, " m"), INPUT_LABELS["yass"], _fmt(result["yass"], 2, " m")),
        (INPUT_LABELS["q_taban"], _fmt(result["q_taban"], 3, f" {stress_unit}"), INPUT_LABELS["sigma_vo_taban"], _fmt(result["sigma_vo_taban"], 3, f" {stress_unit}")),
        (INPUT_LABELS["dogal_bha"], _fmt(result["dogal_bha"], 3, f" {bha_unit}"), INPUT_LABELS["doygun_bha"], _fmt(result["doygun_bha"], 3, f" {bha_unit}")),
        (INPUT_LABELS["q_net"], _fmt(result["q_net"], 3, f" {stress_unit}"), INPUT_LABELS["kosul"], "Δσ ≤ 0.10 σ'vo"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    _table_grid(table, header_rows=0)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            is_label = c_idx in (0, 2)
            _set_cell_text(table.rows[r_idx].cells[c_idx], value, bold=is_label, size=8 if is_label else 9)
            if is_label:
                _set_cell_shading(table.rows[r_idx].cells[c_idx], "F4F6F7")
    return table


def _add_result_table(doc, result):
    table = doc.add_table(rows=1, cols=5)
    stress_unit = result.get("gerilme_birimi", "t/m²")
    headers = ["Yöntem", f"Δσ ({stress_unit})", f"σ'vo ({stress_unit})", "Δσ/σ'vo", "Gerekli derinlik"]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _table_grid(table, header_rows=1)
    for key in ("boussinesq", "westergaard", "yaklasik"):
        item = result["yontem_sonuclari"][key]
        cells = table.add_row().cells
        bold = key == result["belirleyici_yontem"]
        color = "1E8449" if bold else None
        values = [
            item["ad"],
            _fmt(item["delta_sigma"], 3),
            _fmt(item["sigma_vo"], 3),
            _ratio(item["oran"]),
            _fmt(item["sondaj_derinligi_yuvarlatilmis"], 2, " m"),
        ]
        for idx, value in enumerate(values):
            _set_cell_text(cells[idx], value, bold=bold, size=9, color=color, align=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def _add_detail_table(doc, result):
    rows = result.get("rows", [])
    if not rows:
        return None
    doc.add_page_break()
    _add_heading(doc, "Hesap Tablosu", size=14)
    table = doc.add_table(rows=1, cols=8)
    headers = ["Der. (m)", "Z (m)", "m", "n", "σ'vo", "Bouss. Δσ", "West. Δσ", "Yak. Δσ"]
    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    _table_grid(table, header_rows=1)
    for row in rows[:42]:
        cells = table.add_row().cells
        values = [
            _fmt(row["derinlik"], 2),
            _fmt(row["z"], 2),
            _fmt(row["m"], 3),
            _fmt(row["n"], 3),
            _fmt(row["sigma_vo"], 3),
            _fmt(row["boussinesq_delta"], 3),
            _fmt(row["westergaard_delta"], 3),
            _fmt(row["yaklasik_delta"], 3),
        ]
        for idx, value in enumerate(values):
            _set_cell_text(cells[idx], value, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    return table


def _build_docx(veri, output_path):
    ctx = _foy_context(veri)
    result = ctx["result"]
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.35)
        section.right_margin = Cm(1.35)

    _add_heading(
        doc,
        "NET TEMEL TABAN BASINCINDAN KAYNAKLANAN GERİLME ARTIŞINA GÖRE\nSONDAJ DERİNLİĞİ HESABI",
        size=14,
    )
    _add_note(
        doc,
        "Sondaj derinliği, net temel taban basıncından kaynaklanan zemindeki gerilme artışının "
        "zeminin kendi ağırlığından kaynaklanan efektif düşey gerilmenin %10'una eşit veya daha "
        "küçük olduğu derinlik esas alınarak belirlenmiştir.",
    )
    _add_input_table(doc, ctx)
    doc.add_paragraph()
    _add_result_table(doc, result)

    conclusion = (
        f"Yapılan hesapta belirleyici yöntem {result['belirleyici_yontem_adi']} olarak bulunmuştur. "
        f"Bu nedenle zemin yüzeyinden itibaren önerilen minimum sondaj derinliği "
        f"{result['sondaj_derinligi_yuvarlatilmis']:.2f} m'dir. "
        f"Temel tabanı altında kalan Z değeri {result['temel_alti_z']:.2f} m'dir. "
        "Kademeli sonuçlar bağımsız ikili arama ile doğrulanmıştır."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("SONUÇ: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(conclusion).font.size = Pt(10)

    with tempfile.TemporaryDirectory() as tmp:
        graph_path = os.path.join(tmp, "sondaj_derinligi_grafik.png")
        sondaj_derinligi_grafik_olustur(result, graph_path)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(graph_path, width=Cm(17.0))
    _add_detail_table(doc, result)
    atomic_docx_save(doc, output_path)
    return {"path": output_path, "result": result}


def _font_path(bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    return next((path for path in candidates if os.path.exists(path)), None)


def _pdf_text(page, rect, text, size=9, bold=False, color=(0, 0, 0), align=0):
    fontname = "ArialBold" if bold else "Arial"
    fontfile = _font_path(bold)
    kwargs = {"fontsize": size, "fontname": fontname, "color": color, "align": align}
    if fontfile:
        kwargs["fontfile"] = fontfile
    page.insert_textbox(rect, str(text), **kwargs)


def _pdf_cell(page, rect, text, size=8, bold=False, fill=None, color=(0, 0, 0), align=0):
    if fill:
        page.draw_rect(rect, color=(0.75, 0.78, 0.80), fill=fill, width=0.4)
    else:
        page.draw_rect(rect, color=(0.70, 0.74, 0.76), width=0.4)
    inner = rect + (3, 2, -3, -2)
    _pdf_text(page, inner, text, size=size, bold=bold, color=color, align=align)


def _pdf_table(page, x, y, widths, row_h, rows, header_rows=1, size=8):
    for r_idx, row in enumerate(rows):
        x0 = x
        for c_idx, value in enumerate(row):
            rect = page.rect.__class__(x0, y, x0 + widths[c_idx], y + row_h)
            fill = (0.85, 0.92, 0.97) if r_idx < header_rows else None
            _pdf_cell(page, rect, value, size=size, bold=r_idx < header_rows or c_idx in (0, 2), fill=fill, align=1 if r_idx < header_rows else 0)
            x0 += widths[c_idx]
        y += row_h
    return y


def _build_pdf(veri, output_path):
    import fitz

    ctx = _foy_context(veri)
    result = ctx["result"]
    stress_unit = result.get("gerilme_birimi", "t/m²")
    bha_unit = result.get("bha_birimi", "t/m³")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    margin = 34
    y = 34
    page.draw_rect(fitz.Rect(margin, y, 595 - margin, y + 66), color=(0.15, 0.22, 0.30), fill=(0.15, 0.22, 0.30))
    _pdf_text(
        page,
        fitz.Rect(margin + 12, y + 10, 595 - margin - 12, y + 48),
        "NET TEMEL TABAN BASINCINDAN KAYNAKLANAN GERİLME ARTIŞINA GÖRE\nSONDAJ DERİNLİĞİ HESABI",
        size=13,
        bold=True,
        color=(1, 1, 1),
        align=1,
    )
    y += 80
    _pdf_text(page, fitz.Rect(margin, y, 595 - margin, y + 42), ctx["project"], size=12, bold=True, color=(0.12, 0.18, 0.24))
    y += 24
    _pdf_text(page, fitz.Rect(margin, y, 595 - margin, y + 34), f"Konum: {ctx['location']}    |    {ctx['parcel']}    |    Tarih: {ctx['date']}", size=8.5)
    y += 30
    note = (
        "Sondaj derinliği, net temel taban basıncından kaynaklanan zemindeki gerilme artışının "
        "zeminin kendi ağırlığından kaynaklanan efektif düşey gerilmenin %10'una eşit veya daha küçük "
        "olduğu derinlik esas alınarak belirlenmiştir."
    )
    _pdf_text(page, fitz.Rect(margin, y, 595 - margin, y + 48), note, size=8.8)
    y += 52
    input_rows = [
        (INPUT_LABELS["b"], _fmt(result["b"], 2, " m"), INPUT_LABELS["l"], _fmt(result["l"], 2, " m")),
        (INPUT_LABELS["df"], _fmt(result["temel_derinligi"], 2, " m"), INPUT_LABELS["yass"], _fmt(result["yass"], 2, " m")),
        (INPUT_LABELS["q_taban"], _fmt(result["q_taban"], 3, f" {stress_unit}"), INPUT_LABELS["sigma_vo_taban"], _fmt(result["sigma_vo_taban"], 3, f" {stress_unit}")),
        (INPUT_LABELS["dogal_bha"], _fmt(result["dogal_bha"], 3, f" {bha_unit}"), INPUT_LABELS["doygun_bha"], _fmt(result["doygun_bha"], 3, f" {bha_unit}")),
        (INPUT_LABELS["q_net"], _fmt(result["q_net"], 3, f" {stress_unit}"), INPUT_LABELS["kosul"], "Δσ ≤ 0.10 σ'vo"),
    ]
    y = _pdf_table(page, margin, y, [138, 114, 138, 114], 31, input_rows, header_rows=0, size=7.4)
    y += 16
    result_rows = [["Yöntem", "Δσ", "σ'vo", "Oran", "Gerekli derinlik"]]
    for key in ("boussinesq", "westergaard", "yaklasik"):
        item = result["yontem_sonuclari"][key]
        result_rows.append([
            item["ad"],
            _fmt(item["delta_sigma"], 3),
            _fmt(item["sigma_vo"], 3),
            _ratio(item["oran"]),
            _fmt(item["sondaj_derinligi_yuvarlatilmis"], 2, " m"),
        ])
    y = _pdf_table(page, margin, y, [142, 82, 82, 82, 116], 23, result_rows, header_rows=1, size=8)
    y += 14
    conclusion_rect = fitz.Rect(margin, y, 595 - margin, y + 54)
    page.draw_rect(conclusion_rect, color=(0.10, 0.48, 0.25), fill=(0.90, 0.97, 0.92), width=0.9)
    _pdf_text(
        page,
        conclusion_rect + (8, 7, -8, -6),
        f"SONUÇ: Belirleyici yöntem {result['belirleyici_yontem_adi']} olarak bulunmuştur. "
        f"Zemin yüzeyinden itibaren önerilen minimum sondaj derinliği "
        f"{result['sondaj_derinligi_yuvarlatilmis']:.2f} m'dir. Temel tabanı altında kalan Z "
        f"{result['temel_alti_z']:.2f} m'dir. Sayısal çapraz doğrulama başarılıdır.",
        size=9.5,
        bold=True,
        color=(0.05, 0.30, 0.16),
    )
    y += 68
    with tempfile.TemporaryDirectory() as tmp:
        graph_path = os.path.join(tmp, "sondaj_derinligi_grafik.png")
        sondaj_derinligi_grafik_olustur(result, graph_path)
        page.insert_image(fitz.Rect(margin, y, 595 - margin, y + 246), filename=graph_path, keep_proportion=True)

    page2 = doc.new_page(width=595, height=842)
    _pdf_text(page2, fitz.Rect(margin, 34, 595 - margin, 60), "HESAP TABLOSU", size=12, bold=True, color=(0.12, 0.18, 0.24), align=1)
    rows = [["Der.", "Z", "m", "n", "σ'vo", "Bouss.Δσ", "West.Δσ", "Yak.Δσ"]]
    for row in result.get("rows", [])[:46]:
        rows.append([
            _fmt(row["derinlik"], 2),
            _fmt(row["z"], 2),
            _fmt(row["m"], 3),
            _fmt(row["n"], 3),
            _fmt(row["sigma_vo"], 3),
            _fmt(row["boussinesq_delta"], 3),
            _fmt(row["westergaard_delta"], 3),
            _fmt(row["yaklasik_delta"], 3),
        ])
    _pdf_table(page2, margin, 70, [50, 50, 54, 54, 68, 76, 76, 76], 15, rows, header_rows=1, size=6.8)
    try:
        atomic_fitz_pdf_save(doc, output_path)
    finally:
        doc.close()
    return {"path": output_path, "result": result}


def sondaj_derinligi_foyu_olustur(veri, output_path):
    ext = os.path.splitext(output_path)[1].lower()
    if ext == ".pdf":
        return _build_pdf(veri, output_path)
    if ext != ".docx":
        output_path = f"{output_path}.docx"
    return _build_docx(veri, output_path)

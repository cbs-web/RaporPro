# Dosya: RaporPro/raporlama_arazi.py
from __future__ import annotations

import re
import unicodedata

from docx.shared import RGBColor


PMT_STANDARD_TEXT = (
    "Deneyler TS EN ISO 22476-4 ve ASTM D4719-00 standartlarına uygun olarak "
    "yapılmıştır."
)
KAROT_ACIKLAMA_TEXT = (
    "Kesilen birimlerin çakıl içeriği, basınçlı sondaj suyu ile temas ettiğinde "
    "dağılma ve erime özelliği göstermesinin karot yüzdesinin düşmesine neden "
    "olduğu düşünülmektedir."
)


def _text(value):
    return str(value or "").strip()


def _present(value):
    return _text(value) not in ("", "-", "None", "none", "nan", "NaN")


def _number(value):
    text = _text(value).replace("%", "").replace(",", ".")
    if not text:
        return None
    match = re.search(r"[-+]?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _row_value(row, index, key):
    if isinstance(row, dict):
        return row.get(key, "")
    try:
        return row[index]
    except (TypeError, IndexError):
        return ""


def _natural_key(value):
    return [
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r"(\d+)", _text(value))
    ]


def _unique_natural(values):
    unique = {}
    for value in values:
        clean = _text(value)
        if clean:
            unique.setdefault(clean.casefold(), clean)
    return sorted(unique.values(), key=_natural_key)


def _format_number(value):
    value = float(value)
    if value.is_integer():
        return str(int(value))
    return f"{value:.1f}".rstrip("0").rstrip(".").replace(".", ",")


def _turkish_list(values):
    values = list(values or [])
    if not values:
        return ""
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} ve {values[1]}"
    return f"{', '.join(values[:-1])} ve {values[-1]}"


def arazi_deney_rapor_verileri(sondajlar):
    """SPT, PMT ve karot kayıtlarını rapora uygun, doğrulanmış satırlara ayır."""
    spt_data = []
    pmt_data = []
    kaya_data = []
    pmt_sondajlari = []
    tcr_degerleri = []

    for sondaj in sondajlar or []:
        sondaj_no = _text(sondaj.get("no"))
        for row in sondaj.get("spt", []) or []:
            values = [_row_value(row, index, key) for index, key in enumerate(
                ("der", "n15", "n30_1", "n45", "n30")
            )]
            if isinstance(row, dict):
                row_is_valid = any(_present(value) for value in values)
            else:
                try:
                    row_is_valid = len(row) >= 5
                except TypeError:
                    row_is_valid = False
            if row_is_valid:
                spt_data.append([sondaj_no, *values])

        sondajda_pmt_var = False
        for row in sondaj.get("pmt", []) or []:
            depth = _row_value(row, 0, "der")
            em = _row_value(row, 1, "em")
            pl = _row_value(row, 2, "pl")
            if not (_present(depth) and _present(em) and _present(pl)):
                continue
            pmt_data.append([sondaj_no, depth, em, pl])
            sondajda_pmt_var = True
        if sondajda_pmt_var and sondaj_no:
            pmt_sondajlari.append(sondaj_no)

        for row in sondaj.get("kaya", []) or []:
            depth = _row_value(row, 0, "der")
            tcr = _row_value(row, 1, "tcr")
            scr = _row_value(row, 2, "scr")
            rqd = _row_value(row, 3, "rqd")
            tcr_number = _number(tcr)
            if not _present(depth) or tcr_number is None or not 0 <= tcr_number <= 100:
                continue
            kaya_data.append([sondaj_no, depth, tcr, scr, rqd])
            tcr_degerleri.append(tcr_number)

    return {
        "spt_data": spt_data,
        "pmt_data": pmt_data,
        "kaya_data": kaya_data,
        "pmt_sondajlari": _unique_natural(pmt_sondajlari),
        "tcr_degerleri": tcr_degerleri,
    }


def pmt_rapor_cumlesi(sondajlar, table_number=14):
    sondajlar = _unique_natural(sondajlar)
    if not sondajlar:
        return ""
    wells = _turkish_list(sondajlar)
    location = "sondajında" if len(sondajlar) == 1 else "sondajlarında"
    return (
        f"Çalışma alanında {wells} {location} presiyometre deneyi yapılmıştır. "
        f"{PMT_STANDARD_TEXT} Presiyometre deney sonuçları Tablo "
        f"{int(table_number)}'te verilmiştir."
    )


def tcr_rapor_cumlesi(tcr_values):
    values = [float(value) for value in tcr_values or [] if value is not None]
    if not values:
        return ""
    minimum = min(values)
    maximum = max(values)
    if abs(minimum - maximum) < 1e-9:
        value_text = f"%{_format_number(minimum)}'dir"
    else:
        value_text = f"%{_format_number(minimum)}-%{_format_number(maximum)} arasındadır"
    return (
        f"Çalışma alanında yapılan sondajlarda karot yüzdeleri {value_text}. "
        f"{KAROT_ACIKLAMA_TEXT}"
    )


def _normalized(value):
    text = _text(value).translate(str.maketrans({
        "ı": "i", "İ": "I", "ş": "s", "Ş": "S", "ğ": "g", "Ğ": "G",
        "ç": "c", "Ç": "C", "ö": "o", "Ö": "O", "ü": "u", "Ü": "U",
    }))
    text = unicodedata.normalize("NFKD", text.casefold())
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]+", "", text)


def _remove_paragraph(paragraph):
    element = getattr(paragraph, "_element", None)
    parent = element.getparent() if element is not None else None
    if parent is not None:
        parent.remove(element)


def _nearby_body_paragraphs(doc, anchor, before=4, after=4):
    paragraphs = list(doc.paragraphs)
    try:
        index = next(idx for idx, item in enumerate(paragraphs) if item._p is anchor._p)
    except (StopIteration, AttributeError):
        return [anchor] if anchor is not None else []
    return paragraphs[max(0, index - before): min(len(paragraphs), index + after + 1)]


def _is_heading_or_caption(paragraph):
    style_name = _normalized(getattr(getattr(paragraph, "style", None), "name", ""))
    return style_name.startswith("heading") or style_name == "caption"


def _is_pmt_paragraph(text):
    normalized = _normalized(text)
    return (
        "presiyometredeneysonuclari" in normalized
        or "presiyometredeneyiyapilmistir" in normalized
        or normalized == "pmt"
    )


def _is_karot_paragraph(text):
    normalized = _normalized(text)
    return (
        "karotyuzdeleri" in normalized
        or normalized == "kayatablo"
    )


def _find_near(anchor_paragraphs, predicate, narrative=False):
    matches = [paragraph for paragraph in anchor_paragraphs if predicate(paragraph.text)]
    if narrative:
        for paragraph in matches:
            if _is_heading_or_caption(paragraph):
                continue
            normalized = _normalized(paragraph.text)
            if "calismaalaninda" in normalized or "yapilmistir" in normalized:
                return paragraph
        for paragraph in matches:
            if not _is_heading_or_caption(paragraph) and _normalized(paragraph.text) != "pmt":
                return paragraph
    return matches[0] if matches else None


def _caption_number_update(paragraphs, predicate, number):
    for paragraph in paragraphs:
        if not predicate(paragraph.text):
            continue
        if re.search(r"\bTablo\s*\d+\s*:", paragraph.text, flags=re.IGNORECASE):
            updated_text = re.sub(
                r"\bTablo\s*\d+",
                f"Tablo {int(number)}",
                paragraph.text,
                count=1,
                flags=re.IGNORECASE,
            )
            for run in paragraph.runs:
                match = re.search(r"\bTablo\s*\d+", run.text, flags=re.IGNORECASE)
                if match:
                    run.text = re.sub(
                        r"\bTablo\s*\d+",
                        f"Tablo {int(number)}",
                        run.text,
                        count=1,
                        flags=re.IGNORECASE,
                    )
                    return paragraph
            for run in paragraph.runs:
                if re.fullmatch(r"\s*\d+\s*", run.text):
                    run.text = str(int(number))
                    return paragraph
            paragraph.text = updated_text
            return paragraph
    return None


def _paragraph_clear(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _pmt_paragraph_write(paragraph, sondajlar, table_number):
    sondajlar = _unique_natural(sondajlar)
    wells = _turkish_list(sondajlar)
    location = "sondajında" if len(sondajlar) == 1 else "sondajlarında"
    try:
        paragraph.style = "Normal"
    except (KeyError, ValueError):
        pass
    _paragraph_clear(paragraph)
    highlighted = paragraph.add_run(
        f"Çalışma alanında {wells} {location} presiyometre deneyi yapılmıştır. "
    )
    highlighted.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    paragraph.add_run(
        f"{PMT_STANDARD_TEXT} Presiyometre deney sonuçları "
    )
    table_reference = paragraph.add_run(f"Tablo {int(table_number)}")
    table_reference.bold = True
    paragraph.add_run("'te verilmiştir.")


def _tcr_paragraph_write(paragraph, tcr_values):
    values = [float(value) for value in tcr_values or [] if value is not None]
    if not values:
        return
    minimum = min(values)
    maximum = max(values)
    if abs(minimum - maximum) < 1e-9:
        value_text = f"%{_format_number(minimum)}'dir"
    else:
        value_text = f"%{_format_number(minimum)}-%{_format_number(maximum)} arasındadır"
    _paragraph_clear(paragraph)
    paragraph.add_run("Çalışma alanında yapılan sondajlarda ")
    highlighted = paragraph.add_run(f"karot yüzdeleri {value_text}.")
    highlighted.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    paragraph.add_run(f" {KAROT_ACIKLAMA_TEXT}")


def arazi_deney_word_bolumlerini_uygula(doc, paragraph_index, report_data):
    """PMT ve karot Word bloklarını veri varlığına göre güncelle veya kaldır."""
    pmt_rows = report_data.get("pmt_data", [])
    kaya_rows = report_data.get("kaya_data", [])
    pmt_anchor = paragraph_index.get("[PMT]")
    kaya_anchor = paragraph_index.get("[KAYA_TABLO]")
    pmt_near = _nearby_body_paragraphs(doc, pmt_anchor, before=4, after=1)
    kaya_near = _nearby_body_paragraphs(doc, kaya_anchor, before=2, after=5)

    if kaya_rows:
        narrative = _find_near(kaya_near, _is_karot_paragraph, narrative=True)
        if narrative is not None:
            _tcr_paragraph_write(narrative, report_data.get("tcr_degerleri", []))
    else:
        for paragraph in list(kaya_near):
            if _is_karot_paragraph(paragraph.text):
                _remove_paragraph(paragraph)

    pmt_table_number = 14 if kaya_rows else 13
    if pmt_rows:
        narrative = _find_near(pmt_near, _is_pmt_paragraph, narrative=True)
        if narrative is not None:
            _pmt_paragraph_write(
                narrative,
                report_data.get("pmt_sondajlari", []),
                pmt_table_number,
            )
        _caption_number_update(pmt_near, _is_pmt_paragraph, pmt_table_number)
    else:
        for paragraph in list(pmt_near):
            if _is_pmt_paragraph(paragraph.text):
                _remove_paragraph(paragraph)

    return {
        "pmt_var": bool(pmt_rows),
        "kaya_var": bool(kaya_rows),
        "pmt_table_number": pmt_table_number if pmt_rows else None,
    }


__all__ = [
    "arazi_deney_rapor_verileri",
    "arazi_deney_word_bolumlerini_uygula",
    "pmt_rapor_cumlesi",
    "tcr_rapor_cumlesi",
]

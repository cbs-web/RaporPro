# Dosya: RaporPro/rapor_metin_revizyon.py
from __future__ import annotations

import json
import os
import re
from typing import Any

from docx import Document
from performans import gizli_bilgileri_maskele
from yardimcilar import atomic_docx_save


MAX_AI_UNIT_COUNT = 70
MAX_AI_UNIT_CHARS = 900


def _temiz_metin(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _kisalt(text, limit=260):
    text = _temiz_metin(text)
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 3)].rstrip() + "..."


def _unit_ekle(units, unit_id, label, text, **loc):
    text = str(text or "")
    if not text.strip():
        return
    item = {
        "unit_id": unit_id,
        "label": label,
        "text": text,
        "preview": _kisalt(text),
    }
    item.update(loc)
    units.append(item)


def _header_footer_parts(doc):
    parts = [
        ("h", "Üst bilgi", "header"),
        ("fh", "İlk sayfa üst bilgi", "first_page_header"),
        ("eh", "Çift sayfa üst bilgi", "even_page_header"),
        ("f", "Alt bilgi", "footer"),
        ("ff", "İlk sayfa alt bilgi", "first_page_footer"),
        ("ef", "Çift sayfa alt bilgi", "even_page_footer"),
    ]
    for section_index, section in enumerate(doc.sections):
        for prefix, label, attr in parts:
            try:
                part = getattr(section, attr)
            except Exception:
                continue
            yield prefix, label, section_index, part


def word_metin_birimleri_oku(word_path):
    """Word raporundaki düzenlenebilir metin birimlerini okur."""
    if not word_path or not os.path.exists(word_path):
        raise FileNotFoundError("Okunacak Word raporu bulunamadı.")

    doc = Document(word_path)
    units = []
    for idx, paragraph in enumerate(doc.paragraphs):
        _unit_ekle(units, f"p:{idx}", f"Paragraf {idx + 1}", paragraph.text, kind="paragraph", index=idx)

    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    _unit_ekle(
                        units,
                        f"t:{table_index}:{row_index}:{cell_index}:{paragraph_index}",
                        f"Tablo {table_index + 1} R{row_index + 1} C{cell_index + 1}",
                        paragraph.text,
                        kind="table",
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                        paragraph_index=paragraph_index,
                    )

    for prefix, label, section_index, part in _header_footer_parts(doc):
        for paragraph_index, paragraph in enumerate(part.paragraphs):
            _unit_ekle(
                units,
                f"{prefix}:{section_index}:{paragraph_index}",
                f"{label} {section_index + 1}/{paragraph_index + 1}",
                paragraph.text,
                kind=prefix,
                section_index=section_index,
                paragraph_index=paragraph_index,
            )
    return units


def _sk_formatla(text):
    match = re.search(r"SK\s*-?\s*(\d+)", str(text or ""), flags=re.I)
    return f"SK-{int(match.group(1))}" if match else _temiz_metin(text)


def _sk_regex(text):
    match = re.search(r"SK\s*-?\s*(\d+)", str(text or ""), flags=re.I)
    if not match:
        return None
    number = str(int(match.group(1)))
    return re.compile(rf"(?<![A-Z0-9])SK\s*-?\s*0*{re.escape(number)}(?![A-Z0-9])", flags=re.I)


def _literal_regex(old_text):
    old_text = str(old_text or "")
    sk_pattern = _sk_regex(old_text)
    if sk_pattern is not None:
        return sk_pattern
    return re.compile(re.escape(old_text), flags=re.I)


def _ilk_eslesen_ifade(text, old_text):
    if not old_text:
        return ""
    if old_text in text:
        return old_text
    pattern = _literal_regex(old_text)
    match = pattern.search(text)
    return match.group(0) if match else ""


def _degisim_onizleme(current_text, old_text, new_text):
    match_text = _ilk_eslesen_ifade(current_text, old_text)
    if not match_text:
        return current_text, False
    pattern = _literal_regex(match_text)
    updated, count = pattern.subn(str(new_text or ""), current_text, count=1)
    return updated, bool(count)


def _temiz_ifade(text):
    text = _temiz_metin(text)
    text = text.strip(" '\"“”‘’")
    text = re.sub(r"^(raporda|rapor içinde|metinde|cümlede)\s+", "", text, flags=re.I).strip()
    return text.strip(" '\"“”‘’")


def _degisim_ciftleri(note):
    note = str(note or "")
    pairs = []

    quote = r"[\"'“”‘’]"
    for pattern in [
        rf"{quote}(?P<old>.+?){quote}\s*(?:ifadesi\s*)?yerine\s*{quote}(?P<new>.+?){quote}",
        rf"{quote}(?P<old>.+?){quote}\s*(?:ifadesi\s*)?(?:,?\s*)?(?P<new>.+?)\s+olarak\s+(?:düzeltilsin|duzeltilsin|değiştirilsin|degistirilsin|yazılsın|yazilsin)",
    ]:
        for match in re.finditer(pattern, note, flags=re.I | re.S):
            pairs.append((_temiz_ifade(match.group("old")), _temiz_ifade(match.group("new")), 88, "Tırnak içi ifade düzeltmesi"))

    clauses = [part.strip() for part in re.split(r"[\n\r.;]+", note) if part.strip()]
    for clause in clauses:
        for pattern, reason in [
            (r"(?P<old>.+?)\s+ifadesi\s+(?P<new>.+?)\s+olarak\s+(?:düzeltilsin|duzeltilsin|değiştirilsin|degistirilsin|yazılsın|yazilsin)", "İfade düzeltmesi"),
            (r"(?P<old>.+?)\s+yerine\s+(?P<new>.+?)(?:\s+yazılsın|\s+yazilsin|\s+düzeltilsin|\s+duzeltilsin|\s+değiştirilsin|\s+degistirilsin|$)", "Yerine yazma düzeltmesi"),
        ]:
            match = re.search(pattern, clause, flags=re.I)
            if match:
                pairs.append((_temiz_ifade(match.group("old")), _temiz_ifade(match.group("new")), 78, reason))

        sk_ifade = re.search(r"(?P<old>SK\s*-?\s*\d+)\s+ifadesi\s+(?P<new>SK\s*-?\s*\d+)\s+olarak", clause, flags=re.I)
        if sk_ifade:
            pairs.append((_sk_formatla(sk_ifade.group("old")), _sk_formatla(sk_ifade.group("new")), 96, "Sondaj numarası ifade düzeltmesi"))
            continue

        old_match = re.search(r"(?P<old>SK\s*-?\s*\d+)\s+olarak\s+(?:görünen|gorunen|görüntülenen|goruntulenen|yazılan|yazilan|işlenen|islenen)", clause, flags=re.I)
        new_match = re.search(r"(?:aslında|aslinda|gerçekte|gercekte|doğrusu|dogrusu)\s+(?P<new>SK\s*-?\s*\d+)", clause, flags=re.I)
        if old_match and new_match:
            pairs.append((_sk_formatla(old_match.group("old")), _sk_formatla(new_match.group("new")), 95, "Sondaj numarası düzeltmesi"))
            continue

    sk_tokens = [_sk_formatla(match.group(0)) for match in re.finditer(r"SK\s*-?\s*\d+", note, flags=re.I)]
    unique_sk = []
    for token in sk_tokens:
        if token not in unique_sk:
            unique_sk.append(token)
    if len(unique_sk) >= 2:
        pairs.append((unique_sk[0], unique_sk[-1], 82, "Sondaj numarası düzeltmesi"))

    cleaned = []
    seen = set()
    for old_text, new_text, confidence, reason in pairs:
        old_text = _temiz_ifade(old_text)
        new_text = _temiz_ifade(new_text)
        if not old_text or not new_text or old_text == new_text:
            continue
        key = (old_text.lower(), new_text.lower())
        if key in seen:
            continue
        seen.add(key)
        cleaned.append((old_text, new_text, confidence, reason))
    return cleaned


def metin_revizyon_kural_analiz_et(note, units):
    """Düzeltme notundan güvenli eski/yeni ifade önerileri üretir."""
    pairs = _degisim_ciftleri(note)
    warnings = []
    if not pairs:
        return {
            "source": "kural",
            "items": [],
            "warnings": ["Düzeltme notundan doğrudan eski/yeni ifade çifti çıkarılamadı."],
            "unit_count": len(units or []),
        }

    items = []
    seen = set()
    for old_text, new_text, confidence, reason in pairs:
        matched = 0
        for unit in units or []:
            current = unit.get("text", "")
            match_text = _ilk_eslesen_ifade(current, old_text)
            if not match_text:
                continue
            preview, _ok = _degisim_onizleme(current, match_text, new_text)
            key = (unit.get("unit_id"), match_text.lower(), str(new_text).lower())
            if key in seen:
                continue
            seen.add(key)
            matched += 1
            items.append(
                {
                    "unit_id": unit.get("unit_id"),
                    "label": unit.get("label", ""),
                    "old_text": match_text,
                    "new_text": new_text,
                    "current_text": current,
                    "preview_text": preview,
                    "guven": confidence,
                    "reason": reason,
                    "source": "kural",
                }
            )
        if not matched:
            warnings.append(f"'{old_text}' ifadesi rapor metninde bulunamadı.")

    return {"source": "kural", "items": items, "warnings": warnings, "unit_count": len(units or [])}


def _aday_birimler(note, units):
    note_norm = _temiz_metin(note).lower()
    tokens = {tok for tok in re.findall(r"[A-Za-zÇĞİÖŞÜçğıöşü0-9_-]{3,}", note_norm)}
    sk_tokens = {_sk_formatla(match.group(0)).lower() for match in re.finditer(r"SK\s*-?\s*\d+", note, flags=re.I)}
    scored = []
    for unit in units or []:
        text = unit.get("text", "")
        norm = text.lower()
        score = 0
        for token in tokens:
            if token in norm:
                score += 1
        for token in sk_tokens:
            if token and _sk_regex(token) and _sk_regex(token).search(text):
                score += 8
        if score:
            scored.append((score, unit))
    if not scored:
        scored = [(0, unit) for unit in (units or [])[:MAX_AI_UNIT_COUNT]]
    scored.sort(key=lambda item: (-item[0], item[1].get("label", "")))
    return [unit for _score, unit in scored[:MAX_AI_UNIT_COUNT]]


def _ai_prompt(note, candidate_units):
    unit_lines = []
    for unit in candidate_units:
        text = _temiz_metin(unit.get("text", ""))[:MAX_AI_UNIT_CHARS]
        unit_lines.append({"unit_id": unit.get("unit_id"), "label": unit.get("label"), "text": text})
    return (
        "Sen Zemin Rapor Pro icin calisan guvenli bir Word metin revizyon asistanisin.\n"
        "Belediye/kontrolor duzeltme notunu oku. Asagidaki rapor metin birimlerinde yapilmasi gereken "
        "cumle veya ifade duzeltmelerini bul.\n"
        "Cok onemli kurallar:\n"
        "- Sadece verilen metin birimlerinde acikca bulunan ifadeleri degistir.\n"
        "- old_text mevcut birim metninde birebir bulunabilir bir ifade olsun.\n"
        "- Emin degilsen items listesini bos birak.\n"
        "- Cikti yalnizca JSON olsun.\n"
        "Beklenen JSON:\n"
        "{\"items\":[{\"unit_id\":\"p:1\",\"old_text\":\"eski ifade\",\"new_text\":\"yeni ifade\",\"reason\":\"neden\",\"guven\":0}]}\n\n"
        f"Duzeltme notu:\n{note}\n\n"
        f"Rapor metin birimleri:\n{json.dumps(unit_lines, ensure_ascii=False)}"
    )


def _ai_metin_revizyonu(note, units, motor="otomatik", timeout=45):
    try:
        import requests
    except Exception as exc:
        raise RuntimeError(f"requests yüklenemedi: {exc}") from exc

    from ai_motoru import _aktif_motor_sec, _api_key_kontrol, _json_from_text
    from spt_okuma_motoru import openai_model_sec, spt_ayarlarini_yukle

    ayarlar = spt_ayarlarini_yukle()
    aktif = _aktif_motor_sec(ayarlar, motor)
    if aktif == "kural":
        return metin_revizyon_kural_analiz_et(note, units)
    _api_key_kontrol(aktif, ayarlar)

    prompt = _ai_prompt(note, _aday_birimler(note, units))
    if aktif == "openai":
        url = "https://api.openai.com/v1/chat/completions"
        api_key = ayarlar["openai_api_key"]
        model_name = openai_model_sec(ayarlar, "revizyon")
        payload = {
            "model": model_name,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.05,
            "response_format": {"type": "json_object"},
        }
        payload = {key: value for key, value in payload.items() if value is not None}
        response = requests.post(url, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}, json=payload, timeout=timeout)
        if response.status_code != 200:
            detail = gizli_bilgileri_maskele(response.text[:500], (api_key,))
            raise RuntimeError(f"{aktif.upper()} hata kodu {response.status_code}: {detail}")
        raw = response.json()["choices"][0]["message"]["content"]
    elif aktif in ("gemini", "gemini_pro"):
        model_id = "gemini-2.5-pro" if aktif == "gemini_pro" else "gemini-2.5-flash"
        api_key = ayarlar["gemini_api_key"]
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_id}:generateContent"
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.05, "response_mime_type": "application/json"},
        }
        response = requests.post(
            url,
            headers={"Content-Type": "application/json", "x-goog-api-key": api_key},
            json=payload,
            timeout=timeout,
        )
        if response.status_code != 200:
            try:
                msg = response.json().get("error", {}).get("message", response.text)
            except Exception:
                msg = response.text
            detail = gizli_bilgileri_maskele(str(msg)[:500], (api_key,))
            raise RuntimeError(f"GEMINI hata kodu {response.status_code}: {detail}")
        raw = response.json()["candidates"][0]["content"]["parts"][0]["text"]
    else:
        raise RuntimeError(f"Desteklenmeyen AI motoru: {aktif}")

    parsed = _json_from_text(raw)
    raw_items = parsed.get("items", []) if isinstance(parsed, dict) else parsed if isinstance(parsed, list) else []
    return _ai_sonucunu_temizle(raw_items, units, aktif, raw)


def _ai_sonucunu_temizle(raw_items, units, source, raw_response=""):
    by_id = {unit.get("unit_id"): unit for unit in units or []}
    items = []
    warnings = []
    seen = set()
    for raw in raw_items or []:
        unit_id = str((raw or {}).get("unit_id") or "").strip()
        unit = by_id.get(unit_id)
        old_text = _temiz_ifade((raw or {}).get("old_text") or (raw or {}).get("eski") or "")
        new_text = _temiz_ifade((raw or {}).get("new_text") or (raw or {}).get("yeni") or "")
        if not unit or not old_text or not new_text or old_text == new_text:
            continue
        match_text = _ilk_eslesen_ifade(unit.get("text", ""), old_text)
        if not match_text:
            warnings.append(f"AI önerisi atlandı; '{old_text}' ifadesi {unit.get('label', unit_id)} içinde bulunamadı.")
            continue
        try:
            confidence = int(round(float(str((raw or {}).get("guven", (raw or {}).get("confidence", 70))).replace(",", "."))))
        except Exception:
            confidence = 70
        confidence = max(0, min(100, confidence))
        preview, _ok = _degisim_onizleme(unit.get("text", ""), match_text, new_text)
        key = (unit_id, match_text.lower(), new_text.lower())
        if key in seen:
            continue
        seen.add(key)
        items.append(
            {
                "unit_id": unit_id,
                "label": unit.get("label", ""),
                "old_text": match_text,
                "new_text": new_text,
                "current_text": unit.get("text", ""),
                "preview_text": preview,
                "guven": confidence,
                "reason": str((raw or {}).get("reason") or (raw or {}).get("neden") or "AI önerisi").strip(),
                "source": source,
            }
        )
    return {"source": source, "items": items, "warnings": warnings, "unit_count": len(units or []), "raw_response": raw_response}


def metin_revizyon_analiz_et(word_path, note, motor="otomatik", timeout=45, ai_kullan=True):
    units = word_metin_birimleri_oku(word_path)
    note = str(note or "").strip()
    if not note:
        return {"source": "kural", "items": [], "warnings": ["Düzeltme notu boş."], "unit_count": len(units)}
    if not ai_kullan:
        return metin_revizyon_kural_analiz_et(note, units)
    try:
        result = _ai_metin_revizyonu(note, units, motor=motor, timeout=timeout)
        if result.get("items"):
            return result
        fallback = metin_revizyon_kural_analiz_et(note, units)
        fallback["source"] = f"{result.get('source', motor)}+kural"
        fallback["warnings"] = list(result.get("warnings") or []) + ["AI doğrudan metin düzeltmesi üretemedi; kural tabanlı sonuç gösteriliyor."] + list(fallback.get("warnings") or [])
        return fallback
    except Exception as exc:
        fallback = metin_revizyon_kural_analiz_et(note, units)
        fallback["warnings"].insert(0, f"AI metin analizi kullanılamadı, kural tabanlı analiz gösteriliyor: {exc}")
        return fallback


def _paragraph_by_unit_id(doc, unit_id):
    parts = str(unit_id or "").split(":")
    if not parts:
        return None
    kind = parts[0]
    try:
        if kind == "p" and len(parts) == 2:
            return doc.paragraphs[int(parts[1])]
        if kind == "t" and len(parts) == 5:
            table_index, row_index, cell_index, paragraph_index = [int(part) for part in parts[1:]]
            return doc.tables[table_index].rows[row_index].cells[cell_index].paragraphs[paragraph_index]
        if kind in {"h", "fh", "eh", "f", "ff", "ef"} and len(parts) == 3:
            section_index = int(parts[1])
            paragraph_index = int(parts[2])
            attr = {
                "h": "header",
                "fh": "first_page_header",
                "eh": "even_page_header",
                "f": "footer",
                "ff": "first_page_footer",
                "ef": "even_page_footer",
            }[kind]
            return getattr(doc.sections[section_index], attr).paragraphs[paragraph_index]
    except Exception:
        return None
    return None


def _tum_paragraflar(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for _prefix, _label, _section_index, part in _header_footer_parts(doc):
        for paragraph in part.paragraphs:
            yield paragraph


def _paragraf_degistir(paragraph, old_text, new_text):
    if paragraph is None:
        return False
    current = paragraph.text
    match_text = _ilk_eslesen_ifade(current, old_text)
    if not match_text:
        return False
    pattern = _literal_regex(match_text)
    for run in paragraph.runs:
        if _ilk_eslesen_ifade(run.text, match_text):
            run.text = pattern.sub(str(new_text or ""), run.text, count=1)
            return True
    paragraph.text = pattern.sub(str(new_text or ""), current, count=1)
    return True


def metin_revizyonlari_uygula(word_path, revisions, output_path):
    if not word_path or not os.path.exists(word_path):
        return {"success": False, "message": "Revize edilecek Word raporu bulunamadı.", "applied": [], "skipped": []}
    if not output_path:
        return {"success": False, "message": "Kaydedilecek çıktı yolu seçilmedi.", "applied": [], "skipped": []}

    doc = Document(word_path)
    applied = []
    skipped = []
    for revision in revisions or []:
        unit_id = revision.get("unit_id")
        paragraph = _paragraph_by_unit_id(doc, unit_id)
        old_text = revision.get("old_text", "")
        new_text = revision.get("new_text", "")
        if _paragraf_degistir(paragraph, old_text, new_text):
            applied.append(revision)
            continue

        fallback_ok = False
        for candidate in _tum_paragraflar(doc):
            if candidate is paragraph:
                continue
            if _paragraf_degistir(candidate, old_text, new_text):
                fallback_ok = True
                break
        if fallback_ok:
            applied.append(revision)
        else:
            skipped.append(revision)

    if not applied:
        return {
            "success": False,
            "message": "Seçili düzeltmeler rapor içinde uygulanamadı. Eski ifadeler raporda bulunamamış olabilir.",
            "applied": [],
            "skipped": skipped,
        }

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    atomic_docx_save(doc, output_path)
    message = f"Metin revizyonlu rapor oluşturuldu. Uygulanan düzeltme: {len(applied)}."
    if skipped:
        message += f" Atlanan düzeltme: {len(skipped)}."
    return {
        "success": True,
        "message": message,
        "applied": applied,
        "skipped": skipped,
        "output_path": output_path,
    }

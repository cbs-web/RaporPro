import datetime
import os
import re
import shutil

import pandas as pd
from PIL import Image
from tutarlilik_motoru import (
    bulgu_ekle,
    kontrol_ekle,
    kontrol_raporunu_tamamla,
    proje_tutarlilik_raporu,
)

from docx import Document
from ekler import uygun_ek_sablonu
from jeofizik_sheet_motoru import jeofizik_sheet_rows_to_ss_list, jeofizik_sheet_var_mi
from rapor_sablonu import etkin_rapor_sablonu_yolu


KNOWN_TAGS = {
    "[ADA]", "[ALAN_BOYLAM]", "[ALAN_ENLEM]", "[BINA_BILGILERI]",
    "[EGIM_YONU]", "[EGIM_YUZDE]", "[IL]", "[ILCE]", "[IMAR_ALANI]",
    "[IMAR_DURUMU]", "[JEO_KOOR]", "[JEO_PARAMETRE]", "[JEO_SONUC]",
    "[JEO_TARIH]", "[KATEGORI]", "[KATEGORI_ZEMIN]", "[KAYA_TABLO]",
    "[KOT_MAX]", "[KOT_MIN]", "[KOT_ORT]", "[LAB_FIZIK]",
    "[LAB_MEKANIK]", "[LITOLOJI_DAGILIM]", "[MAHALLE]", "[MASW]",
    "[MEVKI]", "[MT_TABLO]", "[PAFTA]", "[PARSEL]", "[PGA]", "[PMT]",
    "[PROJE_ADI]", "[RESIM:PGA]", "[RESIM:SONDAJ]", "[RESIM:TKGM]",
    "[RESIM:MJH]", "[RESIM_JEOFIZIK]", "[RESIM_MJH]", "[RESIM_SONDAJ]", "[RESIM_YERBULDURUR]",
    "[SAYI_MT]", "[SAYI_SS]", "[SONDAJ_BILGISI]", "[SPT]", "[Sondaj]",
    "[VP]", "[YASS_ONERI]", "[YASS_TABLO]", "[HIDROJEOLOJI_DURUM]", "[YEREL_ZEMIN]",
    "[ZEMIN_OZET]", "RESIM:MJH", "RESIM:PGA", "RESIM:TKGM",
    "RESIM:Yerbuldurur",
}

PREFIXED_TAG_RE = re.compile(
    r"^\[S[1-5]_(PROJE_ADI|IL|ILCE|MAHALLE|MEVKI|PAFTA|ADA|PARSEL)\]$"
)

IMAGE_TAGS = {
    "[RESIM_YERBULDURUR]": "img_yer",
    "RESIM:Yerbuldurur": "img_yer",
    "[RESIM:PGA]": "img_pga",
    "RESIM:PGA": "img_pga",
    "[RESIM:TKGM]": "img_tkgm",
    "RESIM:TKGM": "img_tkgm",
    "[RESIM_JEOFIZIK]": "word_img_jeofizik",
    "RESIM:MJH": "img_mjh",
    "[RESIM:MJH]": "img_mjh",
    "[RESIM_MJH]": "img_mjh",
    "[RESIM_SONDAJ]": "word_img_sondaj",
    "[RESIM:SONDAJ]": "word_img_sondaj",
}

TAG_DESCRIPTIONS = {
    "[PROJE_ADI]": "Proje sahibi/adi bilgisini yazar.",
    "[IL]": "Il bilgisini yazar.",
    "[ILCE]": "Ilce bilgisini yazar.",
    "[MAHALLE]": "Mahalle bilgisini yazar.",
    "[MEVKI]": "Mevkii bilgisini yazar.",
    "[PAFTA]": "Pafta bilgisini yazar.",
    "[ADA]": "Ada bilgisini yazar.",
    "[PARSEL]": "Parsel bilgisini yazar.",
    "[BINA_BILGILERI]": "Bina bilgileri ve temel yükleri tablosunu ekler.",
    "[KATEGORI]": "Arazi kategori bilgisini yazar.",
    "[KATEGORI_ZEMIN]": "Kategori zemin/aciklama bilgisini yazar.",
    "[PGA]": "PGA degerini yazar.",
    "[JEO_TARIH]": "Jeofizik calisma tarihini yazar.",
    "[SAYI_SS]": "Sismik serim sayisini yazar.",
    "[SAYI_MT]": "Mikrotremor olcu sayisini yazar.",
    "[YEREL_ZEMIN]": "Yerel zemin sinifini yazar.",
    "[KOT_ORT]": "Ortalama kot bilgisini yazar.",
    "[KOT_MAX]": "Maksimum kot bilgisini yazar.",
    "[KOT_MIN]": "Minimum kot bilgisini yazar.",
    "[EGIM_YUZDE]": "Egim yuzdesini yazar.",
    "[EGIM_YONU]": "Egim yonunu yazar.",
    "[IMAR_ALANI]": "Imar alani bilgisini yazar.",
    "[IMAR_DURUMU]": "Imar durumu/aciklamasini yazar.",
    "[ALAN_ENLEM]": "Calisma alani merkez enlemini yazar.",
    "[ALAN_BOYLAM]": "Calisma alani merkez boylamini yazar.",
    "[SONDAJ_BILGISI]": "Sondaj sayısı ve derinlik özet cümlesini yazar.",
    "[Sondaj]": "Sondaj konum, tarih, kot, derinlik ve litoloji tablosunu ekler.",
    "[YASS_TABLO]": "Yeraltisuyu olcumleri tablosunu ekler.",
    "[YASS_ONERI]": "Yeraltısuyu durumuna göre öneri paragrafını yazar.",
    "[HIDROJEOLOJI_DURUM]": "Dere, taşkın, deniz ve yeraltı suyu durumunu yazar.",
    "[LAB_FIZIK]": "Laboratuvar fiziksel deney ozet tablosunu ekler.",
    "[LAB_MEKANIK]": "Laboratuvar mekanik deney ozet tablosunu ekler.",
    "[ZEMIN_OZET]": "Laboratuvar ve SPT verilerine gore zemin ozet metnini yazar.",
    "[LITOLOJI_DAGILIM]": "Litoloji birimlerinin sondajlara gore derinlik dagilimini yazar.",
    "[SPT]": "SPT deney tablosunu ekler.",
    "[PMT]": "Presiyometre deney tablosunu ekler.",
    "[KAYA_TABLO]": "Kaya/karot deney tablosunu ekler.",
    "[JEO_PARAMETRE]": "Jeofizik parametre tablosunu ekler.",
    "[MASW]": "MASW/Vs tablosunu ekler.",
    "[VP]": "Vp tablosunu ekler.",
    "[JEO_KOOR]": "Jeofizik ve mikrotremor koordinat tablosunu ekler.",
    "[MT_TABLO]": "Mikrotremor olcu tablosunu ekler.",
    "[JEO_SONUC]": "Vs30 ve To degerlerinden jeofizik sonuc cumlesi uretir.",
    "[RESIM_YERBULDURUR]": "Yerbuldurur haritası görselini ekler.",
    "RESIM:Yerbuldurur": "Yerbuldurur haritası görselini ekler.",
    "RESIM:TKGM": "TKGM haritası görselini ekler.",
    "[RESIM:TKGM]": "TKGM haritası görselini ekler.",
    "[RESIM:PGA]": "PGA haritası görselini ekler.",
    "RESIM:PGA": "PGA haritası görselini ekler.",
    "[RESIM_JEOFIZIK]": "Jeofizik/serim haritası görselini ekler.",
    "RESIM:MJH": "Mühendislik jeolojisi haritası görselini ekler.",
    "[RESIM:MJH]": "Mühendislik jeolojisi haritası görselini ekler.",
    "[RESIM_MJH]": "Mühendislik jeolojisi haritası görselini ekler.",
    "[RESIM_SONDAJ]": "Sondaj vaziyet planı görselini ekler.",
    "[RESIM:SONDAJ]": "Sondaj vaziyet planı görselini ekler.",
}

RECOMMENDED_TAGS = [
    "[PROJE_ADI]", "[IL]", "[ILCE]", "[MAHALLE]", "[SONDAJ_BILGISI]",
    "[Sondaj]", "[SPT]", "[YASS_TABLO]", "[HIDROJEOLOJI_DURUM]", "[LAB_FIZIK]", "[JEO_PARAMETRE]",
    "[JEO_KOOR]", "[RESIM_SONDAJ]", "[RESIM_JEOFIZIK]",
]


def is_blank(value):
    return value is None or str(value).strip() in {"", "-", "None", "null"}


def number_or_none(value):
    if is_blank(value):
        return None
    try:
        return float(str(value).strip().replace(",", "."))
    except Exception:
        return None

def image_path_for_tag(app_instance, attr_name):
    return getattr(app_instance, attr_name, None)


def backup_project_file(project_path, keep=10):
    if not project_path or not os.path.exists(project_path):
        return None, None

    try:
        project_dir = os.path.dirname(os.path.abspath(project_path))
        backup_dir = os.path.join(project_dir, "backups")
        os.makedirs(backup_dir, exist_ok=True)

        stem, ext = os.path.splitext(os.path.basename(project_path))
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{stem}_{timestamp}{ext or '.json'}"
        backup_path = os.path.join(backup_dir, backup_name)

        counter = 2
        while os.path.exists(backup_path):
            backup_name = f"{stem}_{timestamp}_{counter}{ext or '.json'}"
            backup_path = os.path.join(backup_dir, backup_name)
            counter += 1

        shutil.copy2(project_path, backup_path)

        backups = [
            os.path.join(backup_dir, name)
            for name in os.listdir(backup_dir)
            if name.startswith(f"{stem}_") and name.endswith(ext or ".json")
        ]
        backups.sort(key=lambda path: os.path.getmtime(path), reverse=True)
        for old_path in backups[keep:]:
            try:
                os.remove(old_path)
            except OSError:
                pass

        return backup_path, None
    except Exception as exc:
        return None, str(exc)


def validate_project_data(veri):
    """Proje verisini yapılandırılmış tutarlılık motoruyla denetle."""
    return proje_tutarlilik_raporu(veri)


def read_word_tags(word_path):
    doc = Document(word_path)
    texts = []
    for paragraph in _iter_doc_paragraphs(doc):
        texts.append(paragraph.text or "")
    joined = "\n".join(texts)
    return sorted(set(re.findall(r"\[[^\]]+\]|RESIM:[A-Za-z0-9_:-]+", joined)))


def _iter_doc_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        containers = [
            section.header, section.first_page_header, section.even_page_header,
            section.footer, section.first_page_footer, section.even_page_footer,
        ]
        for container in containers:
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def is_known_tag(tag):
    return tag in KNOWN_TAGS or bool(PREFIXED_TAG_RE.match(tag))


def get_supported_tags():
    tags = set(KNOWN_TAGS) | set(TAG_DESCRIPTIONS)
    items = []
    for tag in sorted(tags, key=lambda value: value.upper()):
        if PREFIXED_TAG_RE.match(tag):
            category = "Sablon kopya"
        elif "RESIM" in tag:
            category = "Gorsel"
        elif tag in {"[Sondaj]", "[SPT]", "[PMT]", "[KAYA_TABLO]", "[LAB_FIZIK]", "[LAB_MEKANIK]", "[JEO_PARAMETRE]", "[MASW]", "[VP]", "[JEO_KOOR]", "[MT_TABLO]", "[YASS_TABLO]", "[BINA_BILGILERI]"}:
            category = "Tablo"
        else:
            category = "Metin"
        items.append({
            "tag": tag,
            "category": category,
            "description": TAG_DESCRIPTIONS.get(tag, "Desteklenen etiket."),
        })
    items.append({
        "tag": "[S1_PROJE_ADI] ... [S5_PARSEL]",
        "category": "Şablon kopya",
        "description": "Aynı proje künye bilgisini farklı şablon bölümlerinde kullanmak için S1_ ile S5_ arası ön ekler desteklenir.",
    })
    return items


def analyze_word_template(word_path):
    result = {
        "path": word_path,
        "tags": [],
        "known": [],
        "unknown": [],
        "missing_recommended": [],
        "error": None,
    }
    if is_blank(word_path):
        result["error"] = "Word şablonu seçilmemiş."
        return result
    if not os.path.exists(word_path):
        result["error"] = f"Word şablonu bulunamadı: {word_path}"
        return result
    try:
        tags = read_word_tags(word_path)
    except Exception as exc:
        result["error"] = f"Word şablonu okunamadı: {exc}"
        return result

    result["tags"] = tags
    result["known"] = [tag for tag in tags if is_known_tag(tag)]
    result["unknown"] = [tag for tag in tags if not is_known_tag(tag)]
    result["missing_recommended"] = [tag for tag in RECOMMENDED_TAGS if tag not in tags]
    return result


def format_template_analysis(analysis):
    lines = ["WORD ETİKET ANALİZİ", "=" * 19, ""]
    if analysis.get("path"):
        lines.append(f"Şablon: {analysis['path']}")
    if analysis.get("error"):
        lines.append(f"Hata: {analysis['error']}")
        return "\n".join(lines)

    lines.append(f"Toplam etiket: {len(analysis['tags'])}")
    lines.append(f"Bilinen etiket: {len(analysis['known'])}")
    lines.append(f"Bilinmeyen etiket: {len(analysis['unknown'])}")
    lines.append("")

    lines.append("BİLİNMEYEN ETİKETLER")
    if analysis["unknown"]:
        lines.extend(f"- {tag}" for tag in analysis["unknown"])
    else:
        lines.append("- Yok")
    lines.append("")

    lines.append("ŞABLONDA BULUNAN ETİKETLER")
    if analysis["tags"]:
        for tag in analysis["tags"]:
            marker = "OK" if is_known_tag(tag) else "UYARI"
            desc = TAG_DESCRIPTIONS.get(tag, "Kod karşılığı bilinmiyor.")
            lines.append(f"- {marker}: {tag} - {desc}")
    else:
        lines.append("- Etiket bulunamadı.")
    lines.append("")

    lines.append("ÖNERİLEN AMA EKSİK OLAN ETİKETLER")
    if analysis["missing_recommended"]:
        for tag in analysis["missing_recommended"]:
            lines.append(f"- {tag}: {TAG_DESCRIPTIONS.get(tag, '')}")
    else:
        lines.append("- Yok")
    return "\n".join(lines)


def _preflight_file_map(app_instance):
    if hasattr(app_instance, "_dosya_map"):
        try:
            return app_instance._dosya_map()
        except Exception:
            pass
    return {
        key: getattr(app_instance, key, None)
        for key in (
            "word_path", "lab_excel_path", "jeo_excel_path", "kml_path",
            "img_yer", "img_tkgm", "img_pga", "img_mjh",
            "word_img_sondaj", "word_img_jeofizik",
        )
    }


def _preflight_lab_rows(app_instance):
    internal_rows = app_instance.veri.get("lab_sheet", {}).get("rows", [])
    if any(any(not is_blank(cell) for cell in row) for row in internal_rows or []):
        return internal_rows, None
    path = getattr(app_instance, "lab_excel_path", None)
    if is_blank(path) or not os.path.exists(path):
        return [], None
    try:
        frame = pd.read_excel(path, header=None)
        frame = frame.where(pd.notna(frame), "")
        return frame.values.tolist(), None
    except Exception as exc:
        return [], str(exc)


def _image_quality_check(report, tag, path, attr_name):
    check_id = f"rapor.gorsel.{attr_name}.{re.sub(r'[^a-z0-9]+', '', tag.casefold())}"
    exists = bool(path and os.path.isfile(path))
    kontrol_ekle(
        report,
        check_id,
        "Rapor görselleri",
        tag,
        exists,
        os.path.basename(path) if exists else f"{tag} için görsel seçilmemiş veya dosya bulunamıyor.",
        "haritalar",
        "Haritalar sekmesinde ilgili görseli yeniden oluşturup Word için aktarın.",
        failure_level="warning",
        weight=1,
    )
    if not exists:
        return
    try:
        with Image.open(path) as image:
            width, height = image.size
            image.verify()
        if width * height < 600_000 or min(width, height) < 400:
            bulgu_ekle(
                report,
                f"{check_id}.cozunurluk",
                "warning",
                "Rapor görselleri",
                "Görsel çözünürlüğü",
                f"{tag} görseli düşük çözünürlüklü görünüyor ({width}x{height} px).",
                "haritalar",
                "Haritayı Word için yeniden ve daha yüksek çözünürlükle aktarın.",
            )
    except Exception as exc:
        bulgu_ekle(
            report,
            f"{check_id}.okuma",
            "error",
            "Rapor görselleri",
            "Görsel dosyası",
            f"{tag} görseli açılamıyor: {exc}",
            "haritalar",
            "Görseli yeniden oluşturun veya geçerli JPG/PNG dosyası seçin.",
        )


def build_preflight_report(app_instance):
    """Veri, dosya, şablon ve çıktı kaynaklarını tek raporda denetle."""
    file_map = _preflight_file_map(app_instance)
    lab_rows, lab_read_error = _preflight_lab_rows(app_instance)
    report = proje_tutarlilik_raporu(app_instance.veri, file_map, lab_rows=lab_rows)

    if lab_read_error:
        bulgu_ekle(
            report,
            "rapor.lab.okuma",
            "error",
            "Laboratuvar",
            "Lab Excel",
            f"Lab Excel okunamadı: {lab_read_error}",
            "rapor",
            "Lab Excel dosyasını kontrol edin veya veriyi LAB Sheet'e yapıştırın.",
        )

    sondajlar = app_instance.veri.get("sondaj", []) or []
    kesit_options = app_instance.veri.get("kesit_ayarlari", {}) or {}
    selected_names = kesit_options.get("selected_sondajlar") or []
    if selected_names:
        selected_set = set(selected_names)
        selected = [item for item in sondajlar if item.get("no") in selected_set]
    else:
        selected = sondajlar
    if len(selected) >= 2:
        try:
            from kesit_kalite import build_section_quality_report

            section_report = build_section_quality_report(selected, kesit_options)
            section_errors = section_report.get("errors", []) or []
            section_warnings = section_report.get("warnings", []) or []
            kontrol_ekle(
                report,
                "kesit.kalite",
                "Kesit",
                "Kesit kalite kontrolü",
                not section_errors and not section_warnings,
                (
                    "Kesit kalite kontrolü temiz."
                    if not section_errors and not section_warnings
                    else f"{len(section_errors)} hata, {len(section_warnings)} uyarı"
                ),
                "kesit",
                "Kesit kalite ekranında tabaka eşleşmelerini ve koordinatları kontrol edin.",
                failure_level="error" if section_errors else "warning",
                weight=1,
            )
            for idx, detail in enumerate(section_errors):
                bulgu_ekle(
                    report,
                    f"kesit.hata.{idx}",
                    "error",
                    "Kesit",
                    "Kesit hatası",
                    str(detail),
                    "kesit",
                    "Kesit seçim ve kalite ekranından ilgili veriyi düzeltin.",
                )
            for idx, detail in enumerate(section_warnings):
                bulgu_ekle(
                    report,
                    f"kesit.uyari.{idx}",
                    "warning",
                    "Kesit",
                    "Kesit uyarısı",
                    str(detail),
                    "kesit",
                    "Kesit seçim ve kalite ekranından ilgili veriyi kontrol edin.",
                )
            report.setdefault("stats", {})["kesit_sondaj"] = len(selected)
        except Exception as exc:
            bulgu_ekle(
                report,
                "kesit.kalite.calistirma",
                "warning",
                "Kesit",
                "Kesit kalite kontrolü",
                f"Kesit kalite kontrolü çalıştırılamadı: {exc}",
                "kesit",
                "Kesit verilerini kontrol edin.",
            )
    else:
        kontrol_ekle(
            report,
            "kesit.secim",
            "Kesit",
            "Kesit seçimi",
            False,
            "Kesit için en az iki sondaj seçilmemiş.",
            "kesit",
            "Kesit seçim ekranından çizilecek sondajları seçin.",
            failure_level="warning",
            weight=1,
        )

    tags = []
    word_path = etkin_rapor_sablonu_yolu(getattr(app_instance, "word_path", None))
    if word_path and os.path.isfile(word_path):
        try:
            tags = read_word_tags(word_path)
            bulgu_ekle(
                report,
                "rapor.sablon.etiket_bilgisi",
                "info",
                "Rapor şablonu",
                "Word etiketleri",
                f"Word şablonunda {len(tags)} farklı etiket bulundu.",
                "rapor",
                blocking=False,
            )
        except Exception as exc:
            bulgu_ekle(
                report,
                "rapor.sablon.okuma",
                "error",
                "Rapor şablonu",
                "Word şablonu",
                f"Word şablonu okunamadı: {exc}",
                "rapor",
                "Geçerli bir DOCX şablonu seçin.",
            )

    for tag in tags:
        if not is_known_tag(tag):
            bulgu_ekle(
                report,
                f"rapor.sablon.bilinmeyen.{re.sub(r'[^a-z0-9]+', '', tag.casefold())}",
                "warning",
                "Rapor şablonu",
                "Bilinmeyen etiket",
                f"Word şablonunda kod karşılığı bilinmeyen etiket var: {tag}",
                "rapor",
                "Etiketi Etiket Yöneticisi ile karşılaştırın.",
            )

    for tag, attr_name in IMAGE_TAGS.items():
        if tag in tags:
            _image_quality_check(report, tag, image_path_for_tag(app_instance, attr_name), attr_name)

    lab_tags = {"[LAB_FIZIK]", "[LAB_MEKANIK]", "[ZEMIN_OZET]"}
    if lab_tags.intersection(tags):
        lab_ready = bool(lab_rows)
        kontrol_ekle(
            report,
            "rapor.kaynak.lab",
            "Rapor veri kaynakları",
            "Laboratuvar verisi",
            lab_ready,
            "Laboratuvar verisi okunabildi." if lab_ready else "Laboratuvar etiketleri var ancak LAB verisi yok.",
            "rapor",
            "LAB Sheet doldurun veya Lab Excel seçin.",
            failure_level="warning",
            weight=1,
        )

    jeo_tags = {"[JEO_PARAMETRE]", "[MASW]", "[VP]", "[JEO_KOOR]", "[JEO_SONUC]"}
    if jeo_tags.intersection(tags):
        jeo_rows = app_instance.veri.get("jeofizik_sheet", {}).get("rows", [])
        jeo_sheet_ready = jeofizik_sheet_var_mi(app_instance.veri) and bool(jeofizik_sheet_rows_to_ss_list(jeo_rows))
        manual_ready = any(
            item.get("layers") for item in app_instance.veri.get("jeofizik", {}).get("ss_list", [])
        )
        jeo_path = getattr(app_instance, "jeo_excel_path", None)
        jeo_ready = jeo_sheet_ready or manual_ready or bool(jeo_path and os.path.isfile(jeo_path))
        kontrol_ekle(
            report,
            "rapor.kaynak.jeofizik",
            "Rapor veri kaynakları",
            "Jeofizik parametreleri",
            jeo_ready,
            "Jeofizik parametre kaynağı hazır." if jeo_ready else "Jeofizik etiketleri var ancak parametre verisi yok.",
            "jeofizik",
            "Jeofizik Sheet doldurun, Excel bağlayın veya manuel tabaka verisi girin.",
            failure_level="warning",
            weight=1,
        )

    if "[SPT]" in tags:
        has_spt = any(item.get("spt") for item in sondajlar)
        kontrol_ekle(
            report,
            "rapor.kaynak.spt",
            "Rapor veri kaynakları",
            "SPT tablosu",
            has_spt,
            "SPT kayıtları hazır." if has_spt else "Şablonda [SPT] var ancak SPT kaydı yok.",
            "workbook",
            "Workbook SPT sayfasını doldurun.",
            failure_level="warning",
            weight=1,
            sheet="spt",
        )

    try:
        ek_label, ek_path = uygun_ek_sablonu(app_instance.veri)
        ek_ok = bool(ek_path and os.path.isfile(ek_path))
        kontrol_ekle(
            report,
            "rapor.ek_sablonu",
            "Rapor ekleri",
            ek_label,
            ek_ok,
            os.path.basename(ek_path) if ek_ok else f"Ek dosyası bulunamadı: {ek_path or '-'}",
            "rapor",
            "Rapor sekmesinde doğru ek setini veya şablonunu seçin.",
            failure_level="warning",
            weight=1,
        )
    except Exception as exc:
        bulgu_ekle(
            report,
            "rapor.ek_sablonu.okuma",
            "warning",
            "Rapor ekleri",
            "Ek seçimi",
            f"Ek seçimi kontrol edilemedi: {exc}",
            "rapor",
            "Ek ayarlarını kontrol edin.",
        )

    report.setdefault("stats", {})["word_tags"] = len(tags)
    return kontrol_raporunu_tamamla(report)


def format_preflight_report(report):
    lines = ["RAPOR ÖN KONTROL", "=" * 18, ""]
    sections = [
        ("HATALAR", report.get("errors", [])),
        ("UYARILAR", report.get("warnings", [])),
        ("BİLGİ", report.get("info", [])),
    ]
    for title, items in sections:
        lines.append(title)
        if items:
            for item in items:
                lines.append(f"- {item}")
        else:
            lines.append("- Yok")
        lines.append("")
    return "\n".join(lines).strip()

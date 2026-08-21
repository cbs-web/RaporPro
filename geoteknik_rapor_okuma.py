# Dosya: RaporPro/geoteknik_rapor_okuma.py
"""Geoteknik raporlardan kontrollü yapı bilgisi önerileri üretir."""

from __future__ import annotations

import contextlib
import io
import os
import re
import unicodedata
from dataclasses import asdict, dataclass
from pathlib import Path

import fitz
from docx import Document


class GeoteknikRaporOkumaHatasi(RuntimeError):
    """Geoteknik rapor açılamadığında veya yapı bilgisi bulunamadığında kullanılır."""


@dataclass(frozen=True)
class GeoteknikRaporAlani:
    bolum: str
    anahtar: str
    etiket: str
    deger: str
    kaynak: str
    belge_turu: str
    guven: float
    blok_adi: str = ""
    alternatifler: tuple[str, ...] = ()
    uyari: str = ""
    kanit: str = ""


_FIELD_LABELS = {
    "kul": "Bina Kullanım Amacı",
    "sinif": "Bina Kullanım Sınıfı",
    "onem": "Bina Önem Katsayısı",
    "malz": "Yapı Malzemesi",
    "bod": "Bodrum Kat Adedi",
    "kat": "Toplam Kat Adedi",
    "plan": "Plan Boyutları",
    "yukseklik": "Yapı Yüksekliği (Hn)",
    "yukseklik_sinif": "Bina Yükseklik Sınıfı",
    "temel_alan": "Temel Alanı (m²)",
    "ins": "Toplam İnşaat Alanı (m²)",
    "der": "Olası Kazı Derinliği (m)",
    "tem": "Temel Tipi (proje verisinde saklanır)",
    "ysinif": "Yerel Zemin Sınıfı (proje verisinde saklanır)",
    "gqe_min": "G+Q+E Minimum (t/m²)",
    "gqe_ort": "G+Q+E Ortalama (t/m²)",
    "gqe_max": "G+Q+E Maksimum (t/m²)",
    "comb_min": "1.4G+1.6Q Minimum (t/m²)",
    "comb_ort": "1.4G+1.6Q Ortalama (t/m²)",
    "comb_max": "1.4G+1.6Q Maksimum (t/m²)",
}

_FIELD_ORDER = tuple(_FIELD_LABELS)

_IMAR_FIELD_LABELS = {
    "imar_alani": "İmar kullanım kararı / alan türü",
    "parsel_tipi": "Parsel tipi",
    "yol_cephe_sayisi": "Yol cephesi sayısı",
    "yol_yonleri": "Yol cephe yönleri",
    "yol_cepheleri": "Yol ve cephe açıklaması",
    "komsu_parseller": "Komşu parsel numaraları",
    "mevcut_yapilar": "Komşu parsellerdeki yapı durumu",
    "mevcut_kullanim": "Parselin mevcut kullanımı",
    "yol_kaplama": "Yol kaplama türü",
    "yaya_trafik": "Yaya trafiğine açıklık",
    "tasit_trafik": "Taşıt trafiğine açıklık",
    "ulasim_durumu": "Ulaşım durumu",
    "altyapi_durumu": "Altyapı durumu",
    "dogalgaz_hatti": "Doğalgaz hattı",
    "elektrik_hatti": "Elektrik hattı",
    "kanalizasyon_hatti": "Kanalizasyon hattı",
    "temiz_su_hatti": "Temiz/içme suyu hattı",
    "telekom_hatti": "Telekom hattı",
    "yagmur_suyu_hatti": "Yağmur suyu hattı",
    "egim": "Genel eğim",
    "yon": "Eğim yönü",
}

_TURKISH_NUMBER_WORDS = {
    "bir": "1",
    "iki": "2",
    "üç": "3",
    "uc": "3",
    "dört": "4",
    "dort": "4",
}


def _ascii_key(value):
    text = str(value or "").replace("ı", "i").replace("İ", "I")
    text = unicodedata.normalize("NFKD", text)
    text = "".join(char for char in text if not unicodedata.combining(char))
    return re.sub(r"[^a-z0-9]+", " ", text.casefold()).strip()


def _clean_text(value):
    text = str(value or "").replace("\xa0", " ")
    text = text.replace("−", "-").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", text).strip()


def _number_text(value):
    match = re.search(r"[-+]?\d+(?:[.,]\d+)*", _clean_text(value))
    return match.group(0).lstrip("+") if match else ""


def _positive_number_text(value):
    return _number_text(value).lstrip("-")


def _numeric_value(value):
    text = _number_text(value)
    if not text:
        return None
    if "," in text:
        text = text.replace(".", "").replace(",", ".")
    elif text.count(".") > 1:
        text = text.replace(".", "")
    try:
        return float(text)
    except ValueError:
        return None


def _same_value(first, second):
    left = _numeric_value(first)
    right = _numeric_value(second)
    if left is not None and right is not None:
        return abs(left - right) <= 0.005
    return _ascii_key(first) == _ascii_key(second)


def _block_name(value):
    key = _ascii_key(value)
    match = re.fullmatch(r"([a-z0-9]+) blok", key)
    if not match:
        match = re.fullmatch(r"blok ([a-z0-9]+)", key)
    if not match:
        return ""
    token = match.group(1)
    token = token.upper() if len(token) <= 3 else token.title()
    return f"{token} Blok"


def _declared_blocks(text):
    flat = _clean_text(text)
    blocks = []

    for match in re.finditer(
        r"\b([A-ZÇĞİÖŞÜ])\s*,\s*([A-ZÇĞİÖŞÜ])\s*(?:,|ve)\s*"
        r"([A-ZÇĞİÖŞÜ])\s+Blok\b",
        flat,
        flags=re.I,
    ):
        blocks.extend(f"{token.upper()} Blok" for token in match.groups())
    for match in re.finditer(
        r"\b([A-ZÇĞİÖŞÜ0-9]+)\s+Blok\b(?!\s+Say[ıi]s[ıi])",
        flat,
        flags=re.I,
    ):
        name = _block_name(f"{match.group(1)} Blok")
        if name:
            blocks.append(name)

    result = []
    for name in blocks:
        if name not in result:
            result.append(name)
    return result


def _row_key(label):
    key = _ascii_key(label)
    if "bodrum kat adedi" in key and "toplam kat adedi" in key:
        return "kat_pair"
    if "temel alani" in key and "toplam insaat alani" in key:
        return "area_pair"
    if "bina kullanim amaci" in key or key == "kullanim amaci":
        return "kul"
    if "bina kullanim sinifi" in key or key == "bks":
        return "sinif"
    if "bina onem katsayisi" in key:
        return "onem"
    if "yapi malzemesi" in key:
        return "malz"
    if key.startswith("bodrum kat adedi"):
        return "bod"
    if "toplam insaat alani" in key:
        return "ins"
    if "temel oturumu" in key or key.startswith("temel alani"):
        return "temel_alan"
    if key == "kat adedi" or key.startswith("toplam kat adedi"):
        return "kat"
    if re.fullmatch(r"b m", key):
        return "b"
    if re.fullmatch(r"l m", key):
        return "l"
    if key == "bys" or "bina yukseklik sinifi" in key:
        return "yukseklik_sinif"
    if "yapi yuksekligi" in key or "bina yuksekligi" in key:
        return "yukseklik"
    if "yerel zemin sinifi" in key:
        return "ysinif"
    if "ortalama kazi derinligi" in key or key.startswith("olasi kazi derinligi"):
        return "der"
    if key.startswith("temel tipi"):
        return "tem"
    if "1 4g 1 6" in key and "zemin gerilme" in key:
        return "comb"
    if "g q e" in key and "zemin gerilme" in key:
        return "gqe"
    return ""


def _row_values(row):
    return [_clean_text(cell) for cell in row if _clean_text(cell)]


def _fit_block_values(values, block_count):
    values = [value for value in values if value]
    if block_count <= 0:
        return values
    if len(values) == block_count:
        return values
    if len(values) > block_count and len(values) % block_count == 0:
        group_size = len(values) // block_count
        return [values[index * group_size] for index in range(block_count)]
    if len(values) > block_count:
        return values[-block_count:]
    return values + [""] * (block_count - len(values))


def _table_candidate(rows, page_no):
    normalized_rows = [list(row or []) for row in rows or []]
    blocks = []
    for row in normalized_rows[:8]:
        for cell in row:
            name = _block_name(cell)
            if name and name not in blocks:
                blocks.append(name)

    recognized = []
    max_value_count = 0
    for row in normalized_rows:
        parts = _row_values(row)
        if not parts:
            continue
        key = _row_key(parts[0])
        if not key:
            continue
        values = parts[1:]
        max_value_count = max(max_value_count, len(values))
        recognized.append((key, values))

    if len(recognized) < 3:
        return None
    if not blocks:
        inferred_count = max(1, min(max_value_count, 12))
        blocks = ["Yapı"] if inferred_count == 1 else [f"Blok {i + 1}" for i in range(inferred_count)]

    row_map = {}
    for key, values in recognized:
        fitted = _fit_block_values(values, len(blocks))
        if any(fitted):
            row_map[key] = fitted
    return {
        "page": page_no,
        "blocks": blocks,
        "rows": row_map,
        "score": len(row_map),
    }


def _ocr_tokens_to_text(tokens):
    if not tokens:
        return ""
    lines = []
    current = []
    current_y = None
    current_height = 1.0
    for token in sorted(tokens, key=lambda item: (item.cy, item.x0)):
        tolerance = max(current_height, token.height) * 0.65
        if current and current_y is not None and abs(token.cy - current_y) > tolerance:
            lines.append(" ".join(item.text for item in sorted(current, key=lambda item: item.x0)))
            current = []
        current.append(token)
        current_y = sum(item.cy for item in current) / len(current)
        current_height = sum(item.height for item in current) / len(current)
    if current:
        lines.append(" ".join(item.text for item in sorted(current, key=lambda item: item.x0)))
    return "\n".join(lines)


def _pdf_pages(path, task_context=None):
    pages = []
    ocr_used = False
    try:
        document = fitz.open(path)
    except Exception as exc:
        raise GeoteknikRaporOkumaHatasi(f"PDF açılamadı: {exc}") from exc

    with document:
        total = document.page_count
        for index, page in enumerate(document):
            if task_context is not None:
                task_context.check_cancelled()
                task_context.report(index, total, f"Geoteknik rapor okunuyor: sayfa {index + 1}/{total}")
            text = page.get_text("text") or ""
            if len(re.sub(r"\W", "", text, flags=re.UNICODE)) < 30:
                try:
                    from evrak_okuma import _ocr_clip

                    text = _ocr_tokens_to_text(_ocr_clip(page, (0.0, 0.0, 1.0, 1.0), scale=2.0))
                    ocr_used = ocr_used or bool(text)
                except Exception:
                    pass
            tables = []
            page_key = _ascii_key(text)
            table_hints = (
                "yapi bilgileri",
                "toplam insaat alani",
                "temel oturumu",
                "g q e yuklemesi zemin gerilme",
            )
            if any(hint in page_key for hint in table_hints):
                try:
                    with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                        tables = [table.extract() for table in page.find_tables().tables]
                except Exception:
                    tables = []
            pages.append({"no": index + 1, "text": text, "tables": tables})
        if task_context is not None:
            task_context.report(total, total, "Geoteknik rapor metni ve tabloları okundu")
    return pages, ocr_used


def _docx_pages(path, task_context=None):
    try:
        document = Document(path)
    except Exception as exc:
        raise GeoteknikRaporOkumaHatasi(f"Word raporu açılamadı: {exc}") from exc
    if task_context is not None:
        task_context.report(0, 1, "Geoteknik Word raporu okunuyor")
        task_context.check_cancelled()
    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    if task_context is not None:
        task_context.report(1, 1, "Geoteknik Word raporu okundu")
    return [{"no": 1, "text": text, "tables": tables}], False


def _first_page_match(pages, pattern, flags=re.I):
    regex = re.compile(pattern, flags)
    for page in pages:
        match = regex.search(_clean_text(page.get("text", "")))
        if match:
            return match, page.get("no", 1)
    return None, None


def _common_fields(pages):
    result = {}

    def add(key, value, page, confidence):
        if value and key not in result:
            result[key] = {
                "value": _clean_text(value),
                "page": page or 1,
                "confidence": confidence,
                "alternatives": [],
                "warning": "",
            }

    full_text = "\n".join(page.get("text", "") for page in pages)
    full_key = _ascii_key(full_text)
    usage_patterns = (
        (r"mesken\s*\+?\s*işyeri\s+niteliğinde", "Mesken + İşyeri"),
        (r"konut\s*(?:ve|\+)\s*işyeri", "Konut + İşyeri"),
        (
            r"(?:kullanım\s+amacı|parselin\s+plan\s+fonksiyonu)\s*[:=]?\s*"
            r"(mesken|konut|ticaret|otel|okul|hastane)",
            None,
        ),
        (r"\b(mesken|konut|ticaret|otel)\s+niteliğinde", None),
    )
    for pattern, fixed_value in usage_patterns:
        match, page = _first_page_match(pages, pattern)
        if match:
            value = fixed_value or match.group(1).title()
            add("kul", value, page, 0.91)
            break

    match, page = _first_page_match(
        pages,
        r"Bina\s+Kullanım\s+Sınıfı\s*(?:\(\s*BKS\s*\))?\s*(?:BKS\s*)?[:=]?\s*(\d+)",
    )
    if match:
        add("sinif", match.group(1), page, 0.96)

    match, page = _first_page_match(
        pages,
        r"Bina\s+Önem\s+Katsayısı\s*(?:\(?\s*I\s*\)?)?\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
    )
    if match:
        add("onem", match.group(1), page, 0.96)

    match, page = _first_page_match(pages, r"betonarme(?:\s+çerçeve)?\s+taşıyıcı")
    if match:
        add("malz", "Betonarme", page, 0.95)
    elif "betonarme" in full_key:
        page = next((item["no"] for item in pages if "betonarme" in _ascii_key(item.get("text", ""))), 1)
        add("malz", "Betonarme", page, 0.86)

    match, page = _first_page_match(pages, r"(\d+)\s*(?:adet\s+)?bodrum\s+kat")
    if match:
        add("bod", match.group(1), page, 0.94)
    else:
        match, page = _first_page_match(
            pages,
            r"(?:Bodrum\s*[,;+]|Bodrum\s+Kat\s*\+)\s*(?:Zemin|Zemin\s+Kat)|bodrum\s+katlar[ıi]\s+bulunmaktadır|bodrum\s+katlı",
        )
        if match:
            add("bod", "1", page, 0.86)

    match, page = _first_page_match(
        pages,
        r"Bina\s+Yüksekliği\s*(?:\(\s*m\s*\))?\s*[:=]?\s*([0-9]+(?:[.,][0-9]+)?)",
    )
    if match:
        add("yukseklik", match.group(1), page, 0.92)

    match, page = _first_page_match(
        pages,
        r"Bina\s+Yükseklik\s+Sınıfı\s*(?:\(\s*BYS\s*\))?\s*(?:BYS\s*)?[:=]?\s*(\d+)",
    )
    if match:
        add("yukseklik_sinif", match.group(1), page, 0.96)

    match, page = _first_page_match(pages, r"Yerel\s+Zemin\s+Sınıfı\s*[:=]?\s*(Z[A-E])")
    if match:
        add("ysinif", match.group(1).upper(), page, 0.97)

    if "radye temel" in full_key:
        page = next((item["no"] for item in pages if "radye temel" in _ascii_key(item.get("text", ""))), 1)
        add("tem", "Radye", page, 0.92)

    match, page = _first_page_match(pages, r"toplam\s+(\d+)\s*(?:ar\s+)?katlı")
    if match:
        add("kat", match.group(1), page, 0.84)
    return result


def _single_structure_fields(pages):
    """Tek yapılı raporlardaki anlatım metninden geometri ve yükleri ayıkla."""
    result = {}

    def add(key, value, page, confidence):
        value = _clean_text(value)
        if value and key not in result:
            result[key] = {
                "value": value,
                "page": page or 1,
                "confidence": confidence,
                "alternatives": [],
                "warning": "",
            }

    match, page = _first_page_match(
        pages,
        r"bodrum\s+kat[ıi]\s+olmay[ıi]p|bodrum\s+kats[ıi]z|bodrumsuz",
    )
    if match:
        add("bod", "0", page, 0.98)

    match, page = _first_page_match(pages, r"\b(\d+)\s+katl[ıi]\b")
    if match:
        add("kat", match.group(1), page, 0.88)

    match, page = _first_page_match(
        pages,
        r"temel\s+taban[ıi]\s+([0-9]+(?:[.,][0-9]+)?)\s*m(?:2|²)",
    )
    if not match:
        match, page = _first_page_match(
            pages,
            r"temel\s+(?:oturumu|alan[ıi])\s*[:=]?\s*"
            r"([0-9]+(?:[.,][0-9]+)?)\s*m(?:2|²)",
        )
    if match:
        add("temel_alan", match.group(1), page, 0.97)

    match, page = _first_page_match(
        pages,
        r"toplam\s+inşaat\s+alan[ıi](?:\s+ise)?\s*[:=]?\s*"
        r"([0-9]+(?:[.,][0-9]+)?)\s*m(?:2|²)",
    )
    if match:
        add("ins", match.group(1), page, 0.97)

    match, page = _first_page_match(
        pages,
        r"Yap[ıi]\s+Ebatlar[ıi]\s*/\s*Temel\s+Ebatlar[ıi]\s*"
        r"([0-9]+(?:[.,][0-9]+)?)\s*[*x×]\s*"
        r"([0-9]+(?:[.,][0-9]+)?)",
    )
    if match:
        add("plan", f"{match.group(1)} × {match.group(2)} m", page, 0.98)

    match, page = _first_page_match(
        pages,
        r"kaz[ıi]\s+(?:yüksekliği|derinliği)\s*"
        r"([0-9]+(?:[.,][0-9]+)?)\s*m",
    )
    if not match:
        match, page = _first_page_match(
            pages,
            r"temel\s+taban\s+kot[uo]\s*"
            r"(-?[0-9]+(?:[.,][0-9]+)?)\s*m",
        )
    if match:
        add("der", _positive_number_text(match.group(1)), page, 0.92)

    load_patterns = (
        (
            "comb",
            r"1\s*[,.]\s*4\s*G\s*\+\s*1\s*[,.]\s*6\s*Q\s+"
            r"yüklemesi\s+alt[ıi]nda\s+minimum\s*-\s*ortalama\s*-\s*maksimum\s*",
        ),
        (
            "gqe",
            r"(?:depremli\s+yüklemeler\s*)?\(\s*G\s*\+\s*Q\s*\+\s*E\s*\)\s*"
            r"alt[ıi]nda(?:\s+ise)?\s+minimum\s*-\s*ortalama\s*-\s*maksimum\s*",
        ),
    )
    triplet_tail = (
        r"([0-9]+(?:[.,][0-9]+)?)\s*t\s*/\s*m(?:2|²)\s*-\s*"
        r"([0-9]+(?:[.,][0-9]+)?)\s*t\s*/\s*m(?:2|²)\s*-\s*"
        r"([0-9]+(?:[.,][0-9]+)?)\s*t\s*/\s*m(?:2|²)"
    )
    for prefix, pattern in load_patterns:
        match, page = _first_page_match(pages, pattern + triplet_tail)
        if not match:
            continue
        for suffix, value in zip(("min", "ort", "max"), match.groups()[-3:]):
            add(f"{prefix}_{suffix}", value, page, 0.98)

    return result


def _triplet(value):
    numbers = re.findall(r"\d+(?:[.,]\d+)*", _clean_text(value))
    if len(numbers) < 3:
        return ("", "", "")
    return tuple(number.lstrip("+") for number in numbers[:3])


def _pair(value):
    parts = re.split(r"\s*/\s*", _clean_text(value), maxsplit=1)
    return tuple(parts) if len(parts) == 2 else ("", "")


def _put_field(block_data, block, key, value, page, confidence=0.97):
    value = _clean_text(value)
    if not value:
        return
    current = block_data.setdefault(block, {}).get(key)
    if current is None:
        block_data[block][key] = {
            "value": value,
            "page": page,
            "confidence": confidence,
            "alternatives": [],
            "warning": "",
        }
        return
    if not _same_value(current["value"], value):
        alternative = f"{value} (Sayfa {page})"
        if alternative not in current["alternatives"]:
            current["alternatives"].append(alternative)
        current["warning"] = "Raporda aynı alan için farklı bir değer bulundu."


def _table_fields(pages):
    candidates = []
    for page in pages:
        for table in page.get("tables", []):
            candidate = _table_candidate(table, page.get("no", 1))
            if candidate:
                candidates.append(candidate)
    if not candidates:
        return [], {}, []

    best_score = max(candidate["score"] for candidate in candidates)
    candidates = [candidate for candidate in candidates if candidate["score"] >= max(3, best_score - 2)]
    block_order = []
    block_data = {}

    for candidate in candidates:
        blocks = candidate["blocks"]
        rows = candidate["rows"]
        page = candidate["page"]
        for block in blocks:
            if block not in block_order:
                block_order.append(block)
        for index, block in enumerate(blocks):
            for key in (
                "kul",
                "sinif",
                "onem",
                "malz",
                "bod",
                "ins",
                "temel_alan",
                "kat",
                "yukseklik",
                "yukseklik_sinif",
                "tem",
                "ysinif",
            ):
                values = rows.get(key, [])
                if index < len(values):
                    value = values[index]
                    if key in {"sinif", "onem", "bod", "kat", "yukseklik", "yukseklik_sinif"}:
                        value = _number_text(value)
                    elif key == "ysinif":
                        match = re.search(r"Z[A-E]", value, flags=re.I)
                        value = match.group(0).upper() if match else value
                    _put_field(block_data, block, key, value, page)
            pair_values = rows.get("kat_pair", [])
            if index < len(pair_values):
                basement, total = _pair(pair_values[index])
                _put_field(block_data, block, "bod", _number_text(basement), page)
                _put_field(block_data, block, "kat", _number_text(total), page)
            pair_values = rows.get("area_pair", [])
            if index < len(pair_values):
                foundation, total = _pair(pair_values[index])
                _put_field(block_data, block, "temel_alan", _number_text(foundation), page)
                _put_field(block_data, block, "ins", _number_text(total), page)
            if index < len(rows.get("der", [])):
                _put_field(block_data, block, "der", _positive_number_text(rows["der"][index]), page)

            b_values = rows.get("b", [])
            l_values = rows.get("l", [])
            if index < len(b_values) and index < len(l_values):
                b_value = _clean_text(b_values[index])
                l_value = _clean_text(l_values[index])
                if b_value and l_value:
                    _put_field(block_data, block, "plan", f"{b_value} × {l_value} m", page)

            for row_key, prefix in (("gqe", "gqe"), ("comb", "comb")):
                values = rows.get(row_key, [])
                if index >= len(values):
                    continue
                minimum, average, maximum = _triplet(values[index])
                _put_field(block_data, block, f"{prefix}_min", minimum, page)
                _put_field(block_data, block, f"{prefix}_ort", average, page)
                _put_field(block_data, block, f"{prefix}_max", maximum, page)
    return block_order, block_data, candidates


def _pressure_cross_checks(pages, block_data, warnings):
    checks = (
        (r"Temel\s+taban\s+basıncı\s*\(\s*Statik\s*\)", "comb_max"),
        (r"Temel\s+taban\s+basıncı\s*\(\s*Depremli\s*\)", "gqe_max"),
    )
    for page in pages:
        text = _clean_text(page.get("text", ""))
        for index, (heading_pattern, key) in enumerate(checks):
            heading = re.search(heading_pattern, text, flags=re.I)
            if heading is None:
                continue
            segment_end = len(text)
            for next_pattern, _ in checks[index + 1:]:
                next_heading = re.search(next_pattern, text[heading.end():], flags=re.I)
                if next_heading is not None:
                    segment_end = heading.end() + next_heading.start()
                    break
            if key == "gqe_max":
                stop = re.search(r"Yeraltı\s+suyu|TBDY", text[heading.end():], flags=re.I)
                if stop is not None:
                    segment_end = min(segment_end, heading.end() + stop.start())
            segment = text[heading.end():segment_end]
            pairs = re.findall(
                r":?\s*([0-9]+(?:[.,][0-9]+)?)\s*t\s*/\s*m2\s*"
                r"(?:\([^)]*\)\s*)?\(([A-ZÇĞİÖŞÜ0-9]+)\s+Blok\)",
                segment,
                flags=re.I,
            )
            for value, token in pairs:
                block = _block_name(f"{token} Blok")
                record = block_data.get(block, {}).get(key)
                if not record or _same_value(record["value"], value):
                    continue
                alternative = f"{value} (Sayfa {page.get('no', 1)})"
                if alternative not in record["alternatives"]:
                    record["alternatives"].append(alternative)
                record["warning"] = "Raporda aynı alan için farklı bir değer bulundu."
                warnings.append(
                    f"{block} {_FIELD_LABELS[key]} için {record['value']} ve {value} değerleri bulundu."
                )


def _page_content(page):
    """Sayfa metni ile DOCX/PDF tablo hücrelerini tek arama metninde birlestirir."""
    parts = [_clean_text(page.get("text", ""))]
    for table in page.get("tables", []) or []:
        for row in table or []:
            parts.append(" ".join(_clean_text(cell) for cell in row or []))
    return _clean_text(" ".join(part for part in parts if part))


def _imar_kaynak_siralama(pages):
    def score(page):
        key = _ascii_key(_page_content(page))
        value = 0
        if "insaat sahasi hakkinda bilgiler" in key:
            value += 100
        if "parselin plan fonksiyonu" in key or "imar durum bilgileri" in key:
            value += 55
        if "10 2 kazi sevi" in key or "kazi sevi guvenligi" in key:
            value += 15
        return value

    return sorted(pages, key=lambda page: (-score(page), page.get("no", 1)))


def _imar_bolum_adi(text):
    key = _ascii_key(text)
    if "insaat sahasi hakkinda bilgiler" in key:
        return "2 - İnşaat Sahası Hakkında Bilgiler"
    if "imar durum bilgileri" in key or "parselin plan fonksiyonu" in key:
        return "Tablo 1 - İmar Durum Bilgileri"
    if "kazi sevi guvenligi" in key or "10 2 kazi" in key:
        return "10.2 - Kazı Şevi Güvenliği"
    return "Rapor metni"


def _imar_kanit(text, match, limit=260):
    clean = _clean_text(text)
    start = max(0, match.start() - 90)
    end = min(len(clean), match.end() + 150)
    evidence = clean[start:end].strip(" ,;:")
    if start > 0:
        evidence = "…" + evidence
    if end < len(clean):
        evidence += "…"
    return evidence[:limit]


def _imar_deger_kucult(value):
    value = _clean_text(value)
    return value.strip(" .,:;()")


def _imar_sayi(value):
    value = _clean_text(value).casefold()
    return _TURKISH_NUMBER_WORDS.get(value, _number_text(value))


def _imar_parsel_tipi(value):
    key = _ascii_key(value)
    if key.startswith("kose parsel"):
        return "Köşe parsel"
    if key.startswith("ara parsel"):
        return "Ara parsel"
    if key.startswith("ada parsel"):
        return "Ada parsel"
    return "Belirtilmedi"


def _imar_ayni_deger(key, first, second):
    """İmar alanlarını türlerine uygun biçimde karşılaştır."""
    if key == "yol_cephe_sayisi":
        return _same_value(first, second)
    if key == "komsu_parseller":
        left = tuple(re.findall(r"\d+", _clean_text(first)))
        right = tuple(re.findall(r"\d+", _clean_text(second)))
        if left and right:
            return left == right
    return _ascii_key(first) == _ascii_key(second)


def _imar_saha_alanlari(pages, dosya_adi):
    """Geoteknik raporun saha bolumunden onaylanabilir duz alan adaylari uretir."""
    records = {}
    warnings = []

    def add(key, value, page, confidence, match, section=None):
        value = _imar_deger_kucult(value)
        if not value:
            return
        page = page or 1
        text = _page_content(page_item_by_no.get(page, {}))
        section = section or _imar_bolum_adi(text)
        evidence = _imar_kanit(text, match)
        current = records.get(key)
        if current is None:
            records[key] = {
                "value": value,
                "page": page,
                "confidence": confidence,
                "alternatives": [],
                "warning": "",
                "evidence": evidence,
                "section": section,
            }
            return
        if _imar_ayni_deger(key, current["value"], value):
            return
        alternative = f"{value} (Sayfa {page})"
        if alternative not in current["alternatives"]:
            current["alternatives"].append(alternative)
        current["warning"] = (
            "Raporda aynı alan için farklı değer bulundu; otomatik seçim yapılmadı."
        )
        warning = (
            f"{_IMAR_FIELD_LABELS.get(key, key)} için {current['value']} "
            f"(Sayfa {current['page']}) ve {value} (Sayfa {page}) değerleri bulundu."
        )
        if warning not in warnings:
            warnings.append(warning)

    source_pages = _imar_kaynak_siralama(pages)
    page_item_by_no = {page.get("no", 1): page for page in pages}

    # Plan fonksiyonu / imar alani: tablo satirinin bir sonraki satira kadar olan
    # degeri alinir; sabit bir "Konut" varsayimi yapilmaz.
    plan_pattern = re.compile(
        r"parselin\s+plan\s+fonksiyonu\s+(.+?)(?=\s+(?:yan|arka|ön|on)\s+bahçe\s+çekme\s+mesafesi|\s+eğim\s+durumu|\s+mimari\b|$)",
        re.I,
    )
    for page in source_pages:
        text = _page_content(page)
        for match in plan_pattern.finditer(text):
            add("imar_alani", match.group(1), page.get("no", 1), 0.98, match)

    # Parsel tipi yalnız raporda açıkça yazan ibareden alınır. Yol geometrisi
    # tek başına Ada/Ara/Köşe parsel sonucu üretmez.
    parcel_pattern = re.compile(r"\b(köşe|kose|ara|ada)\s+parsel\b", re.I)
    for page in source_pages:
        text = _page_content(page)
        for match in parcel_pattern.finditer(text):
            value = _imar_parsel_tipi(match.group(0))
            if value != "Belirtilmedi":
                add("parsel_tipi", value, page.get("no", 1), 0.99, match)

    count_patterns = (
        re.compile(
            r"\b(?P<count>\d+|bir|iki|üç|uc|dört|dort)\s+yola\s+cepheli",
            re.I,
        ),
        re.compile(
            r"\b(?P<count>\d+|bir|iki|üç|uc|dört|dort)\s+tarafı(?:nda)?\s+yol",
            re.I,
        ),
        re.compile(
            r"\b(?P<count>\d+|bir|iki|üç|uc|dört|dort)\s+cephesi\s+yol",
            re.I,
        ),
        re.compile(
            r"parselin\s+(?P<count>\d+|bir|iki|üç|uc|dört|dort)\s+cephesi\s+yol",
            re.I,
        ),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern in count_patterns:
            for match in pattern.finditer(text):
                count = _imar_sayi(match.group("count"))
                if count:
                    add("yol_cephe_sayisi", count, page.get("no", 1), 0.96, match)
                    if "yol_cepheleri" not in records:
                        count_word = {
                            "1": "bir",
                            "2": "iki",
                            "3": "üç",
                            "4": "dört",
                        }.get(count, count)
                        add(
                            "yol_cepheleri",
                            f"{count_word} yola cepheli",
                            page.get("no", 1),
                            0.91,
                            match,
                        )

    neighbor_patterns = (
        re.compile(
            r"(?:komşu|komsu|diğer|diger|cephesinde|cephelerinde|sınırlı|sinirli)"
            r".{0,130}?\b(?P<numbers>\d+(?:\s*(?:ve|,)\s*\d+)*)\s*nolu\s+"
            r"(?:boş\s+|bos\s+)?(?:komşu\s+|komsu\s+)?parsel\w*",
            re.I,
        ),
        re.compile(
            r"\b(?P<numbers>\d+(?:\s*(?:ve|,)\s*\d+)*)\s+nolu\s+"
            r"(?:boş\s+|bos\s+)?komşu\s+parsel\w*",
            re.I,
        ),
        re.compile(
            r"\b(?P<numbers>\d+(?:\s*(?:ve|,)\s*\d+)*)\s+nolu\s+"
            r"parsel\w*\s+(?:ile\s+)?komşu\b",
            re.I,
        ),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern in neighbor_patterns:
            for match in pattern.finditer(text):
                number_prefix = text[max(0, match.start("numbers") - 35) : match.start("numbers")]
                if re.search(r"\bada\s*,?\s*$", number_prefix, flags=re.I):
                    continue
                numbers = re.sub(r"\s*(?:ve|,)\s*", ", ", match.group("numbers"))
                add("komsu_parseller", numbers, page.get("no", 1), 0.90, match)

    # Komşu yapı durumu ve mevcut kullanım; rapor açıkça söylemiyorsa alan
    # oluşturulmaz. "Boş" yalnız komşu parsellerin yapısızlığı açıkça yazılmışsa
    # kullanılır.
    neighbor_building_patterns = (
        re.compile(
            r"içerisinde\s+oturulan\s+[^.]{0,100}?yapılar\s+ile\s+boş\s+parseller\s+bulunmaktadır",
            re.I,
        ),
        re.compile(r"çevresinde\s+binalar\s+ile\s+boş\s+parseller\s+bulunmaktadır", re.I),
        re.compile(r"\bnolu\s+boş\s+parsel\s+bulunmaktadır", re.I),
        re.compile(r"komşu\s+parsellerde\s+mevcut\s+yapı\s+bulunmamaktadır", re.I),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern in neighbor_building_patterns:
            for match in pattern.finditer(text):
                value = match.group(0)
                if (
                    "bulunmamaktadir" in _ascii_key(value)
                    or "nolu\\s+boş" in pattern.pattern
                ):
                    value = "Boş"
                add("mevcut_yapilar", value, page.get("no", 1), 0.88, match)

    current_use_patterns = (
        re.compile(
            r"(?:taşınmaz(?:ın)?|parsel(?:in)?)\s+hal[iı]\s*haz[ıi]rda\s+"
            r"arsa\s+vasf[ıi]ndadır",
            re.I,
        ),
        re.compile(
            r"(?:taşınmaz(?:ın)?|parsel(?:in)?)\s+hal[iı]\s*haz[ıi]rda\s+"
            r"boş\s+durumdadır",
            re.I,
        ),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern in current_use_patterns:
            for match in pattern.finditer(text):
                value = "Arsa" if "arsa" in _ascii_key(match.group(0)) else "Boş"
                add("mevcut_kullanim", value, page.get("no", 1), 0.95, match)

    road_pattern = re.compile(
        r"(?P<surface>sıcak\s+asfalt|asfalt|kilit\s+parke|stabilize|toprak)"
        r"(?:\s+ile)?\s+(?:kaplı|kaplamalı|yol)",
        re.I,
    )
    for page in source_pages:
        text = _page_content(page)
        for match in road_pattern.finditer(text):
            add("yol_kaplama", match.group("surface"), page.get("no", 1), 0.95, match)

    # Yol yönü yalnız açıkça belirtilen yön/cephe ifadesinden alınır. Parsel
    # geometrisinden veya KML şeklinden yön çıkarımı yapılmaz.
    road_direction_pattern = re.compile(
        r"(?P<direction>"
        r"(?:ön|on|arka|kuzey|güney|doğu|batı)"
        r"(?:\s*(?:,|ve|ile|-)+\s*(?:ön|on|arka|kuzey|güney|doğu|batı))*"
        r")\s+(?:(?:bir|iki|üç|uc|dört|dort|\d+)\s+)?"
        r"(?:yön(?:ünde|lerinde)?|cephe(?:de|lerinde)?)"
        r"(?:.{0,45})?\byol\b",
        re.I,
    )
    for page in source_pages:
        text = _page_content(page)
        for match in road_direction_pattern.finditer(text):
            direction = _clean_text(match.group("direction"))
            add("yol_yonleri", direction, page.get("no", 1), 0.90, match)

    traffic_pattern = re.compile(
        r"yaya\s+ve\s+taşıt\s+trafiğine\s+(?P<status>kısmen\s+)?açık",
        re.I,
    )
    for page in source_pages:
        text = _page_content(page)
        for match in traffic_pattern.finditer(text):
            status = "Kısmen açık" if match.group("status") else "Açık"
            add("yaya_trafik", status, page.get("no", 1), 0.96, match)
            add("tasit_trafik", status, page.get("no", 1), 0.96, match)

    infrastructure_patterns = (
        (re.compile(r"(?:alt\s+ve\s+üst|altyapı\s+ve\s+üstyapı)\s+[^.]{0,50}tamamlanmamış", re.I), "Altyapı ve üstyapı çalışmaları tamamlanmamıştır."),
        (re.compile(r"(?:alt\s+ve\s+üst|altyapı\s+ve\s+üstyapı)\s+[^.]{0,50}tamamlanmış", re.I), "Altyapı ve üstyapı çalışmaları tamamlanmıştır."),
        (re.compile(r"ulaşım\s+ve\s+altyapı\s+sorunu\s+(?:yoktur|bulunmamaktadır)", re.I), "Yok"),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern, value in infrastructure_patterns:
            for match in pattern.finditer(text):
                add("altyapi_durumu", value, page.get("no", 1), 0.93, match)
                if value == "Yok":
                    add("ulasim_durumu", "Yok", page.get("no", 1), 0.93, match)

    utility_patterns = (
        ("dogalgaz_hatti", r"doğal\s*gaz"),
        ("elektrik_hatti", r"elektrik"),
        ("kanalizasyon_hatti", r"kanalizasyon"),
        ("temiz_su_hatti", r"(?:temiz|içme)\s+su"),
        ("telekom_hatti", r"telekom(?:ünikasyon)?"),
        ("yagmur_suyu_hatti", r"yağmur\s+suyu"),
    )
    for page in source_pages:
        text = _page_content(page)
        for key, term in utility_patterns:
            for match in re.finditer(term, text, flags=re.I):
                window = text[match.start() : match.start() + 180]
                if re.search(r"\bhat\w*\b|\bgeç\w*\b|\bbulun\w*\b", window, flags=re.I):
                    add(key, "Var", page.get("no", 1), 0.91, match)
                    break

    slope_patterns = (
        re.compile(
            r"eğim(?:\s+durumu)?\s*(?:[:=])?\s*(?:≈|~)?\s*%\s*"
            r"(?P<slope>\d+(?:[.,]\d+)?(?:\s*[-~]\s*\d+(?:[.,]\d+)?)?)",
            re.I,
        ),
        re.compile(
            r"(?:≈|~)\s*%\s*(?P<slope>\d+(?:[.,]\d+)?(?:\s*[-~]\s*\d+(?:[.,]\d+)?)?)\s+eğim",
            re.I,
        ),
    )
    for page in source_pages:
        text = _page_content(page)
        for pattern in slope_patterns:
            for match in pattern.finditer(text):
                add("egim", match.group("slope"), page.get("no", 1), 0.94, match)

    direction_pattern = re.compile(
        r"\b(?P<first>[A-ZÇĞİÖŞÜ])\s*-\s*(?P<second>[A-ZÇĞİÖŞÜ])\s+doğrultusunda",
        re.I,
    )
    for page in source_pages:
        text = _page_content(page)
        for match in direction_pattern.finditer(text):
            add(
                "yon",
                f"{match.group('first').upper()}-{match.group('second').upper()}",
                page.get("no", 1),
                0.91,
                match,
            )

    return records, warnings


def _block_count_warning(pages, block_order):
    if len(block_order) <= 1:
        return ""
    for page in pages:
        text = _clean_text(page.get("text", ""))
        match = re.search(
            r"Blok\s+Sayısı\s*/\s*Kat\s+Sayısı.*?(\d+)\s*/\s*(\d+)",
            text,
            flags=re.I,
        )
        if match and int(match.group(1)) != len(block_order):
            return (
                f"Sayfa {page.get('no', 1)} sonuç tablosunda blok sayısı {match.group(1)}, "
                f"yapı tablosunda ise {len(block_order)} blok görünüyor."
            )
    return ""


def geoteknik_sayfalarindan_alanlari_ayikla(pages, dosya_adi="Geoteknik Rapor"):
    """Okunmuş sayfa ve tablolardan proje alanı önerileri oluştur."""
    if not pages:
        return [], ["Raporda okunabilir sayfa bulunamadı."], []

    block_order, block_data, candidates = _table_fields(pages)
    full_text = "\n".join(page.get("text", "") for page in pages)
    if not block_order:
        block_order = _declared_blocks(full_text) or ["Yapı"]
    common = _common_fields(pages)

    if len(block_order) == 1:
        single_structure = _single_structure_fields(pages)
        for key, record in single_structure.items():
            common.setdefault(key, record)

    for block in block_order:
        block_data.setdefault(block, {})
        for key, record in common.items():
            if key not in block_data[block]:
                block_data[block][key] = dict(record)

    warnings = []
    _pressure_cross_checks(pages, block_data, warnings)
    count_warning = _block_count_warning(pages, block_order)
    if count_warning:
        warnings.append(count_warning)

    multi = len(block_order) > 1
    fields = []
    for block in block_order:
        for key in _FIELD_ORDER:
            record = block_data.get(block, {}).get(key)
            if not record or not record.get("value"):
                continue
            if multi and key == "ysinif" and block != block_order[0]:
                continue
            page = record.get("page", 1)
            source_detail = "Yapı Bilgileri tablosu" if candidates and any(
                candidate.get("page") == page for candidate in candidates
            ) else "rapor metni"
            block_field = multi and key != "ysinif"
            fields.append(
                GeoteknikRaporAlani(
                    bolum="bina_blok" if block_field else "bina",
                    anahtar=key,
                    etiket=f"{block} - {_FIELD_LABELS[key]}" if block_field else _FIELD_LABELS[key],
                    deger=record["value"],
                    kaynak=f"{dosya_adi} - Sayfa {page}, {source_detail}",
                    belge_turu="Geoteknik Rapor",
                    guven=float(record.get("confidence", 0.8)),
                    blok_adi=block if block_field else "",
                    alternatifler=tuple(record.get("alternatives", [])),
                    uyari=record.get("warning", ""),
                )
            )
    imar_records, imar_warnings = _imar_saha_alanlari(pages, dosya_adi)
    warnings.extend(imar_warnings)
    for key, record in imar_records.items():
        page = record.get("page", 1)
        section = record.get("section", "Rapor metni")
        target_section = "arazi" if key in {"imar_alani", "egim", "yon"} else "rapor_bilgileri"
        fields.append(
            GeoteknikRaporAlani(
                bolum=target_section,
                anahtar=key,
                etiket=_IMAR_FIELD_LABELS.get(key, key),
                deger=record["value"],
                kaynak=f"{dosya_adi} - Sayfa {page}, {section}",
                belge_turu="Geoteknik Rapor",
                guven=float(record.get("confidence", 0.8)),
                alternatifler=tuple(record.get("alternatives", [])),
                uyari=record.get("warning", ""),
                kanit=record.get("evidence", ""),
            )
        )
    return [asdict(field) for field in fields], warnings, block_order


def geoteknik_raporu_oku(path, task_context=None):
    """PDF veya DOCX geoteknik raporu okuyup onaylanabilir alanlar döndür."""
    source = Path(path).expanduser().resolve()
    if not source.is_file():
        raise GeoteknikRaporOkumaHatasi(f"Geoteknik rapor bulunamadı: {source}")

    suffix = source.suffix.casefold()
    if suffix == ".pdf":
        pages, ocr_used = _pdf_pages(source, task_context=task_context)
    elif suffix == ".docx":
        pages, ocr_used = _docx_pages(source, task_context=task_context)
    else:
        raise GeoteknikRaporOkumaHatasi("Yalnızca PDF ve DOCX geoteknik raporları desteklenir.")

    fields, warnings, blocks = geoteknik_sayfalarindan_alanlari_ayikla(pages, source.name)
    if ocr_used:
        warnings.append("Bazı sayfalar OCR ile okundu; bu satırları kaynak raporla karşılaştırın.")
    if not fields:
        detail = "\n".join(warnings)
        raise GeoteknikRaporOkumaHatasi(
            "Raporda yapı veya imar adası bilgilerine aktarılabilecek alan bulunamadı."
            + (f"\n\n{detail}" if detail else "")
        )

    return {
        "dosya": os.fspath(source),
        "klasor": os.fspath(source.parent),
        "belgeler": [{"ad": source.name, "tur": "Geoteknik Rapor", "yol": os.fspath(source)}],
        "alanlar": fields,
        "uyarilar": warnings,
        "bloklar": blocks,
        "sonuc_basligi": "Geoteknik Rapor Yapı Bilgileri",
        "pencere_basligi": "Geoteknik Rapordan Veri Aktar",
        "aktarim_anahtari": "geoteknik_rapor_aktarimi",
    }


__all__ = [
    "GeoteknikRaporAlani",
    "GeoteknikRaporOkumaHatasi",
    "geoteknik_raporu_oku",
    "geoteknik_sayfalarindan_alanlari_ayikla",
]

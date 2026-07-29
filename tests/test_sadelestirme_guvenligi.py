import datetime
from pathlib import Path
import shutil
from types import SimpleNamespace
from xml.etree import ElementTree
from zipfile import ZipFile

from docx import Document

from ai_motoru import ALLOWED_TAGS, DUZELTME_ETIKET_KURALLARI
from arayuz import RaporRobotuArayuz
from kalite_kontrol import KNOWN_TAGS
from ortam_kontrolu import OPTIONAL_DEPENDENCIES, REQUIRED_DEPENDENCIES
from rapor_etiketleri import DUZELTME_ETIKET_ADLARI
from raporlama import docx_metadata_nortrle, rapor_baglami_olustur
from ui_rapor import RaporSekmesiMixin
from uygulama_yollari import kullanici_veri_dizini, kullanici_yolu
from yardimcilar import atomic_docx_save


def test_eski_roaming_verisi_local_appdata_altina_kopyalanir(tmp_path, monkeypatch):
    roaming = tmp_path / "Roaming"
    local = tmp_path / "Local"
    legacy = roaming / "RaporPro"
    legacy.mkdir(parents=True)
    (legacy / "ayarlar.json").write_text('{"tema": "koyu"}', encoding="utf-8")
    monkeypatch.delenv("RAPORPRO_DATA_DIR", raising=False)
    monkeypatch.setenv("APPDATA", str(roaming))
    monkeypatch.setenv("LOCALAPPDATA", str(local))

    result = kullanici_veri_dizini()

    assert result == local / "RaporPro"
    assert (result / "ayarlar.json").read_text(encoding="utf-8") == '{"tema": "koyu"}'


def test_yarida_kalan_roaming_kopyasi_local_veriyi_gizlemez(tmp_path, monkeypatch):
    roaming = tmp_path / "Roaming"
    local = tmp_path / "Local"
    legacy = roaming / "RaporPro"
    legacy_autosave = legacy / "autosave"
    legacy_autosave.mkdir(parents=True)
    (legacy_autosave / "kurtarma.json").write_text('{"tam": true}', encoding="utf-8")
    monkeypatch.delenv("RAPORPRO_DATA_DIR", raising=False)
    monkeypatch.setenv("APPDATA", str(roaming))
    monkeypatch.setenv("LOCALAPPDATA", str(local))
    real_copytree = shutil.copytree

    def interrupted_copy(_source, target):
        target = Path(target)
        target.mkdir(parents=True)
        (target / "yarim.txt").write_text("eksik", encoding="utf-8")
        raise OSError("kopya kesildi")

    monkeypatch.setattr("uygulama_yollari.shutil.copytree", interrupted_copy)
    assert kullanici_yolu("autosave") == legacy_autosave
    assert not (local / "RaporPro" / "autosave").exists()

    monkeypatch.setattr("uygulama_yollari.shutil.copytree", real_copytree)
    result = kullanici_yolu("autosave")
    assert result == local / "RaporPro" / "autosave"
    assert (result / "kurtarma.json").read_text(encoding="utf-8") == '{"tam": true}'


def test_docx_kisisel_metadata_nortrlenir():
    document = Document()
    document.core_properties.author = "Kisi Adi"
    document.core_properties.last_modified_by = "Baska Kisi"
    document.core_properties.comments = "Gizli not"
    document.core_properties.keywords = "musteri"

    docx_metadata_nortrle(document)

    assert document.core_properties.author == "RaporPro"
    assert document.core_properties.last_modified_by == "RaporPro"
    assert document.core_properties.comments == ""
    assert document.core_properties.keywords == ""


def test_atomic_docx_kaydi_extended_metadata_ve_zamanlari_temizler(tmp_path):
    template = (
        Path(__file__).resolve().parents[1]
        / "sablonlar"
        / "rapor"
        / "varsayilan_rapor_sablonu.docx"
    )
    document = Document(template)
    document.core_properties.created = datetime.datetime(2020, 1, 2)
    document.core_properties.last_printed = datetime.datetime(2021, 2, 3)
    output = tmp_path / "paylasilabilir.docx"

    atomic_docx_save(document, output)

    reopened = Document(output)
    assert reopened.core_properties.created is None
    assert reopened.core_properties.modified is None
    assert reopened.core_properties.last_printed is None
    assert reopened.core_properties.revision == 1
    with ZipFile(output) as package:
        app_xml = ElementTree.fromstring(package.read("docProps/app.xml"))
    ns = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
    assert (app_xml.find(f"{{{ns}}}Company").text or "") == ""
    manager = app_xml.find(f"{{{ns}}}Manager")
    assert manager is None or (manager.text or "") == ""
    assert app_xml.find(f"{{{ns}}}Application").text == "RaporPro"


def test_rapor_baglami_ui_nesnesini_ve_gereksiz_alanlari_tasimaz():
    source = SimpleNamespace(
        veri={"kunye": {"sahibi": "Proje"}},
        word_path="sablon.docx",
        img_yer="harita.png",
        api_anahtari="gizli",
        root=object(),
        set_status=lambda *_args, **_kwargs: None,
    )

    context = rapor_baglami_olustur(source)

    assert context.word_path == "sablon.docx"
    assert context.img_yer == "harita.png"
    assert not hasattr(context, "api_anahtari")
    assert not hasattr(context, "root")
    source.veri["kunye"]["sahibi"] = "Degisti"
    assert context.veri["kunye"]["sahibi"] == "Proje"


def test_dis_ai_onayi_ayni_oturumda_bir_kez_sorulur(monkeypatch):
    calls = []
    monkeypatch.setattr(
        "ui_rapor.messagebox.askyesno",
        lambda *args, **kwargs: calls.append((args, kwargs)) or True,
    )
    context = SimpleNamespace()

    assert RaporSekmesiMixin.dis_ai_veri_aktarim_onayi(
        context, "openai", "rapor metni"
    )
    assert RaporSekmesiMixin.dis_ai_veri_aktarim_onayi(
        context, "openai", "rapor metni"
    )
    assert len(calls) == 1
    assert "OpenAI" in calls[0][0][1]


def test_kural_motoru_dis_ai_onayi_istemez(monkeypatch):
    monkeypatch.setattr(
        "ui_rapor.messagebox.askyesno",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("sorulmamali")),
    )
    assert RaporSekmesiMixin.dis_ai_veri_aktarim_onayi(
        SimpleNamespace(), "kural", "rapor metni"
    )


def test_zorunlu_ve_opsiyonel_bagimliliklar_dogru_sinifta():
    required = {module for module, _package, _purpose in REQUIRED_DEPENDENCIES}
    optional = {module for module, _package, _purpose in OPTIONAL_DEPENDENCIES}
    assert "openpyxl" in required
    assert "win32com" in optional
    assert "requests" in optional
    assert required.isdisjoint(optional)


def test_duzeltme_etiketleri_kalite_kaydinda_da_bilinir():
    assert set(DUZELTME_ETIKET_ADLARI).issubset(KNOWN_TAGS)


def test_ai_kurallari_ortak_duzeltme_etiketlerini_eksiksiz_kapsar():
    rule_tags = {item["tag"] for item in DUZELTME_ETIKET_KURALLARI}
    assert rule_tags == set(DUZELTME_ETIKET_ADLARI)
    assert set(ALLOWED_TAGS) == set(DUZELTME_ETIKET_ADLARI)


def test_arayuz_mro_icinde_ayni_callable_birden_fazla_tanimlanmiyor():
    owners = {}
    duplicates = set()
    for cls in RaporRobotuArayuz.__mro__:
        for name, value in cls.__dict__.items():
            if name.startswith("__") or not callable(value):
                continue
            if name in owners:
                duplicates.add(name)
            else:
                owners[name] = cls
    assert not duplicates

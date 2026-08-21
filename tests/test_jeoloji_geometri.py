# Dosya: RaporPro/tests/test_jeoloji_geometri.py
from pathlib import Path
import sqlite3
import tempfile
import zipfile

from docx import Document
from docx.shared import Inches
from PIL import Image

import jeoloji_geometri
from jeoloji_geometri import (
    HARITA_MOD_SECILI,
    HARITA_MOD_TUMU,
    HARITA_MOD_YAKINDAKILER,
    aday_geometrisini_sec,
    adayi_yerel_geometriyle_eslestir,
    eksik_geometrileri_tkgmden_tamamla,
    geometri_harita_poligonlari,
    harita_fit_bounds,
    harita_gorunum_modeli,
    harita_kayitlarini_ayir,
    geometri_katalogu_olustur,
    kml_geometrilerini_oku,
    kml_onbellegini_temizle,
    koordinat_poligon_uyari_metni,
    koordinat_poligon_uyusmazligi,
)
from jeoloji_kutuphanesi import (
    JeolojiKutuphane,
    eksik_kutuphane_geometrilerini_tamamla,
    jeoloji_adaylarini_tara,
)
from ui_jeoloji_adaylari import (
    JeolojiAdayPenceresi,
    aday_geometri_durum_metni,
    aday_harita_geometri_kayitlari,
    aday_secilebilir_mi,
)
from ui_jeoloji_kutuphanesi import (
    JeolojiKutuphanePenceresi,
    duplicate_adayi_hazirla,
    kayit_geometri_metni,
)


def _word_report(path, *, header=None, lat=40.1000, lon=26.4000):
    image_path = path.parent / f"{path.stem}.png"
    Image.new("RGB", (30, 30), (45, 110, 175)).save(image_path)
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = header or (
        "İmar Bilgileri: Çanakkale İli, Merkez İlçesi, Işıklar Köyü, "
        "32N-4D Pafta, 12 Ada, 7 Parsel"
    )
    doc.sections[0].footer.paragraphs[0].text = "KAYNAK FOOTER"
    doc.add_heading("1. GİRİŞ", level=1)
    doc.add_paragraph(f"WGS84 Enlem: {lat:.6f} Boylam: {lon:.6f}")
    doc.add_heading("2. JEOLOJİ", level=1)
    doc.add_paragraph("Jeoloji bölüm metni")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Birim"
    table.cell(0, 1).text = "Tmçk"
    doc.add_picture(str(image_path), width=Inches(0.5))
    doc.add_heading("3. ARAZİ ÇALIŞMALARI", level=1)
    doc.add_paragraph("Arazi bölümü korunmalı")
    doc.save(path)


def _coordinates(points):
    return " ".join(f"{lon},{lat},0" for lon, lat in points)


def _polygon_xml(outer, inner=None):
    inner_xml = ""
    if inner:
        inner_xml = (
            "<innerBoundaryIs><LinearRing><coordinates>"
            + _coordinates(inner)
            + "</coordinates></LinearRing></innerBoundaryIs>"
        )
    return (
        "<Polygon><outerBoundaryIs><LinearRing><coordinates>"
        + _coordinates(outer)
        + "</coordinates></LinearRing></outerBoundaryIs>"
        + inner_xml
        + "</Polygon>"
    )


def _kml_text(placemarks):
    body = []
    for name, description, geometry_xml in placemarks:
        body.append(
            f"<Placemark><name>{name}</name><description>{description}</description>"
            f"{geometry_xml}</Placemark>"
        )
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<kml xmlns="http://www.opengis.net/kml/2.2"><Document>'
        + "".join(body)
        + "</Document></kml>"
    )


def _write_kml(path, placemarks):
    path.write_text(_kml_text(placemarks), encoding="utf-8")
    return path


def _square(lon=26.4, lat=40.1, size=0.002):
    return [
        (lon - size, lat - size),
        (lon + size, lat - size),
        (lon + size, lat + size),
        (lon - size, lat + size),
        (lon - size, lat - size),
    ]


def _map_record(record_id, lon, lat, *, geometry=True, record_coordinate=None):
    record = {
        "id": record_id,
        "ada": str(100 + record_id),
        "parsel": str(record_id),
        "lat": (record_coordinate or (lat, lon))[0],
        "lon": (record_coordinate or (lat, lon))[1],
    }
    if geometry:
        ring = _square(lon, lat, size=0.001)
        record.update(
            {
                "geometry_hash": f"geometry-{record_id}",
                "geometry_metadata": {
                    "polygons": [[ring]],
                    "geometry_hash": f"geometry-{record_id}",
                    "centroid": [lat, lon],
                    "bounds": [lat - 0.001, lon - 0.001, lat + 0.001, lon + 0.001],
                },
            }
        )
    return record


def _candidate(path, **updates):
    value = {
        "source_path": str(path),
        "eligible": True,
        "selected": True,
        "il": "Çanakkale",
        "ilce": "Merkez",
        "mahalle": "Işıklar",
        "pafta": "32N-4D",
        "ada": "12",
        "parsel": "7",
        "lat": 40.1,
        "lon": 26.4,
        "coordinate_source": "rapor_wgs84",
        "warnings": [],
        "analysis": {"warnings": []},
    }
    value.update(updates)
    return value


def test_yerel_kml_tam_eslesme_yanlis_parseli_atlar_ve_mahalle_uyusmazligini_baglamaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        report = root / "rapor.docx"
        report.touch()
        right = _write_kml(
            root / "dogru.kml",
            [("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel", "", _polygon_xml(_square()))],
        )
        _write_kml(
            root / "yanlis.kml",
            [("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 99 Ada, 8 Parsel", "", _polygon_xml(_square(26.41)))],
        )
        mismatch = _write_kml(
            root / "mahalle_uyusmaz.kml",
            [("Çanakkale İli, Merkez İlçesi, Kepez Mahallesi, 12 Ada, 7 Parsel", "", _polygon_xml(_square()))],
        )
        catalog = geometri_katalogu_olustur([root])
        candidate = adayi_yerel_geometriyle_eslestir(_candidate(report), catalog["geometries"])
        assert candidate["geometry_status"] == "local_exact"
        assert Path(candidate["geometry"]["source_path"]) == right

        no_coordinate = adayi_yerel_geometriyle_eslestir(
            _candidate(report, lat=None, lon=None, coordinate_source=""),
            kml_geometrilerini_oku(right),
        )
        assert no_coordinate["coordinate_source"] == "kml_centroid"
        assert round(no_coordinate["lat"], 4) == 40.1
        assert round(no_coordinate["lon"], 4) == 26.4

        mismatch_only = geometri_katalogu_olustur([mismatch])
        rejected = adayi_yerel_geometriyle_eslestir(_candidate(report), mismatch_only["geometries"])
        assert rejected["geometry"] is None
        assert rejected["geometry_status"] == "location_mismatch"
        assert any("uyuşmadı" in warning for warning in rejected["warnings"])


def test_kimliksiz_tek_poligon_yakinlikla_uyarili_eslesir_uzak_poligon_baglanmaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        report = root / "rapor.docx"
        report.touch()
        near = _write_kml(root / "sinir.kml", [("Çalışma Alanı", "", _polygon_xml(_square()))])
        near_geometry = kml_geometrilerini_oku(near)
        candidate = adayi_yerel_geometriyle_eslestir(_candidate(report), near_geometry)
        assert candidate["geometry_status"] == "local_proximity"
        assert any("yakınlığıyla" in warning for warning in candidate["warnings"])

        far = _write_kml(root / "uzak_sinir.kml", [("Çalışma Alanı", "", _polygon_xml(_square(29.0, 41.0)))])
        far_candidate = adayi_yerel_geometriyle_eslestir(_candidate(report), kml_geometrilerini_oku(far))
        assert far_candidate["geometry"] is None
        assert far_candidate["geometry_status"] == "missing"


def test_ayni_kimlikli_farkli_sinirlar_belirsiz_kalir_ve_kullanici_secebilir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        report = root / "rapor.docx"
        report.touch()
        label = "Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel"
        first = _write_kml(root / "bir.kml", [(label, "", _polygon_xml(_square()))])
        second = _write_kml(root / "iki.kml", [(label, "", _polygon_xml(_square(26.405)))])
        options = kml_geometrilerini_oku(first) + kml_geometrilerini_oku(second)
        candidate = adayi_yerel_geometriyle_eslestir(_candidate(report), options)
        assert candidate["geometry_status"] == "ambiguous"
        assert candidate["geometry"] is None
        assert len(candidate["geometry_options"]) == 2
        aday_geometrisini_sec(candidate, candidate["geometry_options"][1], status="local_user_selected")
        assert candidate["geometry_status"] == "local_user_selected"
        assert Path(candidate["geometry"]["source_path"]) == second


def test_multipolygon_inner_ring_ve_kmz_ayristirilir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        outer = _square()
        inner = _square(size=0.0005)
        second = _square(26.41, 40.11, size=0.001)
        multi = "<MultiGeometry>" + _polygon_xml(outer, inner) + _polygon_xml(second) + "</MultiGeometry>"
        kml_text = _kml_text([("12 Ada 7 Parsel", "", multi)])
        kml = root / "multi.kml"
        kml.write_text(kml_text, encoding="utf-8")
        kmz = root / "multi.kmz"
        with zipfile.ZipFile(kmz, "w", zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("doc.kml", kml_text)

        parsed_kml = kml_geometrilerini_oku(kml)[0]
        parsed_kmz = kml_geometrilerini_oku(kmz)[0]
        assert len(parsed_kml["polygons"]) == 2
        assert len(parsed_kml["polygons"][0]) == 2
        assert parsed_kml["geometry_hash"] == parsed_kmz["geometry_hash"]
        map_polygons = geometri_harita_poligonlari(parsed_kml)
        assert len(map_polygons) == 2
        assert len(map_polygons[0]) == 2


def test_geometrili_klasor_taramasi_db_ve_cache_yazmaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        reports = root / "raporlar"
        reports.mkdir()
        report = reports / "rapor.docx"
        _word_report(report)
        _write_kml(
            reports / "parsel.kml",
            [("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel", "", _polygon_xml(_square()))],
        )
        store = JeolojiKutuphane(base_dir=root / "appdata")

        candidates = jeoloji_adaylarini_tara([reports], recursive=True)

        assert store.count() == 0
        assert not list(store.cache_dir.glob("*.docx"))
        assert not list(store.geometry_dir.glob("*.kml"))
        assert len(candidates) == 1
        assert candidates[0]["geometry_status"] == "local_exact"
        assert candidates[0]["geometry"]


def test_fake_tkgm_ayni_parseli_tek_sorgular_ve_hata_adayi_engellemez():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        calls = []

        def fake_resolver(kunye, output_dir):
            calls.append((kunye["ada"], kunye["parsel"]))
            if kunye["parsel"] == "999":
                raise RuntimeError("ağ kapalı")
            path = Path(output_dir) / "tkgm.kml"
            _write_kml(
                path,
                [(f"{kunye['ada']} Ada, {kunye['parsel']} Parsel", "", _polygon_xml(_square()))],
            )
            return {"path": str(path), "label": "TKGM 12/7"}

        first = _candidate(root / "a.docx", geometry_status="missing")
        second = _candidate(root / "b.docx", geometry_status="missing")
        failing = _candidate(root / "c.docx", parsel="999", geometry_status="missing")
        result = eksik_geometrileri_tkgmden_tamamla([first, second, failing], fake_resolver)

        assert result == {"completed": 2, "failed": 1, "skipped": 0, "queries": 2}
        assert calls.count(("12", "7")) == 1
        assert first["geometry_status"] == second["geometry_status"] == "tkgm"
        assert first["geometry_hash"] == second["geometry_hash"]
        assert failing["geometry_status"] == "tkgm_error"
        assert aday_secilebilir_mi(failing)


def test_tarama_fake_tkgm_resolveri_paylasir_ama_kutuphaneye_yazmaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        reports = root / "reports"
        reports.mkdir()
        _word_report(reports / "bir.docx")
        _word_report(reports / "iki.docx")
        calls = []

        def fake_resolver(kunye, output_dir):
            calls.append((kunye["ada"], kunye["parsel"]))
            path = Path(output_dir) / "parsel.kml"
            _write_kml(path, [("12 Ada 7 Parsel", "", _polygon_xml(_square()))])
            return {"path": str(path), "label": "TKGM 12/7"}

        store = JeolojiKutuphane(base_dir=root / "appdata")
        candidates = jeoloji_adaylarini_tara(
            [reports],
            geometry_resolver=fake_resolver,
            complete_missing=True,
        )
        assert calls == [("12", "7")]
        assert [item["geometry_status"] for item in candidates] == ["tkgm", "tkgm"]
        assert store.count() == 0
        assert not list(store.cache_dir.glob("*.docx"))
        assert not list(store.geometry_dir.glob("*.kml"))


def test_secili_aday_yalniz_bolum_docx_ve_normalize_kml_cache_yazar_dedup_korur():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        sources = root / "kaynaklar"
        sources.mkdir()
        source = sources / "rapor.docx"
        source_two = sources / "rapor_ikinci.docx"
        kml = sources / "parseller.kml"
        _word_report(source)
        _word_report(source_two)
        second_doc = Document(source_two)
        second_doc.add_paragraph("İkinci kaynak rapora özgü metin")
        second_doc.save(source_two)
        _write_kml(
            kml,
            [
                ("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel", "", _polygon_xml(_square())),
                ("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 99 Ada, 2 Parsel", "", _polygon_xml(_square(26.5))),
            ],
        )
        source_before, source_two_before, kml_before = source.read_bytes(), source_two.read_bytes(), kml.read_bytes()
        candidates = jeoloji_adaylarini_tara([sources])
        candidate = next(item for item in candidates if item["original_filename"] == source.name)
        candidate_two = next(item for item in candidates if item["original_filename"] == source_two.name)
        store = JeolojiKutuphane(base_dir=root / "data")

        first = store.import_candidate(candidate)
        second = store.import_candidate(candidate)
        third = store.import_candidate(candidate_two)

        assert first["duplicate"] is False
        assert second["duplicate"] is True
        assert third["duplicate"] is False
        assert store.count() == 2
        record = second["record"]
        second_record = third["record"]
        assert Path(record["kml_path"]).is_file()
        assert record["geometry_source"] == "local_kml"
        assert record["geometry_hash"]
        assert record["kml_path"] == second_record["kml_path"]
        assert len(list(store.geometry_dir.glob("*.kml"))) == 1
        cached_geometry = kml_geometrilerini_oku(record["kml_path"])
        assert len(cached_geometry) == 1
        assert cached_geometry[0]["polygon_count"] == 1
        section = Document(record["cache_path"])
        text = "\n".join(paragraph.text for paragraph in section.paragraphs)
        assert "2. JEOLOJİ" in text
        assert "1. GİRİŞ" not in text
        assert "3. ARAZİ ÇALIŞMALARI" not in text
        assert len(section.tables) == 1
        assert len(section.inline_shapes) == 1
        assert source.read_bytes() == source_before
        assert source_two.read_bytes() == source_two_before
        assert kml.read_bytes() == kml_before


def test_eski_db_geometri_kolonlarina_migre_edilir_ve_eski_kayit_okunur():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        db = root / "legacy.sqlite3"
        with sqlite3.connect(db) as connection:
            connection.execute(
                """
                CREATE TABLE geology_sources (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    source_hash TEXT NOT NULL UNIQUE,
                    original_filename TEXT NOT NULL DEFAULT '', original_path TEXT NOT NULL DEFAULT '',
                    added_at TEXT NOT NULL, updated_at TEXT NOT NULL, metadata_json TEXT NOT NULL DEFAULT '{}',
                    il TEXT NOT NULL DEFAULT '', ilce TEXT NOT NULL DEFAULT '', mahalle TEXT NOT NULL DEFAULT '',
                    pafta TEXT NOT NULL DEFAULT '', ada TEXT NOT NULL DEFAULT '', parsel TEXT NOT NULL DEFAULT '',
                    lat REAL, lon REAL, coordinate_source TEXT NOT NULL DEFAULT '',
                    start_heading TEXT NOT NULL DEFAULT '', end_heading TEXT NOT NULL DEFAULT '',
                    start_index INTEGER NOT NULL DEFAULT -1, end_index INTEGER NOT NULL DEFAULT -1,
                    heading_level INTEGER NOT NULL DEFAULT 1, paragraph_count INTEGER NOT NULL DEFAULT 0,
                    table_count INTEGER NOT NULL DEFAULT 0, image_count INTEGER NOT NULL DEFAULT 0,
                    warning_json TEXT NOT NULL DEFAULT '[]', cache_path TEXT NOT NULL DEFAULT '',
                    file_size INTEGER NOT NULL DEFAULT 0
                )
                """
            )
            connection.execute(
                """
                INSERT INTO geology_sources (
                    source_hash, original_filename, added_at, updated_at, il, ilce, mahalle, ada, parsel
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    "a" * 64, "legacy.docx", "2026-01-01", "2026-01-01",
                    "Çanakkale", "Merkez", "Işıklar", "0", "7",
                ),
            )
        connection.close()
        store = JeolojiKutuphane(db_path=db, cache_dir=root / "cache", geometry_dir=root / "geometry")
        record = store.list_records()[0]
        assert record["original_filename"] == "legacy.docx"
        assert record["geometry_metadata"] == {}
        assert record["kml_path"] == ""
        with store._connection() as connection:
            columns = {row["name"] for row in connection.execute("PRAGMA table_info(geology_sources)")}
        assert {"kml_path", "geometry_source", "geometry_hash", "geometry_status", "geometry_metadata_json"} <= columns

        def fake_resolver(_kunye, output_dir):
            path = Path(output_dir) / "legacy.kml"
            _write_kml(path, [("0 Ada 7 Parsel", "", _polygon_xml(_square()))])
            return path

        result = eksik_kutuphane_geometrilerini_tamamla(store, resolver=fake_resolver)
        reopened = JeolojiKutuphane(db_path=db, cache_dir=root / "cache", geometry_dir=root / "geometry").get(record["id"])
        assert result["updated"] == 1
        assert reopened["geometry_metadata"]["polygons"]
        assert Path(reopened["kml_path"]).is_file()


def test_ui_geometri_saf_yardimcilari_durum_ve_secimi_yansitir():
    candidate = _candidate("rapor.docx", geometry_status="ambiguous")
    assert "Belirsiz" in aday_geometri_durum_metni(candidate)
    geometry = {
        "polygons": [[[_square()]]][0],
        "geometry_hash": "hash",
        "centroid": [40.1, 26.4],
        "bounds": [40.0, 26.3, 40.2, 26.5],
        "source_type": "local_kml",
        "source_path": "parsel.kml",
        "placemark_name": "12/7",
    }
    aday_geometrisini_sec(candidate, geometry, status="local_user_selected")
    assert "kullanıcı" in aday_geometri_durum_metni(candidate)
    record = {"geometry_source": "local_kml", "geometry_status": "local_user_selected"}
    assert kayit_geometri_metni(record).startswith("Yerel KML")
    map_records = aday_harita_geometri_kayitlari([{}, candidate])
    assert [(index, ambiguous) for index, _geometry, ambiguous in map_records] == [(1, False)]


def test_duplicate_point_only_kayit_yerel_kml_ile_ayni_id_uzerinde_guncellenir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        report = root / "rapor.docx"
        _word_report(report)
        store = JeolojiKutuphane(base_dir=root / "data")
        initial = jeoloji_adaylarini_tara([report], recursive=False)[0]
        first = store.import_candidate(initial)["record"]
        assert first["kml_path"] == ""

        _write_kml(
            root / "parsel.kml",
            [("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel", "", _polygon_xml(_square()))],
        )
        rescanned = jeoloji_adaylarini_tara([report], recursive=False)[0]
        duplicate_adayi_hazirla(rescanned, first)
        assert rescanned["duplicate_geometry_update"] is True
        assert rescanned["selected"] is True

        result = store.import_candidate(rescanned)
        assert result["duplicate"] is True
        assert result["record"]["id"] == first["id"]
        assert result["record"]["geometry_metadata"]["polygons"]
        assert Path(result["record"]["kml_path"]).is_file()
        assert store.count() == 1


def test_point_only_kutuphaneyi_backfill_kalici_yazar_ve_reopen_kml_parse_etmez(monkeypatch):
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        report = root / "eski.docx"
        _word_report(report)
        source_bytes = report.read_bytes()
        store = JeolojiKutuphane(base_dir=root / "data")
        record = store.import_candidate(jeoloji_adaylarini_tara([report], recursive=False)[0])["record"]
        _write_kml(
            root / "12_7.kml",
            [("Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 12 Ada, 7 Parsel", "", _polygon_xml(_square()))],
        )

        def network_forbidden(_kunye, _output_dir):
            raise AssertionError("Yerel sınır varken TKGM çağrılmamalı")

        result = eksik_kutuphane_geometrilerini_tamamla(store, resolver=network_forbidden)
        assert result["updated"] == result["local"] == 1
        assert result["queries"] == 0
        assert report.read_bytes() == source_bytes

        monkeypatch.setattr(jeoloji_geometri, "_kml_bytes_oku", lambda _path: (_ for _ in ()).throw(AssertionError("reopen parse")))
        reopened = JeolojiKutuphane(base_dir=root / "data").get(record["id"])
        geometry_records, marker_records = harita_kayitlarini_ayir([reopened], cache={})
        assert len(geometry_records) == 1
        assert marker_records == []
        assert Path(reopened["kml_path"]).is_file()


def test_ayni_parselde_iki_point_only_kayit_tkgmyi_bir_kez_sorgular_ve_cachei_paylasir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        first_path = root / "bir.docx"
        second_path = root / "iki.docx"
        _word_report(first_path)
        _word_report(second_path)
        second_doc = Document(second_path)
        second_doc.paragraphs[1].add_run(" ikinci")
        second_doc.save(second_path)
        store = JeolojiKutuphane(base_dir=root / "data")
        store.import_candidate(jeoloji_adaylarini_tara([first_path], recursive=False)[0])
        store.import_candidate(jeoloji_adaylarini_tara([second_path], recursive=False)[0])
        calls = []

        def fake_resolver(kunye, output_dir):
            calls.append((kunye["ada"], kunye["parsel"]))
            path = Path(output_dir) / "tkgm.kml"
            _write_kml(path, [("12 Ada 7 Parsel", "", _polygon_xml(_square()))])
            return path

        result = eksik_kutuphane_geometrilerini_tamamla(store, resolver=fake_resolver)
        records = store.list_records()
        assert result["queries"] == 1
        assert result["tkgm"] == 2
        assert calls == [("12", "7")]
        assert len({record["kml_path"] for record in records}) == 1
        assert len(list(store.geometry_dir.glob("*.kml"))) == 1


def test_kml_parse_ve_harita_poligon_donusumu_onbelleklenir(monkeypatch):
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        path = _write_kml(root / "parsel.kml", [("12 Ada 7 Parsel", "", _polygon_xml(_square()))])
        kml_onbellegini_temizle()
        original = jeoloji_geometri._kml_bytes_oku
        calls = []

        def counted(source):
            calls.append(str(source))
            return original(source)

        monkeypatch.setattr(jeoloji_geometri, "_kml_bytes_oku", counted)
        first = kml_geometrilerini_oku(path)[0]
        second = kml_geometrilerini_oku(path)[0]
        assert len(calls) == 1
        cache = {}
        first_polygons = geometri_harita_poligonlari(first, cache=cache)
        second_polygons = geometri_harita_poligonlari(second, cache=cache)
        assert first_polygons is second_polygons


def test_tkgm_acikca_farkli_parsel_dondururse_sinir_kabul_edilmez():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)

        def wrong_resolver(_kunye, output_dir):
            path = Path(output_dir) / "yanlis.kml"
            _write_kml(path, [("99 Ada 8 Parsel", "", _polygon_xml(_square()))])
            return path

        candidate = _candidate(root / "rapor.docx", geometry_status="missing")
        result = eksik_geometrileri_tkgmden_tamamla([candidate], wrong_resolver)
        assert result["completed"] == 0
        assert result["failed"] == 1
        assert candidate.get("geometry") is None


def test_harita_gorunum_modlari_secili_yakin_ve_tum_kayitlari_ayirir():
    selected = _map_record(1, 26.4000, 40.1000)
    near_polygon = _map_record(2, 26.4050, 40.1040)
    near_fallback = _map_record(3, 26.4020, 40.1080, geometry=False)
    far_polygon = _map_record(4, 27.0000, 41.0000)
    records = [selected, near_polygon, near_fallback, far_polygon]

    selected_model = harita_gorunum_modeli(records, 1, mode=HARITA_MOD_SECILI)
    assert [item["key"] for item in selected_model["items"]] == [1]
    assert selected_model["geometry_count"] == 1
    assert selected_model["fallback_count"] == 0

    nearby_model = harita_gorunum_modeli(records, 1, mode=HARITA_MOD_YAKINDAKILER, radius_km=2)
    assert {item["key"] for item in nearby_model["items"]} == {1, 2, 3}
    assert nearby_model["fallback_count"] == 1

    all_model = harita_gorunum_modeli(records, 1, mode=HARITA_MOD_TUMU)
    assert {item["key"] for item in all_model["items"]} == {1, 2, 3, 4}


def test_secili_mod_secim_yokken_tum_kutuphaneyi_fit_etmez_ama_tumu_modu_gosterir():
    records = [_map_record(1, 26.4, 40.1), _map_record(2, 27.0, 41.0)]
    selected_model = harita_gorunum_modeli(records, None, mode=HARITA_MOD_SECILI)
    assert selected_model["items"] == []
    assert selected_model["bounds"] is None
    assert harita_fit_bounds(selected_model) is None

    all_model = harita_gorunum_modeli(records, None, mode=HARITA_MOD_TUMU)
    assert len(all_model["items"]) == 2
    assert harita_fit_bounds(all_model) is not None


def test_yanlis_record_koordinati_secili_poligon_fitini_ve_yakinligi_bozmaz():
    selected = _map_record(
        1,
        26.366321,
        40.087125,
        record_coordinate=(40.128927, 26.255947),
    )
    near = _map_record(2, 26.3700, 40.0900)
    wrong_coordinate_near = _map_record(3, 26.2559, 40.1289)
    model = harita_gorunum_modeli(
        [selected, near, wrong_coordinate_near],
        1,
        mode=HARITA_MOD_SECILI,
    )
    assert model["items"][0]["center"] == (40.087125, 26.366321)
    assert model["bounds"] == tuple(selected["geometry_metadata"]["bounds"])
    fitted = harita_fit_bounds(model)
    assert fitted[2] < 40.10
    assert fitted[1] > 26.35

    nearby = harita_gorunum_modeli(
        [selected, near, wrong_coordinate_near],
        1,
        mode=HARITA_MOD_YAKINDAKILER,
        radius_km=2,
    )
    assert {item["key"] for item in nearby["items"]} == {1, 2}


def test_fallback_secili_kayit_icin_nokta_bounds_ve_yakin_fit_uretir():
    fallback = _map_record(5, 26.4, 40.1, geometry=False)
    model = harita_gorunum_modeli([fallback], 5, mode=HARITA_MOD_SECILI)
    assert model["items"][0]["kind"] == "fallback"
    assert model["fallback_count"] == 1
    fitted = harita_fit_bounds(model)
    assert fitted[0] < 40.1 < fitted[2]
    assert fitted[1] < 26.4 < fitted[3]


def test_koordinat_poligon_uyusmazligi_250_m_esigiyle_uyari_uretir_veriyi_degistirmez():
    close = _map_record(1, 26.4, 40.1, record_coordinate=(40.1010, 26.4))
    far = _map_record(2, 26.4, 40.1, record_coordinate=(40.1030, 26.4))
    original = (far["lat"], far["lon"])
    assert koordinat_poligon_uyusmazligi(close)["mismatch"] is False
    check = koordinat_poligon_uyusmazligi(far)
    assert check["mismatch"] is True
    assert check["distance_km"] > 0.25
    assert "haritada poligon esas alındı" in koordinat_poligon_uyari_metni(far)
    assert (far["lat"], far["lon"]) == original


class _FakeDrawing:
    def delete(self):
        return None


class _FakeMap:
    def __init__(self):
        self.marker_count = 0
        self.polygon_count = 0
        self.fit_calls = []
        self.polygon_options = []

    def set_polygon(self, _points, **kwargs):
        self.polygon_count += 1
        self.polygon_options.append(kwargs)
        return _FakeDrawing()

    def set_path(self, _points, **_kwargs):
        return _FakeDrawing()

    def set_marker(self, *_args, **_kwargs):
        self.marker_count += 1
        return _FakeDrawing()

    def fit_bounding_box(self, *_args, **_kwargs):
        self.fit_calls.append(_args)


class _FakeStatus:
    def set(self, _value):
        return None


class _FakeAliveWindow:
    def winfo_exists(self):
        return True


def test_ana_harita_poligonlu_kayit_icin_centroid_markeri_cizmez():
    geometry = {
        "polygons": [[_square()]],
        "geometry_hash": "geometry-hash",
        "centroid": [40.1, 26.4],
    }
    record = {
        "id": 1,
        "ada": "12",
        "parsel": "7",
        "lat": 40.1,
        "lon": 26.4,
        "geometry_hash": "geometry-hash",
        "geometry_metadata": geometry,
    }
    window = object.__new__(JeolojiKutuphanePenceresi)
    window._closing = False
    window.win = _FakeAliveWindow()
    window.owner = type("Owner", (), {"veri": {}})()
    window.records = [record]
    window.selected_id = 1
    window._map_widget = _FakeMap()
    window._map_canvas = None
    window._map_markers = []
    window._map_polygons = []
    window._map_paths = []
    window._map_project_marker = None
    window._map_geometry_cache = {}
    window._map_status_var = _FakeStatus()
    window._harita_ciz(fit=False)
    assert window._map_widget.polygon_count == 1
    assert window._map_widget.marker_count == 0


def test_ana_harita_secili_modda_yalniz_poligon_boundsuna_fit_eder():
    selected = _map_record(
        1,
        26.366321,
        40.087125,
        record_coordinate=(40.128927, 26.255947),
    )
    distant = _map_record(2, 29.0, 42.0)
    window = object.__new__(JeolojiKutuphanePenceresi)
    window._closing = False
    window.win = _FakeAliveWindow()
    window.owner = type("Owner", (), {"veri": {}})()
    window.records = [selected, distant]
    window.selected_id = 1
    window.filter_vars = {}
    window._map_widget = _FakeMap()
    window._map_canvas = None
    window._map_markers = []
    window._map_polygons = []
    window._map_paths = []
    window._map_project_marker = None
    window._map_geometry_cache = {}
    window._fallback_icon = None
    window._map_status_var = _FakeStatus()
    window._harita_ciz(fit=True)
    assert window._map_widget.polygon_count == 1
    assert window._map_widget.marker_count == 0
    (top_left, bottom_right), = window._map_widget.fit_calls
    assert top_left[0] < 40.10
    assert bottom_right[1] > 26.35
    assert window._map_widget.polygon_options[0]["fill_color"]


def test_aday_haritasi_poligonlu_kayit_icin_centroid_markeri_cizmez():
    class FakeTree:
        def selection(self):
            return ("0",)

        def focus(self):
            return "0"

    geometry = {
        "polygons": [[_square()]],
        "geometry_hash": "candidate-geometry-hash",
        "centroid": [40.1, 26.4],
    }
    dialog = object.__new__(JeolojiAdayPenceresi)
    dialog._closing = False
    dialog.win = _FakeAliveWindow()
    dialog.tree = FakeTree()
    dialog.candidates = [_candidate("rapor.docx", geometry=geometry, geometry_hash="candidate-geometry-hash")]
    dialog._map_widget = _FakeMap()
    dialog._map_canvas = None
    dialog._map_drawings = []
    dialog._map_markers = []
    dialog._map_geometry_cache = {}
    dialog.map_status_var = _FakeStatus()
    dialog._harita_ciz(fit=False)
    assert dialog._map_widget.polygon_count == 1
    assert dialog._map_widget.marker_count == 0


def test_harita_lazy_init_anlamli_secime_kadar_widget_kurmaz_ve_callback_iptal_edilir():
    class FakeVar:
        def __init__(self, value):
            self.value = value

        def get(self):
            return self.value

    class DeferredWindow(_FakeAliveWindow):
        def __init__(self):
            self.scheduled = []
            self.cancelled = []

        def after(self, delay, callback):
            callback_id = f"after#{len(self.scheduled) + 1}"
            self.scheduled.append((callback_id, delay, callback))
            return callback_id

        def after_cancel(self, callback_id):
            self.cancelled.append(callback_id)

    window = object.__new__(JeolojiKutuphanePenceresi)
    window.win = DeferredWindow()
    window._closing = False
    window._map_ready = False
    window._map_initializing = False
    window._map_init_after_id = None
    window._map_mode_var = FakeVar(HARITA_MOD_SECILI)
    window._map_status_var = _FakeStatus()
    window.selected_id = None
    window._harita_yuklemeyi_zamanla()
    assert window.win.scheduled == []

    window.selected_id = 7
    window._harita_yuklemeyi_zamanla()
    assert not hasattr(window, "_map_widget") or window._map_widget is None
    assert window._map_init_after_id == "after#1"
    assert len(window.win.scheduled) == 1
    window._iptal_harita_init_callback()
    assert window.win.cancelled == ["after#1"]
    assert window._map_initializing is False


def test_kutuphane_tree_secimi_ayni_satiri_tekrar_secmez():
    class FakeTree:
        def __init__(self):
            self.selection_set_calls = []

        def selection(self):
            return ("1",)

        def selection_set(self, iid):
            self.selection_set_calls.append(iid)

        def focus(self, _iid):
            return None

        def see(self, _iid):
            return None

    window = object.__new__(JeolojiKutuphanePenceresi)
    window.tree = FakeTree()
    window.record_map = {1: {"id": 1}}
    window._all_records = []
    window.selected_id = None
    window._map_ready = False
    window._onizleme_goster = lambda _record: None
    window._harita_gorunumu_yenile = lambda **_kwargs: None

    window._kayit_sec(1)

    assert window.selected_id == 1
    assert window.tree.selection_set_calls == []


def test_kutuphane_tree_eventi_mevcut_secimi_yeniden_islemez():
    class FakeTree:
        def selection(self):
            return ("1",)

    window = object.__new__(JeolojiKutuphanePenceresi)
    window.tree = FakeTree()
    window.selected_id = 1
    calls = []
    window._kayit_sec = calls.append

    window._liste_secildi()

    assert calls == []


def test_aday_haritasi_lazy_init_secili_satirdan_sonra_zamanlanir():
    class FakeTree:
        def selection(self):
            return ("0",)

        def focus(self):
            return "0"

    class DeferredWindow(_FakeAliveWindow):
        def __init__(self):
            self.scheduled = []
            self.cancelled = []

        def after(self, delay, callback):
            self.scheduled.append((delay, callback))
            return "after#candidate-map"

        def after_cancel(self, callback_id):
            self.cancelled.append(callback_id)

    dialog = object.__new__(JeolojiAdayPenceresi)
    dialog.win = DeferredWindow()
    dialog.tree = FakeTree()
    dialog.candidates = [_candidate("rapor.docx")]
    dialog._closing = False
    dialog._map_ready = False
    dialog._map_initializing = False
    dialog._map_init_after_id = None
    dialog.map_status_var = _FakeStatus()
    dialog._harita_yuklemeyi_zamanla()
    assert dialog._map_init_after_id == "after#candidate-map"
    assert len(dialog.win.scheduled) == 1
    dialog._iptal_harita_init_callback()
    assert dialog.win.cancelled == ["after#candidate-map"]


def test_aday_penceresi_kapanista_bekleyen_harita_callbackini_iptal_eder():
    class FakeWindow(_FakeAliveWindow):
        def __init__(self):
            self.cancelled = []

        def after_cancel(self, callback_id):
            self.cancelled.append(callback_id)

    dialog = object.__new__(JeolojiAdayPenceresi)
    dialog.win = FakeWindow()
    dialog._closing = False
    dialog.busy = True
    dialog._map_after_id = "after#map"
    dialog._map_fit_requested = True
    dialog._map_widget = None
    dialog._map_canvas = None
    dialog._map_drawings = []
    dialog._map_markers = []
    dialog._map_geometry_cache = {("hash", "a"): []}
    dialog.ebeveyn_kapaniyor()
    assert dialog._closing is True
    assert dialog.win.cancelled == ["after#map"]
    assert dialog._map_geometry_cache == {}

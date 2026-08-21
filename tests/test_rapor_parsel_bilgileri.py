# Dosya: RaporPro/tests/test_rapor_parsel_bilgileri.py

from pathlib import Path

from docx import Document

from proje_sema import (
    PROJE_SEMA_SURUMU,
    proje_verisini_migre_et,
    varsayilan_proje_verisi,
)
from rapor_parsel_bilgileri import (
    inceleme_alani_konum_metni,
    imar_adasi_metni,
    imar_plani_metni,
    parsel_cevre_ozeti_metni,
    parsel_tipi_normalize_et,
    rapor_bilgileri_eksikleri,
    rapor_bilgileri_varsayilanlari,
    rapor_metin_degerleri,
)
from raporlama_parsel import (
    rapor_kosullu_bolumlerini_uygula,
    rapor_sabit_tablolarini_uygula,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "sablonlar" / "rapor" / "varsayilan_rapor_sablonu.docx"


def _sample_project():
    veri = varsayilan_proje_verisi()
    veri["kunye"].update(
        {
            "sahibi": "Eski Proje Adı",
            "il": "Balıkesir",
            "ilce": "Edremit",
            "mah": "Altınkum",
            "paf": "J17",
            "ada": "1262",
            "par": "4",
        }
    )
    veri["arazi"].update(
        {
            "alan_y": "39.58",
            "alan_x": "26.82",
            "ort": "18.4",
            "min": "17.2",
            "max": "19.1",
            "egim": "%3",
            "yon": "güney",
            "imar_durumu": "Önlemli Alan 5.1",
            "pga": "0.42",
        }
    )
    veri["rapor_bilgileri"].update(
        {
            "proje_adi": "Altınkum Konut Projesi",
            "yapi_sahibi": "Örnek Yapı A.Ş.",
            "ilgili_idare": "Edremit Belediyesi",
            "rapor_tarihi": "15.07.2026",
            "parsel_alani_m2": "1240",
            "parsel_tipi": "Köşe parsel",
            "yol_cepheleri": "Kuzey ve doğu yönlerinde imar yoluna cephelidir",
            "komsu_parseller": "Batıda 3, güneyde 5 numaralı parsellere komşudur",
            "plan_adi": "1/1000 Ölçekli Uygulama İmar Planı",
            "plan_onay_tarihi": "10.01.2024",
            "plan_karar_no": "24",
            "plan_onay_idaresi": "Edremit Belediye Meclisi",
            "iklim_tipi": "Akdeniz iklimi",
            "heyelan_durumu": "Yok",
            "kaya_dusmesi_durumu": "Yok",
            "cig_durumu": "Yok",
            "cokme_durumu": "Yok",
            "laboratuvar_adi": "Deney Laboratuvarı",
        }
    )
    veri["sondaj"] = [
        {
            "no": "SK-1",
            "der": "15",
            "spt": [["1.5", "2", "3", "4", "7"]],
            "pmt": [],
            "kaya": [],
        }
    ]
    veri["lab_sheet"]["rows"] = [["SK-1", "1.50", "Kil"]]
    veri["jeofizik"]["ss_list"] = [
        {
            "ad": "Serim 1",
            "coords": ["39.5800", "26.8200", "39.5800", "26.8205", "39.5800", "26.8210"],
            "layers": [{"vp": "500", "vs": "250"}],
        }
    ]
    return veri


def _all_doc_text(doc):
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            chunks.extend(paragraph.text for paragraph in header.paragraphs)
            for table in header.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_rapor_defaults_are_independent():
    first = rapor_bilgileri_varsayilanlari()
    second = rapor_bilgileri_varsayilanlari()
    first["aktif_faylar"].append({"ad": "Test Fayı"})
    assert second["aktif_faylar"] == []


def test_v3_project_migration_preserves_legacy_identity_and_settings():
    legacy = {
        "schema_version": 3,
        "kunye": {"sahibi": "Mevcut Proje"},
        "ayarlar": {
            "taahhut_ilgili_idare": "Test Belediyesi",
            "tutanak_jeofizik_cihaz": "Test Cihazı",
        },
    }
    migrated, info = proje_verisini_migre_et(legacy)
    assert info.yeni_surum == PROJE_SEMA_SURUMU
    assert migrated["rapor_bilgileri"]["proje_adi"] == "Mevcut Proje"
    assert migrated["rapor_bilgileri"]["yapi_sahibi"] == "Mevcut Proje"
    assert migrated["rapor_bilgileri"]["ilgili_idare"] == "Test Belediyesi"
    assert migrated["rapor_bilgileri"]["sismik_cihaz"] == "Test Cihazı"


def test_v5_project_migration_adds_flat_imar_fields_and_normalizes_parcel_type():
    legacy = {
        "schema_version": 5,
        "rapor_bilgileri": {"parsel_tipi": "ada"},
        "dosyalar": {"masw_word_paths": []},
    }
    migrated, _info = proje_verisini_migre_et(legacy)

    assert migrated["rapor_bilgileri"]["parsel_tipi"] == "Ada parsel"
    assert "yol_cephe_sayisi" in migrated["rapor_bilgileri"]
    assert "dogalgaz_hatti" in migrated["rapor_bilgileri"]
    assert "imar_adasi" not in migrated["rapor_bilgileri"]


def test_parsel_tipi_aliases_are_limited_to_canonical_options():
    assert parsel_tipi_normalize_et("ara") == "Ara parsel"
    assert parsel_tipi_normalize_et("KÖŞE") == "Köşe parsel"
    assert parsel_tipi_normalize_et("ada parsel") == "Ada parsel"
    assert parsel_tipi_normalize_et("üç tarafı yol") == "Belirtilmedi"


def test_imar_adasi_metni_is_natural_and_omits_unapproved_empty_fields():
    veri = varsayilan_proje_verisi()
    veri["arazi"].update({"imar_alani": "Konut", "egim": "0-1"})
    veri["rapor_bilgileri"].update(
        {
            "parsel_tipi": "Köşe parsel",
            "yol_cephe_sayisi": "2",
            "yol_cepheleri": "iki yola cepheli",
            "komsu_parseller": "2, 4",
            "mevcut_yapilar": "İçerisinde oturulan iki katlı yapılar ile boş parseller bulunmaktadır",
            "mevcut_kullanim": "Arsa",
            "yol_kaplama": "sıcak asfalt",
            "yaya_trafik": "Açık",
            "tasit_trafik": "Açık",
            "kanalizasyon_hatti": "Var",
            "temiz_su_hatti": "Var",
        }
    )

    text = imar_adasi_metni(veri)

    assert "Çalışma alanı 'Konut Alanı' içinde bulunmaktadır." in text
    assert "iki yola cepheli köşe parsel" in text
    assert "Komşu parsellerde oturulan iki katlı yapılar ile boş parseller bulunmaktadır." in text
    assert "kanalizasyon ve temiz su hatları geçmektedir" in text
    assert "Belirtilmedi" not in text
    assert "Doğalgaz" not in text
    assert "Parsel tipi:" not in text


def test_imar_plani_metni_combines_absent_disaster_and_building_ban_sentences():
    veri = varsayilan_proje_verisi()
    veri["rapor_bilgileri"].update(
        {
            "afete_maruz_bolge": "Yok",
            "yapi_yasagi": "Yok",
        }
    )

    text = imar_plani_metni(veri)

    assert text.count("İnceleme alanı için Afete Maruz Bölge kararı ve yapı yasağı yoktur.") == 1
    assert "bulunmadığı belirtilmiştir" not in text
    assert "yapı yasağı bulunmadığı belirtilmiştir" not in text


def test_imar_plani_metni_keeps_separate_sentences_for_mixed_statuses():
    veri = varsayilan_proje_verisi()
    veri["rapor_bilgileri"].update(
        {
            "afete_maruz_bolge": "Yok",
            "yapi_yasagi": "Var",
        }
    )

    text = imar_plani_metni(veri)

    assert "İnceleme alanı için Afete Maruz Bölge kararı bulunmadığı belirtilmiştir." in text
    assert "Proje verilerinde yapı yasağı bulunduğu belirtilmiştir." in text
    assert "İnceleme alanı için Afete Maruz Bölge kararı ve yapı yasağı yoktur." not in text


def test_dynamic_report_texts_use_project_specific_values():
    values = rapor_metin_degerleri(_sample_project())
    assert "Altınkum Konut Projesi" in values["[ETUT_AMAC_KAPSAM]"]
    assert "1262 ada" in values["[PARSEL_TANITIM]"]
    assert "Kuzey ve doğu" in values["[PARSEL_TANITIM]"]
    assert "10.01.2024" in values["[IMAR_PLANI_ACIKLAMA]"]
    assert "Edremit Belediye Meclisi" in values["[IMAR_PLANI_ACIKLAMA]"]


def test_parsel_cevresi_ozeti_mechanical_fields_and_area_unit():
    veri = _sample_project()
    veri["rapor_bilgileri"].update(
        {
            "parsel_alani_m2": "5.011,97 m2",
            "parsel_cevresi_ozeti": (
                "Parsel, üç yola cepheli köşe parsel niteliğinde olup diğer "
                "cephesinde 9 numaralı boş parsel bulunmaktadır."
            ),
        }
    )

    text = rapor_metin_degerleri(veri)["[PARSEL_TANITIM]"]

    assert (
        "Parsel, üç yola cepheli köşe parsel niteliğinde olup diğer cephesinde "
        "9 numaralı boş parsel bulunmaktadır."
    ) in text
    assert "Parsel köşe parsel niteliğindedir" not in text
    assert "Parselin yol cepheleri" not in text
    assert "Komşu parsel bilgileri" not in text
    assert "Yakın çevredeki mevcut yapılar" not in text
    assert "Parselin mevcut kullanımı" not in text
    assert "5.011,97 m²" in text
    assert "m2 m²" not in text


def test_parsel_cevre_fields_are_rendered_as_natural_text():
    veri = _sample_project()
    veri["rapor_bilgileri"].update(
        {
            "parsel_tipi": "Köşe parsel",
            "yol_cepheleri": "Yola cepheli",
            "komsu_parseller": "9",
            "mevcut_yapilar": "Boş",
            "mevcut_kullanim": "Boş",
        }
    )

    text = parsel_cevre_ozeti_metni(veri["rapor_bilgileri"])

    assert text == (
        "Parsel, yola cepheli köşe parsel niteliğinde olup 9 numaralı parsele "
        "komşudur. Komşu parsellerde mevcut yapı bulunmamakta ve parsel "
        "hâlihazırda boş durumdadır."
    )
    assert "Yakın çevrede mevcut yapı" not in text
    for label in (
        "Parselin yol cepheleri",
        "Komşu parsel bilgileri",
        "Yakın çevredeki mevcut yapılar",
        "Parselin mevcut kullanımı",
    ):
        assert label not in text


def test_parsel_cevre_custom_summary_suppresses_only_automatic_fields():
    veri = _sample_project()
    veri["rapor_bilgileri"].update(
        {
            "parsel_cevresi_ozeti": "Parsel üç yola cepheli köşe parseldir.",
            "bitki_ortusu": "Seyrek ot örtüsü",
        }
    )

    text = rapor_metin_degerleri(veri)["[PARSEL_TANITIM]"]

    assert "Parsel üç yola cepheli köşe parseldir." in text
    assert "Parsel köşe parsel niteliğindedir" not in text
    assert "Parselin yol cepheleri" not in text
    assert "Komşu parsel bilgileri" not in text
    assert "Yakın çevredeki mevcut yapılar" not in text
    assert "Parselin mevcut kullanımı" not in text
    assert "Bitki örtüsü: Seyrek ot örtüsü." in text


def test_parsel_cevre_full_sentences_are_not_relabelled():
    veri = _sample_project()
    veri["rapor_bilgileri"].update(
        {
            "yol_cepheleri": "Parselin kuzey cephesinde imar yolu bulunmaktadır.",
            "komsu_parseller": "Batıda 3 numaralı parsele komşudur.",
            "mevcut_yapilar": "Yakın çevrede konut yapıları bulunmaktadır.",
            "mevcut_kullanim": "Parsel hâlihazırda tarımsal amaçla kullanılmaktadır.",
        }
    )

    text = parsel_cevre_ozeti_metni(veri["rapor_bilgileri"])

    assert "Parselin kuzey cephesinde imar yolu bulunmaktadır." in text
    assert "Batıda 3 numaralı parsele komşudur." in text
    assert "Yakın çevrede konut yapıları bulunmaktadır." in text
    assert "Parsel hâlihazırda tarımsal amaçla kullanılmaktadır." in text
    assert "Komşu parsel bilgileri:" not in text


def test_empty_neighbor_building_fallback_uses_neighbor_parcels_context():
    data = _sample_project()["rapor_bilgileri"]
    data.update(
        {
            "parsel_tipi": "Belirtilmedi",
            "yol_cepheleri": "",
            "komsu_parseller": "",
            "mevcut_yapilar": "Boş",
            "mevcut_kullanim": "Tarla",
        }
    )

    text = parsel_cevre_ozeti_metni(data)

    assert "Komşu parsellerde mevcut yapı bulunmamaktadır." in text
    assert "Yakın çevrede mevcut yapı bulunmamaktadır." not in text


def test_shared_location_sentence_is_used_for_parcel_and_result():
    veri = _sample_project()
    expected = (
        "İnceleme alanı Balıkesir ili, Edremit ilçesi, Altınkum Mahallesi, "
        "1262 ada, 4 parsel sınırlarında ve Enlem: 39.58, Boylam: 26.82 "
        "(WGS84) koordinatlarındadır."
    )

    values = rapor_metin_degerleri(veri)
    assert inceleme_alani_konum_metni(veri) == expected
    assert values["[PARSEL_TANITIM]"].startswith(expected)
    assert values["[SONUC_KONUM]"] == expected


def test_laboratory_name_is_not_required_for_static_template():
    veri = varsayilan_proje_verisi()
    veri["rapor_bilgileri"].update(
        {
            "proje_adi": "Proje",
            "yapi_sahibi": "İşveren",
            "ilgili_idare": "İdare",
            "rapor_tarihi": "01.01.2026",
            "plan_adi": "Plan",
            "yol_cepheleri": "Yol",
            "komsu_parseller": "Komşu",
        }
    )
    assert rapor_bilgileri_eksikleri(veri) == []
    veri["lab_sheet"]["rows"] = [["SK-1", "1.50"]]
    assert rapor_bilgileri_eksikleri(veri) == []


def test_static_tables_are_not_changed_by_project_data():
    doc = Document()
    climate = doc.add_table(rows=1, cols=2)
    climate.rows[0].cells[0].text = "CANAKKALE"
    doc.add_paragraph("Tablo 1. Çanakkale Meteoroloji İstasyonu iklim verileri")
    fault = doc.add_table(rows=2, cols=3)
    fault.rows[0].cells[0].text = "Diri Fayların Literatür Adları"
    fault.rows[1].cells[0].text = "SABİT FAY"
    seismic = doc.add_table(rows=2, cols=4)
    seismic.rows[0].cells[1].text = "P Dalgası İçin Toplam Vuruş Sayısı"
    seismic.rows[1].cells[0].text = "SABİT SİSMİK"
    masw = doc.add_table(rows=2, cols=4)
    masw.rows[0].cells[1].text = "SDalgası İçin Toplam Vuruş Sayısı"
    masw.rows[1].cells[0].text = "SABİT MASW"

    veri = _sample_project()
    veri["rapor_bilgileri"].update(
        {
            "aktif_faylar": [
                {"ad": "Örnek Fayı", "uzaklik_km": "12.5", "buyukluk": "6.8"}
            ],
            "sismik_vurus_sayisi": "6",
            "sismik_kayit_uzunlugu": "0.128",
            "masw_vurus_sayisi": "4",
            "masw_kayit_uzunlugu": "1",
        }
    )
    before = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]
    result = rapor_sabit_tablolarini_uygula(doc, veri)
    after = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in doc.tables
    ]
    assert result == {
        "iklim_tablolari": 0,
        "aktif_fay_satiri": 0,
        "sismik_satiri": 0,
        "masw_satiri": 0,
    }
    assert after == before
    assert "Örnek Fayı" not in "\n".join(
        cell for table in after for row in table for cell in row
    )


def test_empty_jeophysics_section_is_removed_without_touching_next_section():
    doc = Document()
    doc.add_heading("3.1. Jeofizik Çalışmalar", level=2)
    doc.add_paragraph("Jeofizik içeriği")
    doc.add_heading("3.1.1. Sismik Kırılma", level=3)
    doc.add_paragraph("Sismik içerik")
    doc.add_heading("3.2. ARAŞTIRMA ÇUKURLARI", level=2)
    doc.add_paragraph("Araştırma çukuru içeriği")
    doc.add_heading("3.3. SONDAJLAR", level=2)
    doc.add_paragraph("Sondaj içeriği")

    removed = rapor_kosullu_bolumlerini_uygula(
        doc,
        {
            "jeofizik": {"ss_list": [], "mt_list": []},
            "sondaj": [{"no": "SK-1", "spt": [], "pmt": [], "kaya": []}],
        },
    )
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "3.1" in removed
    assert "Jeofizik Çalışmalar" not in text
    assert "Sismik içerik" not in text
    assert "3.2. ARAŞTIRMA ÇUKURLARI" in text


def test_built_in_template_keeps_static_sections_and_dynamic_project_tags():
    doc = Document(TEMPLATE)
    text = _all_doc_text(doc)
    for tag in (
        "[ETUT_AMAC_KAPSAM]",
        "[PARSEL_TANITIM]",
        "[IMAR_PLANI_ACIKLAMA]",
    ):
        assert tag in text
    for static_tag in (
        "[IKLIM_ACIKLAMA]",
        "[DON_DURUM_ACIKLAMA]",
        "[DOGAL_AFET_ACIKLAMA]",
        "[AKTIF_TEKTONIK_ACIKLAMA]",
        "[AKTIF_FAY_GIRIS]",
        "[SISMIK_YONTEM_ACIKLAMA]",
        "[MASW_YONTEM_ACIKLAMA]",
        "[MT_YONTEM_ACIKLAMA]",
        "[MT_DEGERLENDIRME_ACIKLAMA]",
        "[ARASTIRMA_CUKURU_ACIKLAMA]",
        "[SPT_GIRIS]",
        "[SPT_TEKNIK_ACIKLAMA]",
        "[LAB_GIRIS]",
        "[SONUC_KAZI]",
        "[SONUC_KAZI_ONLEM]",
        "[SONUC_EK_ACIKLAMA]",
    ):
        assert static_tag not in text
    for static_text in (
        "Bölge Akdeniz iklimine sahip",
        "Sismik çalışmaları için yapılan ölçüler GEOMETRICS Geode",
        "Sondajlarda kuyu çapı 90mm",
        "Arter Mühendislik Laboratuvarında",
        "Çalışma Alanında Kazı sınıfı",
    ):
        assert static_text in text
    for sample_project_text in (
        "Mehmet Hakan ELMALI",
        "Çanakkale Belediyesi",
        "Aralık 2025",
        "komşu 11 ve 20",
    ):
        assert sample_project_text not in text
    assert len(doc.tables[2].rows) == 6
    assert len(doc.tables[3].rows) == 3
    assert len(doc.tables[4].rows) == 3

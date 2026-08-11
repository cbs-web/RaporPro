# Dosya: RaporPro/tests/test_rapor_sablonu_profilleri.py

import os

from docx import Document

from proje_sema import varsayilan_proje_verisi
from rapor_sablonu import (
    RAPOR_SABLON_PROFILI_DARDANOS_CINARLI,
    RAPOR_SABLON_PROFILI_GENEL,
    dahili_rapor_sablonu_yolu,
    etkin_rapor_sablonu_yolu,
    rapor_sablon_profili_normalize,
    rapor_sablonu_durumu,
)


def _document_text(path):
    doc = Document(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs)


def test_dardanos_cinarli_profili_ayri_ve_okunur_bir_sablondur():
    general_path = dahili_rapor_sablonu_yolu(RAPOR_SABLON_PROFILI_GENEL)
    dardanos_path = dahili_rapor_sablonu_yolu(RAPOR_SABLON_PROFILI_DARDANOS_CINARLI)

    assert os.path.isfile(general_path)
    assert os.path.isfile(dardanos_path)
    assert os.path.normcase(general_path) != os.path.normcase(dardanos_path)

    text = _document_text(dardanos_path)
    assert "[PROJE_ADI]" in text
    assert "[PARSEL_TANITIM]" in text
    assert "[Sondaj]" in text
    assert "Bölge Akdeniz iklimine sahip" in text
    assert "Türkiye Don İndeksi ve Don Penetrasyon Derinliği Haritası" in text


def test_dardanos_cinarli_profili_etikette_ve_etkin_yolda_korunur():
    status = rapor_sablonu_durumu(None, "Dardanos-Çınarlı")

    assert status["ready"] is True
    assert status["source"] == "builtin"
    assert status["profile"] == RAPOR_SABLON_PROFILI_DARDANOS_CINARLI
    assert "Dardanos-Çınarlı" in status["label"]
    assert etkin_rapor_sablonu_yolu(None, "dardanos") == status["path"]


def test_bilinmeyen_profil_ve_eski_projeler_genel_sablona_doner():
    assert rapor_sablon_profili_normalize("bilinmeyen") == RAPOR_SABLON_PROFILI_GENEL
    assert varsayilan_proje_verisi()["ayarlar"]["rapor_sablon_profili"] == RAPOR_SABLON_PROFILI_GENEL
    assert etkin_rapor_sablonu_yolu(None, "bilinmeyen") == dahili_rapor_sablonu_yolu()


def test_ozel_sablon_secimi_dahili_profilden_once_gelir(tmp_path):
    custom_path = tmp_path / "ozel.docx"
    Document().save(custom_path)

    status = rapor_sablonu_durumu(custom_path, RAPOR_SABLON_PROFILI_DARDANOS_CINARLI)

    assert status["source"] == "custom"
    assert status["path"] == os.path.abspath(custom_path)
    assert status["fallback"] is False

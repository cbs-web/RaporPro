# Dosya: RaporPro/tests/test_geoteknik_rapor_okuma.py

from pathlib import Path

import pytest
from docx import Document

from geoteknik_rapor_okuma import (
    geoteknik_raporu_oku,
    geoteknik_sayfalarindan_alanlari_ayikla,
)


def _sample_table():
    return [
        ["", "A BLOK", "B BLOK"],
        ["Toplam İnşaat Alanı (m2)", "9.614,75", "5.928,40"],
        ["Temel Oturumu (m2)", "1.938,63", "971,47"],
        ["Kat Adedi", "9", "9"],
        ["Min. - Ort.- Mak. 1,4G+1,6 Q Yüklemesi Zemin Gerilme (t/m2)", "10,77-15,39-19,06", "11,83-20,73-25,73"],
        ["Min. - Ort.- Mak. G+Q+E Yüklemesi Zemin Gerilme (t/m2)", "0,92-14,81-22,52", "3,17-20,37-25,00"],
        ["B (m)", "18,62 - 19,10", "18,20"],
        ["L (m)", "41,05 - 39,80", "38,20"],
        ["Temel Tipi", "Radye", "Radye"],
        ["BYS", "5", "5"],
        ["Ortalama Kazı Derinliği", "-4,74", "-5,13"],
    ]


def _field(fields, block, key):
    return next(
        item
        for item in fields
        if item.get("blok_adi", "") == block and item["anahtar"] == key
    )


def test_geoteknik_yapi_tablosu_bloklara_ayrilir_ve_ortak_bilgiler_tamamlanir():
    pages = [
        {
            "no": 10,
            "text": (
                "Yapılar Bodrum Kat + Zemin Kat + 7 Normal Kat olmak üzere toplam 9 katlıdır. "
                "Konut ve işyeri kullanımına sahip betonarme çerçeve taşıyıcı sistemdir. "
                "Bina Kullanım Sınıfı BKS = 3, Bina Önem Katsayısı I = 1'dir. "
                "Bina Yüksekliği (m) 24,50 ve Yerel Zemin Sınıfı ZD'dir."
            ),
            "tables": [_sample_table()],
        },
        {
            "no": 38,
            "text": (
                "Temel taban basıncı (Depremli): 22,52 t/m2 (225 kPa) (A Blok) "
                ": 26,00 t/m2 (260 kPa) (B Blok) Yeraltı suyu derinliği 6,00 m."
            ),
            "tables": [],
        },
        {
            "no": 69,
            "text": "Blok Sayısı / Kat Sayısı (Bodrum Kat Dahil) 1 / 9 Kat",
            "tables": [],
        },
    ]

    fields, warnings, blocks = geoteknik_sayfalarindan_alanlari_ayikla(
        pages,
        "ornek.pdf",
    )

    assert blocks == ["A Blok", "B Blok"]
    assert _field(fields, "A Blok", "bod")["deger"] == "1"
    assert _field(fields, "A Blok", "kat")["deger"] == "9"
    assert _field(fields, "A Blok", "plan")["deger"] == "18,62 - 19,10 × 41,05 - 39,80 m"
    assert _field(fields, "A Blok", "temel_alan")["deger"] == "1.938,63"
    assert _field(fields, "A Blok", "gqe_ort")["deger"] == "14,81"
    assert _field(fields, "A Blok", "comb_max")["deger"] == "19,06"
    assert _field(fields, "B Blok", "der")["deger"] == "5,13"
    assert _field(fields, "B Blok", "gqe_max")["alternatifler"] == ("26,00 (Sayfa 38)",)
    assert any("blok sayısı 1" in warning for warning in warnings)

    local_soil = next(item for item in fields if item["anahtar"] == "ysinif")
    assert local_soil["bolum"] == "bina"
    assert local_soil["deger"] == "ZD"
    assert sum(item["anahtar"] == "ysinif" for item in fields) == 1


def test_docx_geoteknik_raporu_okunur(tmp_path):
    document = Document()
    document.add_paragraph(
        "Yapı konut niteliğinde betonarme çerçeve taşıyıcı sistemde, "
        "Bodrum Kat + Zemin Kat + 4 Normal Kat olmak üzere toplam 6 katlıdır. "
        "Bina Kullanım Sınıfı BKS = 3 ve Bina Önem Katsayısı I = 1'dir."
    )
    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Bodrum Kat Adedi / Toplam Kat Adedi", "1 / 6"),
        ("Temel Alanı / Toplam İnşaat Alanı", "300 m² / 1800 m²"),
        ("B (m)", "15"),
        ("L (m)", "20"),
    ):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    path = Path(tmp_path) / "geoteknik.docx"
    document.save(path)

    result = geoteknik_raporu_oku(path)

    assert result["bloklar"] == ["Yapı"]
    by_key = {item["anahtar"]: item for item in result["alanlar"]}
    assert by_key["bod"]["deger"] == "1"
    assert by_key["kat"]["deger"] == "6"
    assert by_key["temel_alan"]["deger"] == "300"
    assert by_key["ins"]["deger"] == "1800"
    assert by_key["plan"]["deger"] == "15 × 20 m"


def test_tek_yapi_anlatimindan_alanlar_ve_yukler_okunur():
    pages = [
        {
            "no": 8,
            "text": (
                "3 - YAPI HAKKINDA BİLGİLER "
                "Parsel üzerinde temel tabanı 113,80 m2 (9,60*11,86 m) oturuma sahip, "
                "2 katlı mesken niteliğinde bina yapılacaktır. Yapının zemin kat oturumu "
                "113,80 m2 toplam inşaat alanı ise 252,40 m2 dir. Söz konusu yapının "
                "bodrum katı olmayıp, temel taban kotu -0,90 m de radye temel sistemine "
                "sahip olacaktır. 1,4G+1,6Q yüklemesi altında minimum-ortalama-maksimum "
                "7,27 t/m2 - 8,70 t/m2 - 12,93 t/m2, depremli yüklemeler (G+Q+E) "
                "altında ise minimum - ortalama - maksimum 5,76 t/m2 - 7,84 t/m2 - "
                "12,04 t/m2 olduğu belirtilmiştir."
            ),
            "tables": [],
        },
        {
            "no": 48,
            "text": (
                "RAPOR SONUCU 3 Blok Sayısı / Kat Sayısı (Bodrum Kat Dahil) 1 / 2 Kat "
                "4 Yapı Ebatları / Temel Ebatları 8,80 * 11,06 m2 / 9,60 * 11,86 m2 "
                "5 Kazı Yüksekliği (±0.00 kotuna göre) ±0.00'a göre kazı yüksekliği 1,00 m "
                "10 Yerel Zemin Sınıfı ZD"
            ),
            "tables": [],
        },
    ]

    fields, warnings, blocks = geoteknik_sayfalarindan_alanlari_ayikla(
        pages,
        "tek-yapi.pdf",
    )

    assert blocks == ["Yapı"]
    assert warnings == []
    by_key = {item["anahtar"]: item for item in fields}
    assert by_key["bod"]["deger"] == "0"
    assert by_key["kat"]["deger"] == "2"
    assert by_key["plan"]["deger"] == "8,80 × 11,06 m"
    assert by_key["temel_alan"]["deger"] == "113,80"
    assert by_key["ins"]["deger"] == "252,40"
    assert by_key["der"]["deger"] == "1,00"
    assert by_key["comb_min"]["deger"] == "7,27"
    assert by_key["comb_ort"]["deger"] == "8,70"
    assert by_key["comb_max"]["deger"] == "12,93"
    assert by_key["gqe_min"]["deger"] == "5,76"
    assert by_key["gqe_ort"]["deger"] == "7,84"
    assert by_key["gqe_max"]["deger"] == "12,04"


def _imar_field(fields, key):
    return next(item for item in fields if item["anahtar"] == key)


def test_imar_sahasi_alanlari_kaynak_kanıt_ve_hedef_bolumle_tasinir():
    pages = [
        {
            "no": 6,
            "text": (
                "2 - İNŞAAT SAHASI HAKKINDA BİLGİLER "
                "Parsel iki yola cepheli köşe parsel olup diğer cephesinde 9 nolu boş parsel bulunmaktadır. "
                "Kuzey ve doğu cephelerinde yol bulunmaktadır. Yollar sıcak asfalt ile kaplıdır. "
                "Yollar yaya ve taşıt trafiğine açık durumdadır. "
                "Ön cephede doğalgaz, elektrik, kanalizasyon ve temiz su hatları geçmektedir. "
                "Eğim Durumu ≈% 0-1."
            ),
            "tables": [],
        },
    ]

    fields, warnings, _blocks = geoteknik_sayfalarindan_alanlari_ayikla(
        pages,
        "imar-ornek.pdf",
    )

    parcel_type = _imar_field(fields, "parsel_tipi")
    assert parcel_type["bolum"] == "rapor_bilgileri"
    assert parcel_type["deger"] == "Köşe parsel"
    assert parcel_type["kaynak"] == "imar-ornek.pdf - Sayfa 6, 2 - İnşaat Sahası Hakkında Bilgiler"
    assert parcel_type["kanit"]
    assert _imar_field(fields, "yol_cephe_sayisi")["deger"] == "2"
    assert _imar_field(fields, "yol_yonleri")["deger"] == "Kuzey ve doğu"
    assert _imar_field(fields, "yol_kaplama")["deger"] == "sıcak asfalt"
    assert _imar_field(fields, "dogalgaz_hatti")["deger"] == "Var"
    assert _imar_field(fields, "elektrik_hatti")["deger"] == "Var"
    assert _imar_field(fields, "egim")["bolum"] == "arazi"
    assert warnings == []


def test_imar_sahasi_belirtilmeyen_hat_ve_riskleri_uretmez():
    pages = [
        {
            "no": 7,
            "text": (
                "2 - İNŞAAT SAHASI HAKKINDA BİLGİLER "
                "Parsel ara parsel konumundadır. Yollar yaya ve taşıt trafiğine açıktır."
            ),
            "tables": [],
        },
    ]

    fields, _warnings, _blocks = geoteknik_sayfalarindan_alanlari_ayikla(pages)
    keys = {item["anahtar"] for item in fields}
    assert "dogalgaz_hatti" not in keys
    assert "elektrik_hatti" not in keys
    assert "heyelan_durumu" not in keys
    assert "kaya_dusmesi_durumu" not in keys
    assert "cig_durumu" not in keys
    assert "cokme_durumu" not in keys


def test_arslanca_road_count_conflict_is_alternative_and_not_silently_resolved():
    pages = [
        {
            "no": 7,
            "text": "2 - İNŞAAT SAHASI HAKKINDA BİLGİLER Parsel köşe parsel olup üç tarafında yol bulunmaktadır.",
            "tables": [],
        },
        {
            "no": 53,
            "text": "10.2 Kazı Şevi Güvenliği İçin Gerekli Önlemler Parselin dört cephesi yollar ile sınırlıdır.",
            "tables": [],
        },
    ]

    fields, warnings, _blocks = geoteknik_sayfalarindan_alanlari_ayikla(pages)
    count = _imar_field(fields, "yol_cephe_sayisi")
    assert count["deger"] == "3"
    assert count["alternatifler"] == ("4 (Sayfa 53)",)
    assert count["uyari"]
    assert any("Yol cephesi sayısı" in warning for warning in warnings)


def test_multiple_neighbor_numbers_and_numbered_plan_names_are_compared_fully():
    pages = [
        {
            "no": 6,
            "text": (
                "2 - İNŞAAT SAHASI HAKKINDA BİLGİLER "
                "Parsel 2 ve 4 nolu komşu parseller ile komşudur. "
                "Parselin Plan Fonksiyonu 1. Derece Konut Alanı "
                "Yan Bahçe Çekme Mesafesi 3 m."
            ),
            "tables": [],
        },
        {
            "no": 8,
            "text": (
                "Parsel 2 ve 5 nolu komşu parseller ile komşudur. "
                "Parselin Plan Fonksiyonu 1. Derece Ticaret Alanı "
                "Yan Bahçe Çekme Mesafesi 3 m."
            ),
            "tables": [],
        },
    ]

    fields, warnings, _blocks = geoteknik_sayfalarindan_alanlari_ayikla(pages)

    neighbors = _imar_field(fields, "komsu_parseller")
    assert neighbors["deger"] == "2, 4"
    assert neighbors["alternatifler"] == ("2, 5 (Sayfa 8)",)
    plan = _imar_field(fields, "imar_alani")
    assert plan["deger"] == "1. Derece Konut Alanı"
    assert plan["alternatifler"] == ("1. Derece Ticaret Alanı (Sayfa 8)",)
    assert len(warnings) >= 2


@pytest.mark.parametrize(
    "path, expected",
    [
        (
            Path(r"C:\Users\Bugra Senel\Desktop\Barbaros 113_10\113 ada 10 parsel Geoteknik Rapor.pdf"),
            {"yol_cephe_sayisi": "3", "komsu_parseller": "9", "yaya_trafik": "Kısmen açık"},
        ),
        (
            Path(r"C:\Users\Bugra Senel\Desktop\Arslanca 1109_1\1109 ada 1 parsel Parsel Geoteknik Rapor.pdf"),
            {"yol_cephe_sayisi": "3", "komsu_parseller": "3", "yaya_trafik": "Açık"},
        ),
        (
            Path(r"C:\Users\Bugra Senel\Desktop\Dardanos 118_3\118 ada 3 parsel Geoteknik Rapor .pdf"),
            {"yol_cephe_sayisi": "2", "komsu_parseller": "2, 4", "dogalgaz_hatti": "Var"},
        ),
    ],
)
def test_regresyon_geoteknik_raporlari(path, expected):
    if not path.is_file():
        pytest.skip(f"Regresyon raporu bu makinede yok: {path}")
    result = geoteknik_raporu_oku(path)
    values = {
        item["anahtar"]: item["deger"]
        for item in result["alanlar"]
        if item["bolum"] in {"arazi", "rapor_bilgileri"}
    }
    for key, value in expected.items():
        assert values[key] == value

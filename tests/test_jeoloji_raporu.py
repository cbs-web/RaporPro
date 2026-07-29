# Dosya: RaporPro/tests/test_jeoloji_raporu.py
import unittest

from docx import Document
from docx.oxml.ns import qn

from jeoloji_raporu import (
    JEOLOJI_BIRIM_KATALOGU,
    jeoloji_birimleri,
    jeoloji_kodu_normalize,
    jeoloji_rapor_metinleri,
    jeoloji_varsayilanlari,
)
from proje_sema import PROJE_SEMA_SURUMU, proje_verisini_migre_et
from raporlama import replace_tag_with_report_blocks
from rapor_sablonu import dahili_rapor_sablonu_yolu
from tutarlilik_motoru import proje_tutarlilik_raporu


class JeolojiRaporTestleri(unittest.TestCase):
    def test_bilinen_kodlar_ascii_yazimdan_normalize_edilir(self):
        self.assertEqual(jeoloji_kodu_normalize("Tmcd"), "Tmçd")
        self.assertEqual(jeoloji_kodu_normalize("TMCK"), "Tmçk")
        self.assertEqual(jeoloji_kodu_normalize("tmal"), "Tmal")

    def test_tek_birim_yalniz_kendi_aciklamasini_uretir(self):
        veri = {
            "jeoloji": {
                "birimler": [
                    {
                        "kod": "Tmal",
                        "konum": "her_ikisi",
                        "durum": "reziduel",
                        "kesitte_kullan": True,
                    }
                ]
            }
        }

        texts = jeoloji_rapor_metinleri(veri)
        regional = "\n".join(texts["bolgesel"])

        self.assertIn("Alçıtepe Üyesi (Tmal)", regional)
        self.assertIn("rezidüel zeminler", regional)
        self.assertNotIn("Çamrakdere Üyesi", regional)
        self.assertIn("Alçıtepe Üyesi (Tmal)", "\n".join(texts["kesit"]))

    def test_coklu_birim_sirayi_korur_ve_kesit_secimini_uygular(self):
        veri = {
            "jeoloji": {
                "birimler": [
                    {
                        "kod": "Qal",
                        "konum": "inceleme_alani",
                        "durum": "aluvyon",
                        "kesitte_kullan": True,
                    },
                    {
                        "kod": "Tmçd",
                        "konum": "yakin_cevre",
                        "durum": "ana_kaya",
                        "kesitte_kullan": False,
                    },
                ]
            }
        }

        texts = jeoloji_rapor_metinleri(veri)
        regional = "\n".join(texts["bolgesel"])
        section = "\n".join(texts["kesit"])

        self.assertLess(regional.index("Alüvyon (Qal)"), regional.index("Çamrakdere Üyesi (Tmçd)"))
        self.assertIn("Alüvyon (Qal)", section)
        self.assertNotIn("Çamrakdere Üyesi", section)
        self.assertIn("yakın çevresinde", regional)

    def test_tekrarli_birimler_tek_kayda_birlesir(self):
        records = jeoloji_birimleri(
            {
                "birimler": [
                    {"kod": "Tmçd", "konum": "inceleme_alani"},
                    {"kod": "Tmcd", "konum": "yakin_cevre"},
                ]
            }
        )

        self.assertEqual(len(records), 1)
        self.assertEqual(records[0]["konum"], "her_ikisi")

    def test_birim_secmeden_tahmin_yapilmaz(self):
        texts = jeoloji_rapor_metinleri({"jeoloji": {"birimler": []}})
        combined = "\n".join(text for values in texts.values() for text in values)

        self.assertIn("tanımlanmamıştır", combined)
        for code in JEOLOJI_BIRIM_KATALOGU:
            self.assertNotIn(code, combined)

    def test_v2_harita_formasyonu_rapor_birimi_degil_oneri_olur(self):
        legacy = {
            "schema_version": 2,
            "harita_cizimleri": {
                "jeoloji": {
                    "formasyon": "Tmçd",
                    "img_path": "eski-altlik.jpg",
                }
            },
        }
        defaults = {
            "schema_version": PROJE_SEMA_SURUMU,
            "jeoloji": jeoloji_varsayilanlari(),
        }

        migrated, info = proje_verisini_migre_et(legacy, defaults)

        self.assertEqual(info.onceki_surum, 2)
        self.assertEqual(migrated["schema_version"], PROJE_SEMA_SURUMU)
        self.assertEqual(migrated["jeoloji"]["birimler"], [])
        self.assertEqual(
            migrated["jeoloji"]["harita_formasyon_onerisi"],
            "Tmçd",
        )

    def test_word_bloklari_sirali_ve_birim_basligi_kalin_eklenir(self):
        doc = Document()
        tag = doc.add_paragraph("[BOLGESEL_JEOLOJI]")

        changed = replace_tag_with_report_blocks(
            doc,
            "[BOLGESEL_JEOLOJI]",
            [
                {"tur": "metin", "metin": "Giriş metni."},
                {"tur": "birim_basligi", "metin": "Alçıtepe Üyesi (Tmal)"},
                {"tur": "metin", "metin": "Birim açıklaması."},
            ],
            paragraph_index={"[BOLGESEL_JEOLOJI]": tag},
        )

        self.assertTrue(changed)
        self.assertEqual(
            [paragraph.text for paragraph in doc.paragraphs],
            [
                "Giriş metni.",
                "Alçıtepe Üyesi (Tmal)",
                "Birim açıklaması.",
            ],
        )
        self.assertTrue(doc.paragraphs[1].runs[0].bold)
        self.assertFalse(doc.paragraphs[0].runs[0].bold)
        self.assertFalse(doc.paragraphs[2].runs[0].bold)
        self.assertTrue(
            all(
                paragraph._p.pPr.find(qn("w:keepLines")) is not None
                for paragraph in doc.paragraphs
            )
        )

    def test_word_bloklari_numarali_yer_tutucuyu_bos_birakmaz(self):
        doc = Document()
        tag = doc.add_paragraph("[JEOLOJI_SONUC]", style="List Number")

        replace_tag_with_report_blocks(
            doc,
            "[JEOLOJI_SONUC]",
            [{"tur": "metin", "metin": "Jeoloji sonucu."}],
            paragraph_index={"[JEOLOJI_SONUC]": tag},
        )

        self.assertEqual([paragraph.text for paragraph in doc.paragraphs], ["Jeoloji sonucu."])
        self.assertEqual(doc.paragraphs[0].style.name, "List Number")

    def test_dahili_sablon_dinamik_jeoloji_etiketlerini_icerir(self):
        doc = Document(dahili_rapor_sablonu_yolu())
        texts = [paragraph.text.strip() for paragraph in doc.paragraphs]

        for tag in (
            "[BOLGESEL_JEOLOJI]",
            "[BOLGESEL_JEOLOJI_BIRIMLERI]",
            "[MUHENDISLIK_JEOLOJISI]",
            "[JEOLOJIK_KESIT_ACIKLAMA]",
            "[JEOLOJI_SONUC]",
            "[MT_BIRIM_METNI]",
        ):
            self.assertEqual(texts.count(tag), 1)
        self.assertFalse(
            any(
                text.startswith("İnceleme alanı literatürde Üst Miyosen yaşlı")
                for text in texts
            )
        )

    def test_on_kontrol_jeolojik_birim_eksikligini_uyarir(self):
        report = proje_tutarlilik_raporu(
            {
                "kunye": {},
                "bina": {},
                "arazi": {},
                "sondaj": [],
                "jeofizik": {},
                "jeoloji": jeoloji_varsayilanlari(),
            }
        )
        finding = next(
            item
            for item in report["findings"]
            if item["id"] == "jeoloji.birimler"
        )

        self.assertEqual(finding["level"], "warning")
        self.assertEqual(finding["target"], "haritalar")


if __name__ == "__main__":
    unittest.main()

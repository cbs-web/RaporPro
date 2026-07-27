# Dosya: RaporPro/tests/test_raporlama_arazi.py
import unittest

from docx import Document

from raporlama_arazi import (
    arazi_deney_rapor_verileri,
    arazi_deney_word_bolumlerini_uygula,
    pmt_rapor_cumlesi,
    tcr_rapor_cumlesi,
)


def _ornek_bolum_doc():
    doc = Document()
    doc.add_paragraph("Tablo 13: Karot Yüzdeleri")
    kaya_tag = doc.add_paragraph("[KAYA_TABLO]")
    doc.add_paragraph(
        "Çalışma alanında yapılan sondajlarda karot yüzdeleri %23-66 arasındadır."
    )
    doc.add_paragraph("3.4.2. Presiyometre Deney Sonuçları")
    doc.add_paragraph(
        "Çalışma alanında SK1 ve SK3 sondajlarında presiyometre deneyi yapılmıştır."
    )
    doc.add_paragraph("Tablo 14: Presiyometre Deney Sonuçları")
    pmt_tag = doc.add_paragraph("[PMT]")
    return doc, {"[KAYA_TABLO]": kaya_tag, "[PMT]": pmt_tag}


class AraziDeneyRaporTestleri(unittest.TestCase):
    def test_rapor_verileri_gecerli_pmt_ve_tcr_kayitlarini_ayirir(self):
        sondajlar = [
            {
                "no": "SK-10",
                "spt": [["1.50", "50/10", "-", "-", "R"]],
                "pmt": [["3.00", "100", "5"], ["", "", ""]],
                "kaya": [["1.50-3.00", "23", "", ""]],
            },
            {
                "no": "SK-2",
                "pmt": [["6.00", "120", "7"]],
                "kaya": [["3.00-4.50", "66", "50", "40"]],
            },
            {
                "no": "SK-3",
                "pmt": [["9.00", "", "8"]],
                "kaya": [["", "75", "", ""]],
            },
        ]

        result = arazi_deney_rapor_verileri(sondajlar)

        self.assertEqual(result["pmt_sondajlari"], ["SK-2", "SK-10"])
        self.assertEqual(result["tcr_degerleri"], [23.0, 66.0])
        self.assertEqual(len(result["pmt_data"]), 2)
        self.assertEqual(len(result["kaya_data"]), 2)
        self.assertEqual(result["spt_data"][0][1:], ["1.50", "50/10", "-", "-", "R"])

    def test_dinamik_rapor_cumleleri_veriyi_yansitir(self):
        pmt_text = pmt_rapor_cumlesi(["SK-10", "SK-2", "SK-3"], table_number=14)
        tcr_text = tcr_rapor_cumlesi([66, 23, 45.5])

        self.assertIn("SK-2, SK-3 ve SK-10 sondajlarında", pmt_text)
        self.assertIn("Tablo 14'te", pmt_text)
        self.assertIn("%23-%66 arasındadır", tcr_text)

    def test_veri_yoksa_pmt_ve_karot_bloklari_kaldirilir(self):
        doc, index = _ornek_bolum_doc()

        arazi_deney_word_bolumlerini_uygula(
            doc,
            index,
            {
                "pmt_data": [],
                "kaya_data": [],
                "pmt_sondajlari": [],
                "tcr_degerleri": [],
            },
        )

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertNotIn("Presiyometre", text)
        self.assertNotIn("Karot Yüzdeleri", text)
        self.assertNotIn("[PMT]", text)
        self.assertNotIn("[KAYA_TABLO]", text)

    def test_veri_varsa_anlatimlar_guncellenir(self):
        doc, index = _ornek_bolum_doc()

        result = arazi_deney_word_bolumlerini_uygula(
            doc,
            index,
            {
                "pmt_data": [["SK-1", "3", "100", "5"], ["SK-3", "6", "120", "7"]],
                "kaya_data": [["SK-1", "1.5-3.0", "23", "", ""]],
                "pmt_sondajlari": ["SK-1", "SK-3"],
                "tcr_degerleri": [23, 66],
            },
        )

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertTrue(result["pmt_var"])
        self.assertTrue(result["kaya_var"])
        self.assertIn("SK-1 ve SK-3 sondajlarında", text)
        self.assertIn("%23-%66 arasındadır", text)
        self.assertIn("Tablo 14: Presiyometre Deney Sonuçları", text)
        pmt_paragraph = next(
            paragraph for paragraph in doc.paragraphs
            if "SK-1 ve SK-3 sondajlarında" in paragraph.text
        )
        tcr_paragraph = next(
            paragraph for paragraph in doc.paragraphs
            if "%23-%66 arasındadır" in paragraph.text
        )
        self.assertEqual(str(pmt_paragraph.runs[0].font.color.rgb), "EE0000")
        self.assertEqual(str(tcr_paragraph.runs[1].font.color.rgb), "EE0000")
        self.assertTrue(any(run.bold and "Tablo 14" in run.text for run in pmt_paragraph.runs))

    def test_yalniz_pmt_varsa_tablo_numarasi_13_olur(self):
        doc, index = _ornek_bolum_doc()

        result = arazi_deney_word_bolumlerini_uygula(
            doc,
            index,
            {
                "pmt_data": [["SK-1", "3", "100", "5"]],
                "kaya_data": [],
                "pmt_sondajlari": ["SK-1"],
                "tcr_degerleri": [],
            },
        )

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertEqual(result["pmt_table_number"], 13)
        self.assertNotIn("Karot Yüzdeleri", text)
        self.assertIn("Tablo 13: Presiyometre Deney Sonuçları", text)
        self.assertIn("Tablo 13'te verilmiştir", text)


if __name__ == "__main__":
    unittest.main()

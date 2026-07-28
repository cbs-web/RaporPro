# Dosya: RaporPro/tests/test_hidrojeoloji_raporu.py
import unittest

from docx import Document

from hidrojeoloji_raporu import (
    hidrojeoloji_durum_metni,
    hidrojeoloji_varsayilanlari,
    hidrojeoloji_word_paragrafini_uygula,
    sondaj_yass_seviyeleri,
)
from proje_sema import PROJE_SEMA_SURUMU, proje_verisini_migre_et
from rapor_sablonu import dahili_rapor_sablonu_yolu


def _arazi(**values):
    data = hidrojeoloji_varsayilanlari()
    data.update(values)
    return {"hidrojeoloji": data}


class HidrojeolojiRaporTestleri(unittest.TestCase):
    def test_varsayilan_durum_girilmemis_bilgiler_hakkinda_tahmin_yapmaz(self):
        text = hidrojeoloji_durum_metni({}, [])

        self.assertEqual(text, "Yapılan sondajlarda yeraltı suyuna rastlanmamıştır.")
        self.assertNotIn("dere", text)
        self.assertNotIn("denize", text)
        self.assertNotIn("taşkın", text)

    def test_dere_yok_taskin_yok_deniz_ve_yass_yok_cumlesi(self):
        text = hidrojeoloji_durum_metni(
            _arazi(
                akar_dere="Yok",
                kuru_dere="Yok",
                taskin_riski="Yok",
                deniz_mesafe="1000",
                yass_durumu="Rastlanmadı",
            ),
            [],
        )

        self.assertIn("akar veya kuru dere bulunmamaktadır", text)
        self.assertIn("taşkın riski belirlenmemiştir", text)
        self.assertIn("denize yaklaşık 1.000 m mesafededir", text)
        self.assertIn("yeraltı suyuna rastlanmamıştır", text)

    def test_kuru_dere_ve_yass_araligi_proje_verisinden_yazilir(self):
        text = hidrojeoloji_durum_metni(
            _arazi(
                akar_dere="Yok",
                kuru_dere="Var",
                kuru_dere_mesafe="250",
                kuru_dere_yon="Kuzey",
                taskin_riski="Var",
                deniz_mesafe="1400",
            ),
            [
                {"yass_d1": "2,80", "yass_d2": "3.40"},
                {"yass_d1": "", "yass_d2": ""},
            ],
        )

        self.assertIn("yaklaşık 250 m kuzeyinde kuru dere yatağı", text)
        self.assertIn("akar dere bulunmamaktadır", text)
        self.assertIn("taşkın riski bulunduğu", text)
        self.assertIn("denize yaklaşık 1.400 m", text)
        self.assertIn("2,80-3,40 m derinlikleri arasında", text)

    def test_akar_ve_kuru_dere_ayri_konumlarla_yazilir(self):
        text = hidrojeoloji_durum_metni(
            _arazi(
                akar_dere="Var",
                akar_dere_mesafe="600",
                akar_dere_yon="Doğu",
                kuru_dere="Var",
                kuru_dere_yon="Güneybatı",
                taskin_riski="Belirsiz",
                yass_durumu="Belirlenemedi",
            ),
            [],
        )

        self.assertIn("yaklaşık 600 m doğusunda akar dere", text)
        self.assertIn("güneybatısında kuru dere yatağı", text)
        self.assertIn("kesin olarak değerlendirilmesi için yeterli olmadığından", text)
        self.assertIn("yeraltı suyu seviyesi kesin olarak belirlenememiştir", text)

    def test_rastlandi_seciminde_seviye_yoksa_deger_uydurulmaz(self):
        text = hidrojeoloji_durum_metni(
            _arazi(yass_durumu="Rastlandı"),
            [],
        )

        self.assertEqual(text, "Yapılan sondajlarda yeraltı suyuna rastlanmıştır.")

    def test_yass_degerleri_sifir_dahil_sirali_ve_tekrarsizdir(self):
        values = sondaj_yass_seviyeleri([
            {"yass_d1": "3,5", "yass_d2": "0"},
            {"yass_d1": "3.50", "yass_d2": ""},
        ])

        self.assertEqual(values, [0.0, 3.5])

    def test_word_etiketi_dinamik_metinle_degistirilir(self):
        doc = Document()
        tag = doc.add_paragraph("[HIDROJEOLOJI_DURUM]")

        changed = hidrojeoloji_word_paragrafini_uygula(
            doc,
            {"[HIDROJEOLOJI_DURUM]": tag},
            "Dinamik hidrojeoloji metni.",
        )

        self.assertTrue(changed)
        self.assertEqual(doc.paragraphs[0].text, "Dinamik hidrojeoloji metni.")

    def test_eski_sablondaki_sabit_paragraf_da_degistirilir(self):
        doc = Document()
        doc.add_paragraph("4. HİDROJEOLOJİ", style="Heading 1")
        doc.add_paragraph(
            "İnceleme alanında akar ve kuru dere yoktur. "
            "İnceleme alanı denize 1000 m mesafededir. "
            "Yeraltı suyuna rastlanmamıştır."
        )

        changed = hidrojeoloji_word_paragrafini_uygula(
            doc,
            {},
            "Yeni hidrojeoloji açıklaması.",
        )

        self.assertTrue(changed)
        self.assertEqual(doc.paragraphs[1].text, "Yeni hidrojeoloji açıklaması.")

    def test_dahili_word_sablonunda_hidrojeoloji_etiketi_bulunur(self):
        doc = Document(dahili_rapor_sablonu_yolu())
        tags = [paragraph.text.strip() for paragraph in doc.paragraphs]

        self.assertIn("[HIDROJEOLOJI_DURUM]", tags)

    def test_v1_projesi_hidrojeoloji_alanlariyla_v2ye_tasinir(self):
        legacy = {"schema_version": 1, "arazi": {"egim": "5"}, "sondaj": []}
        defaults = {
            "schema_version": PROJE_SEMA_SURUMU,
            "arazi": {"hidrojeoloji": hidrojeoloji_varsayilanlari()},
        }

        migrated, info = proje_verisini_migre_et(legacy, defaults)

        self.assertEqual(info.onceki_surum, 1)
        self.assertEqual(migrated["schema_version"], PROJE_SEMA_SURUMU)
        self.assertEqual(
            migrated["arazi"]["hidrojeoloji"]["yass_durumu"],
            "Sondajlardan otomatik",
        )
        self.assertEqual(migrated["arazi"]["egim"], "5")


if __name__ == "__main__":
    unittest.main()

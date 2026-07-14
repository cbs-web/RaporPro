import json
import copy
import json
import os
from pathlib import Path
import tempfile
import threading
import unittest
from types import SimpleNamespace
from unittest import mock

from openpyxl import load_workbook

import spt_okuma_motoru as spt_motoru
from ai_motoru import belediye_duzeltme_analiz_et, duzeltme_metnini_kural_ile_analiz_et, duzeltme_yonlendirmeleri_olustur
from arayuz_proje import ArayuzProjeMixin
from cikti_kalite import (
    cikti_dosyalari_denetle,
    dosya_kalite_raporu,
    kalite_manifestosu_dogrula,
    kalite_manifestosu_yaz,
)
from jeofizik_sheet_motoru import (
    jeofizik_sheet_ozeti,
    jeofizik_sheet_rows_to_ss_list,
    jeofizik_sheet_var_mi,
    jeofizik_ss_koordinatlarini_koru,
)
from karot_motoru import derinlik_araligi_coz, derinlik_orta, standart_karot_araliklari, tcr_hesapla
from motor import GeoEngine, log_ornek_derinligi_formatla
from pmt_excel_motoru import pmt_excel_dosyasi_oku, pmt_kayitlarini_veriye_aktar
from rapor_metin_revizyon import (
    metin_revizyon_kural_analiz_et,
    metin_revizyonlari_uygula,
    word_metin_birimleri_oku,
)
from rapor_revizyon import docx_bolumlerini_degistir, revizyon_etiketi_var_mi, revizyon_isaretleri_ekle
from raporlama import (
    bina_bilgileri_tablolari_olustur,
    bina_bloklari_rapor,
    buyuk_basliklari_yeni_sayfaya_al,
    duzeltme_etiket_sablonu_olustur,
    duzeltme_etiketleri_temizle,
    jeo_parametre_degeri_formatla,
    jeofizik_vp_layers_sadelestir,
    lab_sheet_satirlari,
    lab_sheet_verisi_var_mi,
    litoloji_dagilim_birimi,
    litoloji_dagilim_paragraflari,
    mjh_resim_yolu,
)
from kalite_kontrol import analyze_word_template, build_preflight_report, image_path_for_tag, validate_project_data
from rapor_sablonu import dahili_rapor_sablonu_yolu, etkin_rapor_sablonu_yolu, rapor_sablonu_durumu
from ui_icons import IconManager
from spt_okuma_motoru import (
    SPTKaydi,
    _path_unique_key,
    _select_spt_records_for_batch,
    hedef_derinlige_yuvarla,
    kayit_normalize_et,
    n30_hesapla,
    normalize_sondaj_no,
    openai_model_sec,
    spt_ayarlarini_kaydet,
    spt_ayarlarini_yukle,
)
from sondaj_derinlik import (
    boussinesq_katsayisi,
    efektif_dusey_gerilme,
    efektif_gerilme_yass,
    gerilme_artisi_2_1,
    gerilme_artisi_boussinesq,
    gerilme_yuzde_on_derinlik_hesapla,
    sondaj_derinligi_hesapla,
    sondaj_derinligi_kontrol_sonucu,
    sondaj_derinligi_ozet_metni,
    temel_genisligi_bul,
    westergaard_katsayisi,
)
from sondaj_derinlik_foyu import sondaj_derinligi_foyu_olustur
from ui_spt_okuma_yardimci import collect_image_paths, duplicate_keys as spt_duplicate_keys, record_quality as spt_record_quality
from taahhutname import taahhutname_context, taahhutname_olustur, taahhutname_yapi_adresi
from tkgm_kml import geojson_kml_olustur, konum_adi_normalize_et
from tutanaklar import tutanaklari_olustur
from ekler import (
    EK_SET_NORMAL,
    ek_basliklari,
    ek_bloklari_oku,
    ek_olustur,
    ekleri_rapora_ekle,
    ekler_pdf_olustur,
    proje_presiyometre_var_mi,
    uygun_ek_sablonu,
)
from harita_cikti import eski_paylasimli_temp_harita_yolu_mu, yeni_harita_cikti_yolu
from harita_ayarlari import hgm_ortofoto_url_kaydet, hgm_ortofoto_url_yukle
from harita_referans import affine_from_refs, coord_to_pixel, kml_koordinatlari_oku, pixel_to_coord, ss_harita_etiketi
from gizli_depo import gizli_deger_coz, gizli_deger_mi, gizli_deger_sakla
from ui_kesit import KesitCizimMixin, kesit_hatti_sondaj_sirasi, kesit_kayit_dosya_adi
from ui_kontrol import KontrolPaneliMixin
from ui_sondaj import SondajMixin
from ui_proje_surumleri import ProjeSurumleriMixin
from proje_arsiv import (
    arsiv_kaydi_ekle,
    arsiv_kayitlari_yukle,
    biten_isler_kml_yaz,
    kml_sinir_koordinatlari_oku,
    proje_merkez_koordinati,
)
from proje_surumleri import (
    degisiklik_ozeti,
    eski_yedekleri_ice_aktar,
    proje_verilerini_karsilastir,
    surum_deposunu_kopyala,
    surum_kaydi_olustur,
    surum_verisi_yukle,
    surumleri_listele,
)
from proje_sema import (
    PROJE_SEMA_SURUMU,
    ProjeSemaHatasi,
    proje_verisini_migre_et,
)
from yardimcilar import atomic_docx_save, atomic_json_dump, atomic_write_text, litoloji_yazim_uyarilari, safe_float, zemin_sinifi_cevir
from workbook_motoru import apply_rows_to_veri as wb_apply_rows_to_veri
from workbook_motoru import build_initial_rows as wb_build_initial_rows
from workbook_motoru import excel_workbook_yaz as wb_excel_workbook_yaz
from workbook_motoru import validate_rows as wb_validate_rows
from workbook_motoru import WORKBOOK_SHEET_DEFS
from task_engine import TkTaskEngine
from tutarlilik_motoru import laboratuvar_kayitlarini_ayikla, proje_tutarlilik_raporu
from tutarlilik_ortak import refu_mu
from uygulama_yollari import eski_veriyi_kopyala, kullanici_yolu
from yonetmelik_motoru import (
    duzeltme_yonetmelik_dayanaklari,
    resmi_yonetmelik_baglanti_bul,
    resmi_yonetmelik_indir_ve_ekle,
    varsayilan_yonetmelikleri_hazirla,
    yerlesik_yonetmelik_ekle,
    yonetmelik_ara,
    yonetmelik_ekle,
    yonetmelikleri_listele,
)
import yonetmelik_motoru


class _ImmediateTkRoot:
    def after(self, _delay, callback):
        callback()


class TutarlilikMotoruTestleri(unittest.TestCase):
    @staticmethod
    def _base_veri():
        return {
            "kunye": {"sahibi": "Test Projesi", "il": "Çanakkale", "ilce": "Merkez"},
            "bina": {},
            "arazi": {},
            "ayarlar": {},
            "lab_sheet": {"rows": []},
            "jeofizik_sheet": {"rows": []},
            "kesit_ayarlari": {},
            "sondaj": [
                {
                    "no": "SK-1",
                    "der": "10.0",
                    "y": "40.1000",
                    "x": "26.4000",
                    "k": "100.0",
                    "litoloji": [["0", "5", "Kil"], ["5", "10", "Kum"]],
                    "spt": [["1.50", "2", "3", "4", "7"]],
                    "pmt": [],
                    "kaya": [],
                }
            ],
            "jeofizik": {
                "ss_list": [
                    {
                        "ad": "SS-1",
                        "coords": ["40.1", "26.4", "40.11", "26.41", "40.12", "26.42"],
                        "layers": [{"h": "5", "vp": "600", "vs": "250"}, {"h": "", "vp": "1200", "vs": "500"}],
                    }
                ],
                "mt_list": [],
            },
        }

    def test_bos_proje_sifir_hazirlik_ve_bloklayici_bulgu_uretir(self):
        veri = {
            "kunye": {}, "bina": {}, "arazi": {}, "sondaj": [], "jeofizik": {},
            "ayarlar": {}, "lab_sheet": {"rows": []}, "jeofizik_sheet": {"rows": []},
        }
        report = proje_tutarlilik_raporu(veri, {})
        self.assertEqual(report["score"], 0)
        self.assertTrue(report["blocking"])
        self.assertIn("sondaj.kayit", {item["id"] for item in report["findings"]})

    def test_sondaj_litoloji_ve_deney_uyumsuzluklarini_birlikte_yakalar(self):
        veri = self._base_veri()
        sondaj = veri["sondaj"][0]
        sondaj["litoloji"] = [["0", "4", "Kil"], ["5", "10", "Kum"]]
        sondaj["spt"] = [["1.50", "2", "3", "4", "99"]]
        sondaj["pmt"] = [["11.0", "100", "8"]]
        sondaj["kaya"] = [["6.0-9.0", "40", "50", "60"]]

        report = proje_tutarlilik_raporu(veri)
        ids = {item["id"] for item in report["findings"]}
        self.assertIn("sondaj.sk1.litoloji.1.bosluk", ids)
        self.assertIn("sondaj.sk1.spt.0.n30_toplam", ids)
        self.assertIn("sondaj.sk1.pmt.0.kuyu_disi", ids)
        self.assertIn("sondaj.sk1.kaya.0.oran_sirasi", ids)

    def test_laboratuvar_baslik_ve_derinliklerini_ayiklar(self):
        rows = [
            ["Sondaj No", "Örnek Derinliği", "USCS"],
            ["SK-1", "1.50-2.00", "CL"],
            ["", "3.00", "CL"],
        ]
        parsed = laboratuvar_kayitlarini_ayikla(rows)
        self.assertTrue(parsed["header_found"])
        self.assertEqual(len(parsed["records"]), 2)
        self.assertEqual(parsed["records"][1]["sondaj"], "SK-1")
        self.assertEqual(parsed["records"][0]["bottom"], 2.0)

    def test_validate_project_data_yapilandirilmis_raporu_korur(self):
        report = validate_project_data(self._base_veri())
        self.assertIn("findings", report)
        self.assertIn("checks", report)
        self.assertIn("errors", report)
        self.assertIn("warnings", report)

    def test_spt_elli_darbe_penetrasyon_gosterimi_refu_kabul_edilir(self):
        veri = self._base_veri()
        veri["sondaj"][0]["spt"] = [
            ["1.50", "50/5", "-", "-", "R"],
            ["3.00", "20", "50/10", "-", "R"],
        ]

        report = proje_tutarlilik_raporu(veri)
        spt_blow_findings = [
            item for item in report["findings"]
            if item["id"].startswith("sondaj.sk1.spt") and item["label"] == "SPT darbe sayısı"
        ]

        self.assertEqual(spt_blow_findings, [])
        self.assertTrue(refu_mu("50/10"))
        self.assertTrue(refu_mu("50 / 5 cm"))
        self.assertFalse(refu_mu("10/20"))
        self.assertFalse(refu_mu("49/10"))

    def test_jeofizik_sheet_manuel_serim_koordinatlarini_korur(self):
        veri = self._base_veri()
        veri["jeofizik"]["ss_list"][0]["ad"] = "SS-1"
        veri["jeofizik_sheet"]["rows"] = [
            ["Sismik Ölçü ve Hesaplarının Sahibi   :", "Serim 1", "", ""],
            ["Zemin Parametreleri", "Simge", "TABAKA NO", ""],
            ["", "Birim", "I", "II"],
            ["VP = Boyuna Dalga Hızı", "VP (m/sn)", 600, 1200],
            ["VS = Enine Dalga Hızı", "VS (m/sn)", 250, 500],
            ["Tabaka Kalınlığı", "metre", 5, ""],
        ]

        report = proje_tutarlilik_raporu(veri)
        coordinate_findings = [
            item for item in report["findings"]
            if item["id"].startswith("jeofizik.ss.0.koordinat")
        ]

        self.assertEqual(coordinate_findings, [])

    def test_sondaj_numarasi_bicim_farki_mukerrer_sayilir(self):
        veri = self._base_veri()
        ikinci = dict(veri["sondaj"][0])
        ikinci["no"] = "SK1"
        veri["sondaj"].append(ikinci)
        report = proje_tutarlilik_raporu(veri)
        self.assertTrue(any(
            item["level"] == "error" and "Mükerrer sondaj" in item["label"]
            for item in report["findings"]
        ))

    def test_preflight_mjh_icin_baska_haritayi_kullanmaz(self):
        from docx import Document
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            word_path = os.path.join(tmp, "sablon.docx")
            image_path = os.path.join(tmp, "yer.png")
            doc = Document()
            doc.add_paragraph("[PROJE_ADI]")
            doc.add_paragraph("RESIM:MJH")
            doc.save(word_path)
            Image.new("RGB", (1200, 800), "white").save(image_path)

            app = SimpleNamespace(
                veri=self._base_veri(),
                word_path=word_path,
                lab_excel_path=None,
                jeo_excel_path=None,
                kml_path=None,
                img_yer=image_path,
                img_tkgm=None,
                img_pga=None,
                img_mjh=None,
                word_img_sondaj=None,
                word_img_jeofizik=None,
            )
            report = build_preflight_report(app)
            mjh_findings = [
                item for item in report["findings"]
                if item.get("category") == "Rapor görselleri" and "MJH" in item.get("label", "")
            ]
            self.assertTrue(mjh_findings)
            self.assertTrue(any(item["level"] == "warning" for item in mjh_findings))

    def test_sablonda_olmayan_gorsel_uyari_degil_bilgi_olur(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            word_path = os.path.join(tmp, "sablon.docx")
            doc = Document()
            doc.add_paragraph("[PROJE_ADI]")
            doc.save(word_path)
            app = SimpleNamespace(
                veri=self._base_veri(),
                word_path=word_path,
                lab_excel_path=None,
                jeo_excel_path=None,
                kml_path=None,
                img_yer=None,
                img_tkgm=None,
                img_pga=None,
                img_mjh=None,
                word_img_sondaj=None,
                word_img_jeofizik=None,
            )
            report = build_preflight_report(app)
            visual_warnings = [
                item for item in report["findings"]
                if item["level"] == "warning" and item["category"] == "Rapor görselleri"
            ]
            self.assertEqual(visual_warnings, [])


class TaskEngineTestleri(unittest.TestCase):
    def test_basari_sayacini_ve_callbacki_gunceller(self):
        done = threading.Event()
        results = []
        engine = TkTaskEngine(_ImmediateTkRoot(), max_workers=1, log_failures=False)
        try:
            engine.run("deneme", lambda: 42, on_success=lambda result: (results.append(result), done.set()))
            self.assertTrue(done.wait(2))
            snap = engine.snapshot()
            self.assertEqual(results, [42])
            self.assertEqual(snap.active_count, 0)
            self.assertEqual(snap.completed_count, 1)
            self.assertEqual(snap.failed_count, 0)
        finally:
            engine.shutdown(wait=True)

    def test_hata_sayacini_ve_callbacki_gunceller(self):
        done = threading.Event()
        errors = []

        def fail():
            raise RuntimeError("kontrollu hata")

        engine = TkTaskEngine(_ImmediateTkRoot(), max_workers=1, log_failures=False)
        try:
            engine.run("hata", fail, on_error=lambda exc: (errors.append(str(exc)), done.set()))
            self.assertTrue(done.wait(2))
            snap = engine.snapshot()
            self.assertEqual(errors, ["kontrollu hata"])
            self.assertEqual(snap.active_count, 0)
            self.assertEqual(snap.completed_count, 0)
            self.assertEqual(snap.failed_count, 1)
        finally:
            engine.shutdown(wait=True)


class CiktiKaliteTestleri(unittest.TestCase):
    def test_word_ciktisinda_kalan_etiket_hata_olur(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "rapor.docx")
            document = Document()
            document.add_paragraph("Proje: [PROJE_ADI]")
            document.save(path)

            report = dosya_kalite_raporu(path)

        self.assertEqual(report["state"], "HATA")
        self.assertTrue(any("işlenmemiş etiket" in item["detail"] for item in report["errors"]))

    def test_temiz_word_ciktisi_gecerli_sayilir(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "rapor.docx")
            document = Document()
            document.add_heading("Zemin ve Temel Etüdü Veri Raporu", level=1)
            document.add_paragraph("Sondaj ve arazi deneyleri tamamlanmıştır.")
            document.save(path)
            report = dosya_kalite_raporu(path)

        self.assertEqual(report["state"], "TEMİZ")

    def test_tek_renk_gorsel_uyari_olur(self):
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "harita.png")
            Image.new("RGB", (800, 600), "white").save(path)
            report = dosya_kalite_raporu(path)

        self.assertEqual(report["state"], "UYARI")
        self.assertTrue(any(item["label"] == "Boş görsel" for item in report["warnings"]))

    def test_manifesto_proje_degisimini_algilar(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = os.path.join(tmp, "sonuc.txt")
            Path(output).write_text("RaporPro kalite denetimi " * 10, encoding="utf-8")
            first_data = {"schema_version": 1, "kunye": {"sahibi": "İlk Proje"}}
            quality = cikti_dosyalari_denetle([output], veri=first_data)
            manifest = os.path.join(tmp, "kalite.json")
            kalite_manifestosu_yaz(manifest, quality, veri=first_data)

            clean = kalite_manifestosu_dogrula(manifest, veri=first_data)
            stale = kalite_manifestosu_dogrula(
                manifest,
                veri={"schema_version": 1, "kunye": {"sahibi": "Değişen Proje"}},
            )

        self.assertEqual(clean["state"], "TEMİZ")
        self.assertTrue(any(item["label"] == "Güncelliğini yitirmiş çıktı" for item in stale["warnings"]))


class ProjeSemaTestleri(unittest.TestCase):
    def test_legacy_proje_kayipsiz_migre_edilir(self):
        legacy = {
            "kunye": {"sahibi": "Eski Proje"},
            "word_path": r"C:\projeler\rapor.docx",
            "sondaj": [{"no": "SK-1", "litoloji": None}],
            "jeofizik": {"ss": [{"ad": "SS-1"}]},
            "bilinmeyen_alan": {"koru": True},
        }
        defaults = ArayuzProjeMixin().varsayilan_veri_olustur()

        migrated, info = proje_verisini_migre_et(legacy, defaults)

        self.assertEqual(info.onceki_surum, 0)
        self.assertEqual(info.yeni_surum, PROJE_SEMA_SURUMU)
        self.assertTrue(info.degisti)
        self.assertEqual(migrated["schema_version"], PROJE_SEMA_SURUMU)
        self.assertEqual(migrated["dosyalar"]["word_path"], legacy["word_path"])
        self.assertEqual(migrated["jeofizik"]["ss_list"], [{"ad": "SS-1"}])
        self.assertEqual(migrated["sondaj"][0]["litoloji"], [])
        self.assertEqual(migrated["sondaj"][0]["pmt"], [])
        self.assertEqual(migrated["bilinmeyen_alan"], {"koru": True})
        self.assertNotIn("schema_version", legacy)

    def test_guncel_proje_tekrar_migrasyonda_degisimez(self):
        current = ArayuzProjeMixin().varsayilan_veri_olustur()
        migrated, info = proje_verisini_migre_et(current, current)
        self.assertFalse(info.degisti)
        self.assertEqual(migrated, current)

    def test_daha_yeni_proje_surumu_acilmaz(self):
        with self.assertRaisesRegex(ProjeSemaHatasi, "daha yeni"):
            proje_verisini_migre_et({"schema_version": PROJE_SEMA_SURUMU + 1})


class ProjeKayitDurumuTestleri(unittest.TestCase):
    def test_proje_degisti_mi_imza_ile_anlar(self):
        app = ArayuzProjeMixin()
        app.veri = {"kunye": {"sahibi": "İlk"}, "sondaj": []}
        app.aktif_dosya_yolu = None
        app._son_kayit_imzasi = None

        app.kayit_imzasi_guncelle()
        self.assertFalse(app.proje_degisti_mi())

        app.veri["kunye"]["sahibi"] = "Son"
        self.assertTrue(app.proje_degisti_mi())

    def test_kaydedilmemis_degisiklik_kaydedilip_devam_edebilir(self):
        app = ArayuzProjeMixin()
        app.veri = {"kunye": {"sahibi": "İlk"}, "sondaj": []}
        app.aktif_dosya_yolu = None
        app._son_kayit_imzasi = app.proje_kayit_imzasi()
        app.veri["kunye"]["sahibi"] = "Son"
        app.proje_kilitli_mi = lambda: False
        app.veri_kaydet = mock.Mock(return_value=True)

        with mock.patch("arayuz_proje.messagebox.askyesnocancel", return_value=True) as prompt:
            self.assertTrue(app.kaydedilmemis_degisiklik_onayi("başka proje açma"))

        app.veri_kaydet.assert_called_once_with()
        self.assertIn("Başka proje açma", prompt.call_args.args[1])

    def test_proje_acma_iptal_edilince_yukleme_yapilmaz(self):
        app = ArayuzProjeMixin()
        app.kaydedilmemis_degisiklik_onayi = mock.Mock(return_value=False)
        app.proje_dosyasi_yukle = mock.Mock()

        with mock.patch("arayuz_proje.filedialog.askopenfilename", return_value="proje.json"):
            app.proje_ac()

        app.kaydedilmemis_degisiklik_onayi.assert_called_once_with("başka proje açma")
        app.proje_dosyasi_yukle.assert_not_called()

    def test_yeni_proje_iptal_edilince_sihirbaz_acilmaz(self):
        app = ArayuzProjeMixin()
        app.kaydedilmemis_degisiklik_onayi = mock.Mock(return_value=False)
        app.yeni_proje_sihirbazi = mock.Mock()

        app.yeni_proje()

        app.yeni_proje_sihirbazi.assert_not_called()


class ProjeSurumGecmisiTestleri(unittest.TestCase):
    @staticmethod
    def _veri(derinlik="15", n30=12, lab_rows=None):
        return {
            "kunye": {"sahibi": "Surum Testi", "il": "Canakkale", "ilce": "Merkez"},
            "sondaj": [
                {
                    "no": "SK-1",
                    "der": derinlik,
                    "litoloji": [{"bas": "0", "bit": derinlik, "tanim": "Kil"}],
                    "spt": [{"der": "1.5", "n30": n30}],
                }
            ],
            "lab_sheet": {"rows": lab_rows or [["Sondaj No", "Derinlik"], ["SK-1", "1.5"]]},
            "jeofizik": {"ss_list": [], "mt_list": []},
            "dosyalar": {},
        }

    def test_surum_kaydi_ayni_icerigi_art_arda_cogaltmaz(self):
        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "proje.json")
            data = self._veri()
            with open(project_path, "w", encoding="utf-8") as handle:
                json.dump(data, handle, ensure_ascii=False)

            first, first_created = surum_kaydi_olustur(project_path, data, "Ilk kayit", keep=10)
            second, second_created = surum_kaydi_olustur(project_path, data, "Tekrar", keep=10)

            self.assertTrue(first_created)
            self.assertFalse(second_created)
            self.assertEqual(first["hash"], second["hash"])
            self.assertEqual(len(surumleri_listele(project_path, eski_yedekleri_aktar=False)), 1)

    def test_surumler_saklanir_yuklenir_ve_karsilastirilir(self):
        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "proje.json")
            old_data = self._veri()
            new_data = self._veri(derinlik="18", n30=19)

            surum_kaydi_olustur(project_path, old_data, "Baslangic", keep=10)
            latest, created = surum_kaydi_olustur(project_path, new_data, "Derinlik guncellendi", keep=10)
            records = surumleri_listele(project_path, eski_yedekleri_aktar=False)
            loaded = surum_verisi_yukle(project_path, latest)
            changes = proje_verilerini_karsilastir(old_data, loaded)
            summary = degisiklik_ozeti(changes)

            self.assertTrue(created)
            self.assertEqual(len(records), 2)
            self.assertEqual(loaded["sondaj"][0]["der"], "18")
            self.assertGreaterEqual(summary["total"], 2)
            self.assertTrue(any(item["label"].endswith("N30") for item in changes))
            self.assertTrue(any(item["old"] == "15" and item["new"] == "18" for item in changes))

    def test_eski_backup_dosyasi_gecmise_aktarilir(self):
        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "zemin.json")
            backup_dir = os.path.join(tmp, "backups")
            os.makedirs(backup_dir)
            backup_path = os.path.join(backup_dir, "zemin_20260713_101500.json")
            with open(backup_path, "w", encoding="utf-8") as handle:
                json.dump(self._veri(derinlik="12"), handle, ensure_ascii=False)

            imported = eski_yedekleri_ice_aktar(project_path, keep=10)
            imported_again = eski_yedekleri_ice_aktar(project_path, keep=10)
            records = surumleri_listele(project_path, eski_yedekleri_aktar=False)

            self.assertEqual(imported, 1)
            self.assertEqual(imported_again, 0)
            self.assertEqual(len(records), 1)
            self.assertEqual(records[0]["source"], "legacy_backup")
            self.assertEqual(surum_verisi_yukle(project_path, records[0])["sondaj"][0]["der"], "12")

    def test_sheet_icerik_degisimini_tek_anlamli_satirda_gosterir(self):
        old_data = self._veri(lab_rows=[["Sondaj No", "Derinlik"], ["SK-1", "1.5"]])
        new_data = self._veri(lab_rows=[["Sondaj No", "Derinlik"], ["SK-1", "3.0"]])

        changes = proje_verilerini_karsilastir(old_data, new_data)
        lab_changes = [item for item in changes if item["category"] == "Laboratuvar Sheet"]

        self.assertEqual(len(lab_changes), 1)
        self.assertEqual(lab_changes[0]["label"], "Laboratuvar çalışma sayfası")
        self.assertIn("içerik değişti", lab_changes[0]["new"])

    def test_gercek_spt_satir_biciminde_n30_alani_ayri_gosterilir(self):
        old_data = {"sondaj": [{"no": "SK-1", "spt": [["1.50", "2", "3", "4", "7"]]}]}
        new_data = {"sondaj": [{"no": "SK-1", "spt": [["1.50", "2", "3", "5", "8"]]}]}

        changes = proje_verilerini_karsilastir(old_data, new_data)

        self.assertEqual(len(changes), 2)
        self.assertTrue(any(item["label"].endswith("30-45 cm") for item in changes))
        self.assertTrue(any(item["label"].endswith("N30") for item in changes))
        self.assertTrue(all("SPT [1.50]" in item["label"] for item in changes))

    def test_surum_numarasi_eski_kayitlar_budansa_da_sabit_kalir(self):
        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "proje.json")
            for index in range(7):
                surum_kaydi_olustur(
                    project_path,
                    self._veri(derinlik=str(10 + index)),
                    f"Kayit {index + 1}",
                    keep=5,
                )

            records = surumleri_listele(project_path, eski_yedekleri_aktar=False, keep=5)

            self.assertEqual(len(records), 5)
            self.assertEqual([record["number"] for record in records], [7, 6, 5, 4, 3])

    def test_farkli_kaydet_surumu_yeni_proje_dalina_kopyalar(self):
        with tempfile.TemporaryDirectory() as tmp:
            source_path = os.path.join(tmp, "kaynak.json")
            target_path = os.path.join(tmp, "hedef.json")
            surum_kaydi_olustur(source_path, self._veri(derinlik="12"), "V1", keep=10)
            surum_kaydi_olustur(source_path, self._veri(derinlik="15"), "V2", keep=10)

            copied = surum_deposunu_kopyala(source_path, target_path)
            records = surumleri_listele(target_path, eski_yedekleri_aktar=False)

            self.assertTrue(copied)
            self.assertEqual([record["number"] for record in records], [2, 1])
            self.assertEqual(surum_verisi_yukle(target_path, records[0])["sondaj"][0]["der"], "15")

    def test_manuel_kayit_eski_ve_yeni_durumu_gecmiste_korur(self):
        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "proje.json")
            old_data = self._veri(derinlik="12")
            new_data = self._veri(derinlik="18")
            with open(project_path, "w", encoding="utf-8") as handle:
                json.dump(old_data, handle, ensure_ascii=False)

            app = ArayuzProjeMixin()
            app.veri = new_data
            app.aktif_dosya_yolu = project_path
            app.guncelle_veri_objesi = mock.Mock()
            app.proje_kilitli_mi = mock.Mock(return_value=False)
            app.set_status = mock.Mock()
            app.set_save_indicator = mock.Mock()
            app.recent_project_ekle = mock.Mock()

            self.assertTrue(app.veri_kaydet())
            first_count = len(surumleri_listele(project_path, eski_yedekleri_aktar=False))
            self.assertTrue(app.veri_kaydet())
            second_count = len(surumleri_listele(project_path, eski_yedekleri_aktar=False))
            depths = {
                surum_verisi_yukle(project_path, record)["sondaj"][0]["der"]
                for record in surumleri_listele(project_path, eski_yedekleri_aktar=False)
            }

            self.assertEqual(first_count, 2)
            self.assertEqual(second_count, 2)
            self.assertEqual(depths, {"12", "18"})

    def test_eski_surumu_yuklemek_dosyayi_degistirmeden_projeyi_kirli_yapar(self):
        class VersionApp(ArayuzProjeMixin, ProjeSurumleriMixin):
            pass

        with tempfile.TemporaryDirectory() as tmp:
            project_path = os.path.join(tmp, "proje.json")
            old_data = self._veri(derinlik="12")
            current_data = self._veri(derinlik="18")
            with open(project_path, "w", encoding="utf-8") as handle:
                json.dump(current_data, handle, ensure_ascii=False)
            old_record, _created = surum_kaydi_olustur(project_path, old_data, "Eski", keep=10)
            surum_kaydi_olustur(project_path, current_data, "Guncel", keep=10)

            app = VersionApp()
            app.veri = copy.deepcopy(current_data)
            app.aktif_dosya_yolu = project_path
            app._son_kayit_imzasi = app.proje_kayit_imzasi(current_data)
            app.guncelle_veri_objesi = mock.Mock()
            app.doldur_arayuz = mock.Mock()
            app.proje_baslik_guncelle = mock.Mock()
            app.set_status = mock.Mock()
            app.set_save_indicator = mock.Mock()

            loaded = app.proje_surumunu_calisma_alanina_yukle(old_record)

            self.assertEqual(loaded["sondaj"][0]["der"], "12")
            self.assertEqual(app.aktif_dosya_yolu, project_path)
            self.assertTrue(app.proje_degisti_mi())
            with open(project_path, "r", encoding="utf-8") as handle:
                self.assertEqual(json.load(handle)["sondaj"][0]["der"], "18")


class YardimciFonksiyonTestleri(unittest.TestCase):
    def test_buton_ikon_eslemesi_yenile_ve_yeni_metinlerini_ayirir(self):
        icons = IconManager()

        self.assertEqual(icons.guess_key("Yenile"), "refresh")
        self.assertEqual(icons.guess_key("Yeni Proje"), "plus")

    def test_dahili_rapor_sablonu_okunur_ve_temel_etiketleri_icerir(self):
        path = dahili_rapor_sablonu_yolu()

        self.assertTrue(os.path.isfile(path))
        analysis = analyze_word_template(path)
        self.assertIsNone(analysis["error"])
        self.assertIn("[PROJE_ADI]", analysis["tags"])
        self.assertGreater(len(analysis["tags"]), 20)

    def test_gecersiz_ozel_sablon_dahili_sablona_doner(self):
        status = rapor_sablonu_durumu(r"Z:\bulunmayan\rapor_sablonu.docx")

        self.assertTrue(status["ready"])
        self.assertTrue(status["fallback"])
        self.assertEqual(status["source"], "builtin")
        self.assertEqual(status["path"], dahili_rapor_sablonu_yolu())

    def test_gecerli_ozel_sablon_dahili_sablondan_once_gelir(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            custom_path = os.path.join(tmp, "ozel.docx")
            Document().save(custom_path)

            status = rapor_sablonu_durumu(custom_path)
            self.assertEqual(status["source"], "custom")
            self.assertFalse(status["fallback"])
            self.assertEqual(etkin_rapor_sablonu_yolu(custom_path), os.path.abspath(custom_path))

    def test_serim_adi_haritada_ss_bicimine_donusur(self):
        self.assertEqual(ss_harita_etiketi("Serim 4", 0), "SS-4")
        self.assertEqual(ss_harita_etiketi("SS 7", 0), "SS-7")
        self.assertEqual(ss_harita_etiketi("", 2), "SS-3")

    def test_atomic_json_dump_guvenli_yazar_ve_degistirir(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "proje.json")
            atomic_json_dump({"ad": "İlk"}, path, indent=2, ensure_ascii=False)
            atomic_json_dump({"ad": "Son", "sayi": 2}, path, indent=2, ensure_ascii=False)

            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)

            self.assertEqual(data, {"ad": "Son", "sayi": 2})
            self.assertFalse([name for name in os.listdir(tmp) if name.endswith(".tmp")])

    def test_atomic_write_text_gecici_dosya_birakmaz(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "biten_isler.kml")
            atomic_write_text(path, "ilk", encoding="utf-8")
            atomic_write_text(path, "son", encoding="utf-8")

            with open(path, "r", encoding="utf-8") as f:
                self.assertEqual(f.read(), "son")
            self.assertFalse([name for name in os.listdir(tmp) if name.endswith(".tmp")])

    def test_atomic_docx_save_hata_halinde_mevcut_dosyayi_korur(self):
        class BrokenDocument:
            def save(self, path):
                Path(path).write_bytes(b"yarim")
                raise RuntimeError("kontrollu kayit hatasi")

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "rapor.docx"
            path.write_bytes(b"mevcut-rapor")

            with self.assertRaises(RuntimeError):
                atomic_docx_save(BrokenDocument(), path)

            self.assertEqual(path.read_bytes(), b"mevcut-rapor")
            self.assertFalse([name for name in os.listdir(tmp) if name != "rapor.docx"])

    def test_kullanici_verisi_appdata_benzeri_dizine_tasinir(self):
        with tempfile.TemporaryDirectory() as tmp:
            legacy = Path(tmp) / "eski.json"
            legacy.write_text('{"deger": 1}', encoding="utf-8")
            target_root = Path(tmp) / "AppData" / "RaporPro"

            target = kullanici_yolu(
                "recent_projects.json",
                legacy=legacy,
                base_dir=target_root,
            )

            self.assertEqual(target, target_root / "recent_projects.json")
            self.assertEqual(target.read_text(encoding="utf-8"), '{"deger": 1}')
            self.assertFalse(eski_veriyi_kopyala(legacy, target))

    def test_safe_float_virgul_ve_bos_deger(self):
        self.assertEqual(safe_float("12,5"), 12.5)
        self.assertEqual(safe_float(""), 0.0)

    def test_dpapi_gizli_deger_roundtrip(self):
        encrypted = gizli_deger_sakla("test-api-key")
        self.assertTrue(gizli_deger_mi(encrypted))
        self.assertNotIn("test-api-key", encrypted)
        self.assertEqual(gizli_deger_coz(encrypted), "test-api-key")

    def test_hgm_url_kaynak_kod_disinda_sifreli_saklanir(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "harita_ayarlar.json"
            url = "https://ornek.test/{z}/{x}/{y}?apikey=secret"
            hgm_ortofoto_url_kaydet(url, path)

            self.assertEqual(hgm_ortofoto_url_yukle(path), url)
            self.assertNotIn("secret", path.read_text(encoding="utf-8"))

    def test_mjh_etiketi_baska_haritalara_dusmez(self):
        with tempfile.TemporaryDirectory() as tmp:
            yer = os.path.join(tmp, "yer.jpg")
            tkgm = os.path.join(tmp, "tkgm.jpg")
            Path(yer).write_bytes(b"yer")
            Path(tkgm).write_bytes(b"tkgm")
            app = SimpleNamespace(img_mjh=None, img_yer=yer, img_tkgm=tkgm)

            self.assertIsNone(mjh_resim_yolu(app))
            self.assertIsNone(image_path_for_tag(app, "img_mjh"))

            mjh = os.path.join(tmp, "mjh.jpg")
            Path(mjh).write_bytes(b"mjh")
            app.img_mjh = mjh
            self.assertEqual(mjh_resim_yolu(app), mjh)

    def test_lab_sheet_satirlari_temizlenir_ve_algilanir(self):
        app = type("App", (), {})()
        app.veri = {"lab_sheet": {"rows": [["Sondaj No", "", ""], ["SK-1", "CL", ""], ["", "", ""]]}}
        self.assertTrue(lab_sheet_verisi_var_mi(app))
        self.assertEqual(lab_sheet_satirlari(app), [["Sondaj No"], ["SK-1", "CL"]])

    def test_jeofizik_sheet_serim_parametrelerini_okur(self):
        rows = [
            ["Sismik Ölçü ve Hesaplarının Sahibi   :", "Serim 1", "", "", "", ""],
            ["Zemin Parametreleri", "Simge", "TABAKA NO", "", "", ""],
            ["", "Birim", "I", "II", "III", "IV"],
            ["VP = Boyuna Dalga Hızı (Sıkışma Dalgası Hızı)", "VP (m/sn)", 315, 1292, 1292, 1292],
            ["VS = Enine Dalga Hızı (Kayma Dalgası Hızı)", "VS (m/sn)", 145, 280, 363, 477],
            ["Tabaka Kalınlığı", "metre", 3.3, 11.2, 15.5, ""],
            ["Tabaka Yoğunluğu", "γ (gr/cm3)", 1.30, 1.85, 1.88, 1.92],
            ["Poisson Oranı", "Birimsiz", 0.36, 0.46, 0.46, 0.42],
            ["Elastisite Modülü", "ED (kg/cm2)", 748, 3963, 7029, 12631],
            ["Bulk (Sıkışmazlık) Modülü", "Mc (kg/cm2)", 928, 29036, 27716, 25347],
            ["Maksimum Kayma Modülü (SHEAR) (Gmax.)", "Gmax(kg/cm2)", 274, 1454, 2445, 4222],
            ["Zemin Büyütme Katsayısı", "b", 2.87, 1.05, 0.99, 0.92],
            ["Sismik Hız Oranı", "Vp/Vs", 2.17, 4.61, 3.56, 2.71],
            ["Vs30 = Ortalama Kayma Dalgası Hızı", "Vs30(m/sn)", 285, "", "", ""],
        ]
        veri = {"jeofizik_sheet": {"rows": rows}}
        self.assertTrue(jeofizik_sheet_var_mi(veri))
        summary = jeofizik_sheet_ozeti(veri)
        self.assertEqual(summary["serim"], 1)
        self.assertEqual(summary["layers"], 4)

        ss_list = jeofizik_sheet_rows_to_ss_list(rows)
        self.assertEqual(ss_list[0]["ad"], "Serim 1")
        self.assertEqual(ss_list[0]["layers"][0]["vp"], "315")
        self.assertEqual(ss_list[0]["layers"][0]["vs"], "145")
        self.assertEqual(ss_list[0]["layers"][0]["G"], "274")
        self.assertEqual(ss_list[0]["layers"][0]["vs30"], "285")
        self.assertEqual(ss_list[0]["layers"][3]["h"], "-")

        jeofizik_ss_koordinatlarini_koru(ss_list, [{"ad": "SS-1", "coords": ["1", "2", "3", "4", "5", "6"]}])
        self.assertEqual(ss_list[0]["coords"], ["1", "2", "3", "4", "5", "6"])

    def test_buyuk_basliklar_yeni_sayfaya_alinir(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Kapak")
        heading = doc.add_paragraph("1. GENEL BİLGİLER")
        normal = doc.add_paragraph("Bu normal bir paragraftır.")
        styled = doc.add_paragraph("İKİNCİ BÜYÜK BAŞLIK")
        styled.style = "Heading 1"

        count = buyuk_basliklari_yeni_sayfaya_al(doc)

        self.assertEqual(count, 2)
        self.assertTrue(heading.paragraph_format.page_break_before)
        self.assertFalse(bool(normal.paragraph_format.page_break_before))
        self.assertTrue(styled.paragraph_format.page_break_before)

    def test_duzeltme_etiket_sablonu_secili_etiketleri_yazar(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "duzeltme.docx")
            selected = duzeltme_etiket_sablonu_olustur(["[SPT]", " [PMT] ", "[SPT]"], path)

            self.assertEqual(selected, ["[SPT]", "[PMT]"])
            self.assertEqual(duzeltme_etiketleri_temizle(["", "[SPT]", "[SPT]"]), ["[SPT]"])

            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("RAPOR DÜZELTME ETİKET ÇIKTISI", text)
            self.assertIn("SPT tablosu", text)
            self.assertIn("[SPT]", text)
            self.assertIn("[PMT]", text)

    def test_belediye_duzeltme_kural_analizi_sondaj_etiketi_onerir(self):
        metin = "Fotoğraf kayıtlarında SK-15 görünen kuyu aslında SK-13 tür. Sondajlar başlığı altında açıklama eklensin."
        result = duzeltme_metnini_kural_ile_analiz_et(metin)

        self.assertEqual(result["source"], "kural")
        self.assertIn("[Sondaj]", result["tags"])
        self.assertTrue(result["items"])

    def test_belediye_duzeltme_asistani_ai_kapatilinca_kural_kullanir(self):
        metin = "SPT tablosu ve litoloji dağılımında Çakıllı Siltli Kum birimi kontrol edilsin."
        result = belediye_duzeltme_analiz_et(metin, ai_kullan=False)

        self.assertEqual(result["source"], "kural")
        self.assertIn("[SPT]", result["tags"])
        self.assertIn("[LITOLOJI_DAGILIM]", result["tags"])

    def test_belediye_duzeltme_yonlendirmesi_ek_sondaj_ve_lab_algilar(self):
        sondaj = duzeltme_yonlendirmeleri_olustur("Belediye ek sondaj yapılmasını istedi.")
        self.assertTrue(any(item["id"] == "ek_sondaj" and item["action_key"] == "sondaj_hizli" for item in sondaj))

        lab = duzeltme_yonlendirmeleri_olustur("Ek laboratuvar deneyi yükle ve rapora işle.")
        self.assertTrue(any(item["id"] == "ek_laboratuvar" and item["action_key"] == "lab_excel" for item in lab))

    def test_belediye_duzeltme_yonlendirmesi_lab_notunda_sadece_lab_onerir(self):
        metin = (
            "5-) Lab deneyleri; 24.00 ile 28.50m arasi deney yok. "
            "2m de bir deney olmasi gerekmektedir. Dolayisiyla 25.50, 27.00 ve 30.00m "
            "seviyelerinde deney eklenmesi gerekmektedir."
        )
        items = duzeltme_yonlendirmeleri_olustur(metin)
        ids = [item["id"] for item in items]

        self.assertEqual(ids, ["ek_laboratuvar"])
        self.assertIn("24.00", items[0]["source_text"])
        self.assertIn("deney", [kw.lower() for kw in items[0]["matched_keywords"]])

    def test_rapor_revizyon_isaretli_bolumu_degistirir(self):
        from docx import Document

        target = Document()
        target.add_paragraph("Önce")
        target_p = target.add_paragraph("[SPT]")
        target.add_paragraph("Sonra")
        revizyon_isaretleri_ekle(target, {"[SPT]": target_p})
        target_p.text = "Eski SPT tablosu"

        source = Document()
        source_p = source.add_paragraph("[SPT]")
        revizyon_isaretleri_ekle(source, {"[SPT]": source_p})
        source_p.text = "Yeni SPT tablosu"

        self.assertTrue(revizyon_etiketi_var_mi(target, "[SPT]"))
        updated, missing = docx_bolumlerini_degistir(target, source, ["[SPT]"])

        text = "\n".join(paragraph.text for paragraph in target.paragraphs)
        self.assertEqual(updated, ["[SPT]"])
        self.assertEqual(missing, [])
        self.assertIn("Yeni SPT tablosu", text)
        self.assertNotIn("Eski SPT tablosu", text)

    def test_rapor_metin_revizyonu_sondaj_ifadesini_bulup_uygular(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            source_path = os.path.join(tmp, "rapor.docx")
            output_path = os.path.join(tmp, "rapor_revize.docx")

            doc = Document()
            doc.add_paragraph("Sondaj fotoğraflarında SK-15 olarak görülen kuyu rapora işlenmiştir.")
            doc.save(source_path)

            units = word_metin_birimleri_oku(source_path)
            result = metin_revizyon_kural_analiz_et(
                "SK-15 olarak görünen kuyu aslında SK-13 tür.",
                units,
            )

            self.assertEqual(result["source"], "kural")
            self.assertTrue(result["items"])
            self.assertEqual(result["items"][0]["old_text"], "SK-15")
            self.assertEqual(result["items"][0]["new_text"], "SK-13")

            info = metin_revizyonlari_uygula(source_path, result["items"], output_path)
            self.assertTrue(info["success"])

            revised = Document(output_path)
            text = "\n".join(p.text for p in revised.paragraphs)
            self.assertIn("SK-13 olarak görülen", text)
            self.assertNotIn("SK-15", text)

    def test_rapor_metin_revizyonu_tablo_hucresini_okur(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            source_path = os.path.join(tmp, "rapor.docx")
            output_path = os.path.join(tmp, "rapor_revize.docx")
            doc = Document()
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Kontrol notu: eski ifade"
            doc.save(source_path)

            units = word_metin_birimleri_oku(source_path)
            self.assertTrue(any(unit["kind"] == "table" and "eski ifade" in unit["text"] for unit in units))

            result = metin_revizyon_kural_analiz_et('"eski ifade" yerine "yeni ifade" yazılsın', units)
            revisions = [dict(item, unit_id="p:999") for item in result["items"]]
            info = metin_revizyonlari_uygula(source_path, revisions, output_path)
            self.assertTrue(info["success"])

            revised = Document(output_path)
            self.assertIn("yeni ifade", revised.tables[0].cell(0, 0).text)

    def test_zemin_sinifi_yeni_cakil_kisaltmalarini_cevirir(self):
        self.assertEqual(zemin_sinifi_cevir("sasiGrP"), "Kumlu Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("sasiGrW"), "Kumlu Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("sasiGrM"), "Kumlu Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("siGrP"), "Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("siGrW"), "Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("siGrM"), "Siltli Çakıl")
        self.assertEqual(zemin_sinifi_cevir("siclGr"), "Siltli Killi Çakıl")
        self.assertEqual(zemin_sinifi_cevir("saclGr"), "Kumlu Killi Çakıl")
        self.assertEqual(zemin_sinifi_cevir("clGr"), "Killi Çakıl")
        self.assertEqual(zemin_sinifi_cevir("grCIL"), "Çakıllı Kil")
        self.assertEqual(zemin_sinifi_cevir("grCIH"), "Çakıllı Kil")
        self.assertEqual(zemin_sinifi_cevir("grCIM"), "Çakıllı Kil")
        self.assertEqual(zemin_sinifi_cevir("siSaP"), "Siltli Kum")
        self.assertEqual(zemin_sinifi_cevir("siSaW"), "Siltli Kum")
        self.assertEqual(zemin_sinifi_cevir("siSaM"), "Siltli Kum")

    def test_tkgm_konum_adi_normalizasyonu_turkce_ekleri_temizler(self):
        self.assertEqual(konum_adi_normalize_et("Arslanca Mahallesi"), "arslanca")
        self.assertEqual(konum_adi_normalize_et("ÇANAKKALE"), "canakkale")

    def test_tkgm_geojson_polygon_kml_koordinati_uretir(self):
        geometry = {
            "type": "Polygon",
            "coordinates": [
                [
                    [26.0, 40.0],
                    [26.1, 40.0],
                    [26.1, 40.1],
                    [26.0, 40.1],
                    [26.0, 40.0],
                ]
            ],
        }
        kml = geojson_kml_olustur(geometry, name="Deneme Parsel")
        self.assertIn("<Placemark>", kml)
        self.assertIn("26.00000000,40.00000000,0.00", kml)

    def test_litoloji_yazim_uyarisi_yakin_hatalari_bulur(self):
        self.assertTrue(any("kumlu" in item for item in litoloji_yazim_uyarilari("Killi kumluu")))
        self.assertFalse(litoloji_yazim_uyarilari("Az killi kum"))

    def test_workbook_litoloji_tanim_yazim_uyarisi_verir(self):
        rows = {
            "sondajlar": [{"no": "SK-1", "der": "10"}],
            "litoloji": [{"sondaj_no": "SK-1", "top": "0", "bot": "2", "tanim": "Killi kumluu"}],
            "spt": [],
            "pmt": [],
            "kaya": [],
            "numune": [],
        }
        result = wb_validate_rows(rows)
        self.assertIn(("litoloji", 0, "tanim"), result["warnings"])

    def test_workbook_sondaj_turu_ve_delgi_capi_proje_ayarindan_gelir(self):
        veri = {
            "ayarlar": {"sondaj_turu": "Kaya", "delgi_capi": "89mm"},
            "sondaj": [{"no": "SK-1", "der": "12", "kaya": []}],
        }
        initial, source_nos = wb_build_initial_rows(veri)
        columns = [key for _, key in WORKBOOK_SHEET_DEFS["sondajlar"]["columns"]]
        self.assertNotIn("sondaj_turu", columns)
        self.assertNotIn("delgi_capi", columns)
        headers = initial["sondajlar"][0]
        self.assertEqual(len(headers), len(columns))

        rows = {
            "sondajlar": [{"no": "SK-1", "der": "12"}],
            "litoloji": [],
            "spt": [],
            "pmt": [],
            "kaya": [],
            "numune": [],
        }
        sondajlar, warnings = wb_apply_rows_to_veri(veri, rows, source_nos)
        self.assertFalse(warnings)
        self.assertEqual(sondajlar[0]["sondaj_turu"], "Kaya")
        self.assertEqual(sondajlar[0]["delgi_capi"], "89mm")

    def test_workbook_excel_yazici_tum_sayfalari_olusturur(self):
        payloads = [
            {
                "title": "Sondajlar",
                "headers": ["SondajNo", "Derinlik"],
                "rows": [["SK-1", "15.0"]],
                "widths": [110, 85],
            },
            {
                "title": "SPT",
                "headers": ["SondajNo", "Derinlik", "N30"],
                "rows": [["SK-1", "1.50", "12"], ["SK-1", "3.00", "18"]],
                "widths": [110, 90, 70],
            },
        ]
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "workbook.xlsx")
            result = wb_excel_workbook_yaz(path, payloads)
            workbook = load_workbook(path, data_only=True)

        self.assertEqual(workbook.sheetnames, ["Sondajlar", "SPT"])
        self.assertEqual(workbook["Sondajlar"]["A2"].value, "SK-1")
        self.assertEqual(workbook["SPT"]["C3"].value, "18")
        self.assertEqual(result["sheet_count"], 2)
        self.assertEqual(result["row_count"], 3)

    def test_sondaj_kaydet_tabloyu_yeniden_cizmeden_veriyi_gunceller(self):
        mixin = SondajMixin.__new__(SondajMixin)
        mixin.veri = {"sondaj": [{"no": "SK-1", "der": "10"}]}
        mixin.sondaj_ui_rows = [{
            "no": SimpleNamespace(get=lambda: "SK-2"),
            "der": SimpleNamespace(get=lambda: "18.0"),
        }]
        mixin.sondaj_satir_durumlarini_yenile = mock.Mock()
        mixin.sondaj_tablosunu_ciz = mock.Mock()
        mixin.set_status = mock.Mock()

        mixin.sondaj_verilerini_kaydet()

        self.assertEqual(mixin.veri["sondaj"][0]["no"], "SK-2")
        self.assertEqual(mixin.veri["sondaj"][0]["der"], "18.0")
        mixin.sondaj_satir_durumlarini_yenile.assert_called_once_with()
        mixin.sondaj_tablosunu_ciz.assert_not_called()

    def test_sondaj_satir_durumu_litoloji_eksigini_uyari_olarak_gosterir(self):
        mixin = SondajMixin.__new__(SondajMixin)

        state, message = mixin.sondaj_satir_genel_durumu({"no": "SK-1", "der": "15", "litoloji": []})

        self.assertEqual(state, "warning")
        self.assertEqual(message, "Litoloji eksik")

    def test_sondaj_satir_durumu_gecerli_litoloji_ile_hazir_olur(self):
        mixin = SondajMixin.__new__(SondajMixin)
        sondaj = {
            "no": "SK-1",
            "der": "15",
            "litoloji": [["0", "15", "Kil"]],
            "spt": [],
            "pmt": [],
            "kaya": [],
        }

        state, message = mixin.sondaj_satir_genel_durumu(sondaj)

        self.assertEqual(state, "ok")
        self.assertEqual(message, "Hazır")

    def test_sondaj_satir_durumu_gecersiz_koordinati_uyari_olarak_gosterir(self):
        mixin = SondajMixin.__new__(SondajMixin)
        sondaj = {
            "no": "SK-1",
            "der": "15",
            "y": "140.25",
            "litoloji": [["0", "15", "Kil"]],
        }

        state, message = mixin.sondaj_satir_genel_durumu(sondaj)

        self.assertEqual(state, "warning")
        self.assertEqual(message, "Koordinat geçersiz")

    def test_tamamlama_merkezi_eski_final_kontrol_girisini_korur(self):
        mixin = KontrolPaneliMixin.__new__(KontrolPaneliMixin)
        mixin.final_kontrol_penceresi = mock.Mock(return_value="window")

        result = mixin.tamamlama_merkezi_penceresi()

        self.assertEqual(result, "window")
        mixin.final_kontrol_penceresi.assert_called_once_with()

    def test_spt_helper_klasorden_resimleri_toplar(self):
        with tempfile.TemporaryDirectory() as tmp:
            nested = os.path.join(tmp, "alt")
            os.makedirs(nested)
            image_a = os.path.join(tmp, "a.JPG")
            image_b = os.path.join(nested, "b.png")
            text_file = os.path.join(tmp, "not.txt")
            for path in (image_a, image_b, text_file):
                with open(path, "w", encoding="utf-8") as f:
                    f.write("x")

            shallow = collect_image_paths([tmp], recursive=False)
            recursive = collect_image_paths([tmp], recursive=True)

            self.assertEqual([os.path.basename(path) for path in shallow], ["a.JPG"])
            self.assertEqual(sorted(os.path.basename(path) for path in recursive), ["a.JPG", "b.png"])

    def test_spt_helper_tekrar_ve_kalite_uyarisi_uret(self):
        kayit1 = SPTKaydi(sondaj_no="SK-1", derinlik="1.50", v15="2", v30="3", v45="4", n30="7", guven="95")
        kayit2 = SPTKaydi(sondaj_no="SK-1", derinlik="1.50", v15="2", v30="3", v45="4", n30="7", guven="95")
        records = [{"kayit": kayit1, "include": True}, {"kayit": kayit2, "include": True}]

        self.assertIn(("SK-1", 1.5), spt_duplicate_keys(records))
        quality = spt_record_quality(
            {"kayit": kayit1, "include": True},
            duplicate=True,
            current_sondaj_depth=lambda _no: 1.0,
            valid_sondaj_nolari={"SK-1"},
            settings={"guven_esigi": 90},
        )
        self.assertEqual(quality["level"], "warning")
        self.assertIn("aynı derinlik", quality["message"])
        self.assertIn("sondaj derinliğini geçiyor", quality["message"])


class KesitCizimTestleri(unittest.TestCase):
    def test_kesit_kayit_dosya_adi_sondaj_araligini_uret(self):
        self.assertEqual(kesit_kayit_dosya_adi(["SK-1", "SK-2", "SK-3"]), "Kesit SK1-3")
        self.assertEqual(kesit_kayit_dosya_adi(["SK-1", "SK-3", "SK-5"]), "Kesit SK1-SK3-SK5")
        self.assertEqual(kesit_kayit_dosya_adi([{"no": "BH-01"}, {"no": "BH-02"}]), "Kesit BH1-2")

    def test_kesit_imzasi_manual_duzenlemeyi_ayirir(self):
        mixin = KesitCizimMixin()
        options = {
            "mode": "schematic",
            "selected_sondajlar": ["SK-1", "SK-2"],
            "vertical_exaggeration": "1.0",
            "corr_tolerance": "3.0",
            "dx_default": "25.0",
            "well_width": "2.0",
        }
        signature = mixin._kesit_section_signature(options)
        saved = {"manual_edits_by_section": {signature: {"p1": [[0, 0], [1, 0], [1, 1]]}}}
        self.assertIn("p1", mixin._kesit_manual_edits_for_options(saved, options))

        other_options = dict(options)
        other_options["selected_sondajlar"] = ["SK-2", "SK-3"]
        self.assertEqual(mixin._kesit_manual_edits_for_options(saved, other_options), {})

    def test_kesit_motoru_tekrarli_birimleri_ayri_eslestirir(self):
        sondajlar = [
            {
                "no": "SK-1",
                "k": "100",
                "der": "6",
                "litoloji": [["0", "2", "Kil"], ["2", "3", "Kum"], ["3", "6", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-2",
                "k": "100",
                "der": "6",
                "litoloji": [["0", "1.8", "Kil"], ["1.8", "3.1", "Kum"], ["3.1", "6", "Kil"]],
                "spt": [],
            },
        ]
        fig, _ = GeoEngine.kesit_ciz_interaktif(
            sondajlar,
            options={
                "mode": "schematic",
                "show_legend": False,
                "show_yass": False,
                "show_distance_labels": False,
                "show_layer_depth_labels": False,
                "show_consistency_labels": False,
                "corr_tolerance": "3.0",
            },
        )
        tool = getattr(fig, "_geo_tool", None)
        self.assertIsNotNone(tool)
        self.assertIn("İZLEME", tool.info_text.get_text())
        tool.set_edit_mode(True)
        self.assertIn("Polygon seçmek", tool.info_text.get_text())

        edit_ids = {
            getattr(poly, "_geo_edit_id", "")
            for poly in tool.polygons
        }
        self.assertIn("match:SK-1:SK-2:0:0:kl", edit_ids)
        self.assertIn("match:SK-1:SK-2:1:1:k", edit_ids)
        self.assertIn("match:SK-1:SK-2:2:2:kl", edit_ids)
        first_section_poly = next(poly for poly in tool.polygons if getattr(poly, "_geo_poly_kind", "") != "well")
        tool.select_polygon(first_section_poly)
        self.assertIn("DÜZENLEME", tool.info_text.get_text())
        self.assertIn("seçili", tool.info_text.get_text())

    def test_kesit_motoru_kesit_hatti_istasyon_hesaplar(self):
        sondajlar = [
            {"no": "SK-1", "y": "41.0000", "x": "29.0000", "k": "100", "der": "5", "litoloji": [["0", "5", "Kil"]], "spt": []},
            {"no": "SK-2", "y": "41.0001", "x": "29.0002", "k": "99", "der": "5", "litoloji": [["0", "5", "Kum"]], "spt": []},
            {"no": "SK-3", "y": "41.0002", "x": "29.0004", "k": "98", "der": "5", "litoloji": [["0", "5", "Kil"]], "spt": []},
        ]
        messages = []
        fig, _ = GeoEngine.kesit_ciz_interaktif(
            sondajlar,
            log_callback=lambda message, level="info": messages.append((level, message)),
            options={
                "mode": "line_projection",
                "line_start_y": "41.0000",
                "line_start_x": "29.0000",
                "line_end_y": "41.0002",
                "line_end_x": "29.0004",
                "line_start_no": "SK-1",
                "line_end_no": "SK-3",
                "max_offset": "10",
                "show_legend": False,
                "show_yass": False,
                "show_distance_labels": False,
                "show_layer_depth_labels": False,
                "show_consistency_labels": False,
            },
        )
        self.assertIn("Kesit hatti: SK-1 - SK-3", fig.axes[0].get_title())
        self.assertAlmostEqual(sondajlar[0]["_station"], 0.0, places=3)
        self.assertGreater(sondajlar[1]["_station"], sondajlar[0]["_station"])
        self.assertGreater(sondajlar[2]["_station"], sondajlar[1]["_station"])
        self.assertLess(abs(sondajlar[1]["_offset"]), 0.05)
        self.assertFalse(messages)

    def test_haritadan_kesit_hatti_yakin_sondajlari_sirali_secer(self):
        sondajlar = [
            {"no": "SK-1", "y": "41.0000", "x": "29.0000"},
            {"no": "SK-2", "y": "41.0001", "x": "29.0002"},
            {"no": "SK-3", "y": "41.0002", "x": "29.0004"},
            {"no": "SK-UZAK", "y": "41.0015", "x": "29.0000"},
        ]
        selected = kesit_hatti_sondaj_sirasi(
            sondajlar,
            start=(41.0000, 29.0000),
            end=(41.0002, 29.0004),
            max_offset=12,
        )
        self.assertEqual([item["no"] for item in selected], ["SK-1", "SK-2", "SK-3"])
        self.assertLess(abs(selected[1]["offset"]), 0.05)

    def test_kesit_motoru_tek_kuyudaki_ince_tabakayi_mercek_cizer(self):
        sondajlar = [
            {
                "no": "SK-1",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-2",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "2", "Kil"], ["2", "3", "Kum"], ["3", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-3",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
        ]
        common_options = {
            "mode": "schematic",
            "show_legend": False,
            "show_yass": False,
            "show_distance_labels": False,
            "show_layer_depth_labels": False,
            "show_consistency_labels": False,
            "lens_max_thickness": "2.0",
        }
        fig, _ = GeoEngine.kesit_ciz_interaktif(sondajlar, options={**common_options, "auto_lens": True})
        edit_ids = {getattr(poly, "_geo_edit_id", "") for poly in fig._geo_tool.polygons}
        lens_ids = {item for item in edit_ids if item.startswith("lens:")}
        self.assertIn("lens:SK-1:SK-2:SK-3:1:k", edit_ids)
        self.assertIn("lens-host:SK-2:SK-1:left:1:kl", edit_ids)
        self.assertIn("lens-host:SK-2:SK-3:right:1:kl", edit_ids)
        self.assertEqual(lens_ids, {"lens:SK-1:SK-2:SK-3:1:k"})
        self.assertNotIn("match:SK-1:SK-2:0:0:kl", edit_ids)
        self.assertNotIn("match:SK-2:SK-3:0:0:kl", edit_ids)
        self.assertFalse(any(item.startswith("pinch-left:SK-2:SK-3:1:k") for item in edit_ids))
        self.assertFalse(any(item.startswith("pinch-right:SK-1:SK-2:1:k") for item in edit_ids))
        self.assertFalse(any(item.startswith("pinch-right:SK-1:SK-2:2:kl") for item in edit_ids))
        self.assertFalse(any(item.startswith("pinch-left:SK-2:SK-3:2:kl") for item in edit_ids))

        fig_without_lens, _ = GeoEngine.kesit_ciz_interaktif(sondajlar, options={**common_options, "auto_lens": False})
        edit_ids_without_lens = {getattr(poly, "_geo_edit_id", "") for poly in fig_without_lens._geo_tool.polygons}
        self.assertNotIn("lens:SK-1:SK-2:SK-3:1:k", edit_ids_without_lens)

    def test_kesit_motoru_iki_sondajda_yarim_mercek_cizer(self):
        sondajlar = [
            {
                "no": "SK-1",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "2", "Kil"], ["2", "3", "Kum"], ["3", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-2",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
        ]
        common_options = {
            "mode": "schematic",
            "show_legend": False,
            "show_yass": False,
            "show_distance_labels": False,
            "show_layer_depth_labels": False,
            "show_consistency_labels": False,
            "lens_max_thickness": "2.0",
            "auto_lens": True,
        }
        fig, _ = GeoEngine.kesit_ciz_interaktif(sondajlar, options={**common_options, "two_well_lens": True})
        edit_ids = {getattr(poly, "_geo_edit_id", "") for poly in fig._geo_tool.polygons}
        self.assertIn("half-lens:SK-1:SK-2:right:1:k", edit_ids)
        self.assertIn("lens-host:SK-1:SK-2:right:1:kl", edit_ids)
        polygons_by_id = {getattr(poly, "_geo_edit_id", ""): poly for poly in fig._geo_tool.polygons}
        lens_poly = polygons_by_id["half-lens:SK-1:SK-2:right:1:k"]
        host_poly = polygons_by_id["lens-host:SK-1:SK-2:right:1:kl"]
        host_pattern_zorders = [artist.get_zorder() for artist in getattr(host_poly, "_geo_pattern_artists", [])]
        lens_pattern_zorders = [artist.get_zorder() for artist in getattr(lens_poly, "_geo_pattern_artists", [])]
        self.assertGreater(lens_poly.get_zorder(), host_poly.get_zorder())
        self.assertEqual(lens_poly.get_alpha(), 1.0)
        self.assertTrue(host_pattern_zorders)
        self.assertTrue(lens_pattern_zorders)
        self.assertLess(max(host_pattern_zorders), lens_poly.get_zorder())
        self.assertGreater(min(lens_pattern_zorders), max(host_pattern_zorders))
        fig._geo_tool.refresh_pattern(lens_poly)
        refreshed_lens_pattern_zorders = [artist.get_zorder() for artist in getattr(lens_poly, "_geo_pattern_artists", [])]
        self.assertTrue(refreshed_lens_pattern_zorders)
        self.assertGreater(min(refreshed_lens_pattern_zorders), max(host_pattern_zorders))
        self.assertNotIn("match:SK-1:SK-2:0:0:kl", edit_ids)
        self.assertFalse(any(item.startswith("pinch-left:SK-1:SK-2:1:k") for item in edit_ids))
        self.assertFalse(any(item.startswith("pinch-left:SK-1:SK-2:2:kl") for item in edit_ids))

        fig_without_half_lens, _ = GeoEngine.kesit_ciz_interaktif(sondajlar, options={**common_options, "two_well_lens": False})
        edit_ids_without_half_lens = {getattr(poly, "_geo_edit_id", "") for poly in fig_without_half_lens._geo_tool.polygons}
        self.assertNotIn("half-lens:SK-1:SK-2:right:1:k", edit_ids_without_half_lens)
        self.assertTrue(any(item.startswith("pinch-left:SK-1:SK-2:1:k") for item in edit_ids_without_half_lens))

    def test_kesit_hattinda_coklu_sondajda_uc_mercek_cizer(self):
        sondajlar = [
            {
                "no": "SK-1",
                "x": "29.0000",
                "y": "41.0000",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "2", "Kil"], ["2", "3", "Kum"], ["3", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-2",
                "x": "29.0002",
                "y": "41.0001",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-3",
                "x": "29.0004",
                "y": "41.0002",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
        ]
        fig, _ = GeoEngine.kesit_ciz_interaktif(
            sondajlar,
            options={
                "mode": "line_projection",
                "line_start": (41.0000, 29.0000),
                "line_end": (41.0002, 29.0004),
                "max_offset": "10",
                "show_legend": False,
                "show_yass": False,
                "show_distance_labels": False,
                "show_layer_depth_labels": False,
                "show_consistency_labels": False,
                "lens_max_thickness": "2.0",
                "auto_lens": True,
                "two_well_lens": True,
            },
        )
        edit_ids = {getattr(poly, "_geo_edit_id", "") for poly in fig._geo_tool.polygons}
        self.assertIn("half-lens:SK-1:SK-2:right:1:k", edit_ids)
        self.assertIn("lens-host:SK-1:SK-2:right:1:kl", edit_ids)
        self.assertFalse(any(item.startswith("pinch-left:SK-1:SK-2:1:k") for item in edit_ids))

    def test_kesit_motoru_onizleme_canli_ayar_gruplarini_etiketler(self):
        sondajlar = [
            {
                "no": "SK-1",
                "k": "100",
                "der": "5",
                "yass_d1": "2.0",
                "litoloji": [["0", "1", "Kil"], ["1", "2", "Kum"], ["2", "5", "Kil"]],
                "spt": [["1.5", "5", "6", "7", "13"]],
            },
            {
                "no": "SK-2",
                "k": "99.7",
                "der": "5",
                "yass_d1": "2.2",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [["1.5", "8", "9", "10", "19"]],
            },
        ]
        fig, _ = GeoEngine.kesit_ciz_interaktif(
            sondajlar,
            options={
                "mode": "schematic",
                "auto_lens": True,
                "two_well_lens": True,
                "show_legend": True,
                "show_yass": True,
                "show_yass_labels": True,
                "show_distance_labels": True,
                "show_layer_depth_labels": True,
                "show_well_elevation_labels": True,
                "show_consistency_labels": True,
            },
        )
        ax = fig.axes[0]
        groups = {getattr(artist, "_geo_live_group", "") for artist in ax.get_children()}
        self.assertIn("station", groups)
        self.assertIn("well_elevation", groups)
        self.assertIn("layer_depth", groups)
        self.assertIn("distance", groups)
        self.assertIn("yass", groups)
        self.assertIn("yass_label", groups)
        self.assertIn("consistency", groups)
        self.assertTrue(any(getattr(artist, "_geo_export_group", None) == "legend" for artist in ax.get_children()))

    def test_kesit_motoru_mercek_gizleme_ve_kum_tarama_ayarini_uygular(self):
        sondajlar = [
            {
                "no": "SK-1",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "2", "Kil"], ["2", "3", "Kum"], ["3", "5", "Kil"]],
                "spt": [],
            },
            {
                "no": "SK-2",
                "k": "100",
                "der": "5",
                "litoloji": [["0", "5", "Kil"]],
                "spt": [],
            },
        ]
        fig, _ = GeoEngine.kesit_ciz_interaktif(
            sondajlar,
            options={
                "mode": "schematic",
                "show_legend": False,
                "show_yass": False,
                "show_distance_labels": False,
                "show_layer_depth_labels": False,
                "show_consistency_labels": False,
                "auto_lens": True,
                "two_well_lens": True,
                "sand_pattern_density": "12.5",
                "manual_edits": {"half-lens:SK-1:SK-2:right:1:k": {"hidden": True}},
            },
        )
        polygons_by_id = {getattr(poly, "_geo_edit_id", ""): poly for poly in fig._geo_tool.polygons}
        lens_poly = polygons_by_id["half-lens:SK-1:SK-2:right:1:k"]
        self.assertFalse(lens_poly.get_visible())
        self.assertTrue(getattr(lens_poly, "_geo_hidden", False))
        self.assertTrue(all(not artist.get_visible() for artist in getattr(lens_poly, "_geo_pattern_artists", [])))
        _, _, density_scale = getattr(lens_poly, "_geo_pattern_info")
        self.assertEqual(density_scale, 12.5)


class SPTMotorTestleri(unittest.TestCase):
    def test_n30_hesapla_ve_refu(self):
        self.assertEqual(n30_hesapla("8", "12"), "20")
        self.assertEqual(n30_hesapla("50/10", "12"), "R")

    def test_sondaj_no_normalize(self):
        self.assertEqual(normalize_sondaj_no("SK=004"), "SK-4")
        self.assertEqual(normalize_sondaj_no("4", "SK-1"), "SK-4")

    def test_hedef_derinlik_yuvarlama(self):
        self.assertEqual(hedef_derinlige_yuvarla("4.52"), "4.50")
        self.assertEqual(hedef_derinlige_yuvarla("105"), "")

    def test_kayit_normalize_aralik_ve_spt_metin(self):
        kayit = kayit_normalize_et({"sondaj_no": "sk4", "derinlik": "10.50-10.95", "spt": "8-9-10"})
        self.assertEqual(kayit.sondaj_no, "SK-4")
        self.assertEqual(kayit.derinlik, "10.50")
        self.assertEqual((kayit.v15, kayit.v30, kayit.v45, kayit.n30), ("8", "9", "10", "19"))

    def test_fotograf_sirasi_cift_satira_dusmez(self):
        path1 = r"C:\tmp\DSCF0001.JPG"
        path2 = r"C:\tmp\DSCF0002.JPG"
        records_by_path = [
            (_path_unique_key(path1), [
                kayit_normalize_et({"sondaj_no": "SK-1", "derinlik": "1.50", "spt": "2-3-4", "kaynak": "DSCF0001.JPG"}),
                kayit_normalize_et({"sondaj_no": "SK-1", "derinlik": "3.00", "spt": "5-6-7", "kaynak": "DSCF0001.JPG"}),
            ]),
            (_path_unique_key(path2), [
                kayit_normalize_et({"sondaj_no": "SK-1", "derinlik": "1.50", "spt": "2-3-4", "kaynak": "DSCF0002.JPG"}),
                kayit_normalize_et({"sondaj_no": "SK-1", "derinlik": "3.00", "spt": "5-6-7", "kaynak": "DSCF0002.JPG"}),
            ]),
        ]
        selected, removed_by_sequence, merged_by_location = _select_spt_records_for_batch(records_by_path, [path1, path2])
        self.assertEqual([item.derinlik for item in selected], ["1.50", "3.00"])
        self.assertEqual(removed_by_sequence, 2)
        self.assertEqual(merged_by_location, 0)

    def test_openai_model_ayarlari_ayri_secilir(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "ayarlar.json")
            ayarlar = spt_ayarlarini_yukle(path)
            self.assertEqual(openai_model_sec(ayarlar, "spt"), "gpt-4o-mini")
            self.assertEqual(openai_model_sec(ayarlar, "revizyon"), "gpt-5.5")
            self.assertEqual(openai_model_sec({"openai_model": "", "revizyon_openai_model": "  "}, "revizyon"), "gpt-5.5")

            spt_ayarlarini_kaydet(
                {"openai_model": "gpt-4o-mini-test", "revizyon_openai_model": "gpt-5.5-test"},
                path,
            )
            ayarlar = spt_ayarlarini_yukle(path)
            self.assertEqual(openai_model_sec(ayarlar, "spt"), "gpt-4o-mini-test")
            self.assertEqual(openai_model_sec(ayarlar, "revizyon"), "gpt-5.5-test")

    def test_api_anahtari_ayar_dosyasinda_sifreli_saklanir(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "ayarlar.json")
            spt_ayarlarini_kaydet({"openai_api_key": "secret-openai-key"}, path)

            raw = Path(path).read_text(encoding="utf-8")
            self.assertNotIn("secret-openai-key", raw)
            self.assertTrue(gizli_deger_mi(json.loads(raw)["openai_api_key"]))
            self.assertEqual(spt_ayarlarini_yukle(path)["openai_api_key"], "secret-openai-key")

    def test_eski_spt_ayar_dosyasi_merkezi_ayara_tasinir(self):
        with tempfile.TemporaryDirectory() as tmp:
            eski_yol = spt_motoru.LEGACY_SPT_AYARLAR_PATH
            yeni_yol = spt_motoru.SPT_AYARLAR_PATH
            try:
                spt_motoru.LEGACY_SPT_AYARLAR_PATH = Path(tmp) / "SPT Okuma" / "ayarlar.json"
                spt_motoru.SPT_AYARLAR_PATH = Path(tmp) / "RaporPro" / "ayarlar.json"
                spt_motoru.LEGACY_SPT_AYARLAR_PATH.parent.mkdir(parents=True, exist_ok=True)
                atomic_json_dump(
                    {"aktif_motor": "openai", "openai_api_key": "test-key", "revizyon_openai_model": "gpt-5.5-test"},
                    spt_motoru.LEGACY_SPT_AYARLAR_PATH,
                    ensure_ascii=False,
                    indent=2,
                )

                ayarlar = spt_motoru.spt_ayarlarini_yukle()

                self.assertEqual(ayarlar["openai_api_key"], "test-key")
                self.assertEqual(ayarlar["revizyon_openai_model"], "gpt-5.5-test")
                self.assertTrue(spt_motoru.SPT_AYARLAR_PATH.exists())
                raw = spt_motoru.SPT_AYARLAR_PATH.read_text(encoding="utf-8")
                self.assertNotIn("test-key", raw)
                self.assertTrue(gizli_deger_mi(json.loads(raw)["openai_api_key"]))
            finally:
                spt_motoru.LEGACY_SPT_AYARLAR_PATH = eski_yol
                spt_motoru.SPT_AYARLAR_PATH = yeni_yol


class PMTExcelMotorTestleri(unittest.TestCase):
    def test_presiyometre_excelinden_pmt_kaydi_okunur_ve_aktarilir(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "SK-1-9_0M-SIM-KELIME-SILINMIS.xlsm")
            wb = Workbook()
            ws = wb.active
            ws.title = "Presiyometre"
            ws["F7"] = "SK-1"
            ws["F8"] = 9
            ws["I6"] = 790
            ws["C53"] = 1.1
            ws["C54"] = 9.15
            ws["E53"] = 190
            ws["E54"] = 205
            ws["S16"] = '=2.66*($I$6+((E53+E54)/2))*(G53/G54)'
            ws["S19"] = 8.47
            ws["S22"] = 9.57
            wb.save(path)

            record = pmt_excel_dosyasi_oku(path)

        self.assertEqual(record["sondaj_no"], "SK-1")
        self.assertEqual(record["der"], "9.00")
        self.assertEqual(record["em"], "1410")
        self.assertEqual(record["pl"], "9.57")
        self.assertEqual(record["warnings"], [])

        veri = {"sondaj": [{"no": "SK-1", "pmt": [["9.00", "100", "1"]]}]}
        result = pmt_kayitlarini_veriye_aktar(veri, [record], update_existing=True)
        self.assertEqual(result["updated"], 1)
        self.assertEqual(veri["sondaj"][0]["pmt"], [["9.00", "1410", "9.57"]])


class YonetmelikMotorTestleri(unittest.TestCase):
    def test_resmi_yonetmelik_eki_sayfadan_bulunur(self):
        html = """
        <html><body>
            <a href="/duyuru">Duyuru</a>
            <a href="https://webdosya.csb.gov.tr/db/yapiisleri/icerikler/zemin-format.docx">
                Eki indirmek için tıklayınız
            </a>
        </body></html>
        """
        url = resmi_yonetmelik_baglanti_bul(
            html,
            "https://yapiisleri.csb.gov.tr/haber",
            expected_extensions=(".docx",),
            link_text_hints=("eki indir", "tıklayınız"),
        )
        self.assertEqual(url, "https://webdosya.csb.gov.tr/db/yapiisleri/icerikler/zemin-format.docx")

    def test_yerlesik_yonetmelik_kaynagi_eklenir(self):
        with tempfile.TemporaryDirectory() as tmp:
            result = yerlesik_yonetmelik_ekle("zemin_temel_etudu_2019", base_dir=tmp)

            self.assertTrue(result["embedded_fallback"])
            self.assertFalse(result["already_exists"])
            self.assertEqual(len(yonetmelikleri_listele(tmp)), 1)

            hits = yonetmelik_ara("eksik laboratuvar deneyi ve sondaj logu", base_dir=tmp)
            self.assertTrue(hits)
            self.assertTrue(hits[0]["doc_title"].endswith("(Yerleşik)"))

    def test_resmi_yonetmelik_indirme_hatasinda_yerlesik_kaynak_eklenir(self):
        original_download = yonetmelik_motoru._download_bytes
        try:
            yonetmelik_motoru._download_bytes = lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("offline"))
            with tempfile.TemporaryDirectory() as tmp:
                result = resmi_yonetmelik_indir_ve_ekle("zemin_temel_etudu_2019", base_dir=tmp)

                self.assertTrue(result["embedded_fallback"])
                docs = yonetmelikleri_listele(tmp)
                self.assertEqual(len(docs), 1)
                self.assertTrue(docs[0]["embedded_fallback"])
                self.assertIn("offline", docs[0]["download_error"])
        finally:
            yonetmelik_motoru._download_bytes = original_download

    def test_varsayilan_yonetmelik_programa_hazirlanir(self):
        with tempfile.TemporaryDirectory() as tmp:
            results = varsayilan_yonetmelikleri_hazirla(base_dir=tmp)

            self.assertTrue(results)
            self.assertEqual(len(yonetmelikleri_listele(tmp)), 1)
            self.assertTrue(yonetmelikleri_listele(tmp)[0]["embedded_fallback"])

    def test_yonetmelik_eklenir_ve_duzeltme_dayanagi_bulunur(self):
        metin = """
        ZEMIN VE TEMEL ETUDU UYGULAMA ESASLARI

        MADDE 1 - Sondajlar ve arazi calismalari
        Sondaj derinligi, yapi ve zemin kosullarina gore yeterli olacak sekilde belirlenir.

        MADDE 2 - Laboratuvar deneyleri
        Laboratuvar deneyleri zemin birimlerini temsil edecek araliklarda yapilir.
        Deney bulunmayan seviyeler icin ek laboratuvar deneyi istenebilir.
        """
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "yonetmelik.txt")
            atomic_write_text(src, metin)

            record = yonetmelik_ekle(src, title="Test Yönetmeliği", base_dir=tmp)
            self.assertEqual(record["title"], "Test Yönetmeliği")
            self.assertGreater(record["chunk_count"], 0)
            self.assertEqual(len(yonetmelikleri_listele(tmp)), 1)

            hits = yonetmelik_ara("25.50 seviyesinde ek laboratuvar deneyi isteniyor", base_dir=tmp)
            self.assertTrue(hits)
            self.assertEqual(hits[0]["doc_title"], "Test Yönetmeliği")
            self.assertIn("Laboratuvar", hits[0]["chunk_title"])

            result = duzeltme_yonetmelik_dayanaklari(
                "Lab deneyleri eksik. 25.50 ve 27.00 m seviyelerinde deney eklenmesi gerekmektedir.",
                base_dir=tmp,
            )
            self.assertTrue(result["items"])
            self.assertFalse(result["warnings"])

    def test_html_yonetmelik_eklenir(self):
        html = """
        <html>
            <head><style>.x { color: red; }</style><script>var a = 1;</script></head>
            <body>
                <h1>Türkiye Bina Deprem Yönetmeliği</h1>
                <p>Deprem etkisi altında temel zemini ve temellerin tasarımı için özel kurallar.</p>
            </body>
        </html>
        """
        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "tbdy.html")
            atomic_write_text(src, html)

            record = yonetmelik_ekle(src, title="TBDY HTML", base_dir=tmp)
            self.assertEqual(record["ext"], ".html")

            hits = yonetmelik_ara("temel zemini tasarımı", base_dir=tmp)
            self.assertTrue(hits)
            self.assertEqual(hits[0]["doc_title"], "TBDY HTML")


class SondajDerinlikHesabiTestleri(unittest.TestCase):
    def test_temel_genisligi_acik_degerden_bulunur(self):
        width, source = temel_genisligi_bul({"temel_genislik": "12,5", "plan": "20x30"})
        self.assertEqual(width, 12.5)
        self.assertEqual(source, "Temel genişliği B")

    def test_temel_genisligi_plandan_bulunur(self):
        width, source = temel_genisligi_bul({"plan": "18 x 32"})
        self.assertEqual(width, 18.0)
        self.assertIn("Plan", source)

    def test_sondaj_derinligi_onerisi_ve_eksik_sondaj_bulunur(self):
        veri = {
            "arazi": {"kategori": "Kategori 2"},
            "bina": {"der": "3", "temel_genislik": "8", "yukseklik": "18"},
            "sondaj": [{"no": "SK-1", "der": "15"}, {"no": "SK-2", "der": "18"}],
        }
        result = sondaj_derinligi_hesapla(veri)
        self.assertEqual(result["onerilen_sondaj_derinligi"], 18.0)
        self.assertEqual([item["sondaj"] for item in result["eksik_sondajlar"]], ["SK-1"])
        self.assertIn("18.00 m", sondaj_derinligi_ozet_metni(veri))

    def test_coklu_blokta_en_buyuk_derinlik_alinir(self):
        veri = {
            "arazi": {"kategori": "Kategori 3"},
            "bina": {
                "coklu_blok": True,
                "bloklar": [
                    {"blok_adi": "A", "der": "2", "temel_genislik": "5", "yukseklik": "12"},
                    {"blok_adi": "B", "der": "4", "temel_genislik": "12", "yukseklik": "30"},
                ],
            },
            "sondaj": [{"no": "SK-1", "der": "30"}],
        }
        result = sondaj_derinligi_hesapla(veri)
        self.assertEqual(result["onerilen_sondaj_derinligi"], 28.5)
        self.assertEqual(result["bloklar"][1]["belirleyici_kriter"], "2B temel genişliği")

    def test_gerilme_yuzde_on_derinligi_bulunur(self):
        params = {
            "q_net": "200",
            "b": "10",
            "l": "20",
            "temel_derinligi": "2",
            "bha": "18",
            "target_ratio": "0.10",
            "round_step": "0.50",
        }
        result = gerilme_yuzde_on_derinlik_hesapla(params)
        self.assertTrue(result["ok"])
        self.assertLessEqual(result["oran"], 0.1001)
        self.assertGreater(result["sondaj_derinligi_yuvarlatilmis"], result["temel_derinligi"])

        shallower_depth = max(result["temel_derinligi"], result["sondaj_derinligi"] - result["round_step"])
        delta = gerilme_artisi_boussinesq(result["q_net"], result["b"], result["l"], shallower_depth)
        sigma = result["bha"] * shallower_depth
        self.assertGreater(delta / sigma, 0.10)

    def test_yaklasik_yontem_bagimsiz_kok_ile_dogrulanir(self):
        params = {
            "q_net": "100",
            "b": "10",
            "l": "10",
            "temel_derinligi": "0",
            "bha": "10",
            "target_ratio": "0.10",
            "round_step": "1",
            "max_depth": "50",
            "hesap_yontemi": "yaklasik",
        }
        result = gerilme_yuzde_on_derinlik_hesapla(params)

        self.assertTrue(result["ok"])
        self.assertTrue(result["sayisal_dogrulama"]["ok"])
        self.assertAlmostEqual(
            result["yontem_sonuclari"]["yaklasik"]["surekli_kok_derinligi"],
            15.445115,
            places=5,
        )
        self.assertEqual(result["sondaj_derinligi_yuvarlatilmis"], 16.0)
        self.assertEqual(result["gerilme_birimi"], "kPa")

    def test_sondaj_derinligi_tutarsiz_birimleri_reddeder(self):
        result = gerilme_yuzde_on_derinlik_hesapla({
            "temel_taban_gerilmesi": "200",
            "temel_genisligi": "10",
            "temel_uzunlugu": "20",
            "temel_derinligi": "2",
            "yass": "1",
            "dogal_bha": "1.8",
            "doygun_bha": "20",
        })

        self.assertFalse(result["ok"])
        self.assertTrue(any("aynı birim" in item for item in result["errors"]))

    def test_temel_taban_gerilmesinden_qnet_hesaplanir(self):
        params = {
            "temel_taban_gerilmesi": "20",
            "temel_genisligi": "10",
            "temel_uzunlugu": "20",
            "temel_derinligi": "2",
            "yass": "1",
            "dogal_bha": "1.8",
            "doygun_bha": "2.0",
            "target_ratio": "0.10",
            "round_step": "0.50",
        }
        result = gerilme_yuzde_on_derinlik_hesapla(params)

        self.assertTrue(result["ok"])
        self.assertTrue(result["uses_yass_model"])
        self.assertAlmostEqual(efektif_gerilme_yass(2, 1, 1.8, 2.0), 2.8)
        self.assertAlmostEqual(result["sigma_vo_taban"], 2.8)
        self.assertAlmostEqual(result["q_net"], 17.2)
        self.assertLessEqual(result["oran"], 0.1001)

    def test_sondaj_derinligi_excel_ornegiyle_uyumlu_hesaplanir(self):
        params = {
            "temel_taban_gerilmesi": "10",
            "temel_genisligi": "24",
            "temel_uzunlugu": "83",
            "temel_derinligi": "4",
            "yass": "30",
            "dogal_bha": "1.9",
            "doygun_bha": "2.0",
            "target_ratio": "0.10",
            "round_step": "1",
            "max_depth": "50",
        }
        result = gerilme_yuzde_on_derinlik_hesapla(params)

        self.assertTrue(result["ok"])
        self.assertAlmostEqual(result["q_net"], 2.4)
        self.assertAlmostEqual(result["sigma_vo_taban"], 7.6)
        self.assertAlmostEqual(boussinesq_katsayisi(3, 10.375), 0.24650917381504633)
        self.assertAlmostEqual(westergaard_katsayisi(3, 10.375), 0.2116276499243993)
        self.assertEqual(result["yontem_sonuclari"]["boussinesq"]["sondaj_derinligi_yuvarlatilmis"], 11.0)
        self.assertEqual(result["yontem_sonuclari"]["westergaard"]["sondaj_derinligi_yuvarlatilmis"], 9.0)
        self.assertEqual(result["yontem_sonuclari"]["yaklasik"]["sondaj_derinligi_yuvarlatilmis"], 9.0)
        self.assertEqual(result["sondaj_derinligi_yuvarlatilmis"], 11.0)
        self.assertEqual(result["belirleyici_yontem"], "boussinesq")

    def test_sondaj_derinligi_foyu_docx_ve_pdf_olusturulur(self):
        import fitz

        veri = {
            "kunye": {
                "proje_adi": "Deneme Proje",
                "mah": "Merkez",
                "ilce": "Kepez",
                "il": "Çanakkale",
                "ada": "1109",
                "par": "1",
            },
            "sondaj_derinlik_hesabi": {
                "temel_taban_gerilmesi": "10",
                "temel_genisligi": "24",
                "temel_uzunlugu": "83",
                "temel_derinligi": "4",
                "yass": "30",
                "dogal_bha": "1.9",
                "doygun_bha": "2.0",
                "target_ratio": "0.10",
                "round_step": "1",
                "max_depth": "50",
            },
        }
        with tempfile.TemporaryDirectory() as tmp:
            docx_path = os.path.join(tmp, "sondaj_foyu.docx")
            pdf_path = os.path.join(tmp, "sondaj_foyu.pdf")

            docx_info = sondaj_derinligi_foyu_olustur(veri, docx_path)
            pdf_info = sondaj_derinligi_foyu_olustur(veri, pdf_path)

            self.assertTrue(os.path.exists(docx_path))
            self.assertTrue(os.path.getsize(docx_path) > 1000)
            self.assertEqual(docx_info["result"]["sondaj_derinligi_yuvarlatilmis"], 11.0)
            self.assertTrue(os.path.exists(pdf_path))
            self.assertTrue(os.path.getsize(pdf_path) > 1000)
            self.assertEqual(pdf_info["result"]["sondaj_derinligi_yuvarlatilmis"], 11.0)
            with fitz.open(pdf_path) as doc:
                self.assertGreaterEqual(doc.page_count, 2)

    def test_proje_kontrolu_elle_girilen_gerilme_hesabini_kullanir(self):
        veri = {
            "sondaj_derinlik_hesabi": {
                "temel_taban_gerilmesi": "20",
                "temel_genisligi": "10",
                "temel_uzunlugu": "20",
                "temel_derinligi": "2",
                "yass": "1",
                "dogal_bha": "1.8",
                "doygun_bha": "2.0",
                "target_ratio": "0.10",
                "round_step": "0.50",
            },
            "sondaj": [{"no": "SK-1", "der": "10"}],
        }
        result = sondaj_derinligi_kontrol_sonucu(veri)
        self.assertEqual(result["hesap_tipi"], "gerilme_10")
        self.assertTrue(result["eksik_sondajlar"])


class LitolojiDagilimTestleri(unittest.TestCase):
    def test_litoloji_birimi_son_ana_birime_gore_ayrilir(self):
        self.assertEqual(litoloji_dagilim_birimi("Çakıllı Siltli Kum"), "Çakıllı Siltli Kum")
        self.assertEqual(litoloji_dagilim_birimi("Kumlu Siltli Killi Çakıl"), "Kumlu Siltli Killi Çakıl")

    def test_litoloji_paragrafi_sondaj_litolojisinden_uretilir(self):
        sondajlar = [
            {"no": "SK-1", "litoloji": [["0", "2", "Çakıllı Siltli Kum"], ["2", "4", "Kil"]]},
        ]
        paragraphs = litoloji_dagilim_paragraflari(sondajlar)
        self.assertTrue(any("Çakıllı Siltli Kum" in item for item in paragraphs))
        self.assertTrue(any("Kil" in item for item in paragraphs))


class JeofizikRaporTestleri(unittest.TestCase):
    def test_vp_tablosu_ardisik_ayni_hizlari_sadelestirir(self):
        layers = [
            {"vp": "500", "vs": "200"},
            {"vp": "1600", "vs": "450"},
            {"vp": "1600.0", "vs": "460"},
            {"vp": "1800", "vs": "500"},
        ]
        sade = jeofizik_vp_layers_sadelestir(layers)
        self.assertEqual([item["vp"] for item in sade], ["500", "1600", "1800"])


class BinaBlokRaporTestleri(unittest.TestCase):
    def test_coklu_blok_aktifse_blok_tablolari_olusturulur(self):
        from docx import Document

        bina = {
            "coklu_blok": True,
            "bloklar": [
                {"blok_adi": "A Blok", "kul": "Konut", "kat": "5", "der": "3.0", "tem": "Radye", "gqe_min": "10", "gqe_ort": "11", "gqe_max": "12"},
                {"blok_adi": "B Blok", "kul": "Ticaret", "kat": "3", "der": "2.5", "tem": "Mütemadi", "gqe_min": "8", "gqe_ort": "9", "gqe_max": "10"},
            ],
        }
        self.assertEqual([item["blok_adi"] for item in bina_bloklari_rapor(bina)], ["A Blok", "B Blok"])
        doc = Document()
        tables = bina_bilgileri_tablolari_olustur(doc, bina)
        self.assertEqual(len(tables), 1)
        first_text = "\n".join(cell.text for row in tables[0].rows for cell in row.cells)
        self.assertIn("Bina Bilgileri", first_text)
        self.assertIn("A Blok", first_text)
        self.assertIn("Temel Tipi", first_text)
        self.assertIn("Radye", first_text)
        self.assertIn("Binadan Temel Zeminine Aktarılan En Yükler", first_text)
        self.assertIn("Ortalama", first_text)
        self.assertIn("B Blok", first_text)
        self.assertIn("Mütemadi", first_text)
        gqe_row = tables[0].rows[-2]
        self.assertEqual([gqe_row.cells[i].text for i in range(1, 7)], ["10", "11", "12", "8", "9", "10"])

    def test_tek_bina_temel_tipi_ve_yerel_zemin_rapora_yazilir(self):
        from docx import Document

        bina = {"kul": "Konut", "tem": "Radye", "ysinif": "ZC"}
        doc = Document()
        tables = bina_bilgileri_tablolari_olustur(doc, bina)
        text = "\n".join(cell.text for row in tables[0].rows for cell in row.cells)
        self.assertIn("Temel Tipi", text)
        self.assertIn("Radye", text)
        self.assertIn("Yerel Zemin Sınıfı", text)
        self.assertIn("ZC", text)

    def test_tablo_baslik_tekrari_kapatilabilir(self):
        from docx import Document
        from docx.oxml.ns import qn
        from raporlama_tablo import apply_report_table_style

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        apply_report_table_style(table, header_rows=2, repeat_headers=False)
        for row in table.rows[:2]:
            tr_pr = row._tr.trPr
            self.assertIsNone(tr_pr.find(qn("w:tblHeader")) if tr_pr is not None else None)

    def test_word_sayfa_sonu_paragrafi_page_break_uretiyor(self):
        from docx.oxml.ns import qn
        from raporlama import word_sayfa_sonu_paragrafi

        paragraph = word_sayfa_sonu_paragrafi()
        br = paragraph.find(".//" + qn("w:br"))
        self.assertIsNotNone(br)
        self.assertEqual(br.get(qn("w:type")), "page")

    def test_jeo_parametre_modulleri_tam_sayi_formatlanir(self):
        self.assertEqual(jeo_parametre_degeri_formatla("E", "274,16"), "274")
        self.assertEqual(jeo_parametre_degeri_formatla("G", "1454.88"), "1455")
        self.assertEqual(jeo_parametre_degeri_formatla("K", "29036,98"), "29037")
        self.assertEqual(jeo_parametre_degeri_formatla("vp", "1918.25"), "1918,25")


class LogCizimTestleri(unittest.TestCase):
    def test_profesyonel_log_minimal_veriyle_figure_uretir(self):
        sondaj = {
            "no": "SK-1",
            "der": "1.0",
            "litoloji": [["0", "1", "Kil"]],
            "spt": [],
            "pmt": [],
            "kaya": [],
            "numuneler": [],
        }
        figures = GeoEngine.ciz_profesyonel_log(sondaj, {"kunye": {}, "ayarlar": {}})
        self.assertEqual(len(figures), 1)
        self.assertTrue(figures[0].axes)

    def test_log_ust_bilgi_hizasi_ve_buyuk_pmt_fontu_korunur(self):
        from motor_log_kaynak import LOG_ZEMIN_PROFILI_START_RATIO

        proje_adi = "Uzun Proje Adı Gayrimenkul Tarım Sanayi ve Ticaret Anonim Şirketi"
        sondaj = {
            "no": "SK-4",
            "der": "15.0",
            "litoloji": [["0", "15", "Kil"]],
            "spt": [],
            "pmt": [["3.0", "1726", "13.5"]],
            "kaya": [],
            "numuneler": [],
        }
        figure = GeoEngine.ciz_profesyonel_log(
            sondaj,
            {
                "kunye": {
                    "sahibi": proje_adi,
                    "il": "Çanakkale",
                    "ada": "1476",
                    "par": "7",
                },
                "ayarlar": {
                    "sorumlu_muhendis": "GOKALP DOGAN",
                    "sondor_belge": "MURAT ERCELIK 3629",
                },
            },
        )[0]
        axis = figure.axes[0]
        texts = {artist.get_text(): artist for artist in axis.texts}

        self.assertEqual(texts["1726"].get_fontsize(), 6.1)
        self.assertGreater(texts["Proje Adı"].get_fontsize(), 7.0)
        self.assertGreater(texts["1476"].get_fontsize(), 7.0)
        self.assertLessEqual(texts[proje_adi].get_fontsize(), 8.5)
        self.assertEqual(
            texts["Gökalp DOĞAN"].get_fontsize(),
            texts["Çanakkale"].get_fontsize(),
        )
        self.assertEqual(
            texts["Murat ERÇELİK 3629"].get_fontsize(),
            texts["Çanakkale"].get_fontsize(),
        )
        self.assertEqual(texts["Gökalp DOĞAN"].get_fontweight(), "normal")
        self.assertEqual(texts["Murat ERÇELİK 3629"].get_fontweight(), "normal")

        expected_x = 0.035 + (0.965 - 0.035) * LOG_ZEMIN_PROFILI_START_RATIO
        aligned_body_lines = [
            line
            for line in axis.lines
            if len(line.get_xdata()) == 2
            and abs(float(line.get_xdata()[0]) - expected_x) < 1e-9
            and abs(float(line.get_xdata()[1]) - expected_x) < 1e-9
        ]
        self.assertTrue(aligned_body_lines)
        self.assertTrue(
            any(abs(patch.get_x() - expected_x) < 1e-9 for patch in axis.patches)
        )


class KarotTCRTestleri(unittest.TestCase):
    def test_log_ornek_derinligi_tek_haneli_formatlanir(self):
        self.assertEqual(log_ornek_derinligi_formatla("12.00-13.50"), "12.0-13.5")
        self.assertEqual(log_ornek_derinligi_formatla("12-13.5"), "12.0-13.5")
        self.assertEqual(log_ornek_derinligi_formatla("12.0-13.5"), "12.0-13.5")
        self.assertEqual(log_ornek_derinligi_formatla("12"), "12.0")

    def test_standart_karot_araliklari_30_metreye_kadar_olusturulur(self):
        intervals = standart_karot_araliklari()
        self.assertEqual(intervals[0], (1.5, 3.0))
        self.assertEqual(intervals[-1], (28.5, 30.0))
        self.assertEqual(len(intervals), 19)

    def test_derinlik_araligi_ve_tcr_hesabi(self):
        self.assertEqual(derinlik_araligi_coz("10.50-13.50"), (10.5, 13.5))
        self.assertEqual(derinlik_orta("10.50-13.50"), 12.0)

        result = tcr_hesapla(
            10.5,
            13.5,
            [((0, 20), (100, 20)), ((0, 50), (50, 50))],
            [(0, 0), (100, 0)],
            [(0, 100), (100, 100)],
        )
        self.assertAlmostEqual(result["karot"], 1.5, places=6)
        self.assertAlmostEqual(result["tcr"], 50.0, places=6)


class HaritaReferansTestleri(unittest.TestCase):
    def test_harita_cikti_yollari_projeler_arasi_paylasilmaz(self):
        path1 = yeni_harita_cikti_yolu("sondaj_lokasyon")
        path2 = yeni_harita_cikti_yolu("sondaj_lokasyon")
        self.assertNotEqual(path1, path2)
        self.assertTrue(path1.endswith(".jpg"))

    def test_eski_paylasimli_temp_harita_yolu_guvensiz_sayilir(self):
        legacy = os.path.join(tempfile.gettempdir(), "rapor_sondaj.jpg")
        self.assertTrue(eski_paylasimli_temp_harita_yolu_mu(legacy))
        self.assertFalse(eski_paylasimli_temp_harita_yolu_mu(r"C:\projeler\rapor_sondaj.jpg"))

    def test_kml_kose_noktalari_sirayla_okunur(self):
        kml = """<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <Placemark>
      <Polygon>
        <outerBoundaryIs>
          <LinearRing>
            <coordinates>
              26.0,40.0,0 27.0,40.0,0 27.0,41.0,0 26.0,41.0,0 26.0,40.0,0
            </coordinates>
          </LinearRing>
        </outerBoundaryIs>
      </Polygon>
    </Placemark>
  </Document>
</kml>"""
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "sinir.kml")
            with open(path, "w", encoding="utf-8") as f:
                f.write(kml)
            points = kml_koordinatlari_oku(path)
        self.assertEqual([p["label"] for p in points], ["KML-1", "KML-2", "KML-3", "KML-4"])
        self.assertEqual((points[0]["lat"], points[0]["lon"]), (40.0, 26.0))

    def test_affine_referanslardan_ara_nokta_hesaplanir(self):
        refs = [
            {"coord": {"lat": 40.0, "lon": 26.0}, "pixel": {"x": 100.0, "y": 200.0}},
            {"coord": {"lat": 40.0, "lon": 27.0}, "pixel": {"x": 200.0, "y": 200.0}},
            {"coord": {"lat": 41.0, "lon": 26.0}, "pixel": {"x": 100.0, "y": 100.0}},
        ]
        coeff = affine_from_refs(refs)
        x, y = coord_to_pixel(coeff, 40.5, 26.5)
        self.assertAlmostEqual(x, 150.0, places=6)
        self.assertAlmostEqual(y, 150.0, places=6)

    def test_affine_pixelden_koordinata_ters_hesaplanir(self):
        refs = [
            {"coord": {"lat": 40.0, "lon": 26.0}, "pixel": {"x": 100.0, "y": 200.0}},
            {"coord": {"lat": 40.0, "lon": 27.0}, "pixel": {"x": 200.0, "y": 200.0}},
            {"coord": {"lat": 41.0, "lon": 26.0}, "pixel": {"x": 100.0, "y": 100.0}},
        ]
        coeff = affine_from_refs(refs)
        lat, lon = pixel_to_coord(coeff, 150.0, 150.0)
        self.assertAlmostEqual(lat, 40.5, places=6)
        self.assertAlmostEqual(lon, 26.5, places=6)


class ProjeArsivTestleri(unittest.TestCase):
    def test_proje_merkezi_arazi_yoksa_sondaj_ortalamasi_alinir(self):
        veri = {
            "arazi": {"alan_y": "", "alan_x": ""},
            "sondaj": [
                {"y": "40.0", "x": "26.0"},
                {"y": "41.0", "x": "27.0"},
            ],
        }
        lat, lon = proje_merkez_koordinati(veri)
        self.assertAlmostEqual(lat, 40.5)
        self.assertAlmostEqual(lon, 26.5)

    def test_biten_is_kaydi_ve_kml_olusturulur(self):
        veri = {
            "kunye": {"sahibi": "Test Proje", "mah": "Merkez", "ilce": "Merkez", "il": "Çanakkale", "ada": "1", "par": "2"},
            "arazi": {"alan_y": "40.100000", "alan_x": "26.400000"},
            "sondaj": [],
            "proje_durumu": {"tamamlanma_tarihi": "2026-06-09T12:00:00"},
        }
        with tempfile.TemporaryDirectory() as tmp:
            index_path = os.path.join(tmp, "completed.json")
            project_path = os.path.join(tmp, "test.json")
            record = arsiv_kaydi_ekle(veri, project_path, index_path=index_path)
            records = arsiv_kayitlari_yukle(index_path=index_path)
            self.assertEqual(len(records), 1)
            self.assertEqual(record["name"], "Test Proje Merkez 1 2")

            kml_path = os.path.join(tmp, "biten.kml")
            info = biten_isler_kml_yaz(records, kml_path)
            self.assertEqual(info["written"], 1)
            with open(kml_path, "r", encoding="utf-8") as f:
                text = f.read()
            self.assertIn("<Placemark>", text)
            self.assertIn("26.40000000,40.10000000,0", text)

    def test_biten_is_kml_siniri_proje_kmlinden_bilgileri_kunyeden_alir(self):
        veri = {
            "kunye": {"sahibi": "Test Proje", "mah": "Merkez", "ilce": "Merkez", "il": "Çanakkale", "ada": "463", "par": "104"},
            "arazi": {"alan_y": "", "alan_x": ""},
            "sondaj": [],
            "proje_durumu": {"tamamlanma_tarihi": "2026-06-09T12:00:00"},
        }
        kml = """<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <Placemark>
      <name>KML içindeki ada parsel dikkate alınmayacak</name>
      <Polygon>
        <outerBoundaryIs>
          <LinearRing>
            <coordinates>
              26.4000,40.1000,0 26.4010,40.1000,0 26.4010,40.1010,0 26.4000,40.1010,0 26.4000,40.1000,0
            </coordinates>
          </LinearRing>
        </outerBoundaryIs>
      </Polygon>
    </Placemark>
  </Document>
</kml>"""
        with tempfile.TemporaryDirectory() as tmp:
            input_kml = os.path.join(tmp, "parsel.kml")
            with open(input_kml, "w", encoding="utf-8") as f:
                f.write(kml)
            self.assertEqual(len(kml_sinir_koordinatlari_oku(input_kml)), 4)

            index_path = os.path.join(tmp, "completed.json")
            project_path = os.path.join(tmp, "test.json")
            record = arsiv_kaydi_ekle(veri, project_path, index_path=index_path, kml_path=input_kml)
            self.assertEqual(len(record["boundary"]), 4)

            output_kml = os.path.join(tmp, "biten.kml")
            biten_isler_kml_yaz([record], output_kml)
            with open(output_kml, "r", encoding="utf-8") as f:
                text = f.read()
            self.assertIn("<Polygon>", text)
            self.assertIn("<MultiGeometry>", text)
            self.assertIn("Ada/Parsel: 463/104", text)
            self.assertNotIn("KML içindeki ada parsel", text)


class TaahhutnameTestleri(unittest.TestCase):
    def test_yapi_adresi_mahalle_ilce_il_bilgisinden_uretilir(self):
        veri = {"kunye": {"mah": "Namık Kemal", "ilce": "Merkez", "il": "Çanakkale"}, "ayarlar": {}}
        self.assertEqual(taahhutname_yapi_adresi(veri), "Namık Kemal Mah. Merkez Çanakkale")
        ctx = taahhutname_context(veri, "jeoloji")
        self.assertEqual(ctx["yapi_adresi"], "Namık Kemal Mah. Merkez Çanakkale")
        self.assertEqual(ctx["yapi_sahibi_adresi"], "Namık Kemal Mah. Merkez Çanakkale")

    def test_excel_sablonsuz_adres_ve_baski_alani_uretilir(self):
        veri = {
            "kunye": {
                "sahibi": "Cahit SARACOGLU ve Hiss.",
                "mah": "Namik Kemal",
                "ilce": "Merkez",
                "il": "Canakkale",
                "paf": "H16C14A1B",
                "ada": "463",
                "par": "104",
            },
            "ayarlar": {"taahhut_ilgili_idare": "Canakkale Belediyesi", "taahhut_tarih": "20.11.2025"},
        }
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "jeoloji.xlsx")
            taahhutname_olustur(veri, "jeoloji", path)
            wb = load_workbook(path, data_only=False)
            ws = wb["tahhütname"]
            self.assertEqual(ws["D14"].value, "Namik Kemal Mah. Merkez Canakkale")
            self.assertEqual(ws["D16"].value, "Namik Kemal Mah. Merkez Canakkale")
            self.assertIn("JEOLOJİ", ws["C5"].value)
            self.assertIn("müellifliğini üstlenmemde", ws["A20"].value)
            self.assertEqual(ws.print_area, "'tahhütname'!$A$1:$I$47")
            self.assertEqual(ws.page_setup.paperSize, 9)
            self.assertEqual(ws.page_setup.orientation, "portrait")
            self.assertIsNone(ws.page_setup.fitToWidth)
            self.assertIsNone(ws.page_setup.fitToHeight)
            self.assertAlmostEqual(ws.column_dimensions["A"].width, 12.44140625)
            self.assertAlmostEqual(ws.column_dimensions["B"].width, 6.33203125)
            self.assertAlmostEqual(ws.column_dimensions["I"].width, 8.6640625)
            self.assertEqual(ws.sheet_format.defaultRowHeight, 14.4)
            self.assertIsNone(ws.row_dimensions[20].height)
            self.assertEqual(ws["A4"].border.left.style, "thin")
            self.assertEqual(ws["A20"].border.top.style, "thin")
            self.assertIn("Gerçeğe aykırı beyanda", ws["A44"].value)
            self.assertEqual(ws["F30"].alignment.horizontal, "center")
            self.assertEqual(ws["F30"].border.right.style, "thin")
            self.assertIsNone(getattr(ws["D20"].border.left, "style", None))


class EklerTestleri(unittest.TestCase):
    def test_presiyometreye_gore_ek_secimi_yapilir(self):
        veri = {"kunye": {"sahibi": "Test Proje"}, "sondaj": [{"no": "SK-1", "pmt": []}], "ayarlar": {}}
        self.assertFalse(proje_presiyometre_var_mi(veri))
        label, path = uygun_ek_sablonu(veri)
        self.assertEqual(label, "Tutanaklı")
        self.assertTrue(os.path.exists(path))

        veri["sondaj"][0]["pmt"] = [["4.50", "100", "8"]]
        self.assertTrue(proje_presiyometre_var_mi(veri))
        label, path = uygun_ek_sablonu(veri)
        self.assertEqual(label, "Arazi Deneyli")
        self.assertTrue(os.path.exists(path))

    def test_ek_dosyasi_kopyalanir(self):
        veri = {"kunye": {"sahibi": "Test Proje"}, "sondaj": [{"no": "SK-1", "pmt": []}], "ayarlar": {}}
        with tempfile.TemporaryDirectory() as tmp:
            path = ek_olustur(veri, tmp)
            self.assertTrue(os.path.exists(path))
            self.assertTrue(os.path.basename(path).endswith("_Tutanakli.docx"))

    def test_ekler_rapora_ekler_etiketinden_gomulur(self):
        from docx import Document

        veri = {"kunye": {"sahibi": "Test Proje"}, "sondaj": [{"no": "SK-1", "pmt": []}], "ayarlar": {}}
        blocks = ek_bloklari_oku(uygun_ek_sablonu(veri)[1])
        self.assertEqual(blocks[-1]["no"], "10")
        doc = Document()
        doc.add_paragraph("Rapor metni")
        doc.add_paragraph("[EKLER]")
        info = ekleri_rapora_ekle(doc, veri)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertEqual(info["count"], 10)
        self.assertIn("EK-1", text)
        self.assertIn("EK-10", text)

    def test_ekler_pdf_coklu_dosya_ile_olusturulur(self):
        import fitz
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            image_path = os.path.join(tmp, "vaziyet.png")
            Image.new("RGB", (640, 360), "white").save(image_path)

            pdf_path = os.path.join(tmp, "log.pdf")
            source_pdf = fitz.open()
            source_pdf.new_page(width=842, height=595)
            source_pdf.save(pdf_path)
            source_pdf.close()

            output_path = os.path.join(tmp, "ekler.pdf")
            veri = {
                "kunye": {"sahibi": "Test Proje"},
                "sondaj": [{"no": "SK-1", "pmt": []}],
                "ayarlar": {},
                "ek_icerikleri": {EK_SET_NORMAL: {"1": [image_path, pdf_path]}},
            }
            self.assertEqual(len(ek_basliklari(veri, EK_SET_NORMAL)), 10)
            info = ekler_pdf_olustur(veri, output_path, set_key=EK_SET_NORMAL)
            self.assertTrue(os.path.exists(output_path))
            self.assertEqual(info["cover_count"], 10)
            self.assertEqual(info["attached_count"], 2)
            merged = fitz.open(output_path)
            try:
                self.assertGreaterEqual(merged.page_count, 12)
                self.assertGreater(merged[1].rect.width, merged[1].rect.height)
                self.assertGreater(merged[2].rect.width, merged[2].rect.height)
            finally:
                merged.close()


class TutanakTestleri(unittest.TestCase):
    def test_tutanaklar_program_verisinden_uretilir(self):
        from docx import Document
        from docx.oxml.ns import qn
        from PIL import Image

        veri = {
            "kunye": {"mah": "Namık Kemal", "ada": "463", "par": "105", "sahibi": "Test Proje"},
            "ayarlar": {"firma_adi": "UB ZEMİN MÜHENDİSLİK", "sondaj_turu": "Kaya", "delgi_capi": "89mm"},
            "sondaj": [
                {
                    "no": "SK-1",
                    "k": "100.50",
                    "bas_tar": "01.06.2026",
                    "bit_tar": "02.06.2026",
                    "der": "15.00",
                    "y": "40.100000",
                    "x": "26.400000",
                    "sondaj_turu": "Kaya",
                    "delgi_capi": "76mm",
                    "spt": [["1.50", "2", "3", "4", "7"]],
                    "pmt": [["4.50", "100", "8"]],
                    "kaya": [],
                    "numuneler": [["3.00-3.50", "UD"], ["6.00-6.50", "UD-1"], ["9.00", "Örselenmiş"]],
                    "yass_d1": "6.00",
                }
            ],
            "jeofizik": {
                "tarih": "03.06.2026",
                "ss_list": [{"ad": "SS-1", "coords": ["40.1", "26.4", "40.2", "26.5", "40.3", "26.6"]}],
            },
        }
        with tempfile.TemporaryDirectory() as tmp:
            image_path = os.path.join(tmp, "lokasyon.png")
            Image.new("RGB", (900, 500), "white").save(image_path)
            output_path = os.path.join(tmp, "tutanaklar.docx")
            info = tutanaklari_olustur(veri, output_path, image_path)
            self.assertTrue(os.path.exists(output_path))
            self.assertEqual(info["sondaj_count"], 1)
            self.assertEqual(info["jeofizik_count"], 1)
            doc = Document(output_path)
            self.assertEqual(doc.tables[0].rows[2].cells[2].text, "SK-1")
            self.assertEqual(doc.tables[0].rows[4].cells[2].text, "Kaya")
            self.assertEqual(doc.tables[0].rows[11].cells[2].text, "89mm")
            self.assertEqual(doc.tables[0].rows[12].cells[2].text, "10")
            self.assertEqual(doc.tables[0].rows[13].cells[2].text, "2")
            self.assertEqual(doc.tables[0].rows[14].cells[2].text, "1")
            self.assertEqual(doc.tables[0].rows[15].cells[2].text, "1")
            self.assertEqual(doc.tables[-1].rows[3].cells[1].text, "03.06.2026")
            self.assertEqual(len(doc.inline_shapes), 1)
            self.assertLessEqual(doc.sections[0].bottom_margin.cm, 0.8)
            self.assertIsNotNone(doc.tables[0].rows[0]._tr.trPr.find(qn("w:cantSplit")))
            self.assertTrue(doc.tables[0].rows[0].cells[0].paragraphs[0].paragraph_format.keep_with_next)
            self.assertFalse(bool(doc.tables[2].rows[0].cells[0].paragraphs[0].paragraph_format.keep_with_next))


if __name__ == "__main__":
    unittest.main()

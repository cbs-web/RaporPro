# Dosya: RaporPro/tests/test_jeoloji_kutuphanesi.py
from pathlib import Path
from collections import Counter
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import RGBColor
from PIL import Image

from jeoloji_docx import jeoloji_bolumunu_uygula
from jeoloji_kutuphanesi import (
    JeolojiKutuphane,
    docx_analiz_et,
    haversine_km,
    jeoloji_adaylarini_tara,
    kayitlari_filtrele,
    secili_jeoloji_kaydi,
    sha256_dosya,
)
from proje_sema import PROJE_SEMA_SURUMU, proje_verisini_migre_et, varsayilan_proje_verisi
from ui_jeoloji_kutuphanesi import (
    JeolojiKutuphanePenceresi,
    kayit_ada_parsel_metni,
    kayit_konum_metni,
    proje_secili_jeoloji_kaydi,
)
from ui_jeoloji_adaylari import aday_secilebilir_mi, secili_adaylari_filtrele


def _sample_doc(path, *, explicit_coord=True, header_text="KAYNAK HEADER", include_body_metadata=True):
    image_path = path.parent / f"{path.stem}.png"
    Image.new("RGB", (40, 40), (190, 40, 40)).save(image_path)
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = header_text
    doc.sections[0].footer.paragraphs[0].text = "KAYNAK FOOTER"
    source_base = doc.styles.add_style("ConflictBase", WD_STYLE_TYPE.PARAGRAPH)
    source_base.font.color.rgb = RGBColor(190, 40, 40)
    source_body = doc.styles.add_style("ConflictBody", WD_STYLE_TYPE.PARAGRAPH)
    source_body.base_style = source_base
    source_body.font.color.rgb = RGBColor(210, 30, 30)
    for tag_name in ("next", "link"):
        reference = OxmlElement(f"w:{tag_name}")
        reference.set(qn("w:val"), "ConflictBase")
        source_body._element.append(reference)
    doc.add_heading("1. GİRİŞ", level=1)
    doc.add_paragraph("Giriş metni")
    doc.add_heading("2. JEOLOJİ", level=1)
    if include_body_metadata:
        doc.add_paragraph("İl : Çanakkale")
        doc.add_paragraph("İlçesi : Merkez")
        doc.add_paragraph("Mahallesi : Test Mahallesi")
        doc.add_paragraph("Ada No : 12")
        doc.add_paragraph("Parsel No : 7")
    if explicit_coord:
        doc.add_paragraph("İnceleme alanı Enlem: 40.100000 Boylam: 26.400000 (WGS84) konumundadır.")
    doc.add_paragraph("Kaynak 2. JEOLOJİ metni.")
    bookmarked = doc.add_paragraph("Kaynak stilli metin.", style="ConflictBody")
    _add_bookmark(bookmarked, 0, "sourceMark")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Birim"
    table.cell(0, 1).text = "Koordinat"
    table.cell(1, 0).text = "Tmçk"
    table.cell(1, 1).text = "40.100000 / 26.400000"
    doc.add_picture(str(image_path), width=Inches(1))
    doc.add_heading("2.1 Bölgesel Jeoloji", level=2)
    doc.add_paragraph("Alt başlık metni.")
    doc.add_heading("3. ARAZİ ÇALIŞMALARI", level=1)
    doc.add_paragraph("HEDEF 3 BÖLÜMÜ KORUNMALI")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Enlem"
    table.cell(0, 1).text = "Boylam"
    table.cell(1, 0).text = "40.200000"
    table.cell(1, 1).text = "26.500000"
    table.cell(2, 0).text = "40.100000"
    table.cell(2, 1).text = "26.400000"
    doc.save(path)


def _add_header_table(path, text):
    doc = Document(path)
    table = doc.sections[0].header.add_table(rows=1, cols=1, width=Inches(6))
    table.cell(0, 0).text = text
    doc.save(path)


def _target_doc(path):
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "HEDEF HEADER"
    doc.sections[0].footer.paragraphs[0].text = "HEDEF FOOTER"
    target_image_path = path.parent / f"{path.stem}_target.png"
    Image.new("RGB", (40, 40), (40, 120, 40)).save(target_image_path)
    target_base = doc.styles.add_style("ConflictBase", WD_STYLE_TYPE.PARAGRAPH)
    target_base.font.color.rgb = RGBColor(20, 80, 20)
    target_body = doc.styles.add_style("ConflictBody", WD_STYLE_TYPE.PARAGRAPH)
    target_body.base_style = target_base
    target_body.font.color.rgb = RGBColor(20, 100, 20)
    doc.add_heading("1. GİRİŞ", level=1)
    doc.add_paragraph("Hedef giriş")
    doc.add_heading("2. JEOLOJİ", level=1)
    doc.add_paragraph("Eski jeoloji içeriği")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "Eski tablo"
    doc.add_heading("3. ARAZİ ÇALIŞMALARI", level=1)
    target_paragraph = doc.add_paragraph("HEDEF ARAZİ BÖLÜMÜ", style="ConflictBody")
    _add_bookmark(target_paragraph, 0, "targetMark")
    doc.add_picture(str(target_image_path), width=Inches(1))
    doc.save(path)


def _add_bookmark(paragraph, bookmark_id, name):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _make_floating_first_picture(path):
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)
    with zipfile.ZipFile(path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}
    root = ET.fromstring(members["word/document.xml"])
    inline = root.find(".//{%s}inline" % namespaces["wp"])
    assert inline is not None
    anchor = ET.Element(
        "{%s}anchor" % namespaces["wp"],
        {"distT": "0", "distB": "0", "distL": "0", "distR": "0", "simplePos": "0", "relativeHeight": "0", "behindDoc": "0", "locked": "0", "layoutInCell": "1", "allowOverlap": "1"},
    )
    simple_pos = ET.SubElement(anchor, "{%s}simplePos" % namespaces["wp"], {"x": "0", "y": "0"})
    position_h = ET.SubElement(anchor, "{%s}positionH" % namespaces["wp"], {"relativeFrom": "column"})
    ET.SubElement(position_h, "{%s}posOffset" % namespaces["wp"]).text = "0"
    position_v = ET.SubElement(anchor, "{%s}positionV" % namespaces["wp"], {"relativeFrom": "paragraph"})
    ET.SubElement(position_v, "{%s}posOffset" % namespaces["wp"]).text = "0"
    ET.SubElement(anchor, "{%s}wrapNone" % namespaces["wp"])
    for child in list(inline):
        if child.tag.rsplit("}", 1)[-1] in {"extent", "effectExtent", "docPr", "cNvGraphicFramePr", "graphic"}:
            anchor.append(child)
    inline.getparent().replace(inline, anchor) if hasattr(inline, "getparent") else None
    if not hasattr(inline, "getparent"):
        parent = next(parent for parent in root.iter() if inline in list(parent))
        parent.remove(inline)
        parent.append(anchor)
    members["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    temp_path = Path(str(path) + ".floating.tmp")
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, content in members.items():
            target_zip.writestr(name, content)
    temp_path.replace(path)


def _make_vml_image_reference(path, *, remove_drawing=False):
    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "v": "urn:schemas-microsoft-com:vml",
    }
    for prefix, uri in namespaces.items():
        ET.register_namespace(prefix, uri)
    with zipfile.ZipFile(path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}
    root = ET.fromstring(members["word/document.xml"])
    blip = root.find(".//{%s}blip" % namespaces["a"])
    assert blip is not None
    relationship_id = blip.get("{%s}embed" % namespaces["r"])
    assert relationship_id
    if remove_drawing:
        drawing = root.find(".//{%s}drawing" % namespaces["w"])
        assert drawing is not None
        drawing.getparent().remove(drawing) if hasattr(drawing, "getparent") else None
        if not hasattr(drawing, "getparent"):
            parent = next(parent for parent in root.iter() if drawing in list(parent))
            parent.remove(drawing)

    body = root.find(".//{%s}body" % namespaces["w"])
    assert body is not None
    vml_paragraph = ET.Element("{%s}p" % namespaces["w"])
    run = ET.SubElement(vml_paragraph, "{%s}r" % namespaces["w"])
    pict = ET.SubElement(run, "{%s}pict" % namespaces["w"])
    shape = ET.SubElement(pict, "{%s}shape" % namespaces["v"], {"style": "width:1pt;height:1pt"})
    ET.SubElement(shape, "{%s}imagedata" % namespaces["v"], {"{%s}id" % namespaces["r"]: relationship_id})
    heading = next(
        child
        for child in list(body)
        if child.tag.rsplit("}", 1)[-1] == "p"
        and "3. ARAZİ" in "".join(node.text or "" for node in child.iter() if node.tag.rsplit("}", 1)[-1] == "t")
    )
    body.insert(list(body).index(heading), vml_paragraph)
    members["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    temp_path = Path(str(path) + ".vml.tmp")
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, content in members.items():
            target_zip.writestr(name, content)
    temp_path.replace(path)


def _mark_macro_enabled(path):
    normal = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    macro = b"application/vnd.ms-word.document.macroEnabled.main+xml"
    with zipfile.ZipFile(path, "r") as source_zip:
        members = {name: source_zip.read(name) for name in source_zip.namelist()}
    members["[Content_Types].xml"] = members["[Content_Types].xml"].replace(normal, macro)
    temp_path = Path(str(path) + ".docm.tmp")
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for name, content in members.items():
            target_zip.writestr(name, content)
    temp_path.replace(path)


def test_docx_analizi_sinir_metadata_ve_wgs84_onceligi():
    with tempfile.TemporaryDirectory() as tmp:
        source = Path(tmp) / "eski.docx"
        _sample_doc(source)
        analysis = docx_analiz_et(source)
        boundaries = analysis["boundaries"]
        assert boundaries["start_heading"].startswith("2.")
        assert boundaries["end_heading"].startswith("3.")
        assert analysis["ilce"] == "Merkez"
        assert analysis["mahalle"] == "Test Mahallesi"
        assert analysis["ada"] == "12"
        assert analysis["parsel"] == "7"
        assert analysis["lat"] == 40.1
        assert analysis["lon"] == 26.4
        assert analysis["coordinate_source"] == "rapor_wgs84"
        assert analysis["table_count"] == 1
        assert analysis["image_count"] == 1
        assert any("eski projeye" in warning for warning in analysis["warnings"])


def test_header_imar_bilgileri_ve_header_tablosu_metadata_doldurur():
    with tempfile.TemporaryDirectory() as tmp:
        source = Path(tmp) / "headerli.docx"
        _sample_doc(source, explicit_coord=False, include_body_metadata=False, header_text="Proje Adı: Header testi")
        _add_header_table(
            source,
            "İmar Bilgileri: Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 32N-4D Pafta, 0 Ada, 1147 Parsel",
        )
        analysis = docx_analiz_et(source)
        assert analysis["il"] == "Çanakkale"
        assert analysis["ilce"] == "Merkez"
        assert analysis["mahalle"] == "Işıklar"
        assert analysis["ada"] == "0"
        assert analysis["parsel"] == "1147"
        assert analysis["pafta"] == "32N-4D"
        assert analysis["metadata"]["field_sources"] == {
            "il": "imar_bilgileri",
            "ilce": "imar_bilgileri",
            "mahalle": "imar_bilgileri",
            "pafta": "imar_bilgileri",
            "ada": "imar_bilgileri",
            "parsel": "imar_bilgileri",
        }
        assert any("İmar Bilgileri" in text for text in analysis["metadata"]["header_texts"])
        assert analysis["lat"] is not None
        assert analysis["lon"] is not None


def test_baslik_noktalama_bosluk_farklari_bolum_sinirini_bozmaz():
    with tempfile.TemporaryDirectory() as tmp:
        source = Path(tmp) / "baslik_farki.docx"
        doc = Document()
        doc.add_heading("1.GİRİŞ", level=1)
        doc.add_paragraph("Giriş")
        doc.add_heading("2.JEOLOJİ", level=1)
        doc.add_paragraph("Jeoloji içeriği")
        doc.add_heading("2 . 1 BÖLGESEL JEOLOJİ", level=2)
        doc.add_paragraph("Alt başlık")
        doc.add_heading("3 ) ARAZİ ÇALIŞMALARI", level=1)
        doc.add_paragraph("Arazi içeriği")
        doc.save(source)

        analysis = docx_analiz_et(source)
        assert analysis["boundaries"]["start_heading"] == "2.JEOLOJİ"
        assert analysis["boundaries"]["end_heading"] == "3 ) ARAZİ ÇALIŞMALARI"
        assert "Arazi içeriği" not in analysis["section_text"]


def test_klasor_taramasi_aday_dondurur_ve_kutuphaneye_yazmaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        reports = root / "raporlar"
        nested = reports / "alt"
        nested.mkdir(parents=True)
        second_folder = root / "ikinci_klasor"
        second_folder.mkdir()
        valid = nested / "uygun.docx"
        macro_valid = second_folder / "makrolu.docm"
        rejected = reports / "jeoloji_yok.docx"
        _sample_doc(valid, include_body_metadata=False, header_text="İmar Bilgileri: Bursa İli, Nilüfer İlçesi, Özlüce Mahallesi, H22C Pafta, 0 Ada, 18 Parsel")
        _sample_doc(macro_valid, include_body_metadata=False, header_text="İmar Bilgileri: Ankara İli, Çankaya İlçesi, Kızılay Mahallesi, I29 Pafta, 15 Ada, 4 Parsel")
        _mark_macro_enabled(macro_valid)
        invalid_doc = Document()
        invalid_doc.add_heading("1. GİRİŞ", level=1)
        invalid_doc.add_heading("3. ARAZİ ÇALIŞMALARI", level=1)
        invalid_doc.save(rejected)
        store = JeolojiKutuphane(base_dir=root / "appdata")

        candidates = jeoloji_adaylarini_tara([reports, second_folder], recursive=True)

        assert store.count() == 0
        assert not list(store.cache_dir.glob("*.docx"))
        assert {item["original_filename"] for item in candidates} == {"uygun.docx", "makrolu.docm", "jeoloji_yok.docx"}
        suitable = next(item for item in candidates if item["original_filename"] == "uygun.docx")
        unsuitable = next(item for item in candidates if item["original_filename"] == "jeoloji_yok.docx")
        assert aday_secilebilir_mi(suitable)
        assert suitable["il"] == "Bursa"
        assert suitable["pafta"] == "H22C"
        assert suitable["ada"] == "0"
        assert not aday_secilebilir_mi(unsuitable)
        assert "bulunamadı" in unsuitable["status"]
        macro_candidate = next(item for item in candidates if item["original_filename"] == "makrolu.docm")
        assert aday_secilebilir_mi(macro_candidate)
        assert macro_candidate["ilce"] == "Çankaya"

        added = store.import_candidate(macro_candidate)
        duplicate = store.import_candidate(macro_candidate)
        assert added["duplicate"] is False
        assert duplicate["duplicate"] is True
        assert store.count() == 1
        section_doc = Document(added["record"]["cache_path"])
        section_text = "\n".join(paragraph.text for paragraph in section_doc.paragraphs)
        assert "2. JEOLOJİ" in section_text
        assert "1. GİRİŞ" not in section_text
        assert "3. ARAZİ ÇALIŞMALARI" not in section_text


def test_ui_aday_secim_filtresi_reddedileni_ve_ayni_hashi_tekillestirir():
    candidates = [
        {"source_hash": "abc", "source_path": "a.docx", "eligible": True, "selected": True},
        {"source_hash": "abc", "source_path": "b.docx", "eligible": True, "selected": True},
        {"source_hash": "def", "source_path": "c.docx", "eligible": False, "selected": True},
        {"source_hash": "ghi", "source_path": "d.docx", "eligible": True, "selected": False},
    ]
    selected = secili_adaylari_filtrele(candidates)
    assert [item["source_path"] for item in selected] == ["a.docx"]


def test_dosya_adi_ada_parsel_fallbacki_tarihleri_yanlis_eslestirmez():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        expected = {
            "463_105": ("463", "105"),
            "140-2": ("140", "2"),
            "1345_21": ("1345", "21"),
            "1432_3": ("1432", "3"),
            "213_2": ("213", "2"),
            "117_20": ("117", "20"),
        }
        for stem, pair in expected.items():
            source = root / f"rapor_{stem}.docx"
            _sample_doc(source, explicit_coord=False, include_body_metadata=False, header_text="Proje Adı: Adsız")
            analysis = docx_analiz_et(source)
            assert (analysis["ada"], analysis["parsel"]) == pair
            assert analysis["metadata"]["field_sources"] == {"ada": "dosya_adi", "parsel": "dosya_adi"}

        dated = root / "rapor_2024-05.docx"
        _sample_doc(dated, explicit_coord=False, include_body_metadata=False, header_text="Proje Adı: Tarihli")
        analysis = docx_analiz_et(dated)
        assert analysis["ada"] == ""
        assert analysis["parsel"] == ""


def test_sondaj_koordinat_ortalama_ve_filtre_yakinlik():
    with tempfile.TemporaryDirectory() as tmp:
        source = Path(tmp) / "ortalama.docx"
        _sample_doc(source, explicit_coord=False)
        analysis = docx_analiz_et(source)
        assert round(analysis["lat"], 6) == 40.15
        assert round(analysis["lon"], 6) == 26.45
        assert analysis["coordinate_source"] == "sondaj_ortalamasi"
        records = [
            {"id": 1, "ilce": "Merkez", "mahalle": "Test", "metadata": {"formasyonlar": ["Tmçk"]}, "lat": 40.1, "lon": 26.4},
            {"id": 2, "ilce": "Başka", "mahalle": "Uzak", "metadata": {"formasyonlar": ["Qal"]}, "lat": 41.0, "lon": 27.0},
        ]
        filtered = kayitlari_filtrele(records, ilce="merkez", formasyon="tmck", center=(40.1, 26.4), yaricap_km=5)
        assert [record["id"] for record in filtered] == [1]
        assert haversine_km(40.1, 26.4, 40.1, 26.4) == 0


def test_kutuphane_hash_dedup_cache_ve_kaynak_tasininca_secim():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "rapor.docx"
        _sample_doc(source)
        store = JeolojiKutuphane(base_dir=root / "appdata")
        first = store.import_docx(source)
        second = store.import_docx(source)
        assert first["duplicate"] is False
        assert second["duplicate"] is True
        assert first["record"]["id"] == second["record"]["id"]
        cache_path = Path(first["record"]["cache_path"])
        assert cache_path.is_file()
        assert first["record"]["metadata"]["cache_kind"] == "jeoloji_section"
        assert first["record"]["metadata"]["cache_hash"] == sha256_dosya(cache_path)
        assert first["record"]["source_hash"] != sha256_dosya(cache_path)
        cache_doc = Document(cache_path)
        cache_text = "\n".join(paragraph.text for paragraph in cache_doc.paragraphs)
        assert "2. JEOLOJİ" in cache_text
        assert "1. GİRİŞ" not in cache_text
        assert "3. ARAZİ ÇALIŞMALARI" not in cache_text
        assert "HEDEF 3 BÖLÜMÜ KORUNMALI" not in cache_text
        assert len(cache_doc.tables) == 1
        assert len(cache_doc.inline_shapes) == 1
        assert cache_doc.sections[0].header.paragraphs[0].text == ""
        assert cache_doc.sections[0].footer.paragraphs[0].text == ""
        with zipfile.ZipFile(cache_path) as package:
            names = package.namelist()
            assert not any(name.startswith("word/header") for name in names)
            assert not any(name.startswith("word/footer") for name in names)
            assert any(name.startswith("word/media/") for name in names)
        source.unlink()
        selection = {
            "jeoloji_kutuphanesi": {
                "selected_source_id": first["record"]["id"],
                "selected_source_hash": first["record"]["source_hash"],
                "selected_snapshot": {
                    "cache_name": cache_path.name,
                    "cache_kind": first["record"]["metadata"]["cache_kind"],
                    "cache_hash": first["record"]["metadata"]["cache_hash"],
                },
            }
        }
        selected = secili_jeoloji_kaydi(selection, store=store)
        assert selected is not None
        assert Path(selected["cache_path"]).is_file()


def test_duplicate_hash_yeniden_analizle_metadata_ve_sayimlari_gunceller():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "463_105.docx"
        _sample_doc(source, explicit_coord=False, include_body_metadata=False, header_text="Proje Adı: Yeniden indeks")
        _add_header_table(
            source,
            "İmar Bilgileri: Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 32N-4D Pafta, 0 Ada, 1147 Parsel",
        )
        store = JeolojiKutuphane(base_dir=root / "data")
        first = store.import_docx(source)
        record = first["record"]
        with store._connection() as connection:
            connection.execute(
                """
                UPDATE geology_sources
                SET metadata_json = '{}', il = '', ilce = '', mahalle = '', ada = '', parsel = '',
                    paragraph_count = 0, table_count = 0, image_count = 0, warning_json = '[]'
                WHERE id = ?
                """,
                (record["id"],),
            )
        second = store.import_docx(source)
        updated = second["record"]
        assert second["duplicate"] is True
        assert second["analysis"] is not None
        assert store.count() == 1
        assert updated["id"] == record["id"]
        assert updated["il"] == "Çanakkale"
        assert updated["ilce"] == "Merkez"
        assert updated["mahalle"] == "Işıklar"
        assert updated["pafta"] == "32N-4D"
        assert updated["ada"] == "0"
        assert updated["parsel"] == "1147"
        assert updated["paragraph_count"] > 0
        assert updated["table_count"] == 1
        assert updated["image_count"] == 1
        assert updated["warnings"]
        assert updated["cache_path"] == record["cache_path"]


def test_duplicate_reindex_manuel_metadata_ve_koordinati_korur():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "463_105.docx"
        _sample_doc(source, include_body_metadata=False, header_text="Proje Adı: Manuel koruma")
        _add_header_table(
            source,
            "İmar Bilgileri: Çanakkale İli, Merkez İlçesi, Işıklar Köyü, 32N-4D Pafta, 0 Ada, 1147 Parsel",
        )
        store = JeolojiKutuphane(base_dir=root / "data")
        first = store.import_docx(source)["record"]
        manual = store.update_record(
            first["id"],
            {
                "il": "Bursa",
                "ilce": "Nilüfer",
                "mahalle": "Özlüce",
                "pafta": "H22-C",
                "ada": "99",
                "parsel": "12",
                "lat": "40.987654",
                "lon": "29.123456",
            },
        )
        assert all(manual["metadata"]["field_sources"][key] == "manuel" for key in ("il", "ilce", "mahalle", "pafta", "ada", "parsel"))
        assert manual["metadata"]["field_sources"]["lat"] == "manuel"
        assert manual["metadata"]["field_sources"]["lon"] == "manuel"
        assert manual["coordinate_source"] == "manuel"

        with store._connection() as connection:
            connection.execute(
                "UPDATE geology_sources SET paragraph_count = 0, table_count = 0, image_count = 0, warning_json = '[]' WHERE id = ?",
                (first["id"],),
            )

        result = store.import_docx(source)
        updated = result["record"]
        assert result["duplicate"] is True
        assert updated["il"] == "Bursa"
        assert updated["ilce"] == "Nilüfer"
        assert updated["mahalle"] == "Özlüce"
        assert updated["pafta"] == "H22-C"
        assert updated["ada"] == "99"
        assert updated["parsel"] == "12"
        assert updated["lat"] == 40.987654
        assert updated["lon"] == 29.123456
        assert updated["coordinate_source"] == "manuel"
        assert updated["paragraph_count"] > 0
        assert updated["table_count"] == 1
        assert updated["image_count"] == 1
        assert updated["warnings"]
        assert all(updated["metadata"]["field_sources"][key] == "manuel" for key in ("il", "ilce", "mahalle", "pafta", "ada", "parsel"))


def test_tam_oo_xml_aktarimi_header_footer_ve_3_bolumu_korur():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "source.docx"
        target = root / "target.docx"
        output = root / "output.docx"
        _sample_doc(source)
        _target_doc(target)
        target_doc = Document(target)
        source_record = JeolojiKutuphane(base_dir=root / "data").import_docx(source)["record"]
        result = jeoloji_bolumunu_uygula(target_doc, source_record["cache_path"])
        target_doc.save(output)
        reopened = Document(output)
        headings = [paragraph.text for paragraph in reopened.paragraphs if paragraph.style.name.startswith("Heading")]
        assert sum("2. JEOLOJİ" in heading for heading in headings) == 1
        assert sum("3. ARAZİ ÇALIŞMALARI" in heading for heading in headings) == 1
        assert "Kaynak 2. JEOLOJİ metni." in "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        assert "HEDEF ARAZİ BÖLÜMÜ" in "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        assert "Eski jeoloji içeriği" not in "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        assert len(reopened.tables) == 1
        assert len(reopened.inline_shapes) == 2
        source_paragraph = next(item for item in reopened.paragraphs if item.text == "Kaynak stilli metin.")
        target_paragraph = next(item for item in reopened.paragraphs if item.text == "HEDEF ARAZİ BÖLÜMÜ")
        assert source_paragraph.style.style_id != "ConflictBody"
        assert target_paragraph.style.style_id == "ConflictBody"
        assert source_paragraph.style.font.color.rgb == RGBColor(210, 30, 30)
        assert target_paragraph.style.font.color.rgb == RGBColor(20, 100, 20)
        assert reopened.sections[0].header.paragraphs[0].text == "HEDEF HEADER"
        assert reopened.sections[0].footer.paragraphs[0].text == "HEDEF FOOTER"
        assert "KAYNAK HEADER" not in "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        assert result["inserted_elements"] > 0
        with zipfile.ZipFile(output) as package:
            media = [name for name in package.namelist() if name.startswith("word/media/")]
            assert media
            assert any("image" in name for name in media)
            document_xml = ET.fromstring(package.read("word/document.xml"))
            docpr_ids = [element.get("id") for element in document_xml.iter() if element.tag.rsplit("}", 1)[-1] == "docPr"]
            assert len(docpr_ids) == 2
            assert len(docpr_ids) == len(set(docpr_ids))
            bookmark_ids = [
                element.get(qn("w:id"))
                for element in document_xml.iter()
                if element.tag.rsplit("}", 1)[-1] in {"bookmarkStart", "bookmarkEnd"}
            ]
            assert Counter(bookmark_ids) == Counter({"0": 2, "1": 2})
            styles_xml = ET.fromstring(package.read("word/styles.xml"))
            style_elements = {
                element.get(qn("w:styleId")): element
                for element in styles_xml.iter()
                if element.tag.rsplit("}", 1)[-1] == "style"
            }
            source_style_id = source_paragraph.style.style_id
            source_base_id = next(
                element.get(qn("w:val"))
                for element in style_elements[source_style_id]
                if element.tag.rsplit("}", 1)[-1] == "basedOn"
            )
            assert source_base_id != "ConflictBase"
            for tag_name in ("next", "link"):
                reference = next(
                    element
                    for element in style_elements[source_style_id]
                    if element.tag.rsplit("}", 1)[-1] == tag_name
                )
                assert reference.get(qn("w:val")) == source_base_id


def test_floating_drawing_blip_aktarilir_ve_hedef_word_acilir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "floating.docx"
        target = root / "target.docx"
        output = root / "output.docx"
        _sample_doc(source)
        _make_floating_first_picture(source)
        _target_doc(target)
        target_doc = Document(target)
        record = JeolojiKutuphane(base_dir=root / "data").import_docx(source)["record"]
        analysis = docx_analiz_et(source)
        assert analysis["image_count"] == 1
        with zipfile.ZipFile(record["cache_path"]) as package:
            cache_xml = package.read("word/document.xml").decode("utf-8")
            assert "<wp:anchor" in cache_xml or ":anchor" in cache_xml
        jeoloji_bolumunu_uygula(target_doc, record["cache_path"])
        target_doc.save(output)
        assert Document(output).paragraphs
        with zipfile.ZipFile(output) as package:
            xml = package.read("word/document.xml").decode("utf-8")
            assert "<wp:anchor" in xml or ":anchor" in xml


def test_vml_imagedata_sayilir_ve_drawingml_ile_cift_sayilmaz():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        duplicate = root / "duplicate.docx"
        vml_only = root / "vml_only.docx"
        _sample_doc(duplicate)
        _sample_doc(vml_only)
        _make_vml_image_reference(duplicate)
        _make_vml_image_reference(vml_only, remove_drawing=True)
        assert docx_analiz_et(duplicate)["image_count"] == 1
        assert docx_analiz_et(vml_only)["image_count"] == 1
        record = JeolojiKutuphane(base_dir=root / "data").import_docx(vml_only)["record"]
        assert Document(record["cache_path"]).paragraphs
        with zipfile.ZipFile(record["cache_path"]) as package:
            document_xml = package.read("word/document.xml").decode("utf-8")
            rels_xml = package.read("word/_rels/document.xml.rels").decode("utf-8")
            assert "imagedata" in document_xml
            assert "media/" in rels_xml


def test_ui_acilisinda_proje_secimi_ve_filtre_debounce_kapanista_iptal_edilir():
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "selected.docx"
        _sample_doc(source)
        store = JeolojiKutuphane(base_dir=root / "data")
        record = store.import_docx(source)["record"]
        by_id = proje_secili_jeoloji_kaydi(
            {"jeoloji_kutuphanesi": {"selected_source_id": record["id"]}},
            store,
        )
        by_hash = proje_secili_jeoloji_kaydi(
            {"jeoloji_kutuphanesi": {"selected_source_hash": record["source_hash"]}},
            store,
        )
        assert by_id["id"] == record["id"]
        assert by_hash["id"] == record["id"]

    class FakeWindow:
        def __init__(self):
            self.cancelled = []

        def after(self, _delay, _callback):
            return "after#1"

        def after_cancel(self, callback_id):
            self.cancelled.append(callback_id)

        def winfo_exists(self):
            return True

    class FakeOwner:
        def __init__(self):
            self.closed = False

        def pencere_kapat(self, _window, callback=None):
            self.closed = True
            if callback:
                callback()

    window = object.__new__(JeolojiKutuphanePenceresi)
    window.win = FakeWindow()
    window.owner = FakeOwner()
    window._closing = False
    window._filter_after_id = None
    window._map_pick_mode = False
    window._filtre_keyrelease()
    assert window._filter_after_id == "after#1"
    window.kapat()
    assert window._closing is True
    assert window.win.cancelled == ["after#1"]
    assert window.owner.closed is True


def test_eski_proje_defaultlari_ve_ui_pure_metni_geriye_uyumludur():
    legacy = {"schema_version": PROJE_SEMA_SURUMU, "kunye": {"ilce": "Merkez"}}
    migrated, info = proje_verisini_migre_et(legacy)
    assert info.degisti
    assert migrated["jeoloji_kutuphanesi"] == varsayilan_proje_verisi()["jeoloji_kutuphanesi"]
    assert "selected_source_hash" in migrated["jeoloji_kutuphanesi"]
    record = {"il": "Çanakkale", "ilce": "Merkez", "mahalle": "Test", "ada": "12", "parsel": "7"}
    assert kayit_konum_metni(record) == "Çanakkale / Merkez / Test"
    assert kayit_ada_parsel_metni(record) == "12 / 7"

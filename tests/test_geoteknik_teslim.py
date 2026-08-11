import os
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document
from openpyxl import load_workbook

from geoteknik_teslim import (
    etkin_jeofizik_serimleri,
    ham_veri_kaynaklari,
    ham_verileri_kopyala,
    jeofizik_parametre_paketi_olustur,
    sondaj_veri_paketi_olustur,
)
from ui_cikti import CiktiMerkeziMixin, cikti_profili_secimleri


def _ornek_veri():
    layers = [
        {"h": "3", "vp": "500", "vs": "220", "rho": "1.8", "nu": "0.34", "E": "1234.6", "G": "456.4", "K": "1800.7"},
        {"h": "", "vp": "1600", "vs": "700", "rho": "2.0", "nu": "0.38", "E": "4567.8", "G": "1654.2", "K": "9012.3"},
    ]
    return {
        "kunye": {"sahibi": "Örnek Proje", "il": "Çanakkale", "ilce": "Merkez", "mah": "Arslanca"},
        "ayarlar": {"sondaj_turu": "Zemin", "delgi_capi": "76mm"},
        "dosyalar": {},
        "sondaj": [
            {
                "no": "SK-1", "der": "15", "k": "12.5", "y": "40.10", "x": "",
                "bas_tar": "01.01.2026", "bit_tar": "02.01.2026",
                "yass_d1": "3.0", "yass_t1": "02.01.2026", "yass_d2": "", "yass_t2": "",
                "litoloji": [["0", "0.5", "Bitkisel Toprak"], ["0.5", "15", "Kumlu Kil"]],
                "spt": [["1.50", "3", "4", "5", "9"], ["3.00", "20", "50/10", "-", "R"]],
                "pmt": [["6.0", "120", "8"]],
                "kaya": [["12.0-13.5", "80", "60", "40"]],
                "numuneler": [["1.50-1.95", "DS1"]],
                "spt_kaynaklari": [{"derinlik": "3.00", "kaynak": "DSCF0019.JPG"}],
            }
        ],
        "jeofizik": {
            "ss_list": [
                {"ad": f"Serim {index}", "coords": ["" for _ in range(6)], "layers": [dict(item) for item in layers]}
                for index in range(1, 5)
            ]
        },
        "jeofizik_sheet": {"rows": []},
    }


class GeoteknikTeslimTests(unittest.TestCase):
    def test_sondaj_excel_ve_word_paketi_olusturur(self):
        with tempfile.TemporaryDirectory() as tmp:
            paths = sondaj_veri_paketi_olustur(_ornek_veri(), tmp)
            self.assertEqual(len(paths), 2)
            self.assertTrue(all(os.path.isfile(path) for path in paths))

            workbook = load_workbook(os.path.join(tmp, "Sondaj_Verileri.xlsx"), data_only=False)
            self.assertEqual(
                workbook.sheetnames,
                ["Sondaj Özeti", "Litoloji", "SPT", "Presiyometre", "Kaya ve Karot", "Numuneler"],
            )
            self.assertEqual(workbook["Sondaj Özeti"]["E2"].value, "Girilmedi")
            self.assertEqual(workbook["SPT"]["D3"].value, "50/10")
            self.assertEqual(workbook["SPT"]["F3"].value, "R")
            self.assertEqual(workbook["SPT"]["G3"].value, "Refü")
            self.assertEqual(workbook["SPT"]["H3"].value, "DSCF0019.JPG")
            workbook.close()

            document = Document(os.path.join(tmp, "Sondaj_Ozet_Tablolari.docx"))
            self.assertEqual(len(document.tables), 3)
            document_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Sondaj Özet Tabloları", document_text)
            self.assertIn("Yeraltı Suyu Gözlemleri", document_text)

    def test_jeofizik_paketi_ucer_serimlik_ayri_tablolar_olusturur(self):
        with tempfile.TemporaryDirectory() as tmp:
            paths = jeofizik_parametre_paketi_olustur(_ornek_veri(), tmp)
            self.assertEqual(len(paths), 2)

            workbook = load_workbook(os.path.join(tmp, "Jeofizik_Parametreleri.xlsx"), data_only=False)
            sheet = workbook["Jeofizik Parametreleri"]
            self.assertEqual(sheet["A4"].value, "Vp (m/s)")
            self.assertEqual(sheet["B4"].value, 500)
            self.assertEqual(sheet["B8"].value, 1235)
            self.assertEqual(sheet["D1"].value, "Serim 2")
            self.assertEqual(sheet["F1"].value, "Serim 3")
            self.assertTrue(str(sheet["B4"].font.color.rgb).endswith("C00000"))
            self.assertEqual(len(sheet.row_breaks.brk), 1)
            workbook.close()

            document = Document(os.path.join(tmp, "Jeofizik_Parametreleri.docx"))
            self.assertEqual(len(document.tables), 2)
            self.assertIn("Serim 1 - Serim 2 - Serim 3", "\n".join(p.text for p in document.paragraphs))
            self.assertIn("Serim 4", "\n".join(p.text for p in document.paragraphs))

    def test_ham_kaynaklari_bulur_ve_tekrarsiz_kopyalar(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            project = root / "proje.json"
            project.write_text("{}", encoding="utf-8")
            evrak = root / "EVRAKLAR"
            evrak.mkdir()
            (evrak / "imar.pdf").write_bytes(b"ornek evrak")
            lab = root / "lab.xlsx"
            lab.write_bytes(b"ornek lab")
            veri = _ornek_veri()
            veri["dosyalar"]["lab_excel_path"] = str(lab)
            sources = ham_veri_kaynaklari(veri, str(project))
            self.assertTrue(any(category == "Evraklar" for category, _path in sources))
            self.assertTrue(any(category == "Laboratuvar" for category, _path in sources))

            copied, warnings = ham_verileri_kopyala(sources, root / "teslim")
            self.assertFalse(warnings)
            self.assertEqual(len(copied), 2)
            self.assertTrue(all(os.path.isfile(path) for path in copied))

    def test_geoteknik_profili_yalniz_ilgili_gruplari_secer(self):
        selections = cikti_profili_secimleri("Geoteknik Mühendisine Teslim")
        self.assertTrue(selections["logs"])
        self.assertTrue(selections["sondaj_data"])
        self.assertTrue(selections["geophysics_data"])
        self.assertTrue(selections["source_files"])
        self.assertFalse(selections["report"])
        self.assertFalse(selections["taahhutnameler"])

    def test_manuel_jeofizik_serimleri_kullanilir(self):
        self.assertEqual(len(etkin_jeofizik_serimleri(_ornek_veri())), 4)

    def test_cikti_merkezi_yeni_paketleri_numarali_klasorlere_yazar(self):
        with tempfile.TemporaryDirectory() as tmp:
            mixin = CiktiMerkeziMixin.__new__(CiktiMerkeziMixin)
            mixin.veri = _ornek_veri()
            mixin.cikti_merkezi_progress = mock.Mock()
            mixin.cikti_merkezi_bitti = mock.Mock()
            mixin.set_status = mock.Mock()
            mixin.cikti_merkezi_threaded(
                {
                    "base_folder": tmp,
                    "format": "jpg",
                    "dpi": 150,
                    "veri_snapshot": _ornek_veri(),
                    "map_sources": [],
                    "report_image_sources": [],
                    "sondaj_data": True,
                    "geophysics_data": True,
                    "source_files": False,
                },
                {},
                {"cancelled": False},
            )
            self.assertTrue(os.path.isfile(os.path.join(tmp, "07_Sondaj_Verileri", "Sondaj_Verileri.xlsx")))
            self.assertTrue(os.path.isfile(os.path.join(tmp, "08_Jeofizik_Parametreleri", "Jeofizik_Parametreleri.docx")))
            self.assertTrue(os.path.isfile(os.path.join(tmp, "RaporPro_Cikti_Kalite.json")))
            mixin.cikti_merkezi_bitti.assert_called_once()


if __name__ == "__main__":
    unittest.main()

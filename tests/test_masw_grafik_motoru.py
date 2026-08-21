# Dosya: RaporPro/tests/test_masw_grafik_motoru.py

import hashlib
import json
import os
import tempfile
import unittest

from docx import Document
from docx.shared import Cm
from PIL import Image

from masw_grafik_motoru import (
    MASW_GRAFIK_ETIKETI,
    masw_grafik_kaydi_oku,
    masw_grafiklerini_rapora_ekle,
    masw_word_yollari_normalize,
)
from proje_paketi import proje_paketi_olustur
from proje_sema import PROJE_SEMA_SURUMU, proje_verisini_migre_et


def _sha256(path):
    digest = hashlib.sha256()
    with open(path, "rb") as stream:
        for chunk in iter(lambda: stream.read(64 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _resim_olustur(path, color):
    Image.new("RGB", (420, 260), color).save(path)


def _masw_word_olustur(path, eski_resim, hiz_resmi):
    document = Document()
    document.add_paragraph("MASW + MAM DEĞERLENDİRME")
    document.add_picture(eski_resim, width=Cm(10))
    document.add_paragraph("DİSPERSİYON EĞRİSİ")
    document.add_picture(hiz_resmi, width=Cm(14))
    document.save(path)


def _masw_word_ayni_paragrafta_olustur(path, dispersiyon_resmi, hiz_resmi):
    document = Document()
    document.add_paragraph("DİSPERSİYON EĞRİSİ")
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(dispersiyon_resmi, width=Cm(10))
    paragraph.add_run().add_picture(hiz_resmi, width=Cm(10))
    document.save(path)


def _masw_word_ardisik_cizimlerde_olustur(path, ilk_resim, son_resim):
    document = Document()
    document.add_paragraph("DİSPERSİYON EĞRİSİ")
    document.add_paragraph().add_run().add_picture(ilk_resim, width=Cm(10))
    document.add_paragraph()
    document.add_paragraph().add_run().add_picture(son_resim, width=Cm(10))
    document.add_paragraph("Sonraki bölüm")
    document.save(path)


def _inline_resim_hashleri(document):
    hashes = []
    for shape in document.inline_shapes:
        for blip in shape._inline.xpath(".//a:blip"):
            rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = document.part.related_parts.get(rid)
            if part is not None:
                hashes.append(hashlib.sha256(bytes(part.blob)).hexdigest())
    return hashes


class MaswGrafikMotoruTestleri(unittest.TestCase):
    def test_dispersiyon_basligindan_sonraki_grafigi_secer(self):
        with tempfile.TemporaryDirectory() as tmp:
            before = os.path.join(tmp, "once.png")
            expected = os.path.join(tmp, "hiz.png")
            source = os.path.join(tmp, "MASW-1.docx")
            _resim_olustur(before, "red")
            _resim_olustur(expected, "green")
            _masw_word_olustur(source, before, expected)

            record = masw_grafik_kaydi_oku(source)

            self.assertEqual(record.paragraf_no, 3)
            self.assertEqual(len(record.iliskili_resimler), 1)
            self.assertEqual(
                hashlib.sha256(record.iliskili_resimler[0][1]).hexdigest(),
                _sha256(expected),
            )

    def test_ayni_paragrafta_son_cizimi_secer(self):
        with tempfile.TemporaryDirectory() as tmp:
            dispersiyon = os.path.join(tmp, "dispersiyon.png")
            expected = os.path.join(tmp, "hiz.png")
            source = os.path.join(tmp, "MASW-ayni-paragraf.docx")
            _resim_olustur(dispersiyon, "red")
            _resim_olustur(expected, "green")
            _masw_word_ayni_paragrafta_olustur(source, dispersiyon, expected)

            record = masw_grafik_kaydi_oku(source)

            self.assertEqual(len(record.iliskili_resimler), 1)
            self.assertEqual(
                hashlib.sha256(record.iliskili_resimler[0][1]).hexdigest(),
                _sha256(expected),
            )

    def test_ardisik_cizim_paragraflarinda_son_cizimi_secer(self):
        with tempfile.TemporaryDirectory() as tmp:
            dispersiyon = os.path.join(tmp, "dispersiyon.png")
            expected = os.path.join(tmp, "hiz.png")
            source = os.path.join(tmp, "MASW-ardisik.docx")
            _resim_olustur(dispersiyon, "red")
            _resim_olustur(expected, "green")
            _masw_word_ardisik_cizimlerde_olustur(source, dispersiyon, expected)

            record = masw_grafik_kaydi_oku(source)

            self.assertEqual(len(record.iliskili_resimler), 1)
            self.assertEqual(
                hashlib.sha256(record.iliskili_resimler[0][1]).hexdigest(),
                _sha256(expected),
            )

    def test_basliksiz_fallback_son_cizimi_secer(self):
        with tempfile.TemporaryDirectory() as tmp:
            first = os.path.join(tmp, "ilk.png")
            expected = os.path.join(tmp, "son.png")
            source = os.path.join(tmp, "MASW-basliksiz.docx")
            _resim_olustur(first, "red")
            _resim_olustur(expected, "green")
            document = Document()
            document.add_picture(first, width=Cm(10))
            document.add_paragraph("Ara bölüm")
            document.add_picture(expected, width=Cm(10))
            document.save(source)

            record = masw_grafik_kaydi_oku(source)

            self.assertEqual(
                hashlib.sha256(record.iliskili_resimler[0][1]).hexdigest(),
                _sha256(expected),
            )

    def test_sabit_ornekleri_iki_kaynak_grafikle_degistirir(self):
        with tempfile.TemporaryDirectory() as tmp:
            old = os.path.join(tmp, "old.png")
            graph1 = os.path.join(tmp, "graph1.png")
            graph2 = os.path.join(tmp, "graph2.png")
            source1 = os.path.join(tmp, "MASW-1.docx")
            source2 = os.path.join(tmp, "MASW-2.docx")
            output = os.path.join(tmp, "rapor.docx")
            for path, color in ((old, "gray"), (graph1, "red"), (graph2, "blue")):
                _resim_olustur(path, color)
            _masw_word_olustur(source1, old, graph1)
            _masw_word_olustur(source2, old, graph2)
            source_hashes = (_sha256(source1), _sha256(source2))

            report = Document()
            report.add_paragraph("Şekil 8 MASW Ölçüm Grafikleri")
            report.add_picture(old, width=Cm(10))
            report.add_picture(old, width=Cm(10))
            report.add_paragraph("Tablo 8. İnceleme Alanında ölçülmüş Vs hızları")

            result = masw_grafiklerini_rapora_ekle(report, [source2, source1])
            report.save(output)
            reopened = Document(output)

            self.assertEqual(result.eklenen, 2)
            self.assertEqual(result.kaldirilan_sabit_gorsel, 2)
            self.assertFalse(result.hatalar)
            self.assertEqual(len(reopened.inline_shapes), 2)
            expected_hashes = sorted(
                hashlib.sha256(masw_grafik_kaydi_oku(path).iliskili_resimler[0][1]).hexdigest()
                for path in (source1, source2)
            )
            self.assertEqual(sorted(_inline_resim_hashleri(reopened)), expected_hashes)
            self.assertEqual((_sha256(source1), _sha256(source2)), source_hashes)
            self.assertIn(
                "Şekil 8 MASW Ölçüm Grafikleri",
                [paragraph.text for paragraph in reopened.paragraphs],
            )

    def test_kaynak_yokken_sabit_ornek_ve_basligi_kaldirir(self):
        with tempfile.TemporaryDirectory() as tmp:
            old = os.path.join(tmp, "old.png")
            _resim_olustur(old, "gray")
            report = Document()
            report.add_paragraph("Şekil 8 MASW Ölçüm Grafikleri")
            report.add_picture(old, width=Cm(10))
            report.add_paragraph("Tablo 8. İnceleme Alanında ölçülmüş Vs hızları")

            result = masw_grafiklerini_rapora_ekle(report, [])

            self.assertEqual(result.eklenen, 0)
            self.assertEqual(len(report.inline_shapes), 0)
            self.assertFalse(
                any("MASW Ölçüm Grafikleri" in paragraph.text for paragraph in report.paragraphs)
            )
            self.assertTrue(
                any("Tablo 8" in paragraph.text for paragraph in report.paragraphs)
            )

    def test_etiketli_sablona_grafik_ekler(self):
        with tempfile.TemporaryDirectory() as tmp:
            old = os.path.join(tmp, "old.png")
            graph = os.path.join(tmp, "graph.png")
            source = os.path.join(tmp, "MASW-1.docx")
            _resim_olustur(old, "gray")
            _resim_olustur(graph, "blue")
            _masw_word_olustur(source, old, graph)
            report = Document()
            report.add_paragraph(MASW_GRAFIK_ETIKETI)

            result = masw_grafiklerini_rapora_ekle(report, [source])

            self.assertEqual(result.eklenen, 1)
            self.assertEqual(len(report.inline_shapes), 1)
            self.assertFalse(any(MASW_GRAFIK_ETIKETI in p.text for p in report.paragraphs))

    def test_yollar_tekrarsiz_ve_dogal_sirada_tutulur(self):
        result = masw_word_yollari_normalize(
            [r"C:\x\MASW-10.docx", r"C:\x\MASW-2.docx", r"C:\x\MASW-2.docx"]
        )
        self.assertEqual([os.path.basename(path) for path in result], ["MASW-2.docx", "MASW-10.docx"])

    def test_v4_projesine_masw_kaynak_listesi_eklenir(self):
        migrated, info = proje_verisini_migre_et({"schema_version": 4})
        self.assertEqual(info.yeni_surum, PROJE_SEMA_SURUMU)
        self.assertEqual(migrated["dosyalar"]["masw_word_paths"], [])

    def test_tasinabilir_paket_masw_word_listesini_kopyalar(self):
        with tempfile.TemporaryDirectory() as tmp:
            source_project = os.path.join(tmp, "proje.json")
            source_word = os.path.join(tmp, "MASW-1.docx")
            output_dir = os.path.join(tmp, "paketler")
            os.makedirs(output_dir)
            with open(source_project, "w", encoding="utf-8") as stream:
                stream.write("{}")
            Document().save(source_word)
            veri = {
                "schema_version": PROJE_SEMA_SURUMU,
                "kunye": {"sahibi": "MASW Projesi"},
                "dosyalar": {"masw_word_paths": [source_word]},
            }

            info = proje_paketi_olustur(veri, source_project, output_dir)
            with open(info["project_path"], "r", encoding="utf-8") as stream:
                packaged = json.load(stream)

            packaged_path = packaged["dosyalar"]["masw_word_paths"][0]
            self.assertFalse(os.path.isabs(packaged_path))
            self.assertTrue(os.path.isfile(os.path.join(info["folder"], packaged_path)))


if __name__ == "__main__":
    unittest.main()
